from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
import sys
import os

# Add current directory to path for imports
current_dir = os.path.dirname(os.path.abspath(__file__))
if current_dir not in sys.path:
    sys.path.append(current_dir)

from document_templates import DocumentTemplate
from common_data_handler import CommonDataHandler  # Importa il gestore dati

class BaseVerbaleTemplate(DocumentTemplate):
    """Base template semplificato per tutti i verbali di assemblea"""
    
    def get_form_fields(self, extracted_data: dict) -> dict:
        """Crea i campi del form per i dati comuni usando il CommonDataHandler.
        
        Questo metodo centralizza la creazione dei widget di Streamlit per tutti
        i dati standard (società, assemblea, partecipanti, organi di controllo).
        I template specifici possono estendere questo metodo per aggiungere
        i propri campi univoci.
        """
        form_data = {}

        # Usa i metodi del CommonDataHandler per popolare il form
        form_data.update(CommonDataHandler.extract_and_populate_company_data(extracted_data))
        form_data.update(CommonDataHandler.extract_and_populate_assembly_data(extracted_data))
        form_data.update(CommonDataHandler.extract_and_populate_participants_data(extracted_data))
        
        # Aggiunge i campi per l'organo di controllo e il revisore
        if form_data.get("collegio_sindacale"):
            form_data.update(CommonDataHandler.extract_and_populate_organo_controllo(extracted_data))
        
        if form_data.get("revisore"):
            form_data.update(CommonDataHandler.extract_and_populate_revisore(extracted_data))

        return form_data
    
    def _setup_document_styles(self, doc):
        """Configura gli stili base del documento"""
        styles = doc.styles
        
        # Stile per intestazione società
        try:
            if 'TitoloSocieta' not in [s.name for s in styles]:
                title_style = styles.add_style('TitoloSocieta', WD_STYLE_TYPE.PARAGRAPH)
            else:
                title_style = styles['TitoloSocieta']
            
            title_style.font.name = 'Times New Roman'
            title_style.font.size = Pt(14)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(6)
        except Exception as e:
            print(f"Warning: Could not create/configure TitoloSocieta style: {e}")
        
        # Stile per titolo verbale
        try:
            if 'TitoloVerbale' not in [s.name for s in styles]:
                verbale_title_style = styles.add_style('TitoloVerbale', WD_STYLE_TYPE.PARAGRAPH)
            else:
                verbale_title_style = styles['TitoloVerbale']
            
            verbale_title_style.font.name = 'Times New Roman'
            verbale_title_style.font.size = Pt(16)
            verbale_title_style.font.bold = True
            verbale_title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            verbale_title_style.paragraph_format.space_before = Pt(18)
            verbale_title_style.paragraph_format.space_after = Pt(18)
        except Exception as e:
            print(f"Warning: Could not create/configure TitoloVerbale style: {e}")
        
        # Stile per testo normale
        try:
            if 'BodyText' not in [s.name for s in styles]:
                body_style = styles.add_style('BodyText', WD_STYLE_TYPE.PARAGRAPH)
            else:
                body_style = styles['BodyText']
            
            body_style.font.name = 'Times New Roman'
            body_style.font.size = Pt(12)
            body_style.paragraph_format.space_after = Pt(6)
            body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except Exception as e:
            print(f"Warning: Could not create/configure BodyText style: {e}")

    def _add_company_header(self, doc, data):
        """Aggiunge l'intestazione della società"""
        try:
            p = doc.add_paragraph(style='TitoloSocieta')
        except KeyError:
            # Fallback: usa stile normale con formattazione diretta
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run(data.get('denominazione', '[DENOMINAZIONE]'))
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        try:
            p = doc.add_paragraph(style='TitoloSocieta')
        except KeyError:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run(f"Sede in {data.get('sede_legale', '[SEDE]')}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        try:
            p = doc.add_paragraph(style='TitoloSocieta')
        except KeyError:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run(f"Capitale sociale Euro {data.get('capitale_sociale', '[CAPITALE]')} i.v.")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        try:
            p = doc.add_paragraph(style='TitoloSocieta')
        except KeyError:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run(f"Codice fiscale: {data.get('codice_fiscale', '[CF]')}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    def _add_verbale_title(self, doc, data):
        """Aggiunge il titolo del verbale"""
        doc.add_paragraph()
        
        try:
            p = doc.add_paragraph(style='TitoloVerbale')
        except KeyError:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run("Verbale di assemblea dei soci")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        
        data_str = data.get('data_assemblea').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[DATA]'
        
        try:
            p = doc.add_paragraph(style='TitoloVerbale')
        except KeyError:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run(f"del {data_str}")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)

    def _add_signatures(self, doc, data):
        """Aggiunge le firme"""
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Presidente
        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()
        
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("_____________________")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()
        
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(data.get('presidente', '[PRESIDENTE]'))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()
        
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("Il Presidente")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Segretario
        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()
        
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("_____________________")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()
        
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(data.get('segretario', '[SEGRETARIO]'))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()
        
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("Il Segretario")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    def _add_professional_signatures(self, doc, data):
        """Aggiunge le firme con uno stile più "professionale".
        
        Al momento reimpiega la logica di `_add_signatures`, ma è
        definita come metodo separato così i template che la richiedono
        non sollevano AttributeError e, in futuro, potremo personalizzare
        facilmente il layout delle firme professionali in un solo punto.
        """
        # Riutilizza la logica standard per ora
        self._add_signatures(doc, data)

    def _setup_professional_document(self, data):
        """Configura un documento con formattazione professionale"""
        doc = Document()
        self._setup_document_styles(doc)
        return doc

    # Nuovo helper comune per i template che necessitano di tabelle formattate
    def _create_table_with_style(self, doc: Document, rows: int, cols: int, style_name: str = "Table Grid"):
        """Crea una tabella con il numero di righe/colonne indicato ed applica uno stile predefinito.

        Questo metodo centralizza la creazione delle tabelle per evitare duplicazioni nei singoli
        template e garantire una formattazione omogenea. Di default viene applicato lo stile
        "Table Grid" (bordo completo). Se lo stile richiesto non è presente, viene utilizzato lo
        stile di default della libreria python-docx senza generare eccezioni.
        """
        try:
            table = doc.add_table(rows=rows, cols=cols)
            if style_name and style_name in [s.name for s in doc.styles]:
                table.style = style_name
            elif "Table Grid" in [s.name for s in doc.styles]:
                table.style = "Table Grid"
            # In caso lo stile non esista semplicemente si procede senza impostarlo
        except Exception:
            # Fallback di emergenza: crea comunque la tabella senza stile specifico
            table = doc.add_table(rows=rows, cols=cols)
        return table

    # ------------------------------------------------------------------
    # Helper comune: Collegio Sindacale / Revisore tra i partecipanti
    # ------------------------------------------------------------------
    def _add_organi_controllo_paragraphs(self, doc: Document, data: dict):
        """Aggiunge, in fondo all'elenco dei partecipanti, l'organo di controllo
        (Collegio Sindacale / Sindaco Unico) e l'eventuale revisore, se presenti.
        È sufficiente chiamare questo metodo all'interno di _add_participants_section
        dei vari template.
        """
        if data.get("include_collegio_sindacale"):
            tipo_oc = data.get("tipo_organo_controllo", "Collegio Sindacale")
            sindaci = [s for s in data.get("sindaci", []) if s.get("presente")]
            if tipo_oc == "Collegio Sindacale" and sindaci:
                p = doc.add_paragraph("per il Collegio Sindacale:")
                for s in sindaci:
                    doc.add_paragraph(f"- {s.get('nome', '[NOME]')} {s.get('carica', '')}")
            elif tipo_oc == "Sindaco Unico" and sindaci:
                doc.add_paragraph(f"il Sindaco Unico {sindaci[0].get('nome', '[NOME]')}")

        if data.get("include_revisore"):
            nome_rev = data.get("nome_revisore", "[NOME REVISORE]")
            doc.add_paragraph(f"il revisore contabile Dott. {nome_rev}")

    # ------------------------------------------------------------------
    # Helper: formattazione riga socio (testo) riutilizzabile nei template
    # ------------------------------------------------------------------
    def _format_socio_line(self,
                           nome: str,
                           quota_euro: str,
                           quota_perc: str,
                           tipo_partecipazione: str = "Diretto",
                           delegato: str = "",
                           tipo_soggetto: str = "Persona Fisica",
                           rappresentante_legale: str = "") -> str:
        """Restituisce una stringa descrittiva di un socio, usata nelle anteprime.

        Parametri in ingresso già puliti (stringhe).  Nessuna formattazione numerica.
        """
        if tipo_partecipazione == "Delegato" and delegato:
            if tipo_soggetto == "Società":
                line = f"il Sig. {delegato} delegato della società {nome} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
            else:
                line = f"il Sig. {delegato} delegato del socio {nome} recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
        else:
            if tipo_soggetto == "Società":
                if rappresentante_legale:
                    line = f"la società {nome} nella persona del legale rappresentante {rappresentante_legale} recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                else:
                    line = f"la società {nome} recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
            else:
                line = f"il Sig. {nome} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
        return line

    def add_paragraph_with_font(self, doc, text: str,
                                 size: Pt = Pt(12),
                                 font_name: str = "Times New Roman",
                                 bold: bool = False,
                                 underline: bool = False,
                                 alignment=None,
                                 space_before: Pt | None = None,
                                 space_after: Pt | None = None,
                                 left_indent: Inches | None = None):
        """Crea un paragrafo con formattazione semplificata.

        Questo helper evita ripetizione di codice nei template specifici
        e garantisce uniformità dello stile.
        """
        try:
            p = doc.add_paragraph(style='BodyText')
        except Exception:
            p = doc.add_paragraph()

        run = p.add_run(str(text))
        run.font.name = font_name
        run.font.size = size
        run.bold = bold
        run.underline = underline

        if alignment is not None:
            p.alignment = alignment
        if space_before is not None:
            p.paragraph_format.space_before = space_before
        if space_after is not None:
            p.paragraph_format.space_after = space_after
        if left_indent is not None:
            p.paragraph_format.left_indent = left_indent

        return p

    def _add_signature_table(self, doc: Document, data: dict):
        """Crea una tabella di firme standard (Presidente | Segretario) e la restituisce.
        Questo helper è usato da template che desiderano un layout con linee
        e nomi centrati in fondo al documento.
        """
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"

        # Prima riga: linee firma
        for cell in table.rows[0].cells:
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("_____________________")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        # Seconda riga: nomi + ruoli
        presidente = data.get("presidente", "[PRESIDENTE]")
        segretario = data.get("segretario", "[SEGRETARIO]")

        labels = [(presidente, "Il Presidente"), (segretario, "Il Segretario")]
        for idx, cell in enumerate(table.rows[1].cells):
            nome, ruolo = labels[idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{nome}\n{ruolo}")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        return table
