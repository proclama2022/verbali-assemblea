"""
Template per Verbale di Assemblea - Revoca dell'Amministratore Unico e nomina di nuovo Organo Amministrativo
"""

import sys
import os

# Aggiungi il path della cartella src
current_dir = os.path.dirname(os.path.abspath(__file__))
src_path = os.path.join(os.path.dirname(current_dir), 'src')
if src_path not in sys.path:
    sys.path.append(src_path)

from document_templates import DocumentTemplate, DocumentTemplateFactory
from common_data_handler import CommonDataHandler
from base_verbale_template import BaseVerbaleTemplate
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from datetime import date
import streamlit as st

class VerbaleRevocaNominaTemplate(BaseVerbaleTemplate):
    """Template per Verbale di Assemblea - Revoca dell'Amministratore Unico e nomina di nuovo Organo Amministrativo"""
    
    def get_template_name(self) -> str:
        return "Revoca dell'Amministratore Unico e nomina di nuovo Organo Amministrativo"
    
    def get_required_fields(self) -> list:
        return [
            "denominazione", "sede_legale", "capitale_sociale", "codice_fiscale",
            "data_assemblea", "ora_assemblea", "presidente", "segretario", 
            "soci", "amministratore_revocato", "nuovo_amministratore"
        ]
    
    def get_form_fields(self, extracted_data: dict) -> dict:
        """Genera i campi del form per Streamlit."""
        # Chiama il metodo della classe base per ottenere i campi comuni
        form_data = super().get_form_fields(extracted_data)

        st.subheader("📋 Configurazioni Specifiche Revoca e Nomina")
        
        # Sezione Revoca
        st.subheader("Revoca dell'Organo Amministrativo")
        form_data['revoca_organo'] = st.checkbox("Includi sezione di revoca", value=True, key="revoca_organo_revnom")
        if form_data.get('revoca_organo'):
            form_data['motivo_revoca'] = st.selectbox(
                "Motivo della revoca",
                ["Dimissioni", "Scadenza mandato", "Giusta causa", "Altro"],
                key="motivo_revoca_revnom"
            )
            form_data['amministratori_revocati'] = st.text_area(
                "Nomi degli amministratori revocati (separati da virgola)",
                key="amministratori_revocati_revnom"
            )

        # Sezione Nomina
        st.subheader("Nomina del Nuovo Organo Amministrativo")
        form_data['nomina_organo'] = st.checkbox("Includi sezione di nomina", value=True, key="nomina_organo_revnom")
        if form_data.get('nomina_organo'):
            form_data["tipo_amministrazione_nuovo"] = st.selectbox(
                "Tipo di nuova amministrazione",
                ["Amministratore Unico", "Consiglio di Amministrazione"],
                key="tipo_amministrazione_nuovo_revnom"
            )
            # Aggiungere qui i campi per inserire i dati dei nuovi amministratori
            # Esempio: st.data_editor per una lista di nuovi amministratori
            st.info("La sezione per la nomina dei nuovi amministratori è da implementare.")

        return form_data
    
    def show_preview(self, form_data: dict):
        """Mostra l'anteprima del documento"""
        st.markdown("---")
        st.subheader("👁️ Anteprima Documento")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col1:
            show_preview = st.checkbox("Mostra anteprima", value=False, key="preview_checkbox_revoca_nomina")
        
        with col2:
            if show_preview:
                st.info("💡 L'anteprima si aggiorna automaticamente")
        
        if show_preview:
            try:
                preview_text = self._generate_preview_text(form_data)
                st.text(preview_text)
            except Exception as e:
                st.error(f"Errore nell'anteprima: {e}")
    
    def _generate_preview_text(self, data: dict) -> str:
        """Genera il testo di anteprima completo"""
        try:
            preview = ""
            
            # Header azienda formattato in modo uniforme
            denominazione = data.get('denominazione', '[Denominazione]')
            sede_legale = data.get('sede_legale', '[Sede]')
            capitale_sociale_raw = data.get('capitale_versato') or data.get('capitale_deliberato') or data.get('capitale_sociale', '[Capitale]')
            capitale_sociale = CommonDataHandler.format_currency(capitale_sociale_raw)
            codice_fiscale = data.get('codice_fiscale', '[CF]')

            preview += f"""{denominazione}
Sede in {sede_legale}
Capitale sociale Euro {capitale_sociale} i.v.
Codice fiscale: {codice_fiscale}

"""
            
            # Titolo verbale
            preview += f"""Verbale di assemblea dei soci
del {data.get('data_assemblea', '[Data]').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[Data]'}

"""
            
            # Ordine del giorno
            preview += "Ordine del giorno\n"
            preview += "Revoca dell'Amministratore Unico e nomina di nuovo Organo Amministrativo.\n\n"
            
            # Apertura assemblea
            preview += f"""Il giorno {data.get('data_assemblea', '[Data]').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[Data]'} alle ore {data.get('ora_assemblea', '[Ora]')}, presso {data.get('luogo_assemblea', '[Luogo]')}, si è riunita l'assemblea dei soci della società {data.get('denominazione', '[Denominazione]')}.

"""
            
            # Partecipanti
            preview += "Sono presenti:\n"
            soci = data.get('soci', [])
            total_quota_euro = 0.0
            total_quota_percentuale = 0.0

            if soci:
                for socio in soci:
                    if isinstance(socio, dict):
                        nome = socio.get('nome', '[Nome Socio]')
                        quota_euro_raw = socio.get('quota_euro', '')
                        quota_percentuale_raw = socio.get('quota_percentuale', '')

                        try:
                            quota_euro_str = str(quota_euro_raw).replace('.', '').replace(',', '.')
                            total_quota_euro += float(quota_euro_str)
                        except ValueError:
                            pass
                        try:
                            quota_percentuale_str = str(quota_percentuale_raw).replace('.', '').replace(',', '.')
                            total_quota_percentuale += float(quota_percentuale_str)
                        except ValueError:
                            pass

                        quota = '[Quota]' if not quota_euro_raw or str(quota_euro_raw).strip() == '' else str(quota_euro_raw).strip()
                        percentuale = '[%]' if not quota_percentuale_raw or str(quota_percentuale_raw).strip() == '' else str(quota_percentuale_raw).strip()

                        preview += f"- {nome}, titolare di quote per Euro {quota} pari al {percentuale}%\n"

                formatted_total_quota_euro = CommonDataHandler.format_currency(total_quota_euro)
                formatted_total_quota_percentuale = CommonDataHandler.format_percentage(total_quota_percentuale)
                preview += f"Complessivamente, i soci presenti rappresentano una quota pari a nominali euro {formatted_total_quota_euro} pari al {formatted_total_quota_percentuale}% del Capitale Sociale.\n"
            
            preview += f"\nPresente altresì {data.get('rappresentante_legale', '[Rappresentante]')} in qualità di Amministratore Unico.\n\n"
            
            # Costituzione assemblea
            preview += "L'assemblea risulta regolarmente costituita e può validamente deliberare.\n\n"
            
            # Discussione revoca
            preview += "PRIMO PUNTO ALL'ORDINE DEL GIORNO\n"
            preview += "Revoca dell'Amministratore Unico\n\n"
            
            amm_revocato = data.get('amministratore_revocato', {})
            if amm_revocato.get('nome'):
                preview += f"Il Presidente espone che l'Amministratore Unico {amm_revocato.get('nome', '')} "
                if amm_revocato.get('inadempimenti'):
                    preview += f"ha commesso i seguenti inadempimenti: {amm_revocato.get('inadempimenti', '')}.\n\n"
                else:
                    preview += "non ha più i requisiti per ricoprire la carica.\n\n"
            
            preview += "L'assemblea, all'unanimità, delibera di revocare l'Amministratore Unico dalla carica.\n\n"
            
            # Discussione nomina
            preview += "SECONDO PUNTO ALL'ORDINE DEL GIORNO\n"
            preview += "Nomina del nuovo Organo Amministrativo\n\n"
            
            nuovo_amm = data.get('nuovo_amministratore', {})
            if nuovo_amm.get('nome'):
                preview += f"L'assemblea delibera di nominare quale nuovo Amministratore Unico il {nuovo_amm.get('qualifica', 'Sig.')} {nuovo_amm.get('nome', '')}.\n\n"
                
                if nuovo_amm.get('durata_incarico'):
                    preview += f"La durata dell'incarico è: {nuovo_amm.get('durata_incarico', '')}.\n"
                
                if nuovo_amm.get('compenso'):
                    preview += f"Il compenso annuo lordo è stabilito in Euro {nuovo_amm.get('compenso', '')}.\n\n"
            
            # Chiusura
            ora_chiusura = data.get('ora_chiusura', '[Ora chiusura]')
            preview += f"Non essendovi altro da deliberare, l'assemblea viene sciolta alle ore {ora_chiusura}.\n\n"
            
            # Note aggiuntive
            if data.get('note_aggiuntive'):
                preview += f"Note aggiuntive:\n{data.get('note_aggiuntive', '')}\n\n"
            
            # Firme
            preview += "Il Presidente\n"
            preview += f"{data.get('rappresentante_legale', '[Rappresentante]')}\n\n"
            preview += "Il Segretario\n"
            preview += f"{data.get('segretario', '[Segretario]')}"
            
            return preview
            
        except Exception as e:
            return f"Errore nella generazione dell'anteprima: {str(e)}"
    
    def generate_document(self, data: dict) -> Document:
        """Genera il documento Word del verbale"""
        import os
        
        # Utilizza il template .docx esistente per mantenere la formattazione
        template_path = os.path.join(os.path.dirname(__file__), 'template.docx')
        
        try:
            if os.path.exists(template_path):
                doc = Document(template_path)
                # Rimuovi il contenuto esistente del template mantenendo gli stili
                for paragraph in doc.paragraphs[:]:
                    p = paragraph._element
                    p.getparent().remove(p)
            else:
                doc = Document()
                # Setup stili solo se non usiamo template
                self._setup_document_styles(doc)
        except Exception as e:
            # Fallback a documento vuoto se il template non può essere caricato
            doc = Document()
            self._setup_document_styles(doc)
        
        # Setup stili
        self._setup_document_styles(doc)
        
        # Header azienda
        self._add_company_header(doc, data)
        
        # Titolo verbale
        self._add_verbale_title(doc, data)
        
        # Sezioni del verbale
        self._add_opening_section(doc, data)
        self._add_participants_section(doc, data)
        self._add_preliminary_statements(doc, data)
        self._add_revoca_discussion(doc, data)
        self._add_nomina_discussion(doc, data)
        self._add_closing_section(doc, data)
        self._add_signatures(doc, data)
        
        return doc
    
    def _setup_document_styles(self, doc):
        """Definisce e registra gli stili del documento."""
        styles = doc.styles
        try:
            norm_style = styles['Normal']
        except KeyError:
            norm_style = styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
        
        font = norm_style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)
        norm_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        norm_style.paragraph_format.line_spacing = 1.15
        norm_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT # Default, sovrascritto se necessario
        norm_style.paragraph_format.space_before = Pt(0)
        norm_style.paragraph_format.space_after = Pt(0)

        try:
            comp_header_style = styles['CompanyHeader']
        except KeyError:
            comp_header_style = styles.add_style('CompanyHeader', WD_STYLE_TYPE.PARAGRAPH)
        comp_header_style.base_style = norm_style
        comp_header_style.font.size = Pt(10) # Specifico per indirizzo, C.F., etc.
        comp_header_style.font.bold = False
        comp_header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # space_after per CompanyHeader è gestito in _add_company_header

        try:
            verb_title_style = styles['VerbaleTitle']
        except KeyError:
            verb_title_style = styles.add_style('VerbaleTitle', WD_STYLE_TYPE.PARAGRAPH)
        verb_title_style.base_style = norm_style
        verb_title_style.font.size = Pt(16) # Aumentato
        verb_title_style.font.bold = True
        verb_title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # Centrato
        verb_title_style.paragraph_format.space_before = Pt(12)
        verb_title_style.paragraph_format.space_after = Pt(6)

        try:
            verb_subtitle_style = styles['VerbaleSubtitle']
        except KeyError:
            verb_subtitle_style = styles.add_style('VerbaleSubtitle', WD_STYLE_TYPE.PARAGRAPH)
        verb_subtitle_style.base_style = norm_style
        verb_subtitle_style.font.size = Pt(12) # Aumentato
        verb_subtitle_style.font.bold = False
        verb_subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # Centrato
        verb_subtitle_style.paragraph_format.space_after = Pt(18) # Più spazio prima del corpo

        try:
            section_header_style = styles['SectionHeader']
        except KeyError:
            section_header_style = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
        section_header_style.base_style = norm_style
        section_header_style.font.size = Pt(12) # Aumentato
        section_header_style.font.bold = True
        section_header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        section_header_style.paragraph_format.space_before = Pt(12)
        section_header_style.paragraph_format.space_after = Pt(6)

        try:
            body_text_style = styles['BodyText']
        except KeyError:
            body_text_style = styles.add_style('BodyText', WD_STYLE_TYPE.PARAGRAPH)
        body_text_style.base_style = norm_style
        body_text_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Giustificato
        body_text_style.paragraph_format.space_after = Pt(6)

        try:
            list_bullet_style = styles['CustomListBullet']
        except KeyError:
            list_bullet_style = styles.add_style('CustomListBullet', WD_STYLE_TYPE.PARAGRAPH)
        list_bullet_style.base_style = body_text_style # Eredita da BodyText (quindi giustificato)
        # list_bullet_style.paragraph_format.left_indent = Inches(0.5) # Indentazione gestita al momento dell'aggiunta

        try:
            list_number_style = styles['CustomListNumber']
        except KeyError:
            list_number_style = styles.add_style('CustomListNumber', WD_STYLE_TYPE.PARAGRAPH)
        list_number_style.base_style = body_text_style # Eredita da BodyText (quindi giustificato)
        # list_number_style.paragraph_format.left_indent = Inches(0.5) # Indentazione gestita al momento dell'aggiunta
    
    def _add_company_header(self, doc, data):
        """Aggiunge header azienda utilizzando stili"""
        # Denominazione (più grande e grassetto)
        p_denominazione = doc.add_paragraph(style='CompanyHeader')
        run_denominazione = p_denominazione.add_run(data.get('denominazione', '[DENOMINAZIONE SOCIETA\']').upper())
        run_denominazione.font.size = Pt(11) # Leggermente più grande del resto dell'header
        run_denominazione.font.bold = True
        p_denominazione.paragraph_format.space_after = Pt(0) # Nessuno spazio dopo la denominazione

        # Altre informazioni dell'header (dimensione standard, non grassetto)
        sede_text = f"Sede in {data.get('sede_legale', '[Indirizzo Sede Legale]')}"
        doc.add_paragraph(sede_text, style='CompanyHeader').paragraph_format.space_after = Pt(0)
        
        capitale_text = f"Capitale sociale Euro {data.get('capitale_sociale', '[Importo Capitale Sociale]')} i.v."
        doc.add_paragraph(capitale_text, style='CompanyHeader').paragraph_format.space_after = Pt(0)

        codice_fiscale_text = f"Codice fiscale: {data.get('codice_fiscale', '[Partita IVA/Codice Fiscale]')}"
        p_cf = doc.add_paragraph(codice_fiscale_text, style='CompanyHeader')
        p_cf.paragraph_format.space_after = Pt(12) # Spazio dopo l'intero blocco header

    def _add_verbale_title(self, doc, data):
        """Aggiunge il titolo del verbale e la data"""
        doc.add_paragraph("VERBALE DI ASSEMBLEA DEI SOCI", style='VerbaleTitle')
        
        data_assemblea_obj = data.get('data_assemblea', date.today())
        if isinstance(data_assemblea_obj, str):
            try:
                # Prova a convertire la stringa in oggetto data, se necessario
                data_assemblea_obj = datetime.strptime(data_assemblea_obj, '%Y-%m-%d').date()
            except ValueError:
                data_assemblea_str = "[DATA ASSEMBLEA]" # Fallback se la stringa non è parsabile
        
        if isinstance(data_assemblea_obj, date):
            data_assemblea_str = data_assemblea_obj.strftime('%d/%m/%Y')
        else: # Se ancora non è una data (es. None o altro tipo)
            data_assemblea_str = "[DATA ASSEMBLEA]"
            
        doc.add_paragraph(f"del {data_assemblea_str}", style='VerbaleSubtitle')
    
    def _add_opening_section(self, doc, data):
        """Aggiunge la sezione di apertura del verbale."""
        luogo_assemblea = data.get('luogo_assemblea', '[LUOGO ASSEMBLEA]')
        data_assemblea_obj = data.get('data_assemblea', date.today())
        if isinstance(data_assemblea_obj, str):
            try:
                data_assemblea_obj = datetime.strptime(data_assemblea_obj, '%Y-%m-%d').date()
            except ValueError:
                data_assemblea_str = "[DATA ASSEMBLEA]"
        
        if isinstance(data_assemblea_obj, date):
            data_assemblea_str = data_assemblea_obj.strftime('%d/%m/%Y')
        else:
            data_assemblea_str = "[DATA ASSEMBLEA]"
        
        ora_assemblea = data.get('ora_assemblea', '[ORA ASSEMBLEA]')

        opening_text = f"Oggi {data_assemblea_str} alle ore {ora_assemblea} presso {luogo_assemblea}, " \
                       f"si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:"
        p_opening = doc.add_paragraph(opening_text, style='BodyText')
        p_opening.paragraph_format.space_after = Pt(12) # Spazio dopo il paragrafo introduttivo

        # ORDINE DEL GIORNO
        doc.add_paragraph("ORDINE DEL GIORNO", style='SectionHeader') # SectionHeader ha già il suo spacing

        punti_odg = data.get('punti_ordine_del_giorno', [])
        if not punti_odg: 
            punti_odg = [
                "Approvazione del Bilancio al [DATA BILANCIO] e dei documenti correlati",
                "Delibere consequenziali"
            ]
        
        for i, punto in enumerate(punti_odg):
            # Usiamo CustomListNumber per i punti dell'ODG
            # La numerazione effettiva viene dal testo, lo stile applica solo la formattazione (indent, etc)
            p_odg_item = doc.add_paragraph(f"{i+1}. {punto}", style='CustomListNumber')
            p_odg_item.paragraph_format.left_indent = Inches(0.5)
            p_odg_item.paragraph_format.first_line_indent = Inches(-0.25) # Hanging indent per il numero
            if i == len(punti_odg) -1: # Ultimo elemento
                 p_odg_item.paragraph_format.space_after = Pt(12) # Spazio dopo l'intero ODG
            else:
                 p_odg_item.paragraph_format.space_after = Pt(2) # Spazio ridotto tra gli item
    
    def _add_participants_section(self, doc, data):
        """Aggiunge la sezione relativa ai partecipanti."""
        presidente = data.get('presidente_assemblea', '[PRESIDENTE ASSEMBLEA]')
        statuto_art_presidenza = data.get('statuto_art_presidenza', '[ART. STATUTO PRESIDENZA]')
        
        text_presidenza = f"Assume la presidenza ai sensi dell'art. {statuto_art_presidenza} " \
                          f"dello statuto sociale il Sig. {presidente}, " \
                          f"Amministratore Unico, il quale dichiara e constata:"
        p_presidenza = doc.add_paragraph(text_presidenza, style='BodyText')
        p_presidenza.paragraph_format.space_after = Pt(6)

        constatazioni_iniziali = [
            f"che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. {data.get('statuto_art_convocazione', '[ART. STATUTO CONVOCAZIONE]')} dello statuto sociale) l'intervento all'assemblea può avvenire anche in {data.get('modalita_intervento', 'audioconferenza')}",
            "che sono presenti/partecipano all'assemblea:"
        ]
        for i, const_text in enumerate(constatazioni_iniziali):
            p_const = doc.add_paragraph(f"{i+1}- {const_text}", style='CustomListNumber') # Usiamo lo stile per l'indentazione
            p_const.paragraph_format.left_indent = Inches(0.5)
            p_const.paragraph_format.first_line_indent = Inches(-0.25)
            p_const.paragraph_format.space_after = Pt(2)
        
        # Amministratore Unico come partecipante
        p_au = doc.add_paragraph(f"Amministratore Unico nella persona del suddetto Presidente Sig. {presidente}", style='CustomListBullet')
        p_au.paragraph_format.left_indent = Inches(0.75) # Indentazione maggiore per distinguerlo
        p_au.paragraph_format.first_line_indent = Inches(-0.25)
        p_au.paragraph_format.space_after = Pt(2)

        soci_presenti = data.get('soci_presenti', [])
        if not soci_presenti:
            soci_presenti = [
                {'nome': '[NOME SOCIO 1]', 'quota_euro': '[QUOTA SOCIO 1]', 'quota_percentuale': '[% SOCIO 1]'},
                {'nome': '[NOME SOCIO 2]', 'quota_euro': '[QUOTA SOCIO 2]', 'quota_percentuale': '[% SOCIO 2]'}
            ]

        if soci_presenti:
            p_intro_soci = doc.add_paragraph("nonché i seguenti soci o loro rappresentanti:", style='BodyText')
            p_intro_soci.paragraph_format.left_indent = Inches(0.5) # Allineato con le constatazioni
            p_intro_soci.paragraph_format.space_after = Pt(2)

            for socio in soci_presenti:
                nome = socio.get('nome', '[NOME SOCIO]')
                quota_value = socio.get('quota_euro', '')
                percentuale_value = socio.get('quota_percentuale', '')
                
                # Gestione robusta dei valori nulli o vuoti
                quota = '[QUOTA]' if quota_value is None or str(quota_value).strip() == '' else str(quota_value).strip()
                percentuale = '[%]' if percentuale_value is None or str(percentuale_value).strip() == '' else str(percentuale_value).strip()
                
                socio_text = f"il Sig. {nome} socio recante una quota pari a nominali euro {quota} pari al {percentuale}% del Capitale Sociale"
                p_socio = doc.add_paragraph(socio_text, style='CustomListBullet')
                p_socio.paragraph_format.left_indent = Inches(0.75) # Stessa indentazione dell'AU
                p_socio.paragraph_format.first_line_indent = Inches(-0.25)
                p_socio.paragraph_format.space_after = Pt(2)

        # Ulteriori constatazioni (dopo l'elenco dei soci)
        ulteriori_constatazioni = [
            "che gli intervenuti sono legittimati alla presente assemblea;",
            "che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno."
        ]
        # Continua la numerazione dalle constatazioni iniziali
        start_index_ulteriori = len(constatazioni_iniziali) + 1 
        for i, const_text in enumerate(ulteriori_constatazioni, start=start_index_ulteriori):
            p_uconst = doc.add_paragraph(f"{i}- {const_text}", style='CustomListNumber')
            p_uconst.paragraph_format.left_indent = Inches(0.5)
            p_uconst.paragraph_format.first_line_indent = Inches(-0.25)
            p_uconst.paragraph_format.space_after = Pt(2)

        segretario = data.get('segretario_assemblea', '[SEGRETARIO ASSEMBLEA]')
        p_segretario = doc.add_paragraph(f"I presenti all'unanimità chiamano a fungere da segretario il signor {segretario}, che accetta l'incarico.", style='BodyText')
        p_segretario.paragraph_format.space_after = Pt(12) # Spazio dopo questa sezione
    
    def _add_preliminary_statements(self, doc, data):
        """Aggiunge le dichiarazioni preliminari."""
        # presidente = data.get('presidente_assemblea', '[PRESIDENTE ASSEMBLEA]') # Non usato direttamente qui, ma nel contesto
        capitale_sociale_totale = data.get('capitale_sociale_totale', '[CAPITALE SOCIALE TOTALE]')
        percentuale_presente = data.get('percentuale_capitale_presente', '[% CAPITALE PRESENTE]')
        numero_soci_presenti = data.get('numero_soci_presenti', '[NUMERO SOCI PRESENTI]')
        numero_soci_totali = data.get('numero_soci_totali', '[NUMERO SOCI TOTALI]')

        text1 = f"Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati " \
                f"mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere " \
                f"e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante."
        p1 = doc.add_paragraph(text1, style='BodyText')
        p1.paragraph_format.space_after = Pt(6)

        text2_base = f"Il Presidente constata e fa constatare che l'assemblea risulta regolarmente convocata " \
                     f"e debitamente costituita ai sensi di legge e di statuto, essendo presenti tanti soci " \
                     f"rappresentanti il {percentuale_presente}% del capitale sociale pari ad Euro {capitale_sociale_totale}"
        
        # Aggiungere dettagli sul numero di soci se disponibili e significativi
        if numero_soci_presenti and numero_soci_totali and numero_soci_presenti != '[NUMERO SOCI PRESENTI]': # Controlla anche il placeholder
            text2 = f"{text2_base}, e quindi atta a deliberare sugli argomenti posti all'ordine del giorno (presenti n. {numero_soci_presenti} soci su n. {numero_soci_totali} soci)."
        else:
            text2 = f"{text2_base}, e quindi atta a deliberare sugli argomenti posti all'ordine del giorno."
        
        p2 = doc.add_paragraph(text2, style='BodyText')
        p2.paragraph_format.space_after = Pt(6)

        p3 = doc.add_paragraph(f"Il Presidente dichiara quindi l'assemblea validamente costituita e atta a deliberare.", style='BodyText')
        p3.paragraph_format.space_after = Pt(12) # Spazio maggiore dopo questa sezione
    
    def _add_revoca_discussion(self, doc, data):
        """Aggiunge la discussione sulla revoca."""
        # Il numero "1." è già parte del titolo ODG, quindi qui usiamo solo il testo del punto
        # Se questo è un sotto-punto specifico, la numerazione dovrebbe rifletterlo.
        # Assumiamo che sia il primo punto principale della discussione.
        doc.add_paragraph("REVOCA DELL'AMMINISTRATORE UNICO", style='SectionHeader')

        # presidente = data.get('presidente_assemblea', '[PRESIDENTE ASSEMBLEA]') # Non usato direttamente
        amministratore_da_revocare = data.get('amministratore_da_revocare', '[NOME AMMINISTRATORE DA REVOCARE]')
        motivazione_revoca = data.get('motivazione_revoca', '[MOTIVAZIONE DELLA REVOCA]')

        text1 = f"Il Presidente illustra ai presenti la necessità di procedere alla revoca " \
                f"dell'attuale Amministratore Unico, Sig. {amministratore_da_revocare}, {motivazione_revoca}."
        p1 = doc.add_paragraph(text1, style='BodyText')
        p1.paragraph_format.space_after = Pt(6)

        # Dati per la delibera (da rendere dinamici o configurabili)
        percentuale_voti_favorevoli_revoca = data.get('percentuale_voti_favorevoli_revoca', '[PERCENTUALE VOTI FAVOREVOLI]')
        numero_voti_favorevoli_revoca = data.get('numero_voti_favorevoli_revoca', '[NUMERO VOTI FAVOREVOLI]')
        numero_voti_totali_revoca = data.get('numero_voti_totali_revoca', '[NUMERO VOTI TOTALI]')

        text2 = f"Dopo ampia ed esauriente discussione, l'assemblea, udita la relazione del Presidente, " \
                f"con il voto favorevole di {percentuale_voti_favorevoli_revoca}% del capitale sociale " \
                f"avente diritto al voto, pari a n. {numero_voti_favorevoli_revoca} voti su n. {numero_voti_totali_revoca} voti totali,"
        p2 = doc.add_paragraph(text2, style='BodyText')
        p2.paragraph_format.space_after = Pt(3) # Spazio minore prima di DELIBERA
        
        p_delibera_intro = doc.add_paragraph(style='BodyText') # Paragrafo per DELIBERA
        p_delibera_intro.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_delibera = p_delibera_intro.add_run("DELIBERA")
        run_delibera.bold = True
        run_delibera.font.size = Pt(12) # Leggermente più grande
        p_delibera_intro.paragraph_format.space_after = Pt(6)

        delibera_revoca_text = f"Di revocare, con effetto immediato, dalla carica di Amministratore Unico " \
                               f"della società il Sig. {amministratore_da_revocare}, ringraziandolo per l'opera sin qui svolta."
        
        p_delibera_item = doc.add_paragraph(delibera_revoca_text, style='CustomListBullet') # O CustomListNumber se preferito
        p_delibera_item.paragraph_format.left_indent = Inches(0.5)
        p_delibera_item.paragraph_format.first_line_indent = Inches(-0.25)
        p_delibera_item.paragraph_format.space_after = Pt(12) # Spazio dopo la delibera di revoca
    
    def _add_nomina_discussion(self, doc, data):
        """Aggiunge la discussione sulla nomina."""
        doc.add_paragraph("NOMINA DEL NUOVO ORGANO AMMINISTRATIVO", style='SectionHeader')

        tipo_organo_amministrativo = data.get('tipo_organo_amministrativo', 'Amministratore Unico')
        text1 = f"Il Presidente invita quindi l'assemblea a deliberare in merito alla nomina " \
                f"del nuovo organo amministrativo, proponendo la nomina di un {tipo_organo_amministrativo}."
        p1 = doc.add_paragraph(text1, style='BodyText')
        p1.paragraph_format.space_after = Pt(6)

        # Dati generici per la delibera di nomina
        percentuale_voti_favorevoli_nomina = data.get('percentuale_voti_favorevoli_nomina', '[PERCENTUALE VOTI FAVOREVOLI]')
        numero_voti_favorevoli_nomina = data.get('numero_voti_favorevoli_nomina', '[NUMERO VOTI FAVOREVOLI]')
        numero_voti_totali_nomina = data.get('numero_voti_totali_nomina', '[NUMERO VOTI TOTALI]')

        # Amministratore Unico
        if tipo_organo_amministrativo == 'Amministratore Unico':
            nome_nuovo_amministratore = data.get('nome_nuovo_amministratore', '[NOME NUOVO AMMINISTRATORE]')
            compenso_nuovo_amministratore = data.get('compenso_nuovo_amministratore', '0,00')
            durata_incarico_anni = data.get('durata_incarico_anni', '3')
            scadenza_incarico_data = data.get('scadenza_incarico', '[DATA APPROVAZIONE BILANCIO ESERCIZIO]')

            text_nomina_au_proposta = f"Viene proposto quale Amministratore Unico il Sig. {nome_nuovo_amministratore}."
            p_prop_au = doc.add_paragraph(text_nomina_au_proposta, style='BodyText')
            p_prop_au.paragraph_format.space_after = Pt(6)

            text_delibera_intro_au = f"L'assemblea, dopo breve discussione, con il voto favorevole di {percentuale_voti_favorevoli_nomina}% " \
                                     f"del capitale sociale avente diritto al voto, pari a n. {numero_voti_favorevoli_nomina} voti " \
                                     f"su n. {numero_voti_totali_nomina} voti totali,"
            p_delib_intro_au = doc.add_paragraph(text_delibera_intro_au, style='BodyText')
            p_delib_intro_au.paragraph_format.space_after = Pt(3)

            p_delibera_au_title = doc.add_paragraph(style='BodyText')
            p_delibera_au_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_delibera_au = p_delibera_au_title.add_run("DELIBERA")
            run_delibera_au.bold = True
            run_delibera_au.font.size = Pt(12)
            p_delibera_au_title.paragraph_format.space_after = Pt(6)

            delibere_au_items = [
                f"Di nominare Amministratore Unico della società il Sig. {nome_nuovo_amministratore}, il quale, presente all'assemblea, dichiara di accettare la carica e di non trovarsi in alcuna delle cause di ineleggibilità e/o incompatibilità previste dalla legge e dallo statuto.",
                f"Di stabilire che l'incarico avrà durata di {durata_incarico_anni} esercizi sociali, e quindi fino all'approvazione del bilancio relativo all'esercizio che chiuderà il {scadenza_incarico_data}.",
                f"Di attribuire all'Amministratore Unico un compenso annuo lordo di Euro {compenso_nuovo_amministratore}, oltre al rimborso delle spese sostenute per l'esercizio del suo ufficio.",
                "Di conferire all'Amministratore Unico i più ampi poteri di ordinaria e straordinaria amministrazione, con facoltà di nominare procuratori speciali per singoli atti o categorie di atti."
            ]
            for item_text in delibere_au_items:
                p_item = doc.add_paragraph(item_text, style='CustomListBullet')
                p_item.paragraph_format.left_indent = Inches(0.5)
                p_item.paragraph_format.first_line_indent = Inches(-0.25)
                p_item.paragraph_format.space_after = Pt(3)
            if delibere_au_items: # Aggiungi spazio extra dopo l'ultimo elemento della lista
                doc.paragraphs[-1].paragraph_format.space_after = Pt(12)
        
        # Consiglio di Amministrazione (CdA)
        elif tipo_organo_amministrativo == 'Consiglio di Amministrazione':
            membri_cda_input = data.get('membri_cda', []) # Lista di dict [{'nome_consigliere': '...', 'ruolo_cda': '...'}]
            presidente_cda = data.get('presidente_cda', '[NOME PRESIDENTE CDA]')
            durata_incarico_cda_anni = data.get('durata_incarico_cda_anni', '3')
            scadenza_incarico_cda_data = data.get('scadenza_incarico_cda', '[DATA APPROVAZIONE BILANCIO ESERCIZIO]')
            compenso_cda_descr = data.get('compenso_cda_descrizione', 'un compenso annuo lordo di Euro [IMPORTO] per il Presidente e Euro [IMPORTO] per ciascun Consigliere')

            num_membri_cda = len(membri_cda_input) if membri_cda_input else '[NUMERO]'
            text_nomina_cda_proposta = f"Viene proposto di nominare un Consiglio di Amministrazione composto da {num_membri_cda} membri."
            p_prop_cda = doc.add_paragraph(text_nomina_cda_proposta, style='BodyText')
            p_prop_cda.paragraph_format.space_after = Pt(6)
            
            text_delibera_intro_cda = f"L'assemblea, dopo breve discussione, con il voto favorevole di {percentuale_voti_favorevoli_nomina}% " \
                                      f"del capitale sociale avente diritto al voto, pari a n. {numero_voti_favorevoli_nomina} voti " \
                                      f"su n. {numero_voti_totali_nomina} voti totali,"
            p_delib_intro_cda = doc.add_paragraph(text_delibera_intro_cda, style='BodyText')
            p_delib_intro_cda.paragraph_format.space_after = Pt(3)

            p_delibera_cda_title = doc.add_paragraph(style='BodyText')
            p_delibera_cda_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_delibera_cda = p_delibera_cda_title.add_run("DELIBERA")
            run_delibera_cda.bold = True
            run_delibera_cda.font.size = Pt(12)
            p_delibera_cda_title.paragraph_format.space_after = Pt(6)

            delibere_cda_items = [
                "Di nominare quali membri del Consiglio di Amministrazione i seguenti signori, i quali, presenti all'assemblea, dichiarano di accettare la carica e di non trovarsi in alcuna delle cause di ineleggibilità e/o incompatibilità previste dalla legge e dallo statuto:"
            ]
            for item_text in delibere_cda_items:
                # Questo è un'introduzione alla lista dei membri
                p_item_intro = doc.add_paragraph(item_text, style='BodyText') # Non è un bullet item di per sé
                p_item_intro.paragraph_format.space_after = Pt(3)

            if not membri_cda_input:
                membri_cda_input = [{'nome_consigliere': '[NOME CONSIGLIERE 1]', 'ruolo_cda': 'Consigliere'}]
            for membro in membri_cda_input:
                membro_text = f"Sig. {membro.get('nome_consigliere', '[NOME CONSIGLIERE]')}, {membro.get('ruolo_cda', 'Consigliere')}"
                p_membro = doc.add_paragraph(membro_text, style='CustomListBullet')
                p_membro.paragraph_format.left_indent = Inches(0.75) # Indentazione maggiore per i membri
                p_membro.paragraph_format.first_line_indent = Inches(-0.25)
                p_membro.paragraph_format.space_after = Pt(2)
            if membri_cda_input: # Spazio dopo l'elenco dei membri
                 doc.paragraphs[-1].paragraph_format.space_after = Pt(6)
            
            altre_delibere_cda = [
                f"Di nominare Presidente del Consiglio di Amministrazione il Sig. {presidente_cda}.",
                f"Di stabilire che l'incarico per tutti i membri del Consiglio di Amministrazione avrà durata di {durata_incarico_cda_anni} esercizi sociali, e quindi fino all'approvazione del bilancio relativo all'esercizio che chiuderà il {scadenza_incarico_cda_data}.",
                f"Di attribuire ai membri del Consiglio di Amministrazione {compenso_cda_descr}, oltre al rimborso delle spese sostenute per l'esercizio del loro ufficio.",
                "Di conferire al Consiglio di Amministrazione i più ampi poteri di ordinaria e straordinaria amministrazione, con facoltà di nominare procuratori speciali per singoli atti o categorie di atti. Al Presidente del Consiglio di Amministrazione spetta la legale rappresentanza della società."
            ]
            for item_text in altre_delibere_cda:
                p_item = doc.add_paragraph(item_text, style='CustomListBullet')
                p_item.paragraph_format.left_indent = Inches(0.5)
                p_item.paragraph_format.first_line_indent = Inches(-0.25)
                p_item.paragraph_format.space_after = Pt(3)
            if altre_delibere_cda:
                doc.paragraphs[-1].paragraph_format.space_after = Pt(12)

        # Fallback se tipo_organo_amministrativo non è gestito
        else:
            doc.add_paragraph(f"Discussione sulla nomina di {tipo_organo_amministrativo} non implementata.", style='BodyText').paragraph_format.space_after = Pt(12)
    
    def _add_closing_section(self, doc, data):
        """Aggiunge la sezione di chiusura."""
        ora_fine_assemblea = data.get('ora_fine_assemblea', '[ORA FINE ASSEMBLEA]')
        
        closing_text = f"Null'altro essendovi a deliberare e nessuno chiedendo la parola, il Presidente dichiara sciolta l'assemblea alle ore {ora_fine_assemblea}, previa redazione, lettura e approvazione unanime del presente verbale, che consta di [NUMERO PAGINE] pagine."
        # Nota: [NUMERO PAGINE] è un placeholder che andrebbe gestito dinamicamente se possibile, o rimosso/modificato.
        p_closing = doc.add_paragraph(closing_text, style='BodyText')
        p_closing.paragraph_format.space_before = Pt(12)
        p_closing.paragraph_format.space_after = Pt(12)

        p_lcs = doc.add_paragraph("Letto, confermato e sottoscritto.", style='BodyText')
        p_lcs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT # O CENTER se preferito
        p_lcs.paragraph_format.space_before = Pt(18) # Più spazio prima delle firme
        p_lcs.paragraph_format.space_after = Pt(6)
    
    def _add_signatures(self, doc, data):
        """Aggiunge le firme."""
        # Utilizziamo una tabella per allineare le firme
        # Questo è un approccio comune per layout di firme affiancate.
        # Se le firme devono essere una sotto l'altra, la tabella non è strettamente necessaria
        # ma può aiutare a controllare la larghezza e lo spazio.

        # Decidiamo se usare una tabella o paragrafi semplici
        # Per ora, usiamo paragrafi semplici con allineamento, che è più facile da gestire
        # per firme che appaiono una dopo l'altra verticalmente o allineate a sinistra/destra.

        presidente_nome = data.get('presidente_assemblea', '[NOME PRESIDENTE ASSEMBLEA]')
        segretario_nome = data.get('segretario_assemblea', '[NOME SEGRETARIO ASSEMBLEA]')

        # Spazio prima delle firme
        # doc.add_paragraph().paragraph_format.space_before = Pt(24) # Spazio abbondante

        # Firma del Presidente
        p_pres_label = doc.add_paragraph(style='BodyText')
        p_pres_label.paragraph_format.space_before = Pt(36) # Aumentato spazio prima
        p_pres_label.paragraph_format.space_after = Pt(0)
        run_pres_label = p_pres_label.add_run("Il Presidente")
        # run_pres_label.bold = True # Opzionale, se lo stile 'BodyText' non è già bold
        # p_pres_label.alignment = WD_ALIGN_PARAGRAPH.LEFT # O CENTER/RIGHT
        
        p_pres_name = doc.add_paragraph(presidente_nome, style='BodyText')
        p_pres_name.paragraph_format.space_before = Pt(0) 
        p_pres_name.paragraph_format.space_after = Pt(24) # Spazio dopo il nome del presidente
        # p_pres_name.alignment = WD_ALIGN_PARAGRAPH.LEFT # O CENTER/RIGHT

        # Firma del Segretario
        p_secr_label = doc.add_paragraph(style='BodyText')
        p_secr_label.paragraph_format.space_before = Pt(12) # Spazio prima del segretario
        p_secr_label.paragraph_format.space_after = Pt(0)
        run_secr_label = p_secr_label.add_run("Il Segretario")
        # run_secr_label.bold = True # Opzionale
        # p_secr_label.alignment = WD_ALIGN_PARAGRAPH.LEFT # O CENTER/RIGHT

        p_secr_name = doc.add_paragraph(segretario_nome, style='BodyText')
        p_secr_name.paragraph_format.space_before = Pt(0)
        p_secr_name.paragraph_format.space_after = Pt(12) # Spazio finale
        # p_secr_name.alignment = WD_ALIGN_PARAGRAPH.LEFT # O CENTER/RIGHT

        # Se si volessero le firme affiancate, una tabella sarebbe più indicata:
        # table = doc.add_table(rows=2, cols=2)
        # table.cell(0, 0).text = 'Il Presidente'
        # table.cell(1, 0).text = presidente_nome
        # table.cell(0, 1).text = 'Il Segretario'
        # table.cell(1, 1).text = segretario_nome
        # # Applicare stili e formattazione alle celle/paragrafi della tabella

# Registra il template
DocumentTemplateFactory.register_template('verbale_assemblea_revoca_nomina', VerbaleRevocaNominaTemplate)