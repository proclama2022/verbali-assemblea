"""
Template per Verbale di Assemblea - Ratifica Operato dell'Organo Amministrativo
Questo template è specifico per verbali di ratifica dell'operato degli amministratori.
"""

import sys
import os

# Aggiungi il path della cartella src (relativo alla root del progetto)
current_dir = os.path.dirname(os.path.abspath(__file__))
src_path = os.path.join(os.path.dirname(current_dir), 'src')
if src_path not in sys.path:
    sys.path.append(src_path)

from document_templates import DocumentTemplate, DocumentTemplateFactory
from common_data_handler import CommonDataHandler
from base_verbale_template import BaseVerbaleTemplate
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date
import streamlit as st
import pandas as pd
import re

class VerbaleRatificaOperatoTemplate(BaseVerbaleTemplate):
    """Template per Verbale di Assemblea dei Soci - Ratifica Operato dell'Organo Amministrativo"""
    
    def get_template_name(self) -> str:
        return "Ratifica Operato dell'Organo Amministrativo"
    
    def get_required_fields(self) -> list:
        return [
            "denominazione", "sede_legale", "capitale_sociale", "codice_fiscale",
            "data_assemblea", "ora_assemblea", "presidente", "segretario", 
            "soci", "data_scadenza_mandato", "attivita_specifiche"
        ]
    
    def get_form_fields(self, extracted_data: dict) -> dict:
        """Genera i campi del form per Streamlit utilizzando il CommonDataHandler"""
        form_data = {}
        
        # Utilizza il CommonDataHandler per i dati standard
        # Dati azienda standardizzati
        form_data.update(CommonDataHandler.extract_and_populate_company_data(extracted_data))
        
        # Dati assemblea standardizzati
        form_data.update(CommonDataHandler.extract_and_populate_assembly_data(extracted_data))
        
        # Campi specifici per ratifica operato
        st.subheader("📋 Configurazioni Specifiche Ratifica Operato")
        
        col1, col2 = st.columns(2)
        with col1:
            form_data["ruolo_presidente"] = st.selectbox("Ruolo del presidente", 
                                                        ["Amministratore Unico", 
                                                         "Presidente del Consiglio di Amministrazione",
                                                         "Altro (come da statuto)"])
            form_data["tipo_organo"] = st.selectbox("Tipo di organo amministrativo", 
                                                   ["Amministratore Unico", 
                                                    "Consiglio di Amministrazione",
                                                    "Altro"])
        with col2:
            form_data["data_scadenza_mandato"] = st.date_input("Data scadenza mandato", 
                                                              value=None)
            form_data["include_collegio_sindacale"] = st.checkbox("Include Collegio Sindacale", value=False)
        
        # Partecipanti standardizzati usando il CommonDataHandler
        participants_data = CommonDataHandler.extract_and_populate_participants_data(
            extracted_data, 
            unique_key_suffix="ratifica_operato"
        )
        form_data.update(participants_data)
        
        # Amministratori uscenti
        st.subheader("👥 Amministratori Uscenti")
        
        num_amministratori = st.number_input("Numero di amministratori uscenti", 
                                           min_value=1, max_value=10, value=1,
                                           key="num_admin_uscenti")
        
        amministratori_uscenti = []
        for i in range(int(num_amministratori)):
            st.write(f"**Amministratore {i+1}**")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                nome = st.text_input(f"Nome e Cognome", key=f"admin_nome_{i}",
                                   placeholder="es. Mario Rossi")
                
            with col2:
                ruolo = st.selectbox(f"Ruolo", 
                                   ["Amministratore Unico", "Presidente", "Amministratore Delegato", 
                                    "Consigliere", "Altro"],
                                   key=f"admin_ruolo_{i}")
            
            with col3:
                data_nomina = st.date_input(f"Data nomina", 
                                          value=None, key=f"admin_data_nomina_{i}")
            
            if nome:
                amministratori_uscenti.append({
                    "nome": nome,
                    "ruolo": ruolo,
                    "data_nomina": data_nomina
                })
        
        form_data["amministratori_uscenti"] = amministratori_uscenti
        
        # Attività Specifiche
        st.subheader("📝 Attività Specifiche da Ratificare")
        
        form_data["attivita_specifiche"] = st.text_area(
            "Attività specifiche da ratificare",
            placeholder="Inserire le attività specifiche che necessitano di ratifica...\n\nEsempio:\n- Operazioni straordinarie concluse\n- Contratti stipulati\n- Decisioni gestionali prese\n- Altri atti amministrativi",
            height=150,
            help="Elencare le principali attività svolte dagli amministratori che necessitano di ratifica"
        )
        
        # Configurazioni aggiuntive
        st.subheader("⚙️ Configurazioni Aggiuntive")
        
        col1, col2 = st.columns(2)
        with col1:
            form_data["include_revisore"] = st.checkbox("Include revisore contabile", value=False)
            form_data["include_traduzione"] = st.checkbox("Include sezione traduzione", value=False)
        
        with col2:
            form_data["include_sanzioni_spese"] = st.checkbox("Include clausola sanzioni/spese legali", value=True)
            form_data["votazione_unanime"] = st.checkbox("Votazione unanime", value=True)
        
        # Dettagli votazione se non unanime
        if not form_data["votazione_unanime"]:
            st.subheader("🗳️ Dettagli Votazione")
            col1, col2 = st.columns(2)
            with col1:
                form_data["voti_contrari"] = st.text_input("Voti contrari (nomi)", 
                                                          placeholder="es. Sig. Rossi, Sig. Bianchi")
            with col2:
                form_data["astensioni"] = st.text_input("Astensioni (nomi)", 
                                                       placeholder="es. Sig. Verdi")
        
        # Lingue per traduzione
        if form_data["include_traduzione"]:
            st.subheader("🌐 Traduzione")
            col1, col2 = st.columns(2)
            with col1:
                form_data["persona_traduzione"] = st.text_input("Persona che necessita traduzione",
                                                               placeholder="es. Sig. Smith")
            with col2:
                form_data["lingua_traduzione"] = st.selectbox("Lingua di traduzione",
                                                             ["Inglese", "Francese", "Tedesco", "Spagnolo", "Altra"])
        
        # Note aggiuntive
        st.subheader("📝 Note Aggiuntive")
        form_data["note_aggiuntive"] = st.text_area("Note aggiuntive", 
                                                  placeholder="Eventuali note specifiche sulla ratifica...",
                                                  height=80)
        
        return form_data
    
    def show_preview(self, form_data: dict):
        """Mostra l'anteprima del documento fuori dal form"""
        # Sezione Anteprima
        st.markdown("---")
        st.subheader("👁️ Anteprima Documento")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col1:
            show_preview = st.checkbox("Mostra anteprima", value=False, key="preview_checkbox_ratifica")
        
        with col2:
            if show_preview:
                st.info("💡 L'anteprima si aggiorna automaticamente con i dati inseriti sopra")
        
        if show_preview:
            try:
                preview_text = self._generate_preview_text(form_data)
                st.text(preview_text)
            except Exception as e:
                st.error(f"Errore nell'anteprima: {e}")
                st.exception(e)
    
    def _generate_preview_text(self, data: dict) -> str:
        """Genera il testo di anteprima del verbale"""
        try:
            # Header dell'azienda
            header = f"""{data.get('denominazione', '[Denominazione]')}
Sede in {data.get('sede_legale', '[Sede]')}
Capitale sociale Euro {data.get('capitale_sociale', '[Capitale]')} i.v.
Codice fiscale: {data.get('codice_fiscale', '[CF]')}

Verbale di assemblea dei soci
del {data.get('data_assemblea', '[Data]').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[Data]'}

Oggi {data.get('data_assemblea', '[Data]').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[Data]'} alle ore {data.get('ora_assemblea', '[Ora]').strftime('%H:%M') if hasattr(data.get('ora_assemblea'), 'strftime') else '[Ora]'} presso la sede sociale {data.get('sede_legale', '[Sede]')}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:

Ordine del giorno
Ratifica operato dell'Organo amministrativo"""
            
            # Sezione presidenza
            ruolo_presidente = data.get('ruolo_presidente', 'Amministratore Unico')
            presidente_section = f"""

Assume la presidenza ai sensi dell'art. […] dello statuto sociale il Sig. {data.get('presidente', '[Presidente]')} {ruolo_presidente}, il quale dichiara e constata:

1 - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. […] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza

2 - che sono presenti/partecipano all'assemblea:
l'{ruolo_presidente} nella persona del suddetto Presidente Sig. {data.get('presidente', '[Presidente]')}"""
            
            # Consiglio di Amministrazione se applicabile
            if data.get('tipo_organo') == 'Consiglio di Amministrazione':
                presidente_section += """
[oppure
per il Consiglio di Amministrazione:
il Sig […]
il Sig […]
il Sig […]
assente giustificato il Sig […] il quale ha tuttavia rilasciato apposita dichiarazione scritta, conservata agli atti della Società, dalla quale risulta che il medesimo è stato informato su tutti gli argomenti posti all'ordine del giorno e che lo stesso non si oppone alla trattazione degli stessi]"""
            
            # Collegio sindacale se presente
            if data.get('include_collegio_sindacale', False):
                presidente_section += """
[eventualmente
per il Collegio Sindacale
il Dott. […]
il Dott. […]
il Dott. […]
[oppure
il Sindaco Unico nella persona del Sig. […]]]"""
            
            # Revisore se presente
            if data.get('include_revisore', False):
                presidente_section += """
[eventualmente, se invitato
il revisore contabile Dott. […] <oppure il dott. […] in rappresentanza della società di revisione incaricata del controllo contabile>"""
            
            # Soci presenti
            soci = data.get('soci', [])
            total_quota_euro = 0.0
            total_quota_percentuale = 0.0

            for socio in soci:
                if isinstance(socio, dict):
                    try:
                        # Rimuovi punti e sostituisci virgola con punto per la conversione a float
                        quota_euro_str = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                        total_quota_euro += float(quota_euro_str)
                    except ValueError:
                        pass # Ignora valori non numerici
                    try:
                        quota_percentuale_str = str(socio.get('quota_percentuale', '0')).replace('.', '').replace(',', '.')
                        total_quota_percentuale += float(quota_percentuale_str)
                    except ValueError:
                        pass # Ignora valori non numerici
            
            # Formatta i totali per la visualizzazione
            formatted_total_quota_euro = f"{total_quota_euro:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') # Formato italiano
            formatted_total_quota_percentuale = f"{total_quota_percentuale:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') # Formato italiano

            soci_section = f"\nnonché i seguenti soci o loro rappresentanti, [eventualmente così come iscritti a libro soci e] recanti complessivamente una quota pari a nominali euro {formatted_total_quota_euro} pari al {formatted_total_quota_percentuale}% del Capitale Sociale:"
            
            for socio in soci:
                if isinstance(socio, dict):
                    nome = socio.get('nome', '[Nome Socio]')
                    quota_value = socio.get('quota_euro', '')
                    percentuale_value = socio.get('quota_percentuale', '')
                    
                    # Gestione robusta dei valori nulli o vuoti
                    quota = '[Quota]' if quota_value is None or str(quota_value).strip() == '' else str(quota_value).strip()
                    percentuale = '[%]' if percentuale_value is None or str(percentuale_value).strip() == '' else str(percentuale_value).strip()
                    
                    tipo_partecipazione = socio.get('tipo_partecipazione', 'Diretta')
                    
                    if tipo_partecipazione == 'Delegato':
                        delegato = socio.get('delegato', '[Delegato]')
                        soci_section += f"\nil Sig {delegato} delegato del socio Sig {nome} recante una quota pari a nominali euro {quota} pari al {percentuale}% del Capitale Sociale"
                    else:
                        soci_section += f"\nil Sig {nome} socio recante una quota pari a nominali euro {quota} pari al {percentuale}% del Capitale Sociale"
                else:
                    soci_section += f"\nil Sig {socio} socio"
            
            soci_section += "\n2 - che gli intervenuti sono legittimati alla presente assemblea;\n3 - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno."
            
            # Segretario e traduzione
            segretario_section = f"""

I presenti all'unanimità chiamano a fungere da segretario il signor {data.get('segretario', '[Segretario]')}, che accetta l'incarico."""
            
            # Traduzione se necessaria
            if data.get('include_traduzione', False):
                persona_traduzione = data.get('persona_traduzione', '[Nome]')
                lingua = data.get('lingua_traduzione', 'inglese').lower()
                segretario_section += f"""
[eventualmente In particolare, preso atto che il Sig. {persona_traduzione} non conosce la lingua italiana ma dichiara di conoscere la lingua {lingua}, il Presidente dichiara che provvederà a tradurre dall'italiano al {lingua} (e viceversa) gli interventi dei partecipanti alla discussione nonché a tradurre dall'italiano al {lingua} il verbale che sarà redatto al termine della riunione.]"""
            
            segretario_section += """

Il Presidente constata e fa constatare che l'assemblea risulta regolarmente convocata [oppure totalitaria] e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno."""
            
            segretario_section += """

Si passa quindi allo svolgimento dell'ordine del giorno."""
            
            # Discussione ratifica
            attivita_specifiche = data.get('attivita_specifiche', '[Attività specifiche da ratificare]')
            data_scadenza = data.get('data_scadenza_mandato')
            if hasattr(data_scadenza, 'strftime'):
                data_scadenza_str = data_scadenza.strftime('%d/%m/%Y')
            else:
                data_scadenza_str = '[Data scadenza]'
            
            ratifica_section = f"""

Su invito a parlare del Presidente, il socio conferma, come già informalmente anticipato, l'opportunità e il desiderio di procedere, per quanto occorrer possa, alla ratifica delle operazioni poste in essere e degli atti e attività (finanche di tipo omissivo), anche impugnabili a qualsivoglia culpa in vigilando, compiuti dagli Amministratori uscenti per scadenza naturale del mandato, dalla data della loro nomina sino alla data odierna, e per i precedenti mandati, ed al relativo scarico di responsabilità, con rinuncia incondizionata e irrevocabile, nei limiti massimi consentiti dalla legge, all'esercizio di azioni di responsabilità e/o di risarcimento danno nei loro confronti, ed in particolare con riferimento alle seguenti attività (Attività Specifiche), ivi incluse a titolo esemplificativo ma non esaustivo:

{attivita_specifiche}"""
            
            # Votazione
            if data.get('votazione_unanime', True):
                votazione_text = "all'unanimità"
            else:
                voti_contrari = data.get('voti_contrari', '')
                astensioni = data.get('astensioni', '')
                votazione_text = f"con il voto contrario dei Sigg. {voti_contrari}"
                if astensioni:
                    votazione_text += f" e l'astensione dei Sigg. {astensioni}"
            
            deliberazione_section = f"""

Si passa quindi alla votazione con voto palese in forza della quale il Presidente constata che, {votazione_text}, l'assemblea

d e l i b e r a:

di ratificare l'operato di tutti gli Amministratori della Società uscenti per scadenza naturale del mandato in data {data_scadenza_str}, dando loro ampio scarico in merito alle operazioni, atti e attività (finanche di tipo omissivo), in particolare con riferimento alle Attività Specifiche, e anche imputabili a qualsivoglia culpa in vigilando, da loro compiuti in qualità di Amministratori (ivi inclusi in qualità di Presidente o Amministratore Delegato) dalla data della loro nomina sino alla data {data_scadenza_str}

di rinunciare incondizionatamente e irrevocabilmente, nei limiti massimi consentiti dalla legge, all'esercizio delle azioni di responsabilità e/o di risarcimento danno nei confronti di tutti gli Amministratori della Società uscenti per scadenza naturale del mandato in data {data_scadenza_str} in relazione a ogni e qualsiasi operazione, attività, atto, potenziale omissione, circostanza o fatto compiuto, anche imputabile a qualsivoglia culpa in vigilando, connesso o occorso nello svolgimento del loro ufficio come membri del Consiglio di Amministrazione (ivi incluso quale Presidente o Amministratore Delegato), dalla data della loro nomina sino alla data {data_scadenza_str}, e per i precedenti mandati"""
            
            # Sanzioni e spese se incluse
            sanzioni_section = ""
            if data.get('include_sanzioni_spese', True):
                sanzioni_section = """

che eventuali sanzioni e/o spese legali, sia in sede civile che penale, amministrativa o tributaria, che i Consiglieri uscenti fossero chiamati a sostenere in conseguenza della carica ricoperta fino alla data odierna, restino a carico della Società, che, pertanto, ne sosterrà integralmente l'onere, con la sola eccezione di quelle conseguenti a fatti riferibili a loro colpa grave o dolo accertato con sentenza passata in giudicato."""
            
            # Chiusura
            chiusura_section = f"""

*     *     *

Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.

Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo [eventualmente unitamente a quanto allegato].

L'assemblea viene sciolta alle ore [{data.get('ora_chiusura', '[Ora]')}].


Il Presidente                    Il Segretario
{data.get('presidente', '[PRESIDENTE]')}            {data.get('segretario', '[SEGRETARIO]')}"""
            
            # Combina tutte le sezioni
            full_text = (header + presidente_section + soci_section + segretario_section + 
                        ratifica_section + deliberazione_section + sanzioni_section + chiusura_section)
            
            return full_text
            
        except Exception as e:
            return f"Errore nella generazione dell'anteprima: {str(e)}"
    
    def generate_document(self, data: dict) -> Document:
        """Genera il documento Word del verbale"""
        import os
        
        # Utilizza il template .docx esistente per mantenere la formattazione
        template_path = os.path.join(os.path.dirname(__file__), 'template.docx')
        
        try:
            if os.path.exists(template_path):
                doc = Document(template_path)
                # Rimuovi il contenuto esistente del template mantenendo gli stili
                for paragraph in doc.paragraphs[:]:
                    p = paragraph._element
                    p.getparent().remove(p)
            else:
                doc = Document()
                # Setup stili solo se non usiamo template
                self._setup_document_styles(doc)
        except Exception as e:
            # Fallback a documento vuoto se il template non può essere caricato
            doc = Document()
            self._setup_document_styles(doc)
        
        # Setup stili del documento
        self._setup_document_styles(doc)
        
        # Aggiungi header azienda
        self._add_company_header(doc, data)
        
        # Aggiungi titolo verbale
        self._add_verbale_title(doc, data)
        
        # Aggiungi sezione di apertura
        self._add_opening_section(doc, data)
        
        # Aggiungi partecipanti
        self._add_participants_section(doc, data)
        
        # Aggiungi dichiarazioni preliminari
        self._add_preliminary_statements(doc, data)
        
        # Aggiungi discussione ratifica
        self._add_ratifica_discussion(doc, data)
        
        # Aggiungi deliberazione
        self._add_deliberation_section(doc, data)
        
        # Aggiungi sezione di chiusura
        self._add_closing_section(doc, data)
        
        # Aggiungi firme
        self._add_signatures(doc, data)
        
        return doc
    
    def _setup_document_styles(self, doc):
        """Imposta gli stili del documento"""
        styles = doc.styles
        
        # Stile per il titolo principale
        if 'Titolo Principale' not in [s.name for s in styles]:
            title_style = styles.add_style('Titolo Principale', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Times New Roman'
            title_style.font.size = Pt(14)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # Stile per il testo normale
        normal_style = styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal_style.paragraph_format.line_spacing = 1.15
    
    def _add_company_header(self, doc, data):
        """Aggiunge l'header dell'azienda"""
        # Nome azienda
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data.get('denominazione', '[DENOMINAZIONE SOCIETÀ]'))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Sede
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Sede in {data.get('sede_legale', '[SEDE]')}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        # Capitale sociale
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Capitale sociale Euro {data.get('capitale_sociale', '[CAPITALE]')} i.v.")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        # Codice fiscale
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Codice fiscale: {data.get('codice_fiscale', '[CODICE FISCALE]')}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        doc.add_paragraph()
    
    def _add_verbale_title(self, doc, data):
        """Aggiunge il titolo del verbale"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Verbale di assemblea dei soci")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Data
        data_assemblea = data.get('data_assemblea')
        if hasattr(data_assemblea, 'strftime'):
            data_str = data_assemblea.strftime('%d/%m/%Y')
        else:
            data_str = '[DATA]'
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"del {data_str}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        doc.add_paragraph()
    
    def _add_opening_section(self, doc, data):
        """Aggiunge la sezione di apertura"""
        # Data e ora
        data_assemblea = data.get('data_assemblea')
        ora_assemblea = data.get('ora_assemblea')
        
        if hasattr(data_assemblea, 'strftime'):
            data_str = data_assemblea.strftime('%d/%m/%Y')
        else:
            data_str = '[DATA]'
        
        if hasattr(ora_assemblea, 'strftime'):
            ora_str = ora_assemblea.strftime('%H:%M')
        else:
            ora_str = '[ORA]'
        
        text = f"Oggi {data_str} alle ore {ora_str} presso la sede sociale {data.get('sede_legale', '[SEDE]')}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:"
        
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Ordine del giorno
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Ordine del giorno")
        run.font.bold = True
        
        p = doc.add_paragraph("Ratifica operato dell'Organo amministrativo")
    
    def _add_participants_section(self, doc, data):
        """Aggiunge la sezione dei partecipanti"""
        doc.add_paragraph()
        
        ruolo_presidente = data.get('ruolo_presidente', 'Amministratore Unico')
        presidente = data.get('presidente', '[PRESIDENTE]')
        
        text = f"Assume la presidenza ai sensi dell'art. […] dello statuto sociale il Sig. {presidente} {ruolo_presidente}, il quale dichiara e constata:"
        
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Punto 1 - audioconferenza
        p = doc.add_paragraph("1 - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. […] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza")
        
        # Punto 2 - presenti
        p = doc.add_paragraph("2 - che sono presenti/partecipano all'assemblea:")
        p = doc.add_paragraph(f"l'{ruolo_presidente} nella persona del suddetto Presidente Sig. {presidente}")
        
        # Consiglio di Amministrazione se applicabile
        if data.get('tipo_organo') == 'Consiglio di Amministrazione':
            p = doc.add_paragraph("[oppure")
            p = doc.add_paragraph("per il Consiglio di Amministrazione:")
            amministratori = data.get('amministratori_uscenti', [])
            for admin in amministratori[:3]:  # Mostra primi 3
                if isinstance(admin, dict):
                    nome = admin.get('nome', '[Nome]')
                    p = doc.add_paragraph(f"il Sig {nome}")
            p = doc.add_paragraph("assente giustificato il Sig […] il quale ha tuttavia rilasciato apposita dichiarazione scritta, conservata agli atti della Società, dalla quale risulta che il medesimo è stato informato su tutti gli argomenti posti all'ordine del giorno e che lo stesso non si oppone alla trattazione degli stessi]")
        
        # Collegio sindacale se presente
        if data.get('include_collegio_sindacale', False):
            p = doc.add_paragraph("[eventualmente")
            p = doc.add_paragraph("per il Collegio Sindacale")
            p = doc.add_paragraph("il Dott. […]")
            p = doc.add_paragraph("il Dott. […]")
            p = doc.add_paragraph("il Dott. […]")
            p = doc.add_paragraph("[oppure")
            p = doc.add_paragraph("il Sindaco Unico nella persona del Sig. […]]]")
        
        # Revisore se presente
        if data.get('include_revisore', False):
            p = doc.add_paragraph("[eventualmente, se invitato")
            p = doc.add_paragraph("il revisore contabile Dott. […] <oppure il dott. […] in rappresentanza della società di revisione incaricata del controllo contabile>")
        
        # Soci
        p = doc.add_paragraph("nonché i seguenti soci o loro rappresentanti, [eventualmente così come iscritti a libro soci e] recanti complessivamente una quota pari a nominali euro […] pari al […]% del Capitale Sociale:")
        
        soci = data.get('soci', [])
        for socio in soci:
            if isinstance(socio, dict):
                nome = socio.get('nome', '[Nome Socio]')
                quota_value = socio.get('quota_euro', '')
                percentuale_value = socio.get('quota_percentuale', '')
                
                # Gestione robusta dei valori nulli o vuoti
                quota = '[Quota]' if quota_value is None or str(quota_value).strip() == '' else str(quota_value).strip()
                percentuale = '[%]' if percentuale_value is None or str(percentuale_value).strip() == '' else str(percentuale_value).strip()
                
                tipo_partecipazione = socio.get('tipo_partecipazione', 'Diretta')
                
                if tipo_partecipazione == 'Delegato':
                    delegato = socio.get('delegato', '[Delegato]')
                    p = doc.add_paragraph(f"il Sig. {delegato} delegato del socio Sig {nome} recante una quota pari a nominali euro {quota} pari al {percentuale}% del Capitale Sociale")
                else:
                    p = doc.add_paragraph(f"il Sig. {nome} socio recante una quota pari a nominali euro {quota} pari al {percentuale}% del Capitale Sociale")
        
        # Verifica legittimazione
        p = doc.add_paragraph("2 - che gli intervenuti sono legittimati alla presente assemblea;")
        p = doc.add_paragraph("3 - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno.")
    
    def _add_preliminary_statements(self, doc, data):
        """Aggiunge le dichiarazioni preliminari"""
        doc.add_paragraph()
        
        segretario = data.get('segretario', '[SEGRETARIO]')
        p = doc.add_paragraph(f"I presenti all'unanimità chiamano a fungere da segretario il signor {segretario}, che accetta l'incarico.")
        
        p = doc.add_paragraph("Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante.")
        
        # Traduzione se necessaria
        if data.get('include_traduzione', False):
            persona_traduzione = data.get('persona_traduzione', '[Nome]')
            lingua = data.get('lingua_traduzione', 'inglese').lower()
            p = doc.add_paragraph(f"[eventualmente In particolare, preso atto che il Sig. {persona_traduzione} non conosce la lingua italiana ma dichiara di conoscere la lingua {lingua}, il Presidente dichiara che provvederà a tradurre dall'italiano al {lingua} (e viceversa) gli interventi dei partecipanti alla discussione nonché a tradurre dall'italiano al {lingua} il verbale che sarà redatto al termine della riunione.]")
        
        p = doc.add_paragraph("Il Presidente constata e fa constatare che l'assemblea risulta regolarmente convocata [oppure totalitaria] e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.")
        
        p = doc.add_paragraph("Si passa quindi allo svolgimento dell'ordine del giorno.")
        
        # Separatore
        p = doc.add_paragraph("*     *     *")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_ratifica_discussion(self, doc, data):
        """Aggiunge la discussione per la ratifica"""
        doc.add_paragraph()
        
        p = doc.add_paragraph("Su invito a parlare del Presidente, il socio conferma, come già informalmente anticipato, l'opportunità e il desiderio di procedere, per quanto occorrer possa, alla ratifica delle operazioni poste in essere e degli atti e attività (finanche di tipo omissivo), anche impugnabili a qualsivoglia culpa in vigilando, compiuti dagli Amministratori uscenti per scadenza naturale del mandato, dalla data della loro nomina sino alla data odierna, e per i precedenti mandati, ed al relativo scarico di responsabilità, con rinuncia incondizionata e irrevocabile, nei limiti massimi consentiti dalla legge, all'esercizio di azioni di responsabilità e/o di risarcimento danno nei loro confronti, ed in particolare con riferimento alle seguenti attività (Attività Specifiche), ivi incluse a titolo esemplificativo ma non esaustivo:")
        
        # Attività specifiche
        attivita_specifiche = data.get('attivita_specifiche', '[Attività specifiche da ratificare]')
        if attivita_specifiche and attivita_specifiche.strip():
            # Dividi per righe e aggiungi ogni riga come paragrafo
            righe = attivita_specifiche.split('\n')
            for riga in righe:
                if riga.strip():
                    p = doc.add_paragraph(riga.strip())
        else:
            p = doc.add_paragraph("[Inserire le attività specifiche]")
    
    def _add_deliberation_section(self, doc, data):
        """Aggiunge la sezione di deliberazione"""
        doc.add_paragraph()
        
        # Votazione
        if data.get('votazione_unanime', True):
            votazione_text = "all'unanimità"
        else:
            voti_contrari = data.get('voti_contrari', '')
            astensioni = data.get('astensioni', '')
            votazione_text = f"con il voto contrario dei Sigg. {voti_contrari}"
            if astensioni:
                votazione_text += f" e l'astensione dei Sigg. {astensioni}"
        
        p = doc.add_paragraph(f"Si passa quindi alla votazione con voto palese in forza della quale il Presidente constata che, {votazione_text}, l'assemblea")
        
        p = doc.add_paragraph("d e l i b e r a:")
        run = p.runs[0]
        run.font.bold = True
        run.font.underline = True
        
        # Data scadenza mandato
        data_scadenza = data.get('data_scadenza_mandato')
        if hasattr(data_scadenza, 'strftime'):
            data_scadenza_str = data_scadenza.strftime('%d/%m/%Y')
        else:
            data_scadenza_str = '[Data scadenza]'
        
        p = doc.add_paragraph(f"di ratificare l'operato di tutti gli Amministratori della Società uscenti per scadenza naturale del mandato in data {data_scadenza_str}, dando loro ampio scarico in merito alle operazioni, atti e attività (finanche di tipo omissivo), in particolare con riferimento alle Attività Specifiche, e anche imputabili a qualsivoglia culpa in vigilando, da loro compiuti in qualità di Amministratori (ivi inclusi in qualità di Presidente o Amministratore Delegato) dalla data della loro nomina sino alla data {data_scadenza_str}")
        
        p = doc.add_paragraph(f"di rinunciare incondizionatamente e irrevocabilmente, nei limiti massimi consentiti dalla legge, all'esercizio delle azioni di responsabilità e/o di risarcimento danno nei confronti di tutti gli Amministratori della Società uscenti per scadenza naturale del mandato in data {data_scadenza_str} in relazione a ogni e qualsiasi operazione, attività, atto, potenziale omissione, circostanza o fatto compiuto, anche imputabile a qualsivoglia culpa in vigilando, connesso o occorso nello svolgimento del loro ufficio come membri del Consiglio di Amministrazione (ivi incluso quale Presidente o Amministratore Delegato), dalla data della loro nomina sino alla data {data_scadenza_str}, e per i precedenti mandati, ivi incluso specificatamente, ogni operazione, attività, atto, potenziale omissione, circostanza o fatto che sia riflesso, riferito, menzionato o comunque desumibile da")
        
        p = doc.add_paragraph("i bilanci e le altre situazioni contabili della Società approvati dall'assemblea o comunque presentati ai soci sino alla data odierna nonché i relativi allegati e documenti collegati,")
        
        p = doc.add_paragraph("i verbali delle adunanze e delle deliberazioni delle assemblee della Società ed i relativi allegati, sino alla data odierna, ivi incluse a titolo esemplificativo ma non esaustivo, le Attività Specifiche, precisando che l'assemblea è pienamente informata a riguardo, avendo i soci potuto prendere visione dei sopra citati documenti esistenti e depositati presso la sede sociale o comunque messi a disposizione del pubblico, precisando, inoltre, che ai fini della presente rinuncia si intenderà riflesso nei bilanci o nelle altre situazioni contabili della Società ogni atto gestionale che abbia concorso alla formazione del risultato di esercizio negli anni di riferimento, senza che assumano rilievo i criteri contabili sottesi alla iscrizione a bilancio;")
        
        # Sanzioni e spese se incluse
        if data.get('include_sanzioni_spese', True):
            p = doc.add_paragraph("che eventuali sanzioni e/o spese legali, sia in sede civile che penale, amministrativa o tributaria, che i Consiglieri uscenti fossero chiamati a sostenere in conseguenza della carica ricoperta fino alla data odierna, restino a carico della Società, che, pertanto, ne sosterrà integralmente l'onere, con la sola eccezione di quelle conseguenti a fatti riferibili a loro colpa grave o dolo accertato con sentenza passata in giudicato.")
        
        # Separatore
        p = doc.add_paragraph("*     *     *")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_closing_section(self, doc, data):
        """Aggiunge la sezione di chiusura"""
        doc.add_paragraph()
        
        p = doc.add_paragraph("Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.")
        
        p = doc.add_paragraph("Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo [eventualmente unitamente a quanto allegato].")
        
        ora_chiusura = data.get('ora_chiusura', '[ORA]')
        if hasattr(ora_chiusura, 'strftime'):
            ora_str = ora_chiusura.strftime('%H:%M')
        else:
            ora_str = '[ORA]'
        
        p = doc.add_paragraph(f"L'assemblea viene sciolta alle ore {ora_str}.")
    
    def _add_signatures(self, doc, data):
        """Aggiunge le firme"""
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Usa la tabella di firme standardizzata
        table = self._add_signature_table(doc, data)
        
        # Centra la tabella
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Registra il template
DocumentTemplateFactory.register_template('verbale_assemblea_ratifica_operato', VerbaleRatificaOperatoTemplate)