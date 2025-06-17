"""
Template per Verbale di Assemblea - Nomina Amministratori
Questo template è specifico per verbali di nomina degli amministratori della società.
"""

import sys
import os

# Aggiungi il path della cartella src (relativo alla root del progetto)
current_dir = os.path.dirname(os.path.abspath(__file__))
src_path = os.path.join(os.path.dirname(current_dir), 'src')
if src_path not in sys.path:
    sys.path.append(src_path)

from document_templates import DocumentTemplate, DocumentTemplateFactory
from common_data_handler import CommonDataHandler
from base_verbale_template import BaseVerbaleTemplate
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date
import streamlit as st
import pandas as pd
import re

class VerbaleNominaAmministratoriTemplate(BaseVerbaleTemplate):
    """Template per Verbale di Assemblea dei Soci - Nomina Amministratori"""
    
    def get_template_name(self) -> str:
        return "Nomina Amministratori"
    
    def get_required_fields(self) -> list:
        return [
            "denominazione", "sede_legale", "capitale_sociale", "codice_fiscale",
            "data_assemblea", "ora_assemblea", "presidente", "segretario", 
            "soci", "amministratori", "nuovi_amministratori"
        ]
    
    def get_form_fields(self, extracted_data: dict) -> dict:
        """Genera i campi del form per Streamlit utilizzando il CommonDataHandler."""
        # Chiama il metodo della classe base per ottenere i campi comuni
        form_data = super().get_form_fields(extracted_data)

        # Campi specifici per nomina amministratori
        st.subheader("👥 Configurazioni Specifiche Nomina Amministratori")
        
        col1, col2 = st.columns(2)
        with col1:
            form_data["ruolo_presidente"] = st.selectbox("Ruolo del presidente", 
                                                        CommonDataHandler.get_standard_ruoli_presidente(),
                                                        key="ruolo_presidente_nomina")
            form_data["motivo_nomina"] = st.selectbox("Motivo della nomina", 
                                                     ["Dimissioni dell'organo in carica", 
                                                      "Scadenza mandato", 
                                                      "Decadenza dall'ufficio",
                                                      "Prima nomina",
                                                      "Altro"],
                                                      key="motivo_nomina")
        with col2:
            form_data["tipo_amministrazione"] = st.selectbox("Tipo di amministrazione", 
                                                            ["Amministrazione disgiunta", 
                                                             "Amministrazione congiunta",
                                                             "Amministratore Unico"],
                                                             key="tipo_amministrazione_nomina")
            form_data["durata_incarico"] = st.selectbox("Durata incarico", 
                                                       ["A tempo indeterminato fino a revoca o dimissioni",
                                                        "Tre esercizi", 
                                                        "Un esercizio", 
                                                        "Altra durata"],
                                                        key="durata_incarico_nomina")
        
        # Ordine del giorno specifico per nomina amministratori
        st.subheader("📋 Ordine del Giorno")
        form_data["include_compensi"] = st.checkbox("Includi attribuzione compensi", value=True, key="include_compensi_nomina")
        
        # Nuovo sezione per amministratori nominandi
        st.subheader("👥 Nuovi Amministratori da Nominare")
        
        # Numero di amministratori
        if form_data.get("tipo_amministrazione") == "Amministratore Unico":
            num_admin = 1
            st.info("ℹ️ Amministratore Unico: verrà nominato 1 amministratore")
        else:
            num_admin = st.number_input("Numero di amministratori da nominare", min_value=1, max_value=10, value=2, key="num_admin_nomina")
        
        form_data["nuovi_amministratori"] = []
        
        for i in range(num_admin):
            st.markdown(f"**👤 Amministratore {i+1}**")
            col1, col2 = st.columns(2)
            
            with col1:
                nome = st.text_input(f"Nome e Cognome", key=f"admin_nome_{i}_nomina",
                                   placeholder="es. Mario Rossi")
                data_nascita = st.date_input(f"Data di nascita", key=f"admin_data_nascita_{i}_nomina",
                                           value=None)
                luogo_nascita = st.text_input(f"Luogo di nascita", key=f"admin_luogo_nascita_{i}_nomina",
                                            placeholder="es. Roma (RM)")
            
            with col2:
                codice_fiscale = st.text_input(f"Codice Fiscale", key=f"admin_cf_{i}_nomina",
                                             placeholder="es. RSSMRA80A01H501Z")
                residenza = st.text_input(f"Residenza", key=f"admin_residenza_{i}_nomina",
                                        placeholder="es. Via Roma 1, Milano (MI)")
                qualifica = st.selectbox(f"Qualifica nella società", 
                                       ["Socio", "Amministratore uscente", "Terzo"], 
                                       key=f"admin_qualifica_{i}_nomina")
            
            form_data["nuovi_amministratori"].append({
                "nome": nome,
                "data_nascita": data_nascita,
                "luogo_nascita": luogo_nascita,
                "codice_fiscale": codice_fiscale,
                "residenza": residenza,
                "qualifica": qualifica
            })
        
        # Compensi
        if form_data.get("include_compensi"):
            st.subheader("💰 Compensi Amministratori")
            col1, col2 = st.columns(2)
            with col1:
                form_data["compenso_annuo"] = st.text_input("Compenso annuo (€)", 
                                                           value="0,00",
                                                           help="Compenso annuo lordo per ciascun amministratore",
                                                           key="compenso_annuo_nomina")
                form_data["tipo_compenso"] = st.selectbox("Tipo di compenso", 
                                                         ["Fisso", "Gettoni di presenza", "Percentuale utili", "Misto"],
                                                         key="tipo_compenso_nomina")
            with col2:
                form_data["rimborso_spese"] = st.checkbox("Rimborso spese", value=True, key="rimborso_spese_nomina")
                form_data["modalita_liquidazione"] = st.selectbox("Modalità liquidazione", 
                                                                 ["Periodicamente", "Annuale", "A fine mandato"],
                                                                 key="modalita_liquidazione_nomina")
        
        # Verifiche e dichiarazioni
        st.subheader("📋 Verifiche e Dichiarazioni")
        form_data["verifica_requisiti"] = st.checkbox("Verifica requisiti di eleggibilità completata", value=True, key="verifica_requisiti_nomina")
        form_data["dichiarazioni_ricevute"] = st.checkbox("Dichiarazioni degli amministratori ricevute", value=True, key="dichiarazioni_ricevute_nomina")
        form_data["verifica_incompatibilita"] = st.checkbox("Verifica incompatibilità effettuata", value=True, key="verifica_incompatibilita_nomina")
        
        # Opzioni redazionali
        st.subheader("⚙️ Opzioni Redazionali")
        # I campi assemblea_totalitaria e votazione_unanimita vengono gestiti dalla base,
        # qui potremmo solo personalizzare il comportamento se necessario.
        
        if not form_data.get("votazione_unanimita", True):
            form_data["soci_contrari"] = st.text_input("Soci contrari (separati da virgola)", placeholder="es. Mario Rossi, Anna Verdi", key="soci_contrari_nomina")
            form_data["soci_astenuti"] = st.text_input("Soci astenuti (separati da virgola)", key="soci_astenuti_nomina")
        else:
            form_data["soci_contrari"] = ""
            form_data["soci_astenuti"] = ""
        
        # Note aggiuntive
        st.subheader("📝 Note Aggiuntive")
        form_data["note_aggiuntive"] = st.text_area("Note aggiuntive", 
                                                  placeholder="Eventuali note specifiche sulla nomina...",
                                                  height=80,
                                                  key="note_aggiuntive_nomina")
        
        return form_data
    
    def show_preview(self, form_data: dict):
        """Mostra l'anteprima del documento fuori dal form"""
        # Sezione Anteprima
        st.markdown("---")
        st.subheader("👁️ Anteprima Documento")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col1:
            show_preview = st.checkbox("Mostra anteprima", value=False, key="preview_checkbox_nomina_admin")
        
        with col2:
            if show_preview:
                st.info("💡 L'anteprima si aggiorna automaticamente con i dati inseriti sopra")
        
        if show_preview:
            try:
                # Debug dettagliato: mostra i dati estratti
                with st.expander("Debug: Dati estratti dal form"):
                    st.json(form_data)
                    
                    # Debug specifico per i soci
                    if 'soci' in form_data and form_data['soci']:
                        st.subheader("Debug Soci:")
                        for i, socio in enumerate(form_data['soci']):
                            st.write(f"Socio {i+1}:")
                            st.write(f"  - Nome: {socio.get('nome', 'N/A')}")
                            st.write(f"  - Tipo Soggetto: {socio.get('tipo_soggetto', 'N/A')}")
                            st.write(f"  - Tipo Partecipazione: {socio.get('tipo_partecipazione', 'N/A')}")
                            st.write(f"  - Delegato: {socio.get('delegato', 'N/A')}")
                            st.write(f"  - Rappresentante Legale: {socio.get('rappresentante_legale', 'N/A')}")
                            st.write(f"  - Presente: {socio.get('presente', 'N/A')}")
                            st.write("---")
                
                preview_text = self._generate_preview_text(form_data)
                st.text_area(
                    "Contenuto del verbale:",
                    value=preview_text,
                    height=600,
                    key="preview_text_nomina_admin"
                )
            except Exception as e:
                st.error(f"Errore nell'anteprima: {e}")
                st.exception(e)
    
    def _generate_preview_text(self, data: dict) -> str:
        """Genera il testo di anteprima del verbale"""
        try:
            # Header dell'azienda
            header = f"""{data.get('denominazione', '[Denominazione]')}
Sede in {data.get('sede_legale', '[Sede]')}
Capitale sociale Euro {data.get('capitale_sociale', '[Capitale]')} i.v.
Codice fiscale: {data.get('codice_fiscale', '[CF]')}

Verbale di assemblea dei soci
del {data.get('data_assemblea', '[Data]').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[Data]'}

Oggi {data.get('data_assemblea', '[Data]').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[Data]'} alle ore {data.get('ora_inizio', '[Ora]').strftime('%H:%M') if hasattr(data.get('ora_inizio'), 'strftime') else '[Ora]'} presso la sede sociale {data.get('sede_legale', '[Sede]')}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:

Ordine del giorno
1. nomina degli amministratori della società"""
            
            if data.get('include_compensi', True):
                header += "\n2. attribuzione di compensi agli amministratori della società"
            
            # Sezione presidenza
            presidente_section = f"""

Assume la presidenza ai sensi dell'art. […] dello statuto sociale il Sig. {data.get('presidente', '[Presidente]')} {data.get('ruolo_presidente', 'Amministratore Unico')}, il quale dichiara e constata:

1 - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. […] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza

2 - che sono presenti/partecipano all'assemblea:
l'{data.get('ruolo_presidente', 'Amministratore Unico')} nella persona del suddetto Presidente Sig. {data.get('presidente', '[Presidente]')}"""
            
            # Aggiunge amministratori se presenti
            amministratori = data.get('amministratori', [])
            if amministratori and len(amministratori) > 1:
                presidente_section += "\nPer il Consiglio di Amministrazione:"
                for admin in amministratori:
                    nome_admin = admin.get('nome', '') if isinstance(admin, dict) else str(admin)
                    presidente_section += f"\nil Sig {nome_admin}"
            
            # Collegio sindacale se presente
            if data.get('include_collegio_sindacale', False):
                presidente_section += "\nPer il Collegio Sindacale:\nil Dott. [Nome]\nil Dott. [Nome]\nil Dott. [Nome]]"
            
            # Soci: gestione uniforme presenti/assenti con calcolo quote complessive
            soci_presenti = data.get('soci_presenti', [])
            soci_assenti = data.get('soci_assenti', [])

            # Sezione soci, inizializzata vuota per garantire la definizione anche in assenza di soci presenti
            soci_section = ""

            if soci_presenti:
                total_quota_euro = 0.0
                total_quota_percentuale = 0.0

                for socio in soci_presenti:
                    try:
                        total_quota_euro += float(str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.'))
                    except (ValueError, TypeError):
                        pass
                    try:
                        total_quota_percentuale += float(str(socio.get('quota_percentuale', '0')).replace('.', '').replace(',', '.'))
                    except (ValueError, TypeError):
                        pass

                formatted_euro = f"{total_quota_euro:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
                formatted_perc = f"{total_quota_percentuale:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

                soci_section = f"\nnonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_euro} pari al {formatted_perc}% del Capitale Sociale:"

                for socio in soci_presenti:
                    nome = socio.get('nome', '[Nome Socio]')
                    quota_raw = socio.get('quota_euro', '')
                    perc_raw = socio.get('quota_percentuale', '')
                    quota = '[Quota]' if not quota_raw or str(quota_raw).strip() == '' else str(quota_raw).strip()
                    perc = '[%]' if not perc_raw or str(perc_raw).strip() == '' else str(perc_raw).strip()

                    tipo_partecipazione = socio.get('tipo_partecipazione', 'Diretto')
                    delegato = socio.get('delegato', '').strip()
                    tipo_soggetto = socio.get('tipo_soggetto', 'Persona Fisica')
                    rappresentante_legale = socio.get('rappresentante_legale', '').strip()

                    if tipo_partecipazione == 'Delegato' and delegato:
                        if tipo_soggetto == 'Società':
                            line = f"il Sig. {delegato} delegato della società {nome}"
                            if rappresentante_legale:
                                line += f" (nella persona del legale rappresentante {rappresentante_legale})"
                        else:
                            line = f"il Sig. {delegato} delegato del socio {nome}"
                        line += f" recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"
                    else:
                        if tipo_soggetto == 'Società':
                            if rappresentante_legale:
                                line = f"la società {nome} nella persona del legale rappresentante {rappresentante_legale} recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"
                            else:
                                line = f"la società {nome} recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"
                        else:
                            line = f"il Sig. {nome} socio recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"

                    soci_section += f"\n{line}"

            if soci_assenti:
                soci_section += "\nRisultano invece assenti i seguenti soci:"
                for socio in soci_assenti:
                    nome = socio.get('nome', '[Nome Socio]')
                    soci_section += f"\n- Sig. {nome}"
            
            soci_section += "\n2 - che gli intervenuti sono legittimati alla presente assemblea;\n3 - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno."
            
            # Determina la formula con cui descrivere la validità di convocazione
            convocata_phrase = "risulta totalitaria" if data.get('assemblea_totalitaria') else "risulta regolarmente convocata"

            # Segretario
            segretario_section = f"""

I presenti all'unanimità chiamano a fungere da segretario il signor {data.get('segretario', '[Segretario]')}, che accetta l'incarico.

Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante.

Il Presidente constata e fa constatare che l'assemblea {convocata_phrase} e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.

Si passa quindi allo svolgimento dell'ordine del giorno.

*     *     *"""
            
            # Discussione nomina amministratori
            motivo_nomina = data.get('motivo_nomina', 'Dimissioni dell\'organo in carica')
            
            nomina_section = f"""

Il Presidente informa l'assemblea che si rende necessaria la nomina di un nuovo organo amministrativo [{motivo_nomina.lower()}].

Il Presidente ricorda all'assemblea quanto previsto dall'art. 2475 del Codice Civile e dall'atto costitutivo della società [verificare quanto previsto dall'atto costitutivo in tema di amministrazione disgiunta].

Prende la parola il socio sig. […] che propone di affidare l'{data.get('tipo_amministrazione', 'amministrazione disgiunta').lower()} della società ai Sigg.:\n"""
            
            # Lista nuovi amministratori
            nuovi_admin = data.get('nuovi_amministratori', [])
            for admin in nuovi_admin:
                nome = admin.get('nome', '[Nome]')
                data_nascita = admin.get('data_nascita', '[Data nascita]')
                luogo_nascita = admin.get('luogo_nascita', '[Luogo nascita]')
                codice_fiscale = admin.get('codice_fiscale', '[CF]')
                residenza = admin.get('residenza', '[Residenza]')
                
                if hasattr(data_nascita, 'strftime'):
                    data_nascita_str = data_nascita.strftime('%d/%m/%Y')
                else:
                    data_nascita_str = str(data_nascita) if data_nascita else '[Data nascita]'
                
                nomina_section += f"- {nome}, nato a {luogo_nascita} il {data_nascita_str}, C.F. {codice_fiscale}, residente in {residenza}\n"
            
            nomina_section += """\ndando evidenza della comunicazione scritta con cui i candidati, prima di accettare l'eventuale nomina, hanno dichiarato:
- l'insussistenza a loro carico di cause di ineleggibilità alla carica di amministratore di società ed in particolare di non essere stati dichiarati interdetti, inabilitati o falliti e di non essere stati condannati ad una pena che importa l'interdizione, anche temporanea, dai pubblici uffici o l'incapacità ad esercitare uffici direttivi.
- l'insussistenza a loro carico di interdizioni dal ruolo di amministratore adottate da una Stato membro dell'Unione Europea.

[verificare che l'atto costitutivo non preveda ulteriori requisiti per l'assunzione della carica e quanto previsto da leggi speciali in relazione all'esercizio di particolari attività] [se esiste il collegio sindacale o il revisore, verificare eventuali incompatibilità con i neo amministratori]."""
            
            # Frase votazione (unanimita / maggioranza)
            voto_unanimita = data.get('votazione_unanimita', True)
            votazione_phrase = "all'unanimità"
            if not voto_unanimita:
                contrari = data.get('soci_contrari', '').strip()
                astenuti = data.get('soci_astenuti', '').strip()
                parts = []
                if contrari:
                    parts.append(f"con il voto contrario dei Sigg. {contrari}")
                if astenuti:
                    if parts:
                        parts.append(f"e l'astensione dei Sigg. {astenuti}")
                    else:
                        parts.append(f"con l'astensione dei Sigg. {astenuti}")
                votazione_phrase = " ".join(parts) if parts else "a maggioranza"

            # Compensi se inclusi
            compensi_section = ""
            if data.get('include_compensi', True):
                compensi_section = f"""

Il Presidente invita anche l'assemblea a deliberare il compenso da attribuire all'organo amministrativo che verrà nominato, ai sensi dell'art. […] dello statuto sociale."""
            
            # Deliberazione
            durata_incarico = data.get('durata_incarico', 'A tempo indeterminato fino a revoca o dimissioni')
            compenso_annuo = data.get('compenso_annuo', '0,00')
            
            deliberazione_section = f"""

Segue breve discussione tra i soci al termine della quale si passa alla votazione con voto palese in forza della quale il Presidente constata che, all'unanimità, l'assemblea

d e l i b e r a:

di affidare l'{data.get('tipo_amministrazione', 'amministrazione disgiunta').lower()} della società ai Sigg.:"""
            
            for admin in nuovi_admin:
                nome = admin.get('nome', '[Nome]')
                data_nascita = admin.get('data_nascita', '[Data nascita]')
                luogo_nascita = admin.get('luogo_nascita', '[Luogo nascita]')
                codice_fiscale = admin.get('codice_fiscale', '[CF]')
                residenza = admin.get('residenza', '[Residenza]')
                
                if hasattr(data_nascita, 'strftime'):
                    data_nascita_str = data_nascita.strftime('%d/%m/%Y')
                else:
                    data_nascita_str = str(data_nascita) if data_nascita else '[Data nascita]'
                
                deliberazione_section += f"\n- {nome}, nato a {luogo_nascita} il {data_nascita_str}, C.F. {codice_fiscale}, residente in {residenza}"
            
            deliberazione_section += f"\n\nche gli amministratori restino in carica {durata_incarico.lower()} [verificare che l'atto costitutivo non preveda una durata massima per l'incarico]"
            
            if data.get('include_compensi', True):
                rimborso_text = " oltre al rimborso delle spese sostenute in ragione del suo ufficio" if data.get('rimborso_spese', True) else ""
                modalita_liquidazione = data.get('modalita_liquidazione', 'periodicamente')
                
                deliberazione_section += f"\n\ndi attribuire agli amministratori testè nominati il compenso annuo ed omnicomprensivo pari a nominali euro {compenso_annuo} al lordo di ritenute fiscali e previdenziali{rimborso_text}. Il compenso verrà liquidato {modalita_liquidazione.lower()}, in ragione della permanenza in carica."
            
            # Accettazione
            accettazione_section = "\n\n"
            nomi_admin = [admin.get('nome', '[Nome]') for admin in nuovi_admin]
            if len(nomi_admin) == 1:
                accettazione_section += f"Il sig. {nomi_admin[0]}, presente in assemblea in qualità di [{nuovi_admin[0].get('qualifica', 'socio')}] accetta l'incarico e ringrazia l'assemblea per la fiducia accordata."
            else:
                nomi_str = " e ".join(nomi_admin)
                accettazione_section += f"I sigg. {nomi_str}, presenti in assemblea in qualità di [indicare (socio, amministratore uscente, invitato o altro)] accettano l'incarico e ringraziano l'assemblea per la fiducia accordata."
            
            # Chiusura
            chiusura_section = f"""

*     *     *

Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.

Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo [eventualmente unitamente a quanto allegato].

L'assemblea viene sciolta alle ore [{data.get('ora_chiusura', '[Ora]')}].


Il Presidente                    Il Segretario
{data.get('presidente', '[Presidente]')}            {data.get('segretario', '[Segretario]')}"""
            
            # Combina tutte le sezioni
            full_text = (header + presidente_section + soci_section + segretario_section + 
                        nomina_section + compensi_section + deliberazione_section + 
                        accettazione_section + chiusura_section)
            
            return full_text
            
        except Exception as e:
            return f"Errore nella generazione dell'anteprima: {str(e)}"
    
    def generate_document(self, data: dict) -> Document:
        """Genera il documento Word del verbale"""
        import os
        
        # Utilizza il template .docx esistente per mantenere la formattazione
        template_path = os.path.join(os.path.dirname(__file__), 'template.docx')
        
        try:
            if os.path.exists(template_path):
                doc = Document(template_path)
                # Rimuovi il contenuto esistente del template mantenendo gli stili
                for paragraph in doc.paragraphs[:]:
                    p = paragraph._element
                    p.getparent().remove(p)
            else:
                doc = Document()
                # Setup stili solo se non usiamo template
                self._setup_document_styles(doc)
        except Exception as e:
            # Fallback a documento vuoto se il template non può essere caricato
            doc = Document()
            self._setup_document_styles(doc)
        
        # Setup stili del documento
        self._setup_document_styles(doc)
        
        # Aggiungi header azienda
        self._add_company_header(doc, data)
        
        # Aggiungi titolo verbale
        self._add_verbale_title(doc, data)
        
        # Aggiungi sezione di apertura
        self._add_opening_section(doc, data)
        
        # Aggiungi partecipanti
        self._add_participants_section(doc, data)
        
        # Aggiungi dichiarazioni preliminari
        self._add_preliminary_statements(doc, data)
        
        # Aggiungi discussione nomina amministratori
        self._add_nomination_discussion(doc, data)
        
        # Aggiungi sezione di chiusura
        self._add_closing_section(doc, data)
        
        # Aggiungi firme
        self._add_signatures(doc, data)
        
        return doc
    
    def _setup_document_styles(self, doc):
        """Imposta gli stili del documento"""
        super()._setup_document_styles(doc) # Call base class method first
        styles = doc.styles
        
        # Stile per il titolo principale (aggiuntivo o sovrascritto)
        if 'Titolo Principale' not in [s.name for s in styles]:
            title_style = styles.add_style('Titolo Principale', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Times New Roman'
            title_style.font.size = Pt(14)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # Stile per il testo normale
        normal_style = styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal_style.paragraph_format.line_spacing = 1.15
    
    def _add_company_header(self, doc, data):
        """Aggiunge l'header dell'azienda"""
        # Nome azienda
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data.get('denominazione', '[DENOMINAZIONE SOCIETÀ]'))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Sede
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Sede in {data.get('sede_legale', '[SEDE]')}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        # Capitale sociale
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Capitale sociale Euro {data.get('capitale_sociale', '[CAPITALE]')} i.v.")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        # Codice fiscale
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Codice fiscale: {data.get('codice_fiscale', '[CODICE FISCALE]')}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        doc.add_paragraph()
    
    def _add_verbale_title(self, doc, data):
        """Aggiunge il titolo del verbale"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Verbale di assemblea dei soci")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Data
        data_assemblea = data.get('data_assemblea')
        if hasattr(data_assemblea, 'strftime'):
            data_str = data_assemblea.strftime('%d/%m/%Y')
        else:
            data_str = '[DATA]'
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"del {data_str}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        doc.add_paragraph()
    
    def _add_opening_section(self, doc, data):
        """Aggiunge la sezione di apertura"""
        # Data e ora
        data_assemblea = data.get('data_assemblea')
        ora_assemblea = data.get('ora_assemblea')
        
        if hasattr(data_assemblea, 'strftime'):
            data_str = data_assemblea.strftime('%d/%m/%Y')
        else:
            data_str = '[DATA]'
        
        if hasattr(ora_assemblea, 'strftime'):
            ora_str = ora_assemblea.strftime('%H:%M')
        else:
            ora_str = '[ORA]'
        
        text = f"Oggi {data_str} alle ore {ora_str} presso la sede sociale {data.get('sede_legale', '[SEDE]')}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:"
        
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Ordine del giorno
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Ordine del giorno")
        run.font.bold = True
        
        p = doc.add_paragraph("1. nomina degli amministratori della società")
        
        if data.get('include_compensi', True):
            p = doc.add_paragraph("2. attribuzione di compensi agli amministratori della società")
    
    def _add_participants_section(self, doc, data):
        """Aggiunge la sezione dei partecipanti"""
        doc.add_paragraph()
        
        ruolo_presidente = data.get('ruolo_presidente', 'Amministratore Unico')
        presidente = data.get('presidente', '[PRESIDENTE]')
        
        text = f"Assume la presidenza ai sensi dell'art. […] dello statuto sociale il Sig. {presidente} {ruolo_presidente}, il quale dichiara e constata:"
        
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Punto 1 - audioconferenza
        p = doc.add_paragraph("1 - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. […] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza")
        
        # Punto 2 - presenti
        p = doc.add_paragraph("2 - che sono presenti/partecipano all'assemblea:")
        p = doc.add_paragraph(f"l'{ruolo_presidente} nella persona del suddetto Presidente Sig. {presidente}")
        
        # Soci: gestione uniforme presenti/assenti con calcolo quote complessive
        soci_presenti = data.get('soci_presenti', [])
        soci_assenti = data.get('soci_assenti', [])

        # Fallback retro-compatibilità: se la vecchia chiave 'soci' è presente e le liste non sono valorizzate
        if not soci_presenti and not soci_assenti and 'soci' in data:
            soci_presenti = [s for s in data.get('soci', []) if s.get('presente', True)]
            soci_assenti = [s for s in data.get('soci', []) if not s.get('presente', True)]

        if soci_presenti:
            # Inizializza la variabile soci_section per accumulare il testo dei soci
            soci_section = ""
            total_quota_euro = 0.0
            total_quota_percentuale = 0.0

            for socio in soci_presenti:
                try:
                    total_quota_euro += float(str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.'))
                except (ValueError, TypeError):
                    pass
                try:
                    total_quota_percentuale += float(str(socio.get('quota_percentuale', '0')).replace('.', '').replace(',', '.'))
                except (ValueError, TypeError):
                    pass

            formatted_euro = f"{total_quota_euro:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            formatted_perc = f"{total_quota_percentuale:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

            p = doc.add_paragraph(f"nonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_euro} pari al {formatted_perc}% del Capitale Sociale:")

            for socio in soci_presenti:
                nome = socio.get('nome', '[Nome Socio]')
                quota_raw = socio.get('quota_euro', '')
                perc_raw = socio.get('quota_percentuale', '')
                quota = '[Quota]' if not quota_raw or str(quota_raw).strip() == '' else str(quota_raw).strip()
                perc = '[%]' if not perc_raw or str(perc_raw).strip() == '' else str(perc_raw).strip()

                tipo_partecipazione = socio.get('tipo_partecipazione', 'Diretto')
                delegato = socio.get('delegato', '').strip()
                tipo_soggetto = socio.get('tipo_soggetto', 'Persona Fisica')
                rappresentante_legale = socio.get('rappresentante_legale', '').strip()

                if tipo_partecipazione == 'Delegato' and delegato:
                    if tipo_soggetto == 'Società':
                        line = f"il Sig. {delegato} delegato della società {nome}"
                        if rappresentante_legale:
                            line += f" (nella persona del legale rappresentante {rappresentante_legale})"
                    else:
                        line = f"il Sig. {delegato} delegato del socio {nome}"
                    line += f" recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"
                else:
                    if tipo_soggetto == 'Società':
                        if rappresentante_legale:
                            line = f"la società {nome} nella persona del legale rappresentante {rappresentante_legale} recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"
                        else:
                            line = f"la società {nome} recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"
                    else:
                        line = f"il Sig. {nome} socio recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"

                soci_section += f"\n{line}"

        if soci_assenti:
            soci_section += "\nRisultano invece assenti i seguenti soci:"
            for socio in soci_assenti:
                nome = socio.get('nome', '[Nome Socio]')
                soci_section += f"\n- Sig. {nome}"
        
        # Aggiungi il contenuto della sezione soci al documento se è stata costruita
        if 'soci_section' in locals() and soci_section:
            p = doc.add_paragraph(soci_section.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Verifica legittimazione
        p = doc.add_paragraph("2 - che gli intervenuti sono legittimati alla presente assemblea;")
        p = doc.add_paragraph("3 - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno.")
    
    def _add_preliminary_statements(self, doc, data):
        """Aggiunge le dichiarazioni preliminari"""
        doc.add_paragraph()
        
        segretario = data.get('segretario', '[SEGRETARIO]')
        p = doc.add_paragraph(f"I presenti all'unanimità chiamano a fungere da segretario il signor {segretario}, che accetta l'incarico.")
        
        p = doc.add_paragraph("Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante.")
        
        p = doc.add_paragraph("Il Presidente constata e fa constatare che l'assemblea risulta regolarmente convocata e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.")
        
        p = doc.add_paragraph("Si passa quindi allo svolgimento dell'ordine del giorno.")
        
        # Separatore
        p = doc.add_paragraph("*     *     *")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_nomination_discussion(self, doc, data):
        """Aggiunge la discussione per la nomina degli amministratori"""
        doc.add_paragraph()
        
        motivo_nomina = data.get('motivo_nomina', 'Dimissioni dell\'organo in carica')
        p = doc.add_paragraph(f"Il Presidente informa l'assemblea che si rende necessaria la nomina di un nuovo organo amministrativo [{motivo_nomina.lower()}].")
        
        p = doc.add_paragraph("Il Presidente ricorda all'assemblea quanto previsto dall'art. 2475 del Codice Civile e dall'atto costitutivo della società [verificare quanto previsto dall'atto costitutivo in tema di amministrazione disgiunta].")
        
        tipo_admin = data.get('tipo_amministrazione', 'amministrazione disgiunta')
        p = doc.add_paragraph(f"Prende la parola il socio sig. […] che propone di affidare l'{tipo_admin.lower()} della società ai Sigg.:")
        
        # Lista amministratori nominandi
        nuovi_admin = data.get('nuovi_amministratori', [])
        for admin in nuovi_admin:
            nome = admin.get('nome', '[Nome]')
            data_nascita = admin.get('data_nascita', '[Data nascita]')
            luogo_nascita = admin.get('luogo_nascita', '[Luogo nascita]')
            codice_fiscale = admin.get('codice_fiscale', '[CF]')
            residenza = admin.get('residenza', '[Residenza]')
            
            if hasattr(data_nascita, 'strftime'):
                data_nascita_str = data_nascita.strftime('%d/%m/%Y')
            else:
                data_nascita_str = '[Data nascita]'
            
            p = doc.add_paragraph(f"{nome}, nato a {luogo_nascita} il {data_nascita_str}, C.F. {codice_fiscale}, residente in {residenza}")
        
        # Dichiarazioni
        p = doc.add_paragraph("dando evidenza della comunicazione scritta con cui i candidati, prima di accettare l'eventuale nomina, hanno dichiarato:")
        p = doc.add_paragraph("l'insussistenza a loro carico di cause di ineleggibilità alla carica di amministratore di società ed in particolare di non essere stati dichiarati interdetti, inabilitati o falliti e di non essere stati condannati ad una pena che importa l'interdizione, anche temporanea, dai pubblici uffici o l'incapacità ad esercitare uffici direttivi.")
        p = doc.add_paragraph("l'insussistenza a loro carico di interdizioni dal ruolo di amministratore adottate da una Stato membro dell'Unione Europea.")
        p = doc.add_paragraph("[verificare che l'atto costitutivo non preveda ulteriori requisiti per l'assunzione della carica e quanto previsto da leggi speciali in relazione all'esercizio di particolari attività] [se esiste il collegio sindacale o il revisore, verificare eventuali incompatibilità con i neo amministratori].")
        
        # Compensi
        if data.get('include_compensi', True):
            p = doc.add_paragraph("Il Presidente invita anche l'assemblea a deliberare il compenso da attribuire all'organo amministrativo che verrà nominato, ai sensi dell'art. […] dello statuto sociale.")
        
        # Deliberazione
        p = doc.add_paragraph("Segue breve discussione tra i soci al termine della quale si passa alla votazione con voto palese in forza della quale il Presidente constata che, all'unanimità, l'assemblea")
        
        p = doc.add_paragraph("d e l i b e r a:")
        run = p.runs[0]
        run.font.bold = True
        run.font.underline = True
        
        p = doc.add_paragraph(f"di affidare l'{tipo_admin.lower()} della società ai Sigg.:")
        
        # Lista finale amministratori
        for admin in nuovi_admin:
            nome = admin.get('nome', '[Nome]')
            data_nascita = admin.get('data_nascita', '[Data nascita]')
            luogo_nascita = admin.get('luogo_nascita', '[Luogo nascita]')
            codice_fiscale = admin.get('codice_fiscale', '[CF]')
            residenza = admin.get('residenza', '[Residenza]')
            
            if hasattr(data_nascita, 'strftime'):
                data_nascita_str = data_nascita.strftime('%d/%m/%Y')
            else:
                data_nascita_str = '[Data nascita]'
            
            p = doc.add_paragraph(f"{nome}, nato a {luogo_nascita} il {data_nascita_str}, C.F. {codice_fiscale}, residente in {residenza}")
        
        # Durata incarico
        durata = data.get('durata_incarico', 'a tempo indeterminato fino a revoca o dimissioni')
        p = doc.add_paragraph(f"che gli amministratori restino in carica {durata.lower()} [verificare che l'atto costitutivo non preveda una durata massima per l'incarico]")
        
        # Compensi dettaglio
        if data.get('include_compensi', True):
            compenso = data.get('compenso_annuo', '0,00')
            rimborso_text = " oltre al rimborso delle spese sostenute in ragione del suo ufficio" if data.get('rimborso_spese', True) else ""
            modalita = data.get('modalita_liquidazione', 'periodicamente')
            
            p = doc.add_paragraph(f"di attribuire agli amministratori testè nominati il compenso annuo ed omnicomprensivo pari a nominali euro {compenso} al lordo di ritenute fiscali e previdenziali{rimborso_text}. Il compenso verrà liquidato {modalita.lower()}, in ragione della permanenza in carica.")
        
        # Accettazione
        nomi_admin = [admin.get('nome', '[Nome]') for admin in nuovi_admin]
        if len(nomi_admin) == 1:
            qualifica = nuovi_admin[0].get('qualifica', 'socio') if nuovi_admin else 'socio'
            p = doc.add_paragraph(f"Il sig. {nomi_admin[0]}, presente in assemblea in qualità di {qualifica.lower()} accetta l'incarico e ringrazia l'assemblea per la fiducia accordata.")
        else:
            nomi_str = " e ".join(nomi_admin)
            p = doc.add_paragraph(f"I sigg. {nomi_str}, presenti in assemblea in qualità di [indicare (socio, amministratore uscente, invitato o altro)] accettano l'incarico e ringraziano l'assemblea per la fiducia accordata.")
        
        # Separatore
        p = doc.add_paragraph("*     *     *")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_closing_section(self, doc, data):
        """Aggiunge la sezione di chiusura"""
        doc.add_paragraph()
        
        p = doc.add_paragraph("Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.")
        
        p = doc.add_paragraph("Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo [eventualmente unitamente a quanto allegato].")
        
        ora_chiusura = data.get('ora_chiusura', '[ORA]')
        if hasattr(ora_chiusura, 'strftime'):
            ora_str = ora_chiusura.strftime('%H:%M')
        else:
            ora_str = '[ORA]'
        
        p = doc.add_paragraph(f"L'assemblea viene sciolta alle ore {ora_str}.")
    
    def _add_signatures(self, doc, data):
        """Aggiunge le firme"""
        # Use the standardized signature logic from BaseVerbaleTemplate
        super()._add_signatures(doc, data)
        
        # Nota: La centratura della tabella, se necessaria, dovrebbe essere gestita
        # all'interno del metodo _add_signatures della classe base o qui se si 
        # sovrascrive e si ricrea una tabella specifica.
        # Dato che stiamo usando il metodo della classe base, la logica di centratura 
        # specifica per una tabella non è più direttamente applicabile qui.

# Registra il template
DocumentTemplateFactory.register_template('verbale_assemblea_nomina_amministratori', VerbaleNominaAmministratoriTemplate)