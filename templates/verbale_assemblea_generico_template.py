"""
Template per Verbale di Assemblea Generico - Verbale Generico
Questo template è per verbali di assemblea generici con ordine del giorno personalizzabile.
"""

import sys
import os

# Aggiungi il path della cartella src (relativo alla root del progetto)
current_dir = os.path.dirname(os.path.abspath(__file__))
src_path = os.path.join(os.path.dirname(current_dir), 'src')
if src_path not in sys.path:
    sys.path.append(src_path)

from document_templates import DocumentTemplate, DocumentTemplateFactory
from common_data_handler import CommonDataHandler
from base_verbale_template import BaseVerbaleTemplate
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date
import streamlit as st
import pandas as pd
import re

class VerbaleAssembleaGenericoTemplate(BaseVerbaleTemplate):
    """Template per Verbale di Assemblea dei Soci - Generico"""
    
    def get_template_name(self) -> str:
        return "Verbale Assemblea Generico"
    
    def get_required_fields(self) -> list:
        return [
            "denominazione", "sede_legale", "capitale_sociale", "codice_fiscale",
            "data_assemblea", "ora_assemblea", "presidente", "segretario", 
            "soci", "amministratori", "ordine_del_giorno"
        ]
    
    def get_form_fields(self, extracted_data: dict) -> dict:
        """Genera i campi del form per Streamlit utilizzando il CommonDataHandler."""
        # Chiama il metodo della classe base per ottenere i campi comuni
        form_data = super().get_form_fields(extracted_data)
        
        # Rimuovi i campi specifici gestiti dalla classe base se necessario,
        # per ridefinirli o personalizzarli qui.
        # Esempio:
        # form_data.pop("tipo_assemblea", None)
        
        st.subheader("📋 Configurazioni Specifiche Assemblea Generica")
        
        # Campi specifici per assemblea generica
        col1, col2 = st.columns(2)
        with col1:
            form_data["ruolo_presidente"] = st.selectbox("Ruolo del presidente", 
                                                        CommonDataHandler.get_standard_ruoli_presidente(),
                                                        key="ruolo_presidente_generico")
            # Sovrascrive il tipo di assemblea della base con più opzioni
            form_data["tipo_assemblea"] = st.selectbox("Tipo assemblea", 
                                                      ["Ordinaria", "Straordinaria", "Totalitaria"],
                                                      key="tipo_assemblea_generico")
        with col2:
            form_data["modalita_partecipazione"] = st.selectbox("Modalità di partecipazione", 
                                                               ["Presenza fisica", "Audioconferenza", "Mista"],
                                                               key="modalita_partecipazione_generico")
        
        # Ordine del giorno personalizzabile
        st.subheader("📋 Ordine del Giorno")
        default_odg = "1. [Inserire primo punto]\n2. [Inserire secondo punto]\n3. [Inserire terzo punto]"
        punti_odg = st.text_area("Punti all'ordine del giorno", 
                                default_odg, height=120,
                                help="Inserire i punti dell'ordine del giorno, uno per riga",
                                key="punti_odg_generico")
        form_data["punti_ordine_giorno"] = [p.strip() for p in punti_odg.split("\n") if p.strip()]
        
        # Deliberazioni per ogni punto
        st.subheader("📋 Deliberazioni")
        form_data["deliberazioni"] = []
        
        for i, punto in enumerate(form_data["punti_ordine_giorno"]):
            st.markdown(f"**Deliberazione per: {punto}**")
            col1, col2 = st.columns([3, 1])
            with col1:
                deliberazione = st.text_area(f"Deliberazione punto {i+1}", 
                                           key=f"deliberazione_{i}_generico",
                                           height=100,
                                           placeholder="Inserire la deliberazione presa dall'assemblea per questo punto")
            with col2:
                tipo_voto = st.selectbox(f"Tipo voto", 
                                       ["Unanimità", "Maggioranza", "Con voti contrari", "Con astensioni"],
                                       key=f"voto_{i}_generico")
                
                dettagli_voto = ""
                if tipo_voto in ["Con voti contrari", "Con astensioni"]:
                    dettagli_voto = st.text_input(f"Dettagli voto", 
                                                key=f"dettagli_voto_{i}_generico",
                                                placeholder="es. voto contrario dei Sigg. ...")
            
            form_data["deliberazioni"].append({
                "punto": punto,
                "deliberazione": deliberazione,
                "tipo_voto": tipo_voto,
                "dettagli_voto": dettagli_voto
            })
        
        # Informazioni aggiuntive
        st.subheader("ℹ️ Informazioni Aggiuntive")
        col1, col2 = st.columns(2)
        with col1:
            form_data["presenti_aggiuntivi"] = st.text_area("Altri presenti (se previsti)", 
                                                          placeholder="es. revisore contabile, consulenti, ecc.",
                                                          height=80,
                                                          key="presenti_aggiuntivi_generico")
        with col2:
            form_data["note_linguistiche"] = st.text_area("Note linguistiche/traduzione", 
                                                        placeholder="es. traduzione dall'italiano all'inglese",
                                                        height=80,
                                                        key="note_linguistiche_generico")
        
        # Orario di chiusura
        form_data["ora_chiusura"] = st.time_input("Ora di chiusura assemblea", 
                                                 value=None,
                                                 help="Lasciare vuoto se da compilare manualmente",
                                                 key="ora_chiusura_generico")
        
        return form_data
    
    def show_preview(self, form_data: dict):
        """Mostra l'anteprima del documento fuori dal form"""
        # Sezione Anteprima
        st.markdown("---")
        st.subheader("👁️ Anteprima Documento")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col1:
            show_preview = st.checkbox("Mostra anteprima", value=True, key="preview_checkbox_generico")
        
        with col2:
            if show_preview:
                st.info("💡 L'anteprima si aggiorna automaticamente con i dati inseriti sopra")
        
        if show_preview:
            try:
                preview_text = self._generate_preview_text(form_data)
                st.text(preview_text)
            except Exception as e:
                st.error(f"Errore nell'anteprima: {e}")
                st.exception(e)
    
    def _generate_preview_text(self, data: dict) -> str:
        """Genera il testo di anteprima del verbale"""
        try:
            # Header dell'azienda
            header = f"""{data.get('denominazione', '[Denominazione]')}
Sede in {data.get('sede_legale', '[Sede]')}
Capitale sociale Euro {data.get('capitale_sociale', '[Capitale]')} i.v.
Codice fiscale: {data.get('codice_fiscale', '[CF]')}

Verbale di assemblea dei soci
del {data.get('data_assemblea', '[Data]').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[Data]'}

Oggi {data.get('data_assemblea', '[Data]').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[Data]'} alle ore {data.get('ora_assemblea', '[Ora]').strftime('%H:%M') if hasattr(data.get('ora_assemblea'), 'strftime') else '[Ora]'} presso la sede sociale {data.get('sede_legale', '[Sede]')}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:

Ordine del giorno"""
            
            # Aggiunge i punti dell'ordine del giorno
            for i, punto in enumerate(data.get('punti_ordine_giorno', []), 1):
                header += f"\n{i}. {punto.replace('[', '').replace(']', '')}"
            
            # Sezione presidenza
            presidente_section = f"""

Assume la presidenza ai sensi dell'art. […] dello statuto sociale il Sig. {data.get('presidente', '[Presidente]')} {data.get('ruolo_presidente', 'Amministratore Unico')}, il quale dichiara e constata:

1 - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. […] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza

2 - che sono presenti/partecipano all'assemblea:
l'{data.get('ruolo_presidente', 'Amministratore Unico')} nella persona del suddetto Presidente Sig. {data.get('presidente', '[Presidente]')}"""
            
            # Aggiunge amministratori se presenti
            amministratori = data.get('amministratori', [])
            if amministratori and len(amministratori) > 1:
                presidente_section += "\n[oppure\nper il Consiglio di Amministrazione:"
                for admin in amministratori:
                    nome_admin = admin.get('nome', '') if isinstance(admin, dict) else str(admin)
                    presidente_section += f"\nil Sig {nome_admin}"
            
            # Collegio sindacale se presente
            if data.get('include_collegio_sindacale', False):
                presidente_section += "\n[eventualmente\nper il Collegio Sindacale\nil Dott. [Nome]\nil Dott. [Nome]\nil Dott. [Nome]]"
            
            # Revisore se presente
            if data.get('include_revisore', False):
                nome_rev = data.get('nome_revisore', '[NOME REVISORE]')
                presidente_section += f"\n[eventualmente\nil revisore contabile Dott. {nome_rev}]"
            
            # Soci presenti
            # Fallback per la retrocompatibilità
            if 'soci_presenti' not in data:
                data['soci_presenti'] = data.get('soci', [])
                data['soci_assenti'] = []

            soci_presenti = data.get('soci_presenti', [])
            soci_assenti = data.get('soci_assenti', [])
            total_quota_euro = 0.0
            total_quota_percentuale = 0.0

            for socio in soci_presenti:
                if isinstance(socio, dict):
                    try:
                        quota_euro_str = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                        total_quota_euro += float(quota_euro_str)
                    except (ValueError, TypeError):
                        pass 
                    try:
                        quota_percentuale_str = str(socio.get('quota_percentuale', '0')).replace('.', '').replace(',', '.')
                        total_quota_percentuale += float(quota_percentuale_str)
                    except (ValueError, TypeError):
                        pass
            
            # Formatta i totali per la visualizzazione
            formatted_total_quota_euro = f"{total_quota_euro:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            formatted_total_quota_percentuale = f"{total_quota_percentuale:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

            soci_section = f"\nnonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_total_quota_euro} pari al {formatted_total_quota_percentuale}% del Capitale Sociale:"
            
            for socio in soci_presenti:
                if isinstance(socio, dict):
                    nome = socio.get('nome', '[Nome Socio]')
                    quota_value = socio.get('quota_euro', '')
                    percentuale_value = socio.get('quota_percentuale', '')
                    
                    quota = '[Quota]' if quota_value is None or str(quota_value).strip() == '' else str(quota_value).strip()
                    percentuale = '[%]' if percentuale_value is None or str(percentuale_value).strip() == '' else str(percentuale_value).strip()
                    
                    tipo_partecipazione = socio.get('tipo_partecipazione', 'Diretta')
                    
                    if tipo_partecipazione == 'Delegato':
                        delegato = socio.get('delegato', '')
                        socio_line = f"{delegato} delegato del socio {nome} – quota euro {quota} ({percentuale}%)"
                    else:
                        socio_line = f"{nome} – quota euro {quota} ({percentuale}%)"
                    soci_section += f"\n{socio_line}"
                else:
                    socio_line = f"Sig. {socio}"
                    soci_section += f"\n{socio_line}"

            # Aggiunge soci assenti se presenti
            if soci_assenti:
                soci_section += "\n\nRisultano invece assenti i seguenti soci:"
                for socio in soci_assenti:
                    nome = socio.get('nome', '[Nome Socio]')
                    socio_line = f"Sig. {nome}"
                    soci_section += f"\n{socio_line}"
            
            # Altri presenti se specificati
            altri_presenti_section = ""
            if data.get('presenti_aggiuntivi'):
                altri_presenti_section = f"\n4 - che sono altresì presenti, in qualità di:\n{data.get('presenti_aggiuntivi')}"
            
            # Segretario
            segretario_section = f"""

I presenti all'unanimità chiamano a fungere da segretario il signor {data.get('segretario', '[Segretario]')}, che accetta l'incarico.

Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante."""
            
            # Note linguistiche se presenti
            if data.get('note_linguistiche'):
                segretario_section += f"\n[eventualmente\n{data.get('note_linguistiche')}]"
            
            # Validità assemblea
            validita_section = f"""

Il Presidente constata e fa constatare che l'assemblea risulta regolarmente convocata [oppure totalitaria] e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.

Si passa quindi allo svolgimento dell'ordine del giorno.

*     *     *"""
            
            # Deliberazioni
            deliberazioni_section = ""
            deliberazioni = data.get('deliberazioni', [])
            
            for i, delib in enumerate(deliberazioni, 1):
                punto = delib.get('punto', f'[Punto {i}]')
                deliberazione = delib.get('deliberazione', '[Deliberazione da inserire]')
                tipo_voto = delib.get('tipo_voto', 'Unanimità')
                dettagli_voto = delib.get('dettagli_voto', '')
                
                voto_text = ""
                if tipo_voto == "Unanimità":
                    voto_text = "all'unanimità"
                elif tipo_voto == "Maggioranza":
                    voto_text = "a maggioranza"
                elif tipo_voto == "Con voti contrari":
                    voto_text = f"con il voto contrario dei Sigg. {dettagli_voto}" if dettagli_voto else "con voti contrari"
                elif tipo_voto == "Con astensioni":
                    voto_text = f"con l'astensione dei Sigg. {dettagli_voto}" if dettagli_voto else "con astensioni"
                
                deliberazioni_section += f"""

In relazione al {self._get_ordinal(i)} punto posto all'ordine del giorno [{punto}].

Segue ampia discussione al termine della quale si passa alla votazione con voto palese in forza della quale il Presidente constata che, {voto_text}, l'assemblea

d e l i b e r a:

{deliberazione}

*     *     *"""
            
            # Chiusura
            ora_chiusura = data.get('ora_chiusura', '[Ora]')
            if hasattr(ora_chiusura, 'strftime'):
                ora_chiusura = ora_chiusura.strftime('%H:%M')
            
            chiusura_section = f"""

Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.

Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo [eventualmente unitamente a quanto allegato].

L'assemblea viene sciolta alle ore {ora_chiusura}.


Il Presidente                    Il Segretario
{data.get('presidente', '[Presidente]')}            {data.get('segretario', '[Segretario]')}"""
            
            # Combina tutte le sezioni
            full_text = (header + presidente_section + soci_section + altri_presenti_section + 
                        segretario_section + validita_section + deliberazioni_section + chiusura_section)
            
            return full_text
            
        except Exception as e:
            return f"Errore nella generazione dell'anteprima: {str(e)}"
    
    def _get_ordinal(self, number):
        """Converte un numero in ordinale italiano"""
        ordinals = {
            1: "primo", 2: "secondo", 3: "terzo", 4: "quarto", 5: "quinto",
            6: "sesto", 7: "settimo", 8: "ottavo", 9: "nono", 10: "decimo"
        }
        return ordinals.get(number, f"{number}°")
    
    def generate_document(self, data: dict) -> Document:
        """Genera il documento Word del verbale"""
        import os
        
        # Utilizza il template .docx esistente per mantenere la formattazione
        template_path = os.path.join(os.path.dirname(__file__), 'template.docx')
        
        try:
            if os.path.exists(template_path):
                doc = Document(template_path)
                # Rimuovi il contenuto esistente del template mantenendo gli stili
                for paragraph in doc.paragraphs[:]:
                    p = paragraph._element
                    p.getparent().remove(p)
            else:
                doc = Document()
                # Setup stili solo se non usiamo template
                self._setup_document_styles(doc)
        except Exception as e:
            # Fallback a documento vuoto se il template non può essere caricato
            doc = Document()
            self._setup_document_styles(doc)
        
        # Setup stili del documento
        self._setup_document_styles(doc)
        
        # Aggiungi header azienda
        self._add_company_header(doc, data)
        
        # Aggiungi titolo verbale
        self._add_verbale_title(doc, data)
        
        # Aggiungi sezione di apertura
        self._add_opening_section(doc, data)
        
        # Aggiungi partecipanti
        self._add_participants_section(doc, data)
        
        # Aggiungi dichiarazioni preliminari
        self._add_preliminary_statements(doc, data)
        
        # Aggiungi discussione dei punti
        self._add_points_discussion(doc, data)
        
        # Aggiungi sezione di chiusura
        self._add_closing_section(doc, data)
        
        # Aggiungi firme
        self._add_signatures(doc, data)
        
        return doc
    
    def _setup_document_styles(self, doc):
        """Imposta gli stili del documento"""
        # First call parent's method to setup all base styles
        super()._setup_document_styles(doc)
        
        styles = doc.styles
        
        # Stile per il titolo principale
        if 'Titolo Principale' not in [s.name for s in styles]:
            title_style = styles.add_style('Titolo Principale', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Times New Roman'
            title_style.font.size = Pt(14)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # Stile per i sottotitoli
        if 'Sottotitolo' not in [s.name for s in styles]:
            subtitle_style = styles.add_style('Sottotitolo', WD_STYLE_TYPE.PARAGRAPH)
            subtitle_style.font.name = 'Times New Roman'
            subtitle_style.font.size = Pt(12)
            subtitle_style.font.bold = True
            subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_style.paragraph_format.space_after = Pt(6)
        
        # Stile per il testo normale
        normal_style = styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal_style.paragraph_format.line_spacing = 1.15
    
    def _add_company_header(self, doc, data):
        """Aggiunge l'header dell'azienda"""
        # Nome azienda
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data.get('denominazione', '[DENOMINAZIONE SOCIETÀ]'))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Sede
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Sede in {data.get('sede_legale', '[SEDE]')}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        # Capitale sociale
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Capitale sociale Euro {data.get('capitale_sociale', '[CAPITALE]')} i.v.")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        # Codice fiscale
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Codice fiscale: {data.get('codice_fiscale', '[CODICE FISCALE]')}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        # Spazio
        doc.add_paragraph()
    
    def _add_verbale_title(self, doc, data):
        """Aggiunge il titolo del verbale"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Verbale di assemblea dei soci")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Data
        data_assemblea = data.get('data_assemblea')
        if hasattr(data_assemblea, 'strftime'):
            data_str = data_assemblea.strftime('%d/%m/%Y')
        else:
            data_str = '[DATA]'
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"del {data_str}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        doc.add_paragraph()
    
    def _add_opening_section(self, doc, data):
        font_name = 'Times New Roman'
        
        sede_legale = data.get('sede_legale', '[Sede]')
        data_assemblea = data.get('data_assemblea')
        ora_assemblea = data.get('ora_assemblea')
        
        data_str = data_assemblea.strftime('%d/%m/%Y') if data_assemblea else '[Data]'
        ora_str = ora_assemblea.strftime('%H:%M') if ora_assemblea else '[Ora]'
        
        self.add_paragraph_with_font(doc, f"Oggi {data_str} alle ore {ora_str} presso la sede sociale {sede_legale}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:", size=Pt(12), font_name=font_name)
        
        self.add_paragraph_with_font(doc, "Ordine del giorno", size=Pt(12), font_name=font_name, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        
        for i, punto in enumerate(data.get('punti_ordine_giorno', []), 1):
            self.add_paragraph_with_font(doc, f"{i}. {punto.replace('[', '').replace(']', '')}", size=Pt(12), font_name=font_name, left_indent=Inches(0.5))

    def _add_participants_section(self, doc, data):
        """Aggiunge la sezione dei partecipanti"""
        font_name = 'Times New Roman'
        
        presidente_text = f"Assume la presidenza ai sensi dell'art. […] dello statuto sociale il Sig. {data.get('presidente', '[Presidente]')} {data.get('ruolo_presidente', 'Amministratore Unico')}, il quale dichiara e constata:"
        self.add_paragraph_with_font(doc, presidente_text, size=Pt(12), font_name=font_name, space_after=Pt(6))
        
        statements = [
            "che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. […] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza",
            f"che sono presenti/partecipano all'assemblea:\nl'{data.get('ruolo_presidente', 'Amministratore Unico')} nella persona del suddetto Presidente Sig. {data.get('presidente', '[Presidente]')}"
        ]

        # Aggiunge amministratori se presenti
        amministratori = data.get('amministratori', [])
        if amministratori and len(amministratori) > 1:
            admin_list = "[oppure\nper il Consiglio di Amministrazione:"
            for admin in amministratori:
                nome_admin = admin.get('nome', '') if isinstance(admin, dict) else str(admin)
                admin_list += f"\nil Sig {nome_admin}"
            statements.append(admin_list)
        
        # Collegio sindacale se presente
        if data.get('include_collegio_sindacale', False):
            statements.append("[eventualmente\nper il Collegio Sindacale\nil Dott. [Nome]\nil Dott. [Nome]\nil Dott. [Nome]]")
            
        # Revisore se presente
        if data.get('include_revisore', False):
            nome_rev = data.get('nome_revisore', '[NOME REVISORE]')
            statements.append(f"[eventualmente\nil revisore contabile Dott. {nome_rev}]")
        
        for i, statement in enumerate(statements, 1):
             self.add_paragraph_with_font(doc, f"{i} - {statement}", size=Pt(12), font_name=font_name, left_indent=Inches(0.5))
        
        # Fallback per la retrocompatibilità
        if 'soci_presenti' not in data:
            data['soci_presenti'] = data.get('soci', [])
            data['soci_assenti'] = []

        soci_presenti = data.get('soci_presenti', [])
        soci_assenti = data.get('soci_assenti', [])

        # Soci presenti
        total_quota_euro = 0.0
        total_quota_percentuale = 0.0

        for socio in soci_presenti:
            if isinstance(socio, dict):
                try:
                    quota_euro_str = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                    total_quota_euro += float(quota_euro_str)
                except (ValueError, TypeError):
                    pass
                try:
                    quota_percentuale_str = str(socio.get('quota_percentuale', '0')).replace('.', '').replace(',', '.')
                    total_quota_percentuale += float(quota_percentuale_str)
                except (ValueError, TypeError):
                    pass
        
        formatted_total_quota_euro = f"{total_quota_euro:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        formatted_total_quota_percentuale = f"{total_quota_percentuale:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

        soci_text = f"nonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_total_quota_euro} pari al {formatted_total_quota_percentuale}% del Capitale Sociale:"
        self.add_paragraph_with_font(doc, soci_text, size=Pt(12), font_name=font_name, space_before=Pt(12))
        
        for socio in soci_presenti:
            if isinstance(socio, dict):
                nome = socio.get('nome', '[Nome Socio]')
                quota_value = socio.get('quota_euro', '')
                percentuale_value = socio.get('quota_percentuale', '')
                
                quota = '[Quota]' if quota_value is None or str(quota_value).strip() == '' else str(quota_value).strip()
                percentuale = '[%]' if percentuale_value is None or str(percentuale_value).strip() == '' else str(percentuale_value).strip()
                
                tipo_partecipazione = socio.get('tipo_partecipazione', 'Diretta')
                
                if tipo_partecipazione == 'Delegato':
                    delegato = socio.get('delegato', '')
                    socio_line = f"{delegato} delegato del socio {nome} – quota euro {quota} ({percentuale}%)"
                else:
                    socio_line = f"{nome} – quota euro {quota} ({percentuale}%)"
                p_socio = doc.add_paragraph(socio_line)
                try:
                    p_socio.style = 'List Bullet'
                except Exception:
                    p_socio.text = f"- {p_socio.text}"
            else:
                p_socio = doc.add_paragraph(f"Sig. {socio}")
                try:
                    p_socio.style = 'List Bullet'
                except Exception:
                    p_socio.text = f"- {p_socio.text}"
        
        # Aggiunge soci assenti se presenti
        if soci_assenti:
            self.add_paragraph_with_font(doc, "Risultano invece assenti i seguenti soci:", size=Pt(12), font_name=font_name, space_before=Pt(12))
            for socio in soci_assenti:
                nome = socio.get('nome', '[Nome Socio]')
                p_socio = doc.add_paragraph(f"Sig. {nome}")
                try:
                    p_socio.style = 'List Bullet'
                except Exception:
                    p_socio.text = f"- {p_socio.text}"

        # Altri presenti se specificati
        if data.get('presenti_aggiuntivi'):
            altri_presenti_text = f"4 - che sono altresì presenti, in qualità di:\n{data.get('presenti_aggiuntivi')}"
            self.add_paragraph_with_font(doc, altri_presenti_text, size=Pt(12), font_name=font_name, left_indent=Inches(0.5))

    def _add_preliminary_statements(self, doc, data):
        font_name = 'Times New Roman'
        
        segretario_text = f"""I presenti all'unanimità chiamano a fungere da segretario il signor {data.get('segretario', '[Segretario]')}, che accetta l'incarico.

Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante."""
        self.add_paragraph_with_font(doc, segretario_text, size=Pt(12), font_name=font_name)
        
        if data.get('note_linguistiche'):
            note_text = f"La presente riunione è stata tradotta da un interprete professionista dalla lingua italiana alla lingua inglese e viceversa."
            self.add_paragraph_with_font(doc, note_text, size=Pt(12), font_name=font_name, space_before=Pt(12))

        preliminary_text = """Il Presidente dichiara quindi che l'assemblea è regolarmente costituita e atta a deliberare sugli argomenti posti all'ordine del giorno.

Inizia quindi la discussione sui punti all'ordine del giorno."""
        self.add_paragraph_with_font(doc, preliminary_text, size=Pt(12), font_name=font_name, space_before=Pt(12))

    def _add_points_discussion(self, doc, data):
        """Aggiunge la discussione dei punti all'ordine del giorno"""
        deliberazioni = data.get('deliberazioni', [])
        
        for i, delib in enumerate(deliberazioni, 1):
            doc.add_paragraph()
            
            punto = delib.get('punto', f'[Punto {i}]')
            deliberazione = delib.get('deliberazione', '[Deliberazione da inserire]')
            tipo_voto = delib.get('tipo_voto', 'Unanimità')
            dettagli_voto = delib.get('dettagli_voto', '')
            
            # Introduzione del punto
            ordinal = self._get_ordinal(i)
            p = doc.add_paragraph(f"In relazione al {ordinal} punto posto all'ordine del giorno [{punto}].")
            
            # Voto
            voto_text = ""
            if tipo_voto == "Unanimità":
                voto_text = "all'unanimità"
            elif tipo_voto == "Maggioranza":
                voto_text = "a maggioranza"
            elif tipo_voto == "Con voti contrari":
                voto_text = f"con il voto contrario dei Sigg. {dettagli_voto}" if dettagli_voto else "con voti contrari"
            elif tipo_voto == "Con astensioni":
                voto_text = f"con l'astensione dei Sigg. {dettagli_voto}" if dettagli_voto else "con astensioni"
            
            p = doc.add_paragraph(f"Segue ampia discussione al termine della quale si passa alla votazione con voto palese in forza della quale il Presidente constata che, {voto_text}, l'assemblea")
            
            # Delibera
            p = doc.add_paragraph("d e l i b e r a:")
            run = p.runs[0]
            run.font.bold = True
            run.font.underline = True
            
            # Testo della deliberazione
            p = doc.add_paragraph(deliberazione)
            
            # Separatore
            p = doc.add_paragraph("*     *     *")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_closing_section(self, doc, data):
        """Aggiunge la sezione di chiusura"""
        doc.add_paragraph()
        
        p = doc.add_paragraph("Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.")
        
        p = doc.add_paragraph("Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo [eventualmente unitamente a quanto allegato].")
        
        # Ora di chiusura
        ora_chiusura = data.get('ora_chiusura')
        if hasattr(ora_chiusura, 'strftime'):
            ora_str = ora_chiusura.strftime('%H:%M')
        else:
            ora_str = '[ORA]'
        
        p = doc.add_paragraph(f"L'assemblea viene sciolta alle ore {ora_str}.")
    
    def _add_signatures(self, doc, data):
        """Aggiunge le firme"""
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Usa la tabella di firme standardizzata
        table = self._add_signature_table(doc, data)
        
        # Centra la tabella
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Registra il template
DocumentTemplateFactory.register_template('verbale_assemblea_generico', VerbaleAssembleaGenericoTemplate)