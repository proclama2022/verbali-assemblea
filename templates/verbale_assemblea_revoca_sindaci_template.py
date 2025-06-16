"""
Template per Verbale di Assemblea - Revoca dei sindaci e provvedimenti conseguenti
"""

import sys
import os

# Aggiungi il path della cartella src
current_dir = os.path.dirname(os.path.abspath(__file__))
src_path = os.path.join(os.path.dirname(current_dir), 'src')
if src_path not in sys.path:
    sys.path.append(src_path)

from document_templates import DocumentTemplate, DocumentTemplateFactory
from common_data_handler import CommonDataHandler
from base_verbale_template import BaseVerbaleTemplate
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date
import streamlit as st

class VerbaleRevocaSindaciTemplate(BaseVerbaleTemplate):
    """Template per Verbale di Assemblea - Revoca dei sindaci e provvedimenti conseguenti"""
    
    def get_template_name(self) -> str:
        return "Revoca dei sindaci e provvedimenti conseguenti"
    
    def get_required_fields(self) -> list:
        return [
            "denominazione", "sede_legale", "capitale_sociale", "codice_fiscale",
            "data_assemblea", "ora_assemblea", "presidente", "segretario", 
            "soci", "occasione_irregolarita", "gravi_irregolarita"
        ]
    
    def get_form_fields(self, extracted_data: dict) -> dict:
        """Genera i campi del form per Streamlit."""
        # Chiama il metodo della classe base per ottenere i campi comuni
        form_data = super().get_form_fields(extracted_data)

        st.subheader("⚖️ Configurazioni Specifiche Revoca Sindaci")
        
        col1, col2 = st.columns(2)
        with col1:
            form_data["tipo_organo_controllo"] = st.selectbox(
                "Tipo Organo di Controllo",
                ["Collegio Sindacale", "Sindaco Unico"],
                key="tipo_organo_controllo_revoca"
            )
        with col2:
            form_data["motivo_revoca"] = st.selectbox(
                "Motivo della revoca",
                ["Giusta causa", "Dimissioni volontarie", "Scadenza mandato", "Decesso", "Altro"],
                key="motivo_revoca_sindaci"
            )

        # Dati dei sindaci da revocare
        st.subheader("👥 Sindaci da Revocare")
        sindaci_revocare = st.text_area(
            "Nomi dei sindaci da revocare (separati da virgola)",
            help="Es. Mario Rossi, Luigi Bianchi",
            key="sindaci_revocare_list"
        )
        form_data["sindaci_da_revocare"] = [s.strip() for s in sindaci_revocare.split(",") if s.strip()]

        # Eventuale nomina di nuovi sindaci
        st.subheader("Nomina Nuovi Sindaci (Opzionale)")
        form_data['nomina_nuovi_sindaci'] = st.checkbox("Procedere con la nomina di nuovi sindaci?", key="nomina_nuovi_sindaci_revoca")

        if form_data.get('nomina_nuovi_sindaci'):
            # Qui si potrebbero aggiungere i campi per la nomina dei nuovi sindaci,
            # simile al template di nomina del collegio sindacale.
            st.info("La sezione per la nomina di nuovi sindaci è da implementare.")
            # Esempio:
            # form_data['nuovi_sindaci'] = st.data_editor(...)

        return form_data
    
    def show_preview(self, form_data: dict):
        """Mostra l'anteprima del documento"""
        st.markdown("---")
        st.subheader("👁️ Anteprima Documento")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col1:
            show_preview = st.checkbox("Mostra anteprima", value=False, key="preview_checkbox_revoca_sindaci")
        
        with col2:
            if show_preview:
                st.info("💡 L'anteprima si aggiorna automaticamente")
        
        if show_preview:
            try:
                preview_text = self._generate_preview_text(form_data)
                st.text(preview_text)
            except Exception as e:
                st.error(f"Errore nell'anteprima: {e}")
    
    def _generate_preview_text(self, data: dict) -> str:
        """Genera il testo di anteprima"""
        try:
            # Header uniformato
            denominazione = data.get('denominazione', '[Denominazione]')
            sede_legale = data.get('sede_legale', '[Sede]')
            capitale_sociale_raw = data.get('capitale_versato') or data.get('capitale_deliberato') or data.get('capitale_sociale', '[Capitale]')
            capitale_sociale = CommonDataHandler.format_currency(capitale_sociale_raw)
            codice_fiscale = data.get('codice_fiscale', '[CF]')

            header = f"""{denominazione}
Sede in {sede_legale}
Capitale sociale Euro {capitale_sociale} i.v.
Codice fiscale: {codice_fiscale}

Verbale di assemblea dei soci
del {data.get('data_assemblea', '[Data]').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[Data]'}

Oggi {data.get('data_assemblea', '[Data]').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[Data]'} alle ore {data.get('ora_assemblea', '[Ora]').strftime('%H:%M') if hasattr(data.get('ora_assemblea'), 'strftime') else '[Ora]'} presso la sede sociale {data.get('sede_legale', '[Sede]')}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:

Ordine del giorno
Revoca dei sindaci e provvedimenti conseguenti"""
            
            # Sezione presidenza
            ruolo_presidente = data.get('ruolo_presidente', 'Amministratore Unico')
            presidente_section = f"""

Assume la presidenza ai sensi dell'art. […] dello statuto sociale il Sig. {data.get('presidente', '[Presidente]')} {ruolo_presidente}, il quale dichiara e constata:

1 - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. […] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza

2 - che sono presenti/partecipano all'assemblea:
l'{ruolo_presidente} nella persona del suddetto Presidente Sig. {data.get('presidente', '[Presidente]')}"""
            
            # Consiglio di Amministrazione se applicabile
            if data.get('include_consiglio_amministrazione', False):
                presidente_section += """
[oppure
per il Consiglio di Amministrazione:
il Sig […]
il Sig […]
il Sig […]
assente giustificato il Sig […]]"""
            
            # Collegio sindacale
            sindaci = data.get('sindaci_revocandi', [])
            if sindaci:
                if data.get('tipo_collegio') == "Sindaco Unico":
                    sindaco = sindaci[0] if sindaci else {}
                    nome = sindaco.get('nome', '[Nome]')
                    presidente_section += f"""
[eventualmente
il Sindaco Unico nella persona del Sig. {nome}]"""
                else:
                    presidente_section += """
[eventualmente
per il Collegio Sindacale"""
                    for sindaco in sindaci:
                        nome = sindaco.get('nome', '[Nome]')
                        presidente_section += f"\nil {nome}"
                    presidente_section += "]"
            
            # Revisore se presente
            if data.get('include_revisore', False):
                presidente_section += """
[eventualmente, se invitato
il revisore contabile Dott. […]]"""
            
            # Soci presenti
            soci = data.get('soci', [])
            total_quota_euro = 0.0
            total_quota_percentuale = 0.0

            for socio in soci:
                if isinstance(socio, dict):
                    try:
                        # Rimuovi punti e sostituisci virgola con punto per la conversione a float
                        quota_euro_str = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                        total_quota_euro += float(quota_euro_str)
                    except ValueError:
                        pass # Ignora valori non numerici
                    try:
                        quota_percentuale_str = str(socio.get('quota_percentuale', '0')).replace('.', '').replace(',', '.')
                        total_quota_percentuale += float(quota_percentuale_str)
                    except ValueError:
                        pass # Ignora valori non numerici
            
            # Formatta i totali per la visualizzazione
            formatted_total_quota_euro = CommonDataHandler.format_currency(total_quota_euro)
            formatted_total_quota_percentuale = CommonDataHandler.format_percentage(total_quota_percentuale)

            soci_section = f"\nnonché i seguenti soci o loro rappresentanti, [eventualmente così come iscritti a libro soci e] recanti complessivamente una quota pari a nominali euro {formatted_total_quota_euro} pari al {formatted_total_quota_percentuale}% del Capitale Sociale:"
            
            for socio in soci:
                if isinstance(socio, dict):
                    nome = socio.get('nome', '[Nome Socio]')
                    quota_value = socio.get('quota_euro', '')
                    percentuale_value = socio.get('quota_percentuale', '')
                    
                    # Gestione robusta dei valori nulli o vuoti
                    quota = '[Quota]' if quota_value is None or str(quota_value).strip() == '' else str(quota_value).strip()
                    percentuale = '[%]' if percentuale_value is None or str(percentuale_value).strip() == '' else str(percentuale_value).strip()
                    
                    soci_section += f"\nil Sig {nome} socio recante una quota pari a nominali euro {quota} pari al {percentuale}% del Capitale Sociale"
            
            soci_section += "\n2 - che gli intervenuti sono legittimati alla presente assemblea;\n3 - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno."
            
            # Segretario e traduzione
            segretario_section = f"""

I presenti all'unanimità chiamano a fungere da segretario il signor {data.get('segretario', '[Segretario]')}, che accetta l'incarico.

Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante."""
            
            # Traduzione se necessaria
            if data.get('include_traduzione', False):
                persona_traduzione = data.get('persona_traduzione', '[Nome]')
                lingua = data.get('lingua_traduzione', 'inglese').lower()
                segretario_section += f"""
In particolare, preso atto che il Sig. {persona_traduzione} non conosce la lingua italiana ma dichiara di conoscere la lingua {lingua}, il Presidente dichiara che provvederà a tradurre dall'italiano al {lingua} (e viceversa) gli interventi dei partecipanti alla discussione nonché a tradurre dall'italiano al {lingua} il verbale che sarà redatto al termine della riunione."""
            
            segretario_section += """

Il Presidente constata e fa constatare che l'assemblea risulta regolarmente convocata [oppure totalitaria] e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.

Si passa quindi allo svolgimento dell'ordine del giorno.

*     *     *"""
            
            # Discussione revoca sindaci
            occasione = data.get('occasione_irregolarita', '[occasione delle irregolarità]')
            irregolarita = data.get('gravi_irregolarita', '[gravi irregolarità riscontrate]')
            
            revoca_section = f"""

Il Presidente [o altro soggetto presente in assemblea (amministratore o socio)] prende la parola dichiarando che, in occasione di {occasione} i sindaci attualmente in carica sono incorsi in gravi irregolarità nell'adempimento dei loro doveri. Più precisamente le gravi irregolarità riscontrate sono le seguenti:

{irregolarita}"""
            
            # Dichiarazione aggiuntiva
            soggetto_dichiarante = data.get('soggetto_dichiarante', '[Nome]')
            dichiarazione = data.get('dichiarazione_aggiuntiva', '[dichiarazione]')
            
            if soggetto_dichiarante and soggetto_dichiarante != '[Nome]':
                revoca_section += f"""

Prende la parola il Sig. {soggetto_dichiarante} che dichiara {dichiarazione}."""
            
            # Votazione
            if data.get('votazione_unanime', True):
                votazione_text = "all'unanimità"
            else:
                voti_contrari = data.get('voti_contrari', '')
                astensioni = data.get('astensioni', '')
                votazione_text = f"con il voto contrario dei Sigg. {voti_contrari}"
                if astensioni:
                    votazione_text += f" e l'astensione dei Sigg. {astensioni}"
            
            chi_ricorso = data.get('chi_presenta_ricorso', 'Amministratore Unico')
            
            deliberazione_section = f"""

Esaurita la discussione, si passa alla votazione con voto palese in forza della quale il Presidente constata che, {votazione_text}, l'assemblea

d e l i b e r a:

la revoca dei sindaci della società dal loro ufficio per giusta causa, dando incarico all'{chi_ricorso} di presentare ricorso al Tribunale per l'approvazione della presente delibera."""
            
            # Chiusura
            chiusura_section = f"""

*     *     *

Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.

Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo [eventualmente unitamente a quanto allegato].

L'assemblea viene sciolta alle ore [{data.get('ora_chiusura', '[Ora]')}].


Il Presidente                    Il Segretario
{data.get('presidente', '[PRESIDENTE]')}            {data.get('segretario', '[SEGRETARIO]')}"""
            
            # Combina tutte le sezioni
            full_text = (header + presidente_section + soci_section + segretario_section + 
                        revoca_section + deliberazione_section + chiusura_section)
            
            return full_text
            
        except Exception as e:
            return f"Errore nella generazione dell'anteprima: {str(e)}"
    
    def generate_document(self, data: dict) -> Document:
        """Genera il documento Word del verbale"""
        import os
        
        # Utilizza il template .docx esistente per mantenere la formattazione
        template_path = os.path.join(os.path.dirname(__file__), 'template.docx')
        
        try:
            if os.path.exists(template_path):
                doc = Document(template_path)
                # Rimuovi il contenuto esistente del template mantenendo gli stili
                for paragraph in doc.paragraphs[:]:
                    p = paragraph._element
                    p.getparent().remove(p)
            else:
                doc = Document()
                # Setup stili solo se non usiamo template
                self._setup_document_styles(doc)
        except Exception as e:
            # Fallback a documento vuoto se il template non può essere caricato
            doc = Document()
            self._setup_document_styles(doc)
        
        # Setup stili
        self._setup_document_styles(doc)
        
        # Header azienda
        self._add_company_header(doc, data)
        
        # Titolo verbale
        self._add_verbale_title(doc, data)
        
        # Sezioni del verbale
        self._add_opening_section(doc, data)
        self._add_participants_section(doc, data)
        self._add_preliminary_statements(doc, data)
        self._add_revoca_discussion(doc, data)
        self._add_closing_section(doc, data)
        self._add_signatures(doc, data)
        
        return doc
    
    def _setup_document_styles(self, doc):
        """Imposta gli stili del documento"""
        styles = doc.styles
        
        if 'Titolo Principale' not in [s.name for s in styles]:
            title_style = styles.add_style('Titolo Principale', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Times New Roman'
            title_style.font.size = Pt(14)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        normal_style = styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal_style.paragraph_format.line_spacing = 1.15
    
    def _add_company_header(self, doc, data):
        """Aggiunge header azienda"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data.get('denominazione', '[DENOMINAZIONE]'))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Sede in {data.get('sede_legale', '[SEDE]')}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Capitale sociale Euro {data.get('capitale_sociale', '[CAPITALE]')} i.v.")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Codice fiscale: {data.get('codice_fiscale', '[CF]')}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        doc.add_paragraph()
    
    def _add_verbale_title(self, doc, data):
        """Aggiunge titolo verbale"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Verbale di assemblea dei soci")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        
        data_assemblea = data.get('data_assemblea')
        if hasattr(data_assemblea, 'strftime'):
            data_str = data_assemblea.strftime('%d/%m/%Y')
        else:
            data_str = '[DATA]'
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"del {data_str}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        doc.add_paragraph()
    
    def _add_opening_section(self, doc, data):
        """Aggiunge sezione apertura"""
        data_assemblea = data.get('data_assemblea')
        ora_assemblea = data.get('ora_assemblea')
        
        if hasattr(data_assemblea, 'strftime'):
            data_str = data_assemblea.strftime('%d/%m/%Y')
        else:
            data_str = '[DATA]'
        
        if hasattr(ora_assemblea, 'strftime'):
            ora_str = ora_assemblea.strftime('%H:%M')
        else:
            ora_str = '[ORA]'
        
        text = f"Oggi {data_str} alle ore {ora_str} presso la sede sociale {data.get('sede_legale', '[SEDE]')}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:"
        
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Ordine del giorno")
        run.font.bold = True
        
        p = doc.add_paragraph("Revoca dei sindaci e provvedimenti conseguenti")
    
    def _add_participants_section(self, doc, data):
        """Aggiunge sezione partecipanti"""
        doc.add_paragraph()
        
        ruolo_presidente = data.get('ruolo_presidente', 'Amministratore Unico')
        presidente = data.get('presidente', '[PRESIDENTE]')
        
        text = f"Assume la presidenza ai sensi dell'art. […] dello statuto sociale il Sig. {presidente} {ruolo_presidente}, il quale dichiara e constata:"
        
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        p = doc.add_paragraph("1 - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. […] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza")
        
        p = doc.add_paragraph("2 - che sono presenti/partecipano all'assemblea:")
        p = doc.add_paragraph(f"l'{ruolo_presidente} nella persona del suddetto Presidente Sig. {presidente}")
        
        # Consiglio di Amministrazione se presente
        if data.get('include_consiglio_amministrazione', False):
            p = doc.add_paragraph("[oppure")
            p = doc.add_paragraph("per il Consiglio di Amministrazione:")
            p = doc.add_paragraph("il Sig […]")
            p = doc.add_paragraph("il Sig […]")
            p = doc.add_paragraph("il Sig […]")
            p = doc.add_paragraph("assente giustificato il Sig […]]")
        
        # Collegio sindacale
        sindaci = data.get('sindaci_revocandi', [])
        if sindaci:
            if data.get('tipo_collegio') == "Sindaco Unico":
                sindaco = sindaci[0] if sindaci else {}
                nome = sindaco.get('nome', '[Nome]')
                p = doc.add_paragraph("[eventualmente")
                p = doc.add_paragraph(f"il Sindaco Unico nella persona del Sig. {nome}]")
            else:
                p = doc.add_paragraph("[eventualmente")
                p = doc.add_paragraph("per il Collegio Sindacale")
                for sindaco in sindaci:
                    nome = sindaco.get('nome', '[Nome]')
                    p = doc.add_paragraph(f"il {nome}")
                p = doc.add_paragraph("]")
        
        # Revisore se presente
        if data.get('include_revisore', False):
            p = doc.add_paragraph("[eventualmente, se invitato")
            p = doc.add_paragraph("il revisore contabile Dott. […]]")
        
        # Soci
        p = doc.add_paragraph("nonché i seguenti soci o loro rappresentanti, [eventualmente così come iscritti a libro soci e] recanti complessivamente una quota pari a nominali euro […] pari al […]% del Capitale Sociale:")
        
        soci = data.get('soci', [])
        for socio in soci:
            if isinstance(socio, dict):
                nome = socio.get('nome', '[Nome Socio]')
                quota_value = socio.get('quota_euro', '')
                percentuale_value = socio.get('quota_percentuale', '')
                
                # Gestione robusta dei valori nulli o vuoti
                quota = '[Quota]' if quota_value is None or str(quota_value).strip() == '' else str(quota_value).strip()
                percentuale = '[%]' if percentuale_value is None or str(percentuale_value).strip() == '' else str(percentuale_value).strip()
                
                p = doc.add_paragraph(f"il Sig {nome} socio recante una quota pari a nominali euro {quota} pari al {percentuale}% del Capitale Sociale")
        
        p = doc.add_paragraph("2 - che gli intervenuti sono legittimati alla presente assemblea;")
        p = doc.add_paragraph("3 - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno.")
        
        # Presenti in qualità di
        if data.get('include_presenti_qualita', False):
            p = doc.add_paragraph("[eventualmente")
            p = doc.add_paragraph("4 - che sono altresì presenti, in qualità di […]:")
            p = doc.add_paragraph("il Sig. […]")
            p = doc.add_paragraph("il Sig. […]]")
    
    def _add_preliminary_statements(self, doc, data):
        """Aggiunge dichiarazioni preliminari"""
        doc.add_paragraph()
        
        segretario = data.get('segretario', '[SEGRETARIO]')
        p = doc.add_paragraph(f"I presenti all'unanimità chiamano a fungere da segretario il signor {segretario}, che accetta l'incarico.")
        
        p = doc.add_paragraph("Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante.")
        
        # Traduzione se necessaria
        if data.get('include_traduzione', False):
            persona_traduzione = data.get('persona_traduzione', '[Nome]')
            lingua = data.get('lingua_traduzione', 'inglese').lower()
            p = doc.add_paragraph(f"In particolare, preso atto che il Sig. {persona_traduzione} non conosce la lingua italiana ma dichiara di conoscere la lingua {lingua}, il Presidente dichiara che provvederà a tradurre dall'italiano al {lingua} (e viceversa) gli interventi dei partecipanti alla discussione nonché a tradurre dall'italiano al {lingua} il verbale che sarà redatto al termine della riunione.")
        
        p = doc.add_paragraph("Il Presidente constata e fa constatare che l'assemblea risulta regolarmente convocata [oppure totalitaria] e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.")
        
        p = doc.add_paragraph("Si passa quindi allo svolgimento dell'ordine del giorno.")
        
        # Separatore
        p = doc.add_paragraph("*     *     *")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_revoca_discussion(self, doc, data):
        """Aggiunge discussione revoca sindaci"""
        doc.add_paragraph()
        
        occasione = data.get('occasione_irregolarita', '[occasione delle irregolarità]')
        irregolarita = data.get('gravi_irregolarita', '[gravi irregolarità riscontrate]')
        
        p = doc.add_paragraph(f"Il Presidente [o altro soggetto presente in assemblea (amministratore o socio)] prende la parola dichiarando che, in occasione di {occasione} i sindaci attualmente in carica sono incorsi in gravi irregolarità nell'adempimento dei loro doveri. Più precisamente le gravi irregolarità riscontrate sono le seguenti:")
        
        # Irregolarità specifiche
        if irregolarita and irregolarita.strip():
            righe = irregolarita.split('\n')
            for riga in righe:
                if riga.strip():
                    p = doc.add_paragraph(riga.strip())
        else:
            p = doc.add_paragraph("[Inserire le gravi irregolarità]")
        
        # Dichiarazione aggiuntiva
        soggetto_dichiarante = data.get('soggetto_dichiarante', '')
        dichiarazione = data.get('dichiarazione_aggiuntiva', '')
        
        if soggetto_dichiarante and soggetto_dichiarante.strip():
            p = doc.add_paragraph(f"Prende la parola il Sig. {soggetto_dichiarante} che dichiara {dichiarazione if dichiarazione else '[dichiarazione]'}.")
        
        # Votazione
        if data.get('votazione_unanime', True):
            votazione_text = "all'unanimità"
        else:
            voti_contrari = data.get('voti_contrari', '')
            astensioni = data.get('astensioni', '')
            votazione_text = f"con il voto contrario dei Sigg. {voti_contrari}"
            if astensioni:
                votazione_text += f" e l'astensione dei Sigg. {astensioni}"
        
        p = doc.add_paragraph(f"Esaurita la discussione, si passa alla votazione con voto palese in forza della quale il Presidente constata che, {votazione_text}, l'assemblea")
        
        p = doc.add_paragraph("d e l i b e r a:")
        run = p.runs[0]
        run.font.bold = True
        run.font.underline = True
        
        # Deliberazione
        chi_ricorso = data.get('chi_presenta_ricorso', 'Amministratore Unico')
        
        if data.get('ricorso_tribunale', True):
            p = doc.add_paragraph(f"la revoca dei sindaci della società dal loro ufficio per giusta causa, dando incarico all'{chi_ricorso} di presentare ricorso al Tribunale per l'approvazione della presente delibera.")
        else:
            p = doc.add_paragraph("la revoca dei sindaci della società dal loro ufficio per giusta causa.")
        
        # Separatore
        p = doc.add_paragraph("*     *     *")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_closing_section(self, doc, data):
        """Aggiunge sezione chiusura"""
        doc.add_paragraph()
        
        p = doc.add_paragraph("Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.")
        
        p = doc.add_paragraph("Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo [eventualmente unitamente a quanto allegato].")
        
        ora_chiusura = data.get('ora_chiusura', '[ORA]')
        p = doc.add_paragraph(f"L'assemblea viene sciolta alle ore {ora_chiusura}.")
    
    def _add_signatures(self, doc, data):
        """Aggiunge firme"""
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Usa la tabella di firme standardizzata
        table = self._add_signature_table(doc, data)
        
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Registra il template
DocumentTemplateFactory.register_template('verbale_assemblea_revoca_sindaci', VerbaleRevocaSindaciTemplate)