"""
Template per Verbale di Assemblea - Nomina del Revisore
Questo template è specifico per verbali di nomina del revisore della società.
"""

import sys
import os

# Aggiungi il path della cartella src (relativo alla root del progetto)
current_dir = os.path.dirname(os.path.abspath(__file__))
src_path = os.path.join(os.path.dirname(current_dir), 'src')
if src_path not in sys.path:
    sys.path.append(src_path)

from document_templates import DocumentTemplateFactory
from common_data_handler import CommonDataHandler
from base_verbale_template import BaseVerbaleTemplate
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date, datetime
import streamlit as st
import pandas as pd
import re

class VerbaleNominaRevisoreTemplate(BaseVerbaleTemplate):
    """Template per Verbale di Assemblea dei Soci - Nomina del Revisore"""
    
    def get_template_name(self) -> str:
        return "Nomina del Revisore"
    
    def get_required_fields(self) -> list:
        return [
            "denominazione", "sede_legale", "capitale_sociale", "codice_fiscale",
            "data_assemblea", "ora_assemblea", "presidente", "segretario", 
            "soci", "amministratori", "revisore_dati"
        ]
    
    def get_form_fields(self, extracted_data: dict) -> dict:
        """Genera i campi del form per Streamlit."""
        # Chiama il metodo della classe base per ottenere i campi comuni
        form_data = super().get_form_fields(extracted_data)

        st.subheader("🔍 Configurazioni Nomina Revisore")

        # Campi specifici del template
        col1, col2 = st.columns(2)
        with col1:
            form_data['tipo_revisore'] = st.selectbox(
                "Tipo di Revisore",
                ["Revisore Legale (persona fisica)", "Società di Revisione"],
                key="tipo_revisore_nomina"
            )
        with col2:
            form_data['durata_incarico_revisore'] = st.selectbox(
                "Durata dell'incarico",
                ["Tre esercizi", "Un esercizio", "Altro"],
                key="durata_incarico_revisore"
            )

        # Dati del revisore
        st.subheader("Dati del Nuovo Revisore")
        if form_data.get('tipo_revisore') == "Revisore Legale (persona fisica)":
            form_data['nome_revisore'] = st.text_input("Nome e Cognome Revisore", key="nome_revisore_nomina")
            form_data['cf_revisore'] = st.text_input("Codice Fiscale Revisore", key="cf_revisore_nomina")
        else:
            form_data['ragione_sociale_revisore'] = st.text_input("Ragione Sociale Società di Revisione", key="ragione_sociale_revisore_nomina")
            form_data['piva_revisore'] = st.text_input("Partita IVA Società di Revisione", key="piva_revisore_nomina")

        # Compenso
        st.subheader("💰 Compenso Revisore")
        form_data['compenso_revisore'] = st.text_input("Compenso annuo lordo per il revisore (€)", "0,00", key="compenso_revisore_nomina")
        
        return form_data
    
    def show_preview(self, form_data: dict):
        """Mostra l'anteprima del documento fuori dal form"""
        # Sezione Anteprima
        st.markdown("---")
        st.subheader("👁️ Anteprima Documento")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col1:
            show_preview = st.checkbox("Mostra anteprima", value=False, key="preview_checkbox")
        
        if show_preview:
            try:
                # Debug dettagliato: mostra i dati estratti
                with st.expander("Debug: Dati estratti dal form"):
                    st.json(form_data)
                    
                    # Debug specifico per i soci
                    if 'soci' in form_data and form_data['soci']:
                        st.subheader("Debug Soci:")
                        for i, socio in enumerate(form_data['soci']):
                            st.write(f"Socio {i+1}:")
                            st.write(f"  - Nome: {socio.get('nome', 'N/A')}")
                            st.write(f"  - Tipo Soggetto: {socio.get('tipo_soggetto', 'N/A')}")
                            st.write(f"  - Tipo Partecipazione: {socio.get('tipo_partecipazione', 'N/A')}")
                            st.write(f"  - Delegato: {socio.get('delegato', 'N/A')}")
                            st.write(f"  - Rappresentante Legale: {socio.get('rappresentante_legale', 'N/A')}")
                            st.write(f"  - Presente: {socio.get('presente', 'N/A')}")
                            st.write("---")
                
                preview_text = self._generate_preview_text(form_data)
                st.text_area(
                    "Contenuto del verbale:",
                    value=preview_text,
                    height=600,
                    key="preview_text_nomina_revisore"
                )
            except Exception as e:
                st.error(f"Errore nell'anteprima: {e}")
    
    def _generate_preview_text(self, data: dict) -> str:
        """Genera il testo di anteprima del verbale"""
        
        # Estrazione dati di base
        denominazione = data.get('denominazione', '[DENOMINAZIONE]')
        sede = data.get('sede_legale', '[SEDE]')
        capitale = data.get('capitale_sociale', '[CAPITALE]')
        cf = data.get('codice_fiscale', '[CODICE FISCALE]')
        
        data_assemblea = data.get('data_assemblea', date.today())
        ora_assemblea = data.get('ora_assemblea', '10:00')
        presidente = data.get('presidente', '[PRESIDENTE]')
        segretario = data.get('segretario', '[SEGRETARIO]')
        
        # Dati revisore
        revisore_dati = data.get('revisore_dati', {})
        revisore_nome = revisore_dati.get('nome', '[NOME REVISORE]')
        durata_incarico = data.get('durata_incarico', '[DURATA INCARICO]')
        compenso = data.get('compenso_annuo', '[COMPENSO]')
        motivo_nomina = data.get('motivo_nomina', '[MOTIVO NOMINA]')
        
        if isinstance(data_assemblea, str):
            data_str = data_assemblea
        else:
            data_str = data_assemblea.strftime('%d/%m/%Y')
        
        # Generazione testo anteprima
        text = f"""{denominazione.upper()}
Sede in {sede}
Capitale sociale Euro {capitale} i.v.
Codice fiscale: {cf}

Verbale di assemblea dei soci
del {data_str}

Oggi {data_str} alle ore {ora_assemblea} presso la sede sociale {sede}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:

Ordine del giorno
{', '.join(data.get('punti_ordine_giorno', ['nomina del revisore della società']))}

Assume la presidenza ai sensi dell'art. [...] dello statuto sociale il Sig. {presidente} {data.get('ruolo_presidente', 'Amministratore Unico')}, il quale dichiara e constata:

1 - che l'assemblea risulta {data.get('tipo_assemblea', 'regolarmente convocata')}
"""

        # Aggiungi partecipazione in audioconferenza se prevista
        if data.get('modalita_partecipazione', False):
            text += "2 - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. [...] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza\n"
            next_num = 3
        else:
            next_num = 2
        
        # Aggiungi sezione partecipanti
        text += f"{next_num} - che sono presenti/partecipano all'assemblea:\n"
        text += f"l'Amministratore Unico nella persona del suddetto Presidente Sig. {presidente}\n"
        
        # Collegio Sindacale / Sindaco Unico se presente
        if data.get('include_collegio_sindacale'):
            tipo_oc = data.get('tipo_organo_controllo', 'Collegio Sindacale')
            sindaci = [s for s in data.get('sindaci', []) if s.get('presente')]
            if tipo_oc == 'Collegio Sindacale' and sindaci:
                text += "per il Collegio Sindacale:\n"
                for s in sindaci:
                    text += f"- {s.get('nome', '[NOME]')} {s.get('carica', '')}\n"
            elif tipo_oc == 'Sindaco Unico' and sindaci:
                text += f"il Sindaco Unico {sindaci[0].get('nome', '[NOME]')}\n"
        
        # Revisore se presente e diverso da quello da nominare
        if data.get('include_revisore'):
            nome_rev_esistente = data.get('nome_revisore', '[NOME REVISORE]')
            text += f"il revisore contabile Dott. {nome_rev_esistente}\n"
        
        # Gestione soci presenti e assenti in modo uniforme
        soci_presenti = data.get('soci_presenti', [])
        soci_assenti = data.get('soci_assenti', [])

        # Fallback retrocompatibilità
        if not soci_presenti and not soci_assenti and 'soci' in data:
            soci_presenti = [s for s in data.get('soci', []) if s.get('presente', True)]
            soci_assenti = [s for s in data.get('soci', []) if not s.get('presente', True)]

        if soci_presenti:
            total_quota_euro = 0.0
            total_quota_percentuale = 0.0

            for socio in soci_presenti:
                try:
                    quota_euro_str = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                    total_quota_euro += float(quota_euro_str)
                except (ValueError, TypeError):
                    pass
                try:
                    quota_perc_str = str(socio.get('quota_percentuale', '0')).replace('.', '').replace(',', '.')
                    total_quota_percentuale += float(quota_perc_str)
                except (ValueError, TypeError):
                    pass

            formatted_total_quota_euro = f"{total_quota_euro:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            formatted_total_quota_percentuale = f"{total_quota_percentuale:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

            text += f"nonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_total_quota_euro} pari al {formatted_total_quota_percentuale}% del Capitale Sociale:\n"

            for socio in soci_presenti:
                nome = socio.get('nome', '[NOME SOCIO]')
                quota_raw = socio.get('quota_euro', '')
                perc_raw = socio.get('quota_percentuale', '')

                quota = '[QUOTA]' if not quota_raw or str(quota_raw).strip() == '' else str(quota_raw).strip()
                perc = '[%]' if not perc_raw or str(perc_raw).strip() == '' else str(perc_raw).strip()

                tipo_partecipazione = socio.get('tipo_partecipazione', 'Diretto')
                delegato = socio.get('delegato', '').strip()
                tipo_soggetto = socio.get('tipo_soggetto', 'Persona Fisica')
                rappresentante_legale = socio.get('rappresentante_legale', '').strip()

                if tipo_partecipazione == 'Delegato' and delegato:
                    if tipo_soggetto == 'Società':
                        line = f"il Sig. {delegato} delegato della società {nome}"
                        if rappresentante_legale:
                            line += f" (nella persona del legale rappresentante {rappresentante_legale})"
                    else:
                        line = f"il Sig. {delegato} delegato del socio {nome}"
                    line += f" recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"
                else:
                    if tipo_soggetto == 'Società':
                        if rappresentante_legale:
                            line = f"la società {nome} nella persona del legale rappresentante {rappresentante_legale} recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"
                        else:
                            line = f"la società {nome} recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"
                    else:
                        line = f"il Sig. {nome} socio recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"

                text += line + "\n"

        # Soci assenti
        if soci_assenti:
            text += "Risultano invece assenti i seguenti soci:\n"
            for socio in soci_assenti:
                nome = socio.get('nome', '[NOME SOCIO]')
                text += f"- Sig. {nome}\n"
        
        text += f"""
{next_num+1} - che gli intervenuti sono legittimati alla presente assemblea;
{next_num+2} - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno.

I presenti all'unanimità chiamano a fungere da segretario il signor {segretario}, che accetta l'incarico.

Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante.

Il Presidente constata e fa constatare che l'assemblea risulta {data.get('tipo_assemblea', 'regolarmente convocata')} e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.

Si passa quindi allo svolgimento dell'ordine del giorno.

*     *     *

Il Presidente informa l'assemblea che si rende necessaria la nomina del Revisore della società poiché {motivo_nomina}.

Il Presidente ricorda all'assemblea quanto previsto dall'art. 2477 del Codice Civile e dall'atto costitutivo della società.

Prende la parola il socio sig. {data.get('socio_proponente', '[SOCIO PROPONENTE]')} che propone di affidare l'incarico di Revisore al Dott. {revisore_nome} che, prima dell'accettazione dell'incarico, ha reso noti all'assemblea gli incarichi di amministrazione e di controllo ricoperti presso altre società, confermando la propria indipendenza e l'insussistenza di cause di ineleggibilità.

Segue breve discussione tra i soci al termine della quale si passa alla votazione con voto palese in forza della quale il Presidente constata che, all'unanimità, l'assemblea

DELIBERA:

di nominare Revisore della società, per {durata_incarico} il Dott. {revisore_nome} nato a {revisore_dati.get('nato_a', '[LUOGO NASCITA]')} il {revisore_dati.get('nato_il', '[DATA NASCITA]')}, residente in {revisore_dati.get('residente', '[LUOGO RESIDENZA]')}, codice fiscale {revisore_dati.get('codice_fiscale', '[CODICE FISCALE]')}, Revisore Contabile pubblicato sulla G.U. in data {revisore_dati.get('albo_data_gu', '[DATA GU]')} N. {revisore_dati.get('albo_num_gu', '[NUM GU]')}, {revisore_dati.get('albo_serie_gu', '[SERIE GU]')} serie speciale.

di corrispondere al Revisore un compenso annuo pari a euro {compenso}
"""

        # Aggiunta presenza revisore se applicabile
        if data.get('revisore_presente', False):
            text += f"""
Il Dott. {revisore_nome} presente in assemblea in qualità di {data.get('revisore_qualita', 'invitato')} accetta l'incarico e ringrazia l'assemblea per la fiducia accordata.
"""
        else:
            text += f"""
[L'accettazione della carica da parte del Dott. {revisore_nome} potrà avvenire successivamente]
"""

        text += """
*     *     *

Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.

Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo.

L'assemblea viene sciolta alle ore [...].


Il Presidente                    Il Segretario
_________________            _________________
"""

        return text
    
    def generate_document(self, data: dict) -> Document:
        """Genera il documento Word del verbale di nomina revisore con formattazione professionale"""
        # Usa il metodo della classe base per setup professionale
        doc = self._setup_professional_document(data)
        
        # Aggiungi intestazione azienda (dalla classe base)
        self._add_company_header(doc, data)
        
        # Aggiungi titolo verbale (dalla classe base)
        self._add_verbale_title(doc, data)
        
        # Aggiungi sezione di apertura con stile professionale
        self._add_opening_section_professional(doc, data)
        
        # Aggiungi sezione partecipanti con stile professionale
        self._add_participants_section_professional(doc, data)
        
        # Aggiungi dichiarazioni preliminari con stile professionale
        self._add_preliminary_statements_professional(doc, data)
        
        # Aggiungi discussione nomina revisore
        self._add_nomina_revisore_discussion(doc, data)
        
        # Aggiungi sezione di chiusura con stile professionale
        self._add_closing_section_professional(doc, data)
        
        # Aggiungi firme professionali (dalla classe base)
        self._add_professional_signatures(doc, data)
        
        return doc
    
    def _setup_document_styles(self, doc):
        """Imposta gli stili del documento usando solo quelli della classe base"""
        # Usa solo gli stili standardizzati della classe base
        super()._setup_document_styles(doc)
    
    # Rimuovo il metodo sovrascritto per usare quello della classe base
    
    # Rimuovo il metodo sovrascritto per usare quello della classe base
    
    def _add_opening_section(self, doc, data):
        """Aggiungi sezione di apertura"""
        data_assemblea = data.get('data_assemblea', date.today())
        if isinstance(data_assemblea, str):
            data_str = data_assemblea
        else:
            data_str = data_assemblea.strftime('%d/%m/%Y')
        
        ora = data.get('ora_assemblea', '10:00')
        sede = data.get('sede_legale', '[SEDE]')
        
        p = doc.add_paragraph()
        p.add_run(f"Oggi {data_str} alle ore {ora} presso la sede sociale {sede}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:")
        
        # Ordine del giorno
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Ordine del giorno")
        run.bold = True
        
        punti_odg = data.get('punti_ordine_giorno', ['nomina del revisore della società'])
        for punto in punti_odg:
            p = doc.add_paragraph()
            p.add_run(punto)
        
        doc.add_paragraph()
    
    def _add_participants_section(self, doc, data):
        """Aggiungi sezione partecipanti"""
        presidente = data.get('presidente', '[PRESIDENTE]')
        ruolo_presidente = data.get('ruolo_presidente', 'Amministratore Unico')
        
        # Presidenza
        p = doc.add_paragraph()
        p.add_run(f"Assume la presidenza ai sensi dell'art. [...] dello statuto sociale il Sig. {presidente} {ruolo_presidente}, il quale dichiara e constata:")
        
        # Dichiarazioni
        counter = 1
        
        # Tipo assemblea
        p = doc.add_paragraph()
        p.add_run(f"{counter} - che l'assemblea risulta {data.get('tipo_assemblea', 'regolarmente convocata')}")
        counter += 1
        
        # Audioconferenza se prevista
        if data.get('modalita_partecipazione', False):
            p = doc.add_paragraph()
            p.add_run(f"{counter} - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. [...] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza")
            counter += 1
        
        # Presenti
        p = doc.add_paragraph()
        p.add_run(f"{counter} - che sono presenti/partecipano all'assemblea:")
        
        p = doc.add_paragraph()
        p.add_run(f"l'Amministratore Unico nella persona del suddetto Presidente Sig. {presidente}")
        
        # Organi di controllo
        self._add_organi_controllo_paragraphs(doc, data)
        
        # Gestione soci presenti e assenti in modo uniforme
        soci_presenti = data.get('soci_presenti', [])
        soci_assenti = data.get('soci_assenti', [])

        # Fallback retrocompatibilità
        if not soci_presenti and not soci_assenti and 'soci' in data:
            soci_presenti = [s for s in data.get('soci', []) if s.get('presente', True)]
            soci_assenti = [s for s in data.get('soci', []) if not s.get('presente', True)]

        if soci_presenti:
            total_quota_euro = 0.0
            total_quota_percentuale = 0.0

            for socio in soci_presenti:
                try:
                    quota_euro_str = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                    total_quota_euro += float(quota_euro_str)
                except (ValueError, TypeError):
                    pass
                try:
                    quota_perc_str = str(socio.get('quota_percentuale', '0')).replace('.', '').replace(',', '.')
                    total_quota_percentuale += float(quota_perc_str)
                except (ValueError, TypeError):
                    pass

            formatted_total_quota_euro = f"{total_quota_euro:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            formatted_total_quota_percentuale = f"{total_quota_percentuale:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

            p = doc.add_paragraph()
            p.add_run(f"nonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_total_quota_euro} pari al {formatted_total_quota_percentuale}% del Capitale Sociale:")

            for socio in soci_presenti:
                nome = socio.get('nome', '[NOME SOCIO]')
                quota_raw = socio.get('quota_euro', '')
                perc_raw = socio.get('quota_percentuale', '')

                quota = '[QUOTA]' if not quota_raw or str(quota_raw).strip() == '' else str(quota_raw).strip()
                perc = '[%]' if not perc_raw or str(perc_raw).strip() == '' else str(perc_raw).strip()

                tipo_partecipazione = socio.get('tipo_partecipazione', 'Diretto')
                delegato = socio.get('delegato', '').strip()
                tipo_soggetto = socio.get('tipo_soggetto', 'Persona Fisica')
                rappresentante_legale = socio.get('rappresentante_legale', '').strip()

                p = doc.add_paragraph()

                if tipo_partecipazione == 'Delegato' and delegato:
                    if tipo_soggetto == 'Società':
                        line = f"il Sig. {delegato} delegato della società {nome}"
                        if rappresentante_legale:
                            line += f" (nella persona del legale rappresentante {rappresentante_legale})"
                    else:
                        line = f"il Sig. {delegato} delegato del socio {nome}"
                    line += f" recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"
                else:
                    if tipo_soggetto == 'Società':
                        if rappresentante_legale:
                            line = f"la società {nome} nella persona del legale rappresentante {rappresentante_legale} recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"
                        else:
                            line = f"la società {nome} recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"
                    else:
                        line = f"il Sig. {nome} socio recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"

                p.add_run(line)

        # Soci assenti
        if soci_assenti:
            p = doc.add_paragraph()
            p.add_run("Risultano invece assenti i seguenti soci:")
            for socio in soci_assenti:
                nome = socio.get('nome', '[NOME SOCIO]')
                p = doc.add_paragraph()
                p.add_run(f"- Sig. {nome}")
        
        counter += 1
        
        # Altre dichiarazioni
        p = doc.add_paragraph()
        p.add_run(f"{counter} - che gli intervenuti sono legittimati alla presente assemblea;")
        counter += 1
        
        p = doc.add_paragraph()
        p.add_run(f"{counter} - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno.")
        
        doc.add_paragraph()
    
    def _add_preliminary_statements(self, doc, data):
        """Aggiungi dichiarazioni preliminari"""
        segretario = data.get('segretario', '[SEGRETARIO]')
        
        p = doc.add_paragraph()
        p.add_run(f"I presenti all'unanimità chiamano a fungere da segretario il signor {segretario}, che accetta l'incarico.")
        
        p = doc.add_paragraph()
        p.add_run("Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante.")
        
        p = doc.add_paragraph()
        tipo_assemblea = data.get('tipo_assemblea', 'regolarmente convocata')
        p.add_run(f"Il Presidente constata e fa constatare che l'assemblea risulta {tipo_assemblea} e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.")
        
        p = doc.add_paragraph()
        p.add_run("Si passa quindi allo svolgimento dell'ordine del giorno.")
        
        # Separatore
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("*     *     *")
        
        doc.add_paragraph()
    
    def _add_nomina_revisore_discussion(self, doc, data):
        """Aggiungi sezione discussione nomina revisore"""
        motivo_nomina = data.get('motivo_nomina', '[MOTIVO NOMINA]')
        socio_proponente = data.get('socio_proponente', '[SOCIO PROPONENTE]')
        revisore_dati = data.get('revisore_dati', {})
        revisore_nome = revisore_dati.get('nome', '[NOME REVISORE]')
        
        # Informativa presidente
        p = doc.add_paragraph()
        p.add_run(f"Il Presidente informa l'assemblea che si rende necessaria la nomina del Revisore della società poiché {motivo_nomina}.")
        
        p = doc.add_paragraph()
        p.add_run("Il Presidente ricorda all'assemblea quanto previsto dall'art. 2477 del Codice Civile e dall'atto costitutivo della società.")
        
        # Proposta
        p = doc.add_paragraph()
        p.add_run(f"Prende la parola il socio sig. {socio_proponente} che propone di affidare l'incarico di Revisore al Dott. {revisore_nome} che, prima dell'accettazione dell'incarico, ha reso noti all'assemblea gli incarichi di amministrazione e di controllo ricoperti presso altre società, confermando la propria indipendenza e l'insussistenza di cause di ineleggibilità.")
        
        # Discussione e votazione
        p = doc.add_paragraph()
        p.add_run("Segue breve discussione tra i soci al termine della quale si passa alla votazione con voto palese in forza della quale il Presidente constata che, all'unanimità, l'assemblea")
        
        # DELIBERA
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("D E L I B E R A:")
        run.bold = True
        
        doc.add_paragraph()
        
        # Nomina
        durata_incarico = data.get('durata_incarico', '[DURATA INCARICO]')
        nato_a = revisore_dati.get('nato_a', '[LUOGO NASCITA]')
        nato_il = revisore_dati.get('nato_il', '[DATA NASCITA]')
        if isinstance(nato_il, date):
            nato_il_str = nato_il.strftime('%d/%m/%Y')
        else:
            nato_il_str = str(nato_il)
        
        residente = revisore_dati.get('residente', '[LUOGO RESIDENZA]')
        cf_revisore = revisore_dati.get('codice_fiscale', '[CODICE FISCALE]')
        
        albo_data = revisore_dati.get('albo_data_gu', '[DATA GU]')
        if isinstance(albo_data, date):
            albo_data_str = albo_data.strftime('%d/%m/%Y')
        else:
            albo_data_str = str(albo_data)
            
        albo_num = revisore_dati.get('albo_num_gu', '[NUM GU]')
        albo_serie = revisore_dati.get('albo_serie_gu', '[SERIE GU]')
        
        p = doc.add_paragraph()
        p.add_run(f"di nominare Revisore della società, per {durata_incarico} il Dott. {revisore_nome} nato a {nato_a} il {nato_il_str}, residente in {residente}, codice fiscale {cf_revisore}, Revisore Contabile pubblicato sulla G.U. in data {albo_data_str} N. {albo_num}, {albo_serie} serie speciale.")
        
        # Compenso
        compenso = data.get('compenso_annuo', '[COMPENSO]')
        p = doc.add_paragraph()
        p.add_run(f"di corrispondere al Revisore un compenso annuo pari a euro {compenso}")
        
        doc.add_paragraph()
        
        # Accettazione
        if data.get('revisore_presente', False):
            qualita = data.get('revisore_qualita', 'invitato')
            p = doc.add_paragraph()
            p.add_run(f"Il Dott. {revisore_nome} presente in assemblea in qualità di {qualita} accetta l'incarico e ringrazia l'assemblea per la fiducia accordata.")
        else:
            p = doc.add_paragraph()
            p.add_run(f"[L'accettazione della carica da parte del Dott. {revisore_nome} potrà avvenire successivamente]")
        
        # Separatore
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("*     *     *")
        
        doc.add_paragraph()
    
    def _add_closing_section(self, doc, data):
        """Aggiungi sezione di chiusura"""
        p = doc.add_paragraph()
        p.add_run("Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.")
        
        p = doc.add_paragraph()
        p.add_run("Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo.")
        
        p = doc.add_paragraph()
        p.add_run("L'assemblea viene sciolta alle ore [...].")
        
        doc.add_paragraph()
        doc.add_paragraph()
    
    def _add_opening_section_professional(self, doc, data):
        """Aggiungi sezione di apertura con stile professionale"""
        data_assemblea = data.get('data_assemblea', date.today())
        if isinstance(data_assemblea, str):
            data_str = data_assemblea
        else:
            data_str = data_assemblea.strftime('%d/%m/%Y')
        
        ora = data.get('ora_assemblea', '10:00')
        sede = data.get('sede_legale', '[SEDE]')
        
        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"Oggi {data_str} alle ore {ora} presso la sede sociale {sede}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:")
        
        # Ordine del giorno
        doc.add_paragraph()
        p = doc.add_paragraph(style='Heading1')
        p.add_run("ORDINE DEL GIORNO")
        
        punti_odg = data.get('punti_ordine_giorno', ['nomina del revisore della società'])
        for punto in punti_odg:
            p = doc.add_paragraph(style='BodyText')
            p.add_run(punto)
        
        doc.add_paragraph()

    def _add_participants_section_professional(self, doc, data):
        """Aggiungi sezione partecipanti con stile professionale"""
        presidente = data.get('presidente', '[PRESIDENTE]')
        ruolo_presidente = data.get('ruolo_presidente', 'Amministratore Unico')
        
        # Presidenza
        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"Assume la presidenza ai sensi dell'art. [...] dello statuto sociale il Sig. {presidente} {ruolo_presidente}, il quale dichiara e constata:")
        
        # Dichiarazioni
        counter = 1
        
        # Tipo assemblea
        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"{counter} - che l'assemblea risulta {data.get('tipo_assemblea', 'regolarmente convocata')}")
        counter += 1
        
        # Audioconferenza se prevista
        if data.get('modalita_partecipazione', False):
            p = doc.add_paragraph(style='BodyText')
            p.add_run(f"{counter} - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. [...] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza")
            counter += 1
        
        # Presenti
        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"{counter} - che sono presenti/partecipano all'assemblea:")
        
        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"l'Amministratore Unico nella persona del suddetto Presidente Sig. {presidente}")
        
        # Organi di controllo
        self._add_organi_controllo_paragraphs(doc, data)
        
        # Gestione soci presenti e assenti in modo uniforme
        soci_presenti = data.get('soci_presenti', [])
        soci_assenti = data.get('soci_assenti', [])

        # Fallback retrocompatibilità
        if not soci_presenti and not soci_assenti and 'soci' in data:
            soci_presenti = [s for s in data.get('soci', []) if s.get('presente', True)]
            soci_assenti = [s for s in data.get('soci', []) if not s.get('presente', True)]

        if soci_presenti:
            total_quota_euro = 0.0
            total_quota_percentuale = 0.0

            for socio in soci_presenti:
                try:
                    quota_euro_str = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                    total_quota_euro += float(quota_euro_str)
                except (ValueError, TypeError):
                    pass
                try:
                    quota_perc_str = str(socio.get('quota_percentuale', '0')).replace('.', '').replace(',', '.')
                    total_quota_percentuale += float(quota_perc_str)
                except (ValueError, TypeError):
                    pass

            formatted_total_quota_euro = f"{total_quota_euro:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            formatted_total_quota_percentuale = f"{total_quota_percentuale:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

            p = doc.add_paragraph()
            p.add_run(f"nonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_total_quota_euro} pari al {formatted_total_quota_percentuale}% del Capitale Sociale:")

            for socio in soci_presenti:
                nome = socio.get('nome', '[NOME SOCIO]')
                quota_raw = socio.get('quota_euro', '')
                perc_raw = socio.get('quota_percentuale', '')

                quota = '[QUOTA]' if not quota_raw or str(quota_raw).strip() == '' else str(quota_raw).strip()
                perc = '[%]' if not perc_raw or str(perc_raw).strip() == '' else str(perc_raw).strip()

                tipo_partecipazione = socio.get('tipo_partecipazione', 'Diretto')
                delegato = socio.get('delegato', '').strip()
                tipo_soggetto = socio.get('tipo_soggetto', 'Persona Fisica')
                rappresentante_legale = socio.get('rappresentante_legale', '').strip()

                if tipo_partecipazione == 'Delegato' and delegato:
                    if tipo_soggetto == 'Società':
                        line = f"il Sig. {delegato} delegato della società {nome}"
                        if rappresentante_legale:
                            line += f" (nella persona del legale rappresentante {rappresentante_legale})"
                    else:
                        line = f"il Sig. {delegato} delegato del socio {nome}"
                    line += f" recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"
                else:
                    if tipo_soggetto == 'Società':
                        if rappresentante_legale:
                            line = f"la società {nome} nella persona del legale rappresentante {rappresentante_legale} recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"
                        else:
                            line = f"la società {nome} recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"
                    else:
                        line = f"il Sig. {nome} socio recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"

                p = doc.add_paragraph()

                if tipo_partecipazione == 'Delegato' and delegato:
                    if tipo_soggetto == 'Società':
                        line = f"il Sig. {delegato} delegato della società {nome}"
                        if rappresentante_legale:
                            line += f" (nella persona del legale rappresentante {rappresentante_legale})"
                    else:
                        line = f"il Sig. {delegato} delegato del socio {nome}"
                    line += f" recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"
                else:
                    if tipo_soggetto == 'Società':
                        if rappresentante_legale:
                            line = f"la società {nome} nella persona del legale rappresentante {rappresentante_legale} recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"
                        else:
                            line = f"la società {nome} recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"
                    else:
                        line = f"il Sig. {nome} socio recante una quota pari a nominali euro {quota} pari al {perc}% del Capitale Sociale"

                p.add_run(line)

        # Soci assenti
        if soci_assenti:
            text += "Risultano invece assenti i seguenti soci:\n"
            for socio in soci_assenti:
                nome = socio.get('nome', '[NOME SOCIO]')
                text += f"- Sig. {nome}\n"
        
        counter += 1
        
        # Altre dichiarazioni
        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"{counter} - che gli intervenuti sono legittimati alla presente assemblea;")
        counter += 1
        
        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"{counter} - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno.")
        
        doc.add_paragraph()

    def _add_preliminary_statements_professional(self, doc, data):
        """Aggiungi dichiarazioni preliminari con stile professionale"""
        segretario = data.get('segretario', '[SEGRETARIO]')
        
        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"I presenti all'unanimità chiamano a fungere da segretario il signor {segretario}, che accetta l'incarico.")
        
        p = doc.add_paragraph(style='BodyText')
        p.add_run("Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante.")
        
        p = doc.add_paragraph(style='BodyText')
        tipo_assemblea = data.get('tipo_assemblea', 'regolarmente convocata')
        p.add_run(f"Il Presidente constata e fa constatare che l'assemblea risulta {tipo_assemblea} e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.")
        
        p = doc.add_paragraph(style='BodyText')
        p.add_run("Si passa quindi allo svolgimento dell'ordine del giorno.")
        
        # Separatore
        p = doc.add_paragraph(style='BodyText')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("*     *     *")
        run.bold = True
        
        doc.add_paragraph()

    def _add_closing_section_professional(self, doc, data):
        """Aggiungi sezione di chiusura con stile professionale"""
        p = doc.add_paragraph(style='BodyText')
        p.add_run("Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.")
        
        p = doc.add_paragraph(style='BodyText')
        p.add_run("Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo.")
        
        p = doc.add_paragraph(style='BodyText')
        p.add_run("L'assemblea viene sciolta alle ore [...].")
        
        doc.add_paragraph()

    def _add_organi_controllo_paragraphs(self, doc, data):
        # Implementa la logica per aggiungere i paragrafi relativi agli organi di controllo
        pass

# Registra il template
DocumentTemplateFactory.register_template("nomina_revisore", VerbaleNominaRevisoreTemplate)
