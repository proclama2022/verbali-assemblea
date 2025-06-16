"""
Template semplificato per Verbale di Assemblea - Approvazione Bilancio
"""

import sys
import os

# Aggiungi il path della cartella src
current_dir = os.path.dirname(os.path.abspath(__file__))
src_path = os.path.join(os.path.dirname(current_dir), 'src')
if src_path not in sys.path:
    sys.path.append(src_path)

from document_templates import DocumentTemplate, DocumentTemplateFactory
from base_verbale_template import BaseVerbaleTemplate
from docx import Document
from datetime import date
import streamlit as st
from common_data_handler import CommonDataHandler

class VerbaleStandardTemplate(BaseVerbaleTemplate):
    """Template semplificato per Verbale di Assemblea dei Soci - Approvazione Bilancio"""
    
    def get_template_name(self) -> str:
        return "Verbale Standard - Approvazione Bilancio"
    
    def get_required_fields(self) -> list:
        return [
            "denominazione", "sede_legale", "capitale_sociale", "codice_fiscale",
            "data_assemblea", "presidente", "segretario"
        ]
    
    def get_form_fields(self, extracted_data: dict) -> dict:
        """Genera i campi del form per Streamlit in modo semplificato"""
        form_data = {}
        
        # Dati azienda
        st.subheader("🏢 Dati Società")
        col1, col2 = st.columns(2)
        with col1:
            form_data["denominazione"] = st.text_input("Denominazione", extracted_data.get("denominazione", ""))
            form_data["sede_legale"] = st.text_input("Sede legale", extracted_data.get("sede_legale", ""))
        with col2:
            form_data["capitale_sociale"] = st.text_input("Capitale sociale", extracted_data.get("capitale_sociale", "10.000,00"))
            form_data["codice_fiscale"] = st.text_input("Codice fiscale", extracted_data.get("codice_fiscale", ""))
        
        # Dati assemblea
        st.subheader("📅 Dati Assemblea")
        col1, col2 = st.columns(2)
        with col1:
            form_data["data_assemblea"] = st.date_input("Data assemblea", date.today())
            form_data["ora_inizio"] = st.text_input("Ora inizio", "09:00")
        with col2:
            form_data["data_chiusura_bilancio"] = st.date_input("Data chiusura bilancio", date(date.today().year - 1, 12, 31))
            form_data["ora_fine"] = st.text_input("Ora fine", "10:00")
        
        # Partecipanti
        st.subheader("👥 Partecipanti")
        col1, col2 = st.columns(2)
        with col1:
            form_data["presidente"] = st.text_input("Presidente", extracted_data.get("presidente", ""))
            form_data["segretario"] = st.text_input("Segretario", extracted_data.get("segretario", ""))
        with col2:
            form_data["ruolo_presidente"] = st.selectbox("Ruolo presidente", 
                                                        CommonDataHandler.get_standard_ruoli_presidente())
        
        # Soci (semplificato)
        st.subheader("👥 Soci Presenti")
        soci_text = st.text_area("Elenco soci (uno per riga)", 
                                 extracted_data.get("soci_text", "Mario Rossi - 50%\nLuigi Bianchi - 50%"),
                                 height=100)
        form_data["soci_text"] = soci_text
        
        # Risultato esercizio
        st.subheader("💰 Risultato Esercizio")
        col1, col2 = st.columns(2)
        with col1:
            form_data["utile_esercizio"] = st.text_input("Utile dell'esercizio (€)", "0,00")
            form_data["destinazione_utile"] = st.text_area("Destinazione utile", 
                                                          "- a riserva legale per il 5%\n- a nuovo il residuo")
        with col2:
            form_data["esito_votazione"] = st.selectbox("Esito votazione", 
                                                       ["approvato all'unanimità", "approvato a maggioranza"])
        
        return form_data
    
    def show_preview(self, form_data: dict):
        """Mostra anteprima semplificata"""
        with st.expander("👁️ Anteprima Documento", expanded=False):
            preview_text = self._generate_preview_text(form_data)
            st.text_area("Anteprima", preview_text, height=400, disabled=True)
    
    def _generate_preview_text(self, data: dict) -> str:
        """Genera anteprima testuale semplificata"""
        try:
            denominazione = data.get('denominazione', '[DENOMINAZIONE]')
            sede_legale = data.get('sede_legale', '[SEDE]')
            capitale_sociale_raw = data.get('capitale_versato') or data.get('capitale_deliberato') or data.get('capitale_sociale', '[CAPITALE]')
            capitale_sociale = CommonDataHandler.format_currency(capitale_sociale_raw)
            codice_fiscale = data.get('codice_fiscale', '[CF]')
            
            data_assemblea = data.get('data_assemblea', date.today())
            if hasattr(data_assemblea, 'strftime'):
                data_str = data_assemblea.strftime('%d/%m/%Y')
            else:
                data_str = str(data_assemblea)
            
            data_chiusura = data.get('data_chiusura_bilancio', date.today())
            if hasattr(data_chiusura, 'strftime'):
                data_chiusura_str = data_chiusura.strftime('%d/%m/%Y')
            else:
                data_chiusura_str = str(data_chiusura)
            
            presidente = data.get('presidente', '[PRESIDENTE]')
            segretario = data.get('segretario', '[SEGRETARIO]')
            ruolo_presidente = data.get('ruolo_presidente', 'Amministratore Unico')
            ora_inizio = data.get('ora_inizio', '[ORA]')
            ora_fine = data.get('ora_fine', '[ORA]')
            soci_text = data.get('soci_text', '[SOCI]')
            utile_esercizio = data.get('utile_esercizio', '[UTILE]')
            destinazione_utile = data.get('destinazione_utile', '[DESTINAZIONE]')
            esito_votazione = data.get('esito_votazione', 'approvato all\'unanimità')
            
            preview = f"""{denominazione}
Sede in {sede_legale}
Capitale sociale Euro {capitale_sociale} i.v.
Codice fiscale: {codice_fiscale}

VERBALE DI ASSEMBLEA DEI SOCI
del {data_str}

Oggi {data_str} alle ore {ora_inizio} presso la sede sociale si è tenuta l'assemblea generale dei soci per discutere e deliberare sul seguente:

ORDINE DEL GIORNO
1. Approvazione del Bilancio al {data_chiusura_str}
2. Destinazione del risultato dell'esercizio

Assume la presidenza il Sig. {presidente} {ruolo_presidente}, il quale constata che sono presenti:

SOCI PRESENTI:
{soci_text}

I presenti chiamano a fungere da segretario il Sig. {segretario}.

Il Presidente constata che l'assemblea è validamente costituita e atta a deliberare.

PRIMO PUNTO - APPROVAZIONE BILANCIO
Il Presidente illustra il bilancio al {data_chiusura_str}. Dopo discussione, l'assemblea {esito_votazione}

DELIBERA
di approvare il bilancio di esercizio al {data_chiusura_str}.

SECONDO PUNTO - DESTINAZIONE RISULTATO
Il Presidente propone di destinare l'utile di Euro {utile_esercizio} come segue:
{destinazione_utile}

L'assemblea {esito_votazione}

DELIBERA
di destinare l'utile come proposto.

L'assemblea viene sciolta alle ore {ora_fine}.

                    Il Presidente                    Il Segretario
                 {presidente}                     {segretario}
"""
            return preview
            
        except Exception as e:
            return f"Errore nella generazione dell'anteprima: {str(e)}"
    
    def generate_document(self, data: dict) -> Document:
        """Genera il documento Word semplificato"""
        doc = self._setup_professional_document(data)
        
        # Intestazione società
        self._add_company_header(doc, data)
        
        # Titolo verbale
        self._add_verbale_title(doc, data)
        
        # Contenuto principale
        self._add_main_content(doc, data)
        
        # Firme
        self._add_signatures(doc, data)
        
        return doc
    
    def _add_main_content(self, doc, data):
        """Aggiunge il contenuto principale del verbale"""
        # Apertura assemblea
        data_str = data.get('data_assemblea').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[DATA]'
        ora_inizio = data.get('ora_inizio', '[ORA]')
        
        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"Oggi {data_str} alle ore {ora_inizio} presso la sede sociale si è tenuta l'assemblea generale dei soci per discutere e deliberare sul seguente:")

        # Ordine del giorno
        doc.add_paragraph()
        p = doc.add_paragraph(style='BodyText')
        p.add_run("ORDINE DEL GIORNO").bold = True

        # Partecipanti
        presidente = data.get('presidente', '[PRESIDENTE]')
        ruolo_presidente = data.get('ruolo_presidente', 'Amministratore Unico')
        soci_text = data.get('soci_text', '[SOCI]')
        segretario = data.get('segretario', '[SEGRETARIO]')

        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"Assume la presidenza il Sig. {presidente} {ruolo_presidente}, il quale constata che sono presenti:")

        doc.add_paragraph(style='BodyText').add_run("SOCI PRESENTI:").bold = True
        for socio_line in soci_text.split('\n'):
            if socio_line.strip():
                doc.add_paragraph(socio_line.strip(), style='ListParagraph')

        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"I presenti chiamano a fungere da segretario il Sig. {segretario}.")

        p = doc.add_paragraph(style='BodyText')
        p.add_run("Il Presidente constata che l'assemblea è validamente costituita e atta a deliberare.")

        # Primo punto - Approvazione Bilancio
        doc.add_paragraph()
        p = doc.add_paragraph(style='BodyText')
        p.add_run("PRIMO PUNTO - APPROVAZIONE BILANCIO").bold = True

        data_chiusura_str = data.get('data_chiusura_bilancio').strftime('%d/%m/%Y') if hasattr(data.get('data_chiusura_bilancio'), 'strftime') else '[DATA]'
        esito_votazione = data.get('esito_votazione', 'approvato all\'unanimità')

        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"Il Presidente illustra il bilancio al {data_chiusura_str}. Dopo discussione, l'assemblea {esito_votazione}")

        doc.add_paragraph(style='BodyText').add_run("DELIBERA").bold = True
        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"di approvare il bilancio di esercizio al {data_chiusura_str}.")

        # Secondo punto - Destinazione risultato
        doc.add_paragraph()
        p = doc.add_paragraph(style='BodyText')
        p.add_run("SECONDO PUNTO - DESTINAZIONE RISULTATO").bold = True

        utile_esercizio = data.get('utile_esercizio', '[UTILE]')
        destinazione_utile = data.get('destinazione_utile', '[DESTINAZIONE]')

        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"Il Presidente propone di destinare l'utile di Euro {utile_esercizio} come segue:")

        for line in destinazione_utile.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip(), style='ListParagraph')

        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"L'assemblea {esito_votazione}")

        doc.add_paragraph(style='BodyText').add_run("DELIBERA").bold = True
        p = doc.add_paragraph(style='BodyText')
        p.add_run("di destinare l'utile come proposto.")

        # Chiusura assemblea
        ora_fine = data.get('ora_fine', '[ORA]')
        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"L'assemblea viene sciolta alle ore {ora_fine}.")



    def _add_signatures(self, doc, data):
        """Aggiunge le firme utilizzando una tabella formattata."""
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE

        # Aggiungi spazio prima delle firme
        doc.add_paragraph().paragraph_format.space_before = Pt(36)
        
        table = doc.add_table(rows=2, cols=2)
        table.autofit = False # Permette di controllare le larghezze delle colonne
        table.columns[0].width = Inches(3.0)
        table.columns[1].width = Inches(3.0)
        # table.style = 'TableGrid' # Utilizza lo stile di tabella predefinita o uno personalizzato

        # Riga 1: Nomi
        row_cells = table.rows[0].cells
        p_presidente = row_cells[0].add_paragraph()
        p_presidente.add_run("Il Presidente").bold = True
        p_presidente.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p_segretario = row_cells[1].add_paragraph()
        p_segretario.add_run("Il Segretario").bold = True
        p_segretario.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Riga 2: Firme
        row_cells = table.rows[1].cells
        p_firma_presidente = row_cells[0].add_paragraph()
        p_firma_presidente.add_run(data.get('presidente', '[PRESIDENTE]'))
        p_firma_presidente.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p_firma_segretario = row_cells[1].add_paragraph()
        p_firma_segretario.add_run(data.get('segretario', '[SEGRETARIO]'))
        p_firma_segretario.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Aggiungi un paragrafo vuoto per spaziatura dopo la tabella
        doc.add_paragraph()

# Registra il template nel factory
DocumentTemplateFactory.register_template("verbale_standard", VerbaleStandardTemplate)
