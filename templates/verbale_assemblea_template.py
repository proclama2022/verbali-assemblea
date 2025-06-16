"""
Template per Verbale di Assemblea - Approvazione Bilancio
Questo template è specifico per verbali di approvazione bilancio.
"""

import sys
import os

# Aggiungi il path della cartella src (relativo alla root del progetto)
current_dir = os.path.dirname(os.path.abspath(__file__))
src_path = os.path.join(os.path.dirname(current_dir), 'src')
if src_path not in sys.path:
    sys.path.append(src_path)

from document_templates import DocumentTemplate, DocumentTemplateFactory
from common_data_handler import CommonDataHandler
from base_verbale_template import BaseVerbaleTemplate
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from datetime import date
import streamlit as st
import pandas as pd
import re

class VerbaleApprovazioneBilancioTemplate(BaseVerbaleTemplate):
    """Template per Verbale di Assemblea dei Soci - Approvazione Bilancio"""
    
    def get_template_name(self) -> str:
        return "Approvazione Bilancio di Assemblea"
    
    def get_required_fields(self) -> list:
        return [
            "denominazione", "sede_legale", "capitale_sociale", "codice_fiscale",
            "data_assemblea", "ora_assemblea", "presidente", "segretario", 
            "soci", "amministratori"
        ]
    
    def get_form_fields(self, extracted_data: dict) -> dict:
        """Genera i campi del form per Streamlit utilizzando il CommonDataHandler"""
        form_data = {}
        
        # Utilizza il CommonDataHandler per i dati standard
        # Dati azienda standardizzati
        form_data.update(CommonDataHandler.extract_and_populate_company_data(extracted_data))
        
        # Dati assemblea standardizzati
        form_data.update(CommonDataHandler.extract_and_populate_assembly_data(extracted_data))
        
        # Campi specifici per approvazione bilancio
        st.subheader("📋 Configurazioni Specifiche Approvazione Bilancio")
        
        col1, col2 = st.columns(2)
        with col1:
            form_data["presidente"] = st.text_input("Presidente", form_data.get("presidente", ""))
            form_data["ruolo_presidente"] = st.selectbox("Ruolo del presidente",
                                                          CommonDataHandler.get_standard_ruoli_presidente())
            form_data["data_chiusura"] = st.date_input("Data chiusura bilancio",
                                                      CommonDataHandler.get_default_date_chiusura())
        with col2:
            form_data["segretario"] = st.text_input("Segretario", form_data.get("segretario", ""))
            # Campi specifici per questo template
            form_data["tipo_bilancio"] = st.selectbox("Tipo di bilancio",
                                                              ["Ordinario", "Abbreviato", "Micro"])
            form_data["presenza_nota_integrativa"] = st.checkbox("Nota integrativa presente", value=True)
        
        # Partecipanti standardizzati usando il CommonDataHandler
        participants_data = CommonDataHandler.extract_and_populate_participants_data(
            extracted_data,
            unique_key_suffix="bilancio"
        )
        # Aggiungi validazione per rappresentante_legale per società
        for socio in participants_data.get('soci', []):
            if socio.get('tipo_soggetto') == 'Società' and not socio.get('rappresentante_legale'):
                # Imposta un valore di default per rappresentante_legale se non presente
                socio['rappresentante_legale'] = "Rappresentante non specificato"
        form_data.update(participants_data)
        
        # Ordine del giorno specifico per approvazione bilancio
        st.subheader("📋 Ordine del Giorno")
        default_odg = f"1. Approvazione del Bilancio di esercizio chiuso al {form_data['data_chiusura'].strftime('%d/%m/%Y')} e dei relativi documenti che lo compongono\n2. Destinazione del risultato dell'esercizio\n3. Deliberazioni inerenti e conseguenti"
        punti_odg = st.text_area("Punti all'ordine del giorno (numerati automaticamente)", 
                                default_odg, height=120)
        form_data["punti_ordine_giorno"] = [p.strip() for p in punti_odg.split("\n") if p.strip()]
        
        # Destinazione risultato migliorata
        st.subheader("💰 Destinazione del Risultato")
        
        col1, col2 = st.columns(2)
        with col1:
            utile_esercizio = st.text_input("Utile dell'esercizio (€)", "0,00", 
                                          help="Inserisci l'utile dell'esercizio per calcolare automaticamente le destinazioni")
            form_data["utile_esercizio"] = utile_esercizio
            
            # Calcola automaticamente le percentuali se c'è un utile
            try:
                utile_value = float(utile_esercizio.replace(",", "."))
                riserva_legale_suggerita = utile_value * 0.05  # 5% a riserva legale
                utile_residuo = utile_value - riserva_legale_suggerita
            except:
                riserva_legale_suggerita = 0
                utile_residuo = 0
        with col2:
            form_data["riserva_legale"] = st.text_input("A riserva legale (€)", 
                                                       f"{riserva_legale_suggerita:.2f}".replace(".", ","))
            form_data["altre_riserve"] = st.text_input("Ad altre riserve (€)", "0,00")
        
        col1, col2 = st.columns(2)
        with col1:
            form_data["dividendi"] = st.text_input("A dividendi (€)", "0,00")
            form_data["riporto_nuovo"] = st.text_input("A nuovo (€)", 
                                                      f"{utile_residuo:.2f}".replace(".", ","))
        with col2:
            form_data["tipo_destinazione"] = st.selectbox("Tipo principale di destinazione", 
                                                         ["A nuovo", "A riserve", "A dividendi", "Mista"])
        
        # Campo per Altre Informazioni o Note
        st.subheader("📝 Altre Informazioni o Note")
        form_data["altre_informazioni_note"] = st.text_area(
            "Inserisci qui eventuali altre informazioni, note o clausole specifiche da includere nel verbale:",
            height=150,
            key="altre_info_note_bilancio"
        )
        
        return form_data
    
    def show_preview(self, form_data: dict):
        """Mostra l'anteprima del documento fuori dal form"""
        # Sezione Anteprima
        st.markdown("---")
        st.subheader("👁️ Anteprima Documento")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col1:
            show_preview = st.checkbox("Mostra anteprima", value=False, key="preview_checkbox")
        
        with col2:
            if show_preview:
                st.info("💡 L'anteprima si aggiorna automaticamente con i dati inseriti sopra")
        
        if show_preview:
            try:
                # Debug avanzato: mostra alcuni dati per verificare
                st.markdown("**🔍 Debug - Informazioni sui dati**")
                st.write(f"**Tipo dati ricevuti:** {type(form_data)}")
                st.write(f"**Numero di campi:** {len(form_data) if form_data else 0}")
                st.write("**Campi principali:**")
                st.write(f"- Denominazione: {form_data.get('denominazione', 'NON IMPOSTATA')}")
                st.write(f"- Presidente: {form_data.get('presidente', 'NON IMPOSTATO')}")
                st.write(f"- Data assemblea: {form_data.get('data_assemblea', 'NON IMPOSTATA')}")
                st.write(f"- Soci: {len(form_data.get('soci', []))} presenti")
                
                # Debug specifico per i soci
                soci_debug = form_data.get('soci', [])
                if soci_debug:
                    st.write("**Debug Soci Dettagliato:**")
                    for i, socio in enumerate(soci_debug):
                        st.write(f"  Socio {i+1}: {socio}")
                        # Debug specifico per i nuovi campi
                        tipo_soggetto = socio.get('tipo_soggetto', 'NON TROVATO')
                        rappresentante = socio.get('rappresentante_legale', 'NON TROVATO')
                        st.write(f"    - Tipo Soggetto: {tipo_soggetto}")
                        st.write(f"    - Rappresentante Legale: {rappresentante}")
                        
                        # ⚠️ DEBUG SPECIFICO PER IL PROBLEMA
                        if tipo_soggetto == 'Società':
                            if rappresentante and rappresentante != 'NON TROVATO' and rappresentante.strip():
                                st.success(f"✅ Società '{socio.get('nome', '')}' ha rappresentante legale: '{rappresentante}'")
                            else:
                                st.error(f"❌ Società '{socio.get('nome', '')}' NON ha rappresentante legale specificato!")
                                st.write(f"   Valore grezzo: '{repr(rappresentante)}'")
                        elif tipo_soggetto == 'Persona Fisica':
                            st.info(f"ℹ️ '{socio.get('nome', '')}' è una Persona Fisica (rappresentante legale non necessario)")
                else:
                    st.write("❌ Nessun dato sui soci trovato")
                    
                    # Mostra tutti i campi disponibili
                    st.write("**Tutti i campi disponibili:**")
                    if form_data:
                        for key, value in form_data.items():
                            if isinstance(value, list):
                                st.write(f"- {key}: lista con {len(value)} elementi")
                            else:
                                preview_value = str(value)[:50] + "..." if len(str(value)) > 50 else str(value)
                                st.write(f"- {key}: {preview_value}")
                    else:
                        st.error("❌ form_data è vuoto o None!")
                
                    # Debug UI per i soci nell'anteprima
                    if st.checkbox("🔍 Debug Soci nell'Anteprima", key="debug_soci_anteprima"):
                        soci_debug = form_data.get('soci', [])
                        st.write("**Dati soci che arrivano all'anteprima:**")
                        for i, socio in enumerate(soci_debug):
                            tipo_soggetto = socio.get('tipo_soggetto', 'NON TROVATO')
                            rappresentante = socio.get('rappresentante_legale', 'NON TROVATO')
                            tipo_part = socio.get('tipo_partecipazione', 'NON TROVATO')
                            delegato = socio.get('delegato', 'NON TROVATO')
                            
                            st.write(f"**Socio {i+1}: {socio.get('nome', 'NOME MANCANTE')}**")
                            st.write(f"  - Tipo Partecipazione: '{tipo_part}'")
                            st.write(f"  - Tipo Soggetto: '{tipo_soggetto}'")
                            st.write(f"  - Delegato: '{delegato}'")
                            st.write(f"  - Rappresentante Legale: '{rappresentante}'")
                            
                            # Test delle condizioni
                            if tipo_part == 'Delegato' and delegato:
                                st.success(f"✅ MATCH: È un delegato")
                                if tipo_soggetto == 'Società':
                                    st.success(f"✅ MATCH: È una società")
                                    if rappresentante:
                                        st.success(f"✅ MATCH: Ha rappresentante legale")
                                    else:
                                        st.error(f"❌ NO MATCH: Rappresentante legale vuoto")
                                else:
                                    st.info(f"ℹ️ INFO: Non è una società")
                            else:
                                st.info(f"ℹ️ INFO: Non è un delegato")
                                if tipo_soggetto == 'Società':
                                    st.success(f"✅ MATCH: È una società diretta")
                                    if rappresentante:
                                        st.success(f"✅ MATCH: Ha rappresentante legale")
                                    else:
                                        st.error(f"❌ NO MATCH: Rappresentante legale vuoto")
                    
# Genera anteprima con try/catch dettagliato
try:
    preview_text = self._generate_preview_text(form_data)

    if not preview_text:
        st.error("❌ Anteprima vuota - nessun testo generato")
    elif len(preview_text) < 100:
        st.warning("⚠️ Anteprima molto breve - possibile errore nei dati")
        st.code(preview_text)
    else:
        st.success("✅ Anteprima generata correttamente")

    # Aggiungi debug per soci nell'anteprima
    st.write("**Debug Soci nell'Anteprima:**")
    soci_debug = form_data.get('soci', [])
    for i, socio in enumerate(soci_debug):
        st.write(f"Socio {i+1}: {socio.get('nome', 'NOME MANCANTE')}")
        st.write(f"  Tipo Soggetto: {socio.get('tipo_soggetto', 'NON SPECIFICATO')}")
        st.write(f"  Rappresentante Legale: {socio.get('rappresentante_legale', 'NON SPECIFICATO')}")

except Exception as e:
    st.error(f"❌ Errore nella generazione dell'anteprima: {e}")
    preview_text = "Errore nella generazione dell'anteprima"
    # Aggiungi log dettagliato dell'errore
    st.error(f"Tipo di errore: {type(e).__name__}")
    st.error(f"Dettagli errore: {str(e)}")
    if hasattr(e, '__dict__'):
        st.write("Dettagli aggiuntivi:", e.__dict__)
                    
                    # Campo di testo modificabile per l'anteprima
                    edited_preview_text = st.text_area(
                        "Modifica l'anteprima qui (il documento finale userà questo testo):",
                        value=preview_text,
                        height=400,
                        key="editable_preview_text"
                    )
                    # Salva il testo modificato nello stato della sessione per usarlo nella generazione del documento
                    st.session_state['final_document_text'] = edited_preview_text
                    
                    # Statistiche anteprima
                    if 'preview_text' in locals():
                        word_count = len(preview_text.split())
                        char_count = len(preview_text)
                        st.caption(f"📊 Statistiche: {word_count} parole, {char_count} caratteri")
                        
                        # Pulsante per copiare
                        st.code(preview_text[:200] + "..." if len(preview_text) > 200 else preview_text)
                        st.error(f"❌ Errore nella generazione dell'anteprima: {str(e)}")
                        st.code(f"Errore dettagliato: {repr(e)}")
                        
                        # Debug traceback
                        import traceback
                        st.code(traceback.format_exc())
                        
                        # Mostra i dati per debug
                        st.write("🔍 **Dati completi per debug:**")
                        st.json(form_data)
                        
            except Exception as e:
                st.error(f"❌ Errore generale nell'anteprima: {str(e)}")
                import traceback
                st.code(traceback.format_exc())
    
    def _generate_preview_text(self, data: dict) -> str:
        """Genera un'anteprima testuale del verbale seguendo il formato esatto richiesto"""
        try:
            # Verifica che abbiamo i dati essenziali
            if not data:
                return "❌ Nessun dato disponibile per l'anteprima"
            
            if not isinstance(data, dict):
                return f"❌ Dati non validi per l'anteprima: tipo {type(data)} invece di dict"
            
            # Intestazione società
            denominazione = data.get('denominazione', '[DENOMINAZIONE]')
            sede_legale = data.get('sede_legale', '[SEDE]')
            capitale_deliberato = str(data.get('capitale_deliberato', '10.000,00')).strip()
            capitale_versato = str(data.get('capitale_versato', '10.000,00')).strip()
            capitale_sottoscritto = str(data.get('capitale_sottoscritto', '10.000,00')).strip()
            codice_fiscale = data.get('codice_fiscale', '[CODICE FISCALE]')
            
            # Gestione "i.v." (interamente versato) - solo se versato = deliberato per la preview
            if capitale_versato == capitale_deliberato:
                capitale_preview = f"Capitale sociale Euro {capitale_deliberato} i.v."
            else:
                capitale_preview = f"Capitale sociale deliberato Euro {capitale_deliberato}, sottoscritto Euro {capitale_sottoscritto}, versato Euro {capitale_versato}"
            
            preview = f"{denominazione}\n"
            preview += f"Sede in {sede_legale}\n"
            preview += f"{capitale_preview}\n"
            preview += f"Codice fiscale e Partita IVA: {codice_fiscale}\n\n"
            
            # Titolo
            preview += "Verbale di assemblea dei soci\n"
            
            # Data
            data_assemblea = data.get('data_assemblea')
            if data_assemblea:
                try:
                    if hasattr(data_assemblea, 'strftime'):
                        data_str = data_assemblea.strftime('%d/%m/%Y')
                    else:
                        data_str = str(data_assemblea)
                except:
                    data_str = '[DATA]'
            else:
                data_str = '[DATA]'
            
            preview += f"del {data_str}\n\n"
            
            # Apertura assemblea
            ora_inizio = data.get('ora_inizio', '[ORA]')
            luogo = data.get('luogo_assemblea', sede_legale)
            
            preview += f"Oggi {data_str} alle ore {ora_inizio} presso la sede sociale {luogo}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:\n"
            
            # Ordine del giorno
            preview += "Ordine del giorno\n"
            
            # Data chiusura bilancio
            data_chiusura = data.get('data_chiusura')
            if data_chiusura:
                try:
                    if hasattr(data_chiusura, 'strftime'):
                        data_chiusura_str = data_chiusura.strftime('%d/%m/%Y')
                    else:
                        data_chiusura_str = str(data_chiusura)
                except:
                    data_chiusura_str = '[DATA CHIUSURA]'
            else:
                data_chiusura_str = '[DATA CHIUSURA]'
            
            preview += f"Approvazione del Bilancio al {data_chiusura_str} e dei documenti correlati;\n"
            preview += "Delibere consequenziali.\n"
            
            # Presidenza
            presidente = data.get('presidente', '[PRESIDENTE]')
            ruolo_presidente = data.get('ruolo_presidente', 'Amministratore Unico')
            
            preview += f"Assume la presidenza ai sensi dell'art. […] dello statuto sociale il Sig. {presidente} {ruolo_presidente}, il quale dichiara e constata:\n"
            
            # Dichiarazioni preliminari
            preview += "1 - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. […] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza\n"
            preview += "2 - che sono presenti/partecipano all'assemblea:\n"
            
            # Amministratori
            amministratori = data.get('amministratori', [])
            ruolo_presidente = data.get('ruolo_presidente', 'Amministratore Unico')

            # Calcolo e visualizzazione quote totali soci
            soci = data.get('soci', [])
            total_quota_euro = 0.0
            total_quota_percentuale = 0.0

            for socio in soci:
                try:
                    quota_euro_str = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                    total_quota_euro += float(quota_euro_str)
                except (ValueError, TypeError):
                    pass # Ignora valori non numerici

                try:
                    quota_percentuale_str = str(socio.get('quota_percentuale', '0')).replace(',', '.')
                    total_quota_percentuale += float(quota_percentuale_str)
                except (ValueError, TypeError):
                    pass # Ignora valori non numerici

            # Formattazione per l'output italiano
            formatted_total_quota_euro = CommonDataHandler.format_currency(total_quota_euro)
            formatted_total_quota_percentuale = CommonDataHandler.format_percentage(total_quota_percentuale)

            preview += f"\n- Soci presenti: {len(soci)} per un totale di Euro {formatted_total_quota_euro} ({formatted_total_quota_percentuale}% del capitale sociale)\n"

            # Determina se è Amministratore Unico o CdA
            if ruolo_presidente == 'Amministratore Unico' or len(amministratori) <= 1:
                preview += f"l'Amministratore Unico nella persona del suddetto Presidente Sig. {presidente}\n"
            else:
                preview += "per il Consiglio di Amministrazione:\n"
                for amm in amministratori:
                    if isinstance(amm, dict) and amm.get('nome', '').strip():
                        nome = amm.get('nome', '')
                        carica = amm.get('carica', 'Consigliere')
                        if nome != presidente:  # Evita la duplicazione del presidente
                            preview += f"il Sig {nome} ({carica})\n"
            
            # Collegio sindacale (se presente)
            if data.get('collegio_sindacale'):
                preview += "per il Collegio Sindacale:\n"
                preview += "il Dott. [SINDACO 1]\n"
                preview += "il Dott. [SINDACO 2]\n"
                preview += "il Dott. [SINDACO 3]\n"
            
            # Revisore (se presente)
            if data.get('revisore'):
                nome_revisore = data.get('nome_revisore', '[NOME REVISORE]')
                preview += f"il revisore contabile Dott. {nome_revisore}\n"
            
            # Soci presenti - calcolo automatico dei totali
            soci = data.get('soci', [])
            totale_quote_euro = 0
            totale_quote_perc = 0
            soci_presenti = []
            
            if soci and isinstance(soci, list):
                for socio in soci:
                    if isinstance(socio, dict) and socio.get('nome', '').strip() and socio.get('presente', True):
                        nome = socio.get('nome', '').strip()
                        quota_euro_str = socio.get('quota_euro', '0').replace('€', '').replace('Euro', '').strip()
                        quota_perc_str = socio.get('quota_percentuale', '0').replace('%', '').replace(',', '.').strip()
                        tipo_partecipazione = socio.get('tipo_partecipazione', 'Diretto')
                        delegato = socio.get('delegato', '').strip()
                        tipo_soggetto = socio.get('tipo_soggetto', 'Persona Fisica')
                        rappresentante_legale = socio.get('rappresentante_legale', '').strip()
                        
                        # Debug: mostra i valori dei nuovi campi
                        print(f"DEBUG Socio {nome}: tipo_soggetto={tipo_soggetto}, rappresentante_legale={rappresentante_legale}")
                        
                        # Calcola i totali - gestione CORRETTA del formato italiano
                        try:
                            if quota_euro_str:
                                # Gestione del formato italiano: "10.000,00" dove punto=migliaia, virgola=decimali
                                if ',' in quota_euro_str:
                                    # Se c'è una virgola, è il separatore decimale
                                    quota_clean = quota_euro_str.replace('.', '').replace(',', '.')
                                else:
                                    # Se non c'è virgola, rimuovi solo i punti delle migliaia
                                    quota_clean = quota_euro_str.replace('.', '')
                                totale_quote_euro += float(quota_clean)
                        except:
                            pass
                        
                        try:
                            if quota_perc_str:
                                totale_quote_perc += float(quota_perc_str)
                        except:
                            pass
                        
                        soci_presenti.append({
                            'nome': nome,
                            'quota_euro': socio.get('quota_euro', '0'),
                            'quota_percentuale': socio.get('quota_percentuale', '0%'),
                            'tipo_partecipazione': tipo_partecipazione,
                            'delegato': delegato,
                            'tipo_soggetto': tipo_soggetto,
                            'rappresentante_legale': rappresentante_legale
                        })
            
            formatted_totale_quote_euro = CommonDataHandler.format_currency(totale_quote_euro)
            formatted_totale_quote_perc = CommonDataHandler.format_percentage(totale_quote_perc)
            preview += f"nonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_totale_quote_euro} pari al {formatted_totale_quote_perc} del Capitale Sociale:\n"
            
            for socio_info in soci_presenti:
                nome = socio_info['nome']
                quota_euro = socio_info['quota_euro']
                quota_perc = socio_info['quota_percentuale']
                tipo = socio_info['tipo_partecipazione']
                delegato = socio_info['delegato']
                tipo_soggetto = socio_info['tipo_soggetto']
                rappresentante_legale = socio_info['rappresentante_legale']
                
                # DEBUG: Aggiungiamo output dettagliato per ogni socio
                print(f"DEBUG ANTEPRIMA - Socio: {nome}")
                print(f"  - tipo_partecipazione: '{tipo}'")
                print(f"  - tipo_soggetto: '{tipo_soggetto}'")
                print(f"  - delegato: '{delegato}'")
                print(f"  - rappresentante_legale: '{rappresentante_legale}'")
                print(f"  - delegato truthy: {bool(delegato)}")
                print(f"  - rappresentante_legale truthy: {bool(rappresentante_legale)}")
                
                # Gestione completa: delegati e rappresentanti legali
                if tipo == 'Delegato' and delegato:
                    print(f"  -> BRANCH: Delegato")
                    if tipo_soggetto == 'Società':
                        print(f"  -> SUB-BRANCH: Società delegata")
                        preview += f"il Sig {delegato} delegato della società {nome}"
                        if rappresentante_legale:
                            print(f"  -> ADDING: rappresentante legale ({rappresentante_legale})")
                            preview += f" (nella persona del legale rappresentante {rappresentante_legale})"
                        else:
                            print(f"  -> NOT ADDING: rappresentante legale (vuoto o False)")
                        preview += f" recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale\n"
                    else:
                        print(f"  -> SUB-BRANCH: Persona fisica delegata")
                        preview += f"il Sig {delegato} delegato del socio {nome} recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale\n"
                else:
                    print(f"  -> BRANCH: Diretto")
                    if tipo_soggetto == 'Società':
                        print(f"  -> SUB-BRANCH: Società diretta")
                        if rappresentante_legale:
                            print(f"  -> ADDING: rappresentante legale ({rappresentante_legale})")
                            preview += f"la società {nome} nella persona del legale rappresentante {rappresentante_legale} recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale\n"
                        else:
                            print(f"  -> NOT ADDING: rappresentante legale (vuoto o False)")
                            preview += f"la società {nome} recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale\n"
                    else:
                        print(f"  -> SUB-BRANCH: Persona fisica diretta")
                        preview += f"il Sig {nome} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale\n"
            
            preview += "2 - che gli intervenuti sono legittimati alla presente assemblea;\n"
            preview += "3 - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno.\n"
            
            # Segretario
            segretario = data.get('segretario', '[SEGRETARIO]')
            preview += f"I presenti all'unanimità chiamano a fungere da segretario il signor {segretario}, che accetta l'incarico.\n"
            
            # Identificazione partecipanti
            preview += "Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante.\n"
            
            # Constatazione validità
            tipo_convocazione = data.get('tipo_convocazione', 'regolarmente convocata')
            if 'totalitaria' in tipo_convocazione.lower():
                preview += "Il Presidente constata e fa constatare che l'assemblea risulta totalitaria e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.\n"
            else:
                preview += "Il Presidente constata e fa constatare che l'assemblea risulta regolarmente convocata e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.\n"
            
            preview += "Si passa quindi allo svolgimento dell'ordine del giorno.\n"
            preview += "*     *     *\n"
            
            # Primo punto - Bilancio
            preview += f"In relazione al primo punto il presidente legge il bilancio al {data_chiusura_str} composto da stato patrimoniale, conto economico e nota integrativa (allegati di seguito al presente verbale);\n"
            
            # Collegio sindacale (se presente)
            if data.get('collegio_sindacale'):
                preview += f"Prende la parola il Presidente del Collegio Sindacale che legge la relazione del collegio sindacale al Bilancio chiuso al {data_chiusura_str} (allegata di seguito al presente verbale).\n"
            
            # Revisore (se presente)
            if data.get('revisore'):
                nome_revisore = data.get('nome_revisore', '[NOME REVISORE]')
                preview += f"Prende infine la parola il Dott. {nome_revisore} che legge la relazione del revisore contabile (allegata di seguito al presente verbale).\n"
            
            # Votazione
            esito = data.get('esito_votazione', 'approvato all\'unanimità')
            if 'unanimità' in esito:
                preview += "Segue breve discussione tra i soci al termine della quale si passa alla votazione con voto palese in forza della quale il Presidente constata che, all'unanimità, l'assemblea\n"
            else:
                preview += "Segue breve discussione tra i soci al termine della quale si passa alla votazione con voto palese in forza della quale il Presidente constata che, [ESITO VOTAZIONE], l'assemblea\n"
            
            preview += "delibera\n"
            preview += f"l'approvazione del bilancio di esercizio chiuso al {data_chiusura_str} e dei relativi documenti che lo compongono.\n"
            preview += "*     *     *\n"
            
            # Secondo punto - Destinazione risultato
            preview += "In relazione al secondo punto posto all'ordine del giorno, il Presidente, "
            if data.get('sentito_parere_sindaci'):
                preview += "sentito il parere favorevole del Collegio Sindacale, "
            
            tipo_risultato = data.get('tipo_risultato', 'Utile')
            if tipo_risultato == 'Perdita':
                preview += "propone all'assemblea di ripianare la perdita:\n"
            else:
                preview += "propone all'assemblea di così destinare il risultato d'esercizio:\n"
            
            destinazioni = data.get('destinazioni_risultato', [])
            if destinazioni and isinstance(destinazioni, list):
                for dest in destinazioni:
                    if dest and str(dest).strip():
                        preview += f"{str(dest).strip()}\n"
            else:
                if tipo_risultato == 'Perdita':
                    preview += "- al ripianamento della perdita per Euro [SPECIFICARE]\n"
                else:
                    preview += "- a riserva legale per il 5% dell'utile di esercizio\n- a riserva straordinaria\n- a dividendo\n"
            
            preview += "\n*     *     *\n"
            
            # Chiusura
            preview += "Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.\n"
            preview += "Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo.\n"
            
            ora_fine = data.get('ora_fine', '[ORA FINE]')
            preview += f"L'assemblea viene sciolta alle ore {ora_fine}.\n"

            # Aggiunta di Altre Informazioni o Note all'anteprima
            altre_info = data.get('altre_informazioni_note', '')
            if altre_info and altre_info.strip():
                preview += "\nALTRA INFORMAZIONI O NOTE:\n"
                preview += f"{altre_info.strip()}\n"
            
            return preview
            
        except Exception as e:
            import traceback
            error_msg = f"❌ ERRORE nella generazione dell'anteprima:\n{str(e)}\n\n"
            error_msg += f"Tipo errore: {type(e).__name__}\n"
            error_msg += f"Traceback completo:\n{traceback.format_exc()}\n\n"
            error_msg += f"Dati ricevuti: {type(data)} con {len(data) if data else 0} elementi\n"
            if data:
                error_msg += f"Chiavi disponibili: {list(data.keys())}\n"
                error_msg += f"Primi 3 elementi: {dict(list(data.items())[:3])}\n"
            return error_msg
    
    def generate_document(self, data: dict) -> Document:
        """Genera il documento Word con formattazione professionale"""
        import streamlit as st
        
        # Controlla se c'è un testo modificato dall'utente nell'anteprima
        if hasattr(st.session_state, 'final_document_text') and st.session_state.final_document_text:
            # Usa il testo modificato dall'utente
            return self._create_document_from_text(st.session_state.final_document_text)
        else:
            # Genera il documento normalmente se non c'è testo modificato
            # Utilizza il template .docx esistente per mantenere la formattazione
            import os
            template_path = os.path.join(os.path.dirname(__file__), 'template.docx')
            
            try:
                if os.path.exists(template_path):
                    doc = Document(template_path)
                else:
                    doc = Document()
                    # Configurazione stili del documento solo se non usiamo template
                    self._setup_document_styles(doc)
            except Exception as e:
                # Fallback a documento vuoto se il template non può essere caricato
                doc = Document()
                self._setup_document_styles(doc)
            
            # Intestazione società
            self._add_company_header(doc, data)
            
            # Titolo verbale
            self._add_verbale_title(doc, data)
            
            # Apertura assemblea
            self._add_opening_section(doc, data)
            
            # Partecipanti
            self._add_participants_section(doc, data)
            
            # Constatazioni preliminari
            self._add_preliminary_statements(doc, data)
            
            # Svolgimento ordine del giorno (specifico per approvazione bilancio)
            self._add_bilancio_discussion(doc, data)
            
            # Chiusura
            self._add_closing_section(doc, data)
            
            # Firme
            self._add_signatures(doc, data)
            
            return doc
    
    def _create_document_from_text(self, text: str) -> Document:
        """Crea un documento Word dal testo modificato dall'utente con formattazione automatica"""
        # Utilizza il template .docx esistente per mantenere la formattazione
        import os
        template_path = os.path.join(os.path.dirname(__file__), 'template.docx')
        
        try:
            if os.path.exists(template_path):
                doc = Document(template_path)
                # Rimuovi il contenuto esistente del template mantenendo gli stili
                for paragraph in doc.paragraphs[:]:
                    p = paragraph._element
                    p.getparent().remove(p)
            else:
                doc = Document()
                # Configurazione stili di base solo se non usiamo template
                self._setup_document_styles(doc)
        except Exception as e:
            # Fallback a documento vuoto se il template non può essere caricato
            doc = Document()
            self._setup_document_styles(doc)
        
        # Analizza la struttura del testo e applica la formattazione automatica
        sections = self._analyze_text_structure(text)
        
        for section in sections:
            self._add_formatted_section(doc, section)
        
        return doc
    
    def _setup_document_styles(self, doc):
        """Configura gli stili del documento in modo più robusto e completo."""
        styles = doc.styles

    def _analyze_text_structure(self, text: str) -> list:
        """
        Analizza il testo e restituisce una lista di sezioni con i loro stili
        """
        sections = []
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            # Riga vuota
            if not line_stripped:
                sections.append({
                    'type': 'empty', 
                    'content': '', 
                    'style': None,
                    'original_line': line
                })
                continue
            
            # 1. Intestazione aziendale (prime 5 righe, spesso maiuscolo, contiene nome società)
            if i < 5 and (line_stripped.isupper() or 
                         any(keyword in line_stripped.upper() for keyword in 
                             ['S.R.L.', 'S.P.A.', 'S.R.L', 'SPA', 'SRL', 'SEDE', 'CAPITALE', 'CODICE'])):
                sections.append({
                    'type': 'company_header', 
                    'content': line_stripped, 
                    'style': 'CompanyHeader',
                    'original_line': line
                })
                
            # 2. Titolo principale verbale
            elif ('VERBALE' in line_stripped.upper() and 'ASSEMBLEA' in line_stripped.upper()):
                sections.append({
                    'type': 'main_title', 
                    'content': line_stripped, 
                    'style': 'VerbaleTitle',
                    'original_line': line
                })
                
            # 3. Sottotitoli (parentesi o date)
            elif ((line_stripped.startswith('(') and line_stripped.endswith(')')) or 
                  'del ' in line_stripped or 'DEL ' in line_stripped):
                sections.append({
                    'type': 'subtitle', 
                    'content': line_stripped, 
                    'style': 'VerbaleSubtitle',
                    'original_line': line
                })
                
            # 4. Intestazioni di sezione (maiuscolo, parole chiave specifiche)
            elif (line_stripped.isupper() and len(line_stripped) > 5 and
                  any(keyword in line_stripped for keyword in 
                      ['ORDINE', 'SOCI', 'AMMINISTRATORI', 'PUNTO', 'DELIBERA', 'PRESENTE', 'RAPPRESENTAT'])):
                sections.append({
                    'type': 'section_header', 
                    'content': line_stripped, 
                    'style': 'SectionHeader',
                    'original_line': line
                })
                
            # 5. Separatori (linee di asterischi o trattini)
            elif line_stripped in ['*     *     *', '* * *', '---', '___'] or line_stripped.startswith('_' * 10):
                sections.append({
                    'type': 'separator', 
                    'content': line_stripped, 
                    'style': 'BodyText',
                    'original_line': line,
                    'center': True
                })
                
            # 6. Elenchi puntati
            elif line_stripped.startswith(('•', '-', '*', '- ', '• ')):
                sections.append({
                    'type': 'bullet_list', 
                    'content': line_stripped, 
                    'style': 'List Bullet',
                    'original_line': line
                })
                
            # 7. Elenchi numerati
            elif re.match(r'^\d+[\.\)]\s', line_stripped):
                sections.append({
                    'type': 'numbered_list', 
                    'content': line_stripped, 
                    'style': 'List Number',
                    'original_line': line
                })
                
            # 8. Totale quote (formattazione speciale)
            elif 'Totale quote rappresentate' in line_stripped or 'totale quote' in line_stripped.lower():
                sections.append({
                    'type': 'total_summary', 
                    'content': line_stripped, 
                    'style': 'BodyText',
                    'original_line': line,
                    'bold': True
                })
                
            # 9. Testo normale
            else:
                sections.append({
                    'type': 'body_text', 
                    'content': line_stripped, 
                    'style': 'BodyText',
                    'original_line': line
                })
        
        return sections
    
    def _add_formatted_section(self, doc, section):
        """
        Aggiunge una sezione al documento con la formattazione appropriata
        """
        if section['type'] == 'empty':
            doc.add_paragraph()
            return
            
        content = section['content']
        style = section['style']
        
        try:
            if section['type'] == 'company_header':
                # Intestazione aziendale - centrata
                p = doc.add_paragraph(content, style=style)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            elif section['type'] == 'main_title':
                # Titolo principale - centrato, maiuscolo, grassetto
                p = doc.add_paragraph(content.upper(), style=style)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            elif section['type'] == 'subtitle':
                # Sottotitolo - centrato, grassetto
                p = doc.add_paragraph(content, style=style)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            elif section['type'] == 'section_header':
                # Intestazione sezione - centrata, maiuscolo, grassetto
                p = doc.add_paragraph(content.upper(), style=style)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            elif section['type'] == 'separator':
                # Separatore - centrato
                p = doc.add_paragraph(content, style=style)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            elif section['type'] == 'bullet_list':
                # Elenco puntato
                # Rimuovi il simbolo del punto se presente, sarà aggiunto automaticamente
                clean_content = content
                if content.startswith(('•', '-', '*')):
                    clean_content = content[1:].strip()
                elif content.startswith(('- ', '• ')):
                    clean_content = content[2:].strip()
                    
                p = doc.add_paragraph(clean_content, style='List Bullet')
                
            elif section['type'] == 'numbered_list':
                # Elenco numerato
                # Rimuovi la numerazione se presente, sarà aggiunta automaticamente
                clean_content = re.sub(r'^\d+[\.\)]\s*', '', content)
                p = doc.add_paragraph(clean_content, style='List Number')
                
            elif section['type'] == 'total_summary':
                # Totale - grassetto
                p = doc.add_paragraph(content, style=style)
                p.runs[0].font.bold = True
                
            else:
                # Testo normale - giustificato
                p = doc.add_paragraph(content, style=style)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
        except Exception as e:
            # Fallback a testo normale se c'è un errore
            p = doc.add_paragraph(content, style='BodyText')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Stile 'Normal' (base per il documento)
        try:
            normal_style = styles['Normal']
            normal_style.font.name = 'Times New Roman'
            normal_style.font.size = Pt(11)
            normal_style.paragraph_format.space_after = Pt(6)
            normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        except KeyError:
            # Se 'Normal' non esiste, crealo (improbabile per docx standard)
            normal_style = styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
            normal_style.font.name = 'Times New Roman'
            normal_style.font.size = Pt(11)
            normal_style.paragraph_format.space_after = Pt(6)
            normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Stile per intestazione società
        try:
            company_style = styles.add_style('CompanyHeader', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError: # Lo stile esiste già
            company_style = styles['CompanyHeader']
        company_style.base_style = styles['Normal']
        company_style.font.name = 'Times New Roman'
        company_style.font.size = Pt(12)
        company_style.font.bold = True
        company_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_style.paragraph_format.space_after = Pt(12)
        company_style.paragraph_format.space_before = Pt(6)

        # Stile per titolo verbale
        try:
            title_style = styles.add_style('VerbaleTitle', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            title_style = styles['VerbaleTitle']
        title_style.base_style = styles['Normal']
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.font.all_caps = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_before = Pt(24)
        title_style.paragraph_format.space_after = Pt(18)

        # Stile per sottotitoli o date sotto il titolo principale
        try:
            subtitle_style = styles.add_style('VerbaleSubtitle', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            subtitle_style = styles['VerbaleSubtitle']
        subtitle_style.base_style = styles['Normal']
        subtitle_style.font.name = 'Times New Roman'
        subtitle_style.font.size = Pt(12)
        subtitle_style.font.bold = True
        subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_style.paragraph_format.space_after = Pt(12)
        subtitle_style.paragraph_format.space_before = Pt(6)

        # Stile per intestazioni di sezione (es. ORDINE DEL GIORNO)
        try:
            section_header_style = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            section_header_style = styles['SectionHeader']
        section_header_style.base_style = styles['Normal']
        section_header_style.font.name = 'Times New Roman'
        section_header_style.font.size = Pt(14)  # Aumentata dimensione font
        section_header_style.font.bold = True
        section_header_style.font.all_caps = True
        section_header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrato
        section_header_style.paragraph_format.space_before = Pt(18)
        section_header_style.paragraph_format.space_after = Pt(8)

        # Stile per testo principale del paragrafo
        try:
            body_text_style = styles.add_style('BodyText', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            body_text_style = styles['BodyText']
        body_text_style.base_style = styles['Normal']
        body_text_style.font.name = 'Times New Roman'
        body_text_style.font.size = Pt(12)  # Aumentata dimensione font per leggibilità
        body_text_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_text_style.paragraph_format.first_line_indent = Inches(0.0)
        body_text_style.paragraph_format.line_spacing = 1.15
        body_text_style.paragraph_format.space_after = Pt(6)
        body_text_style.paragraph_format.space_before = Pt(0)

        # Stili per elenchi puntati e numerati (se non si usano quelli built-in)
        try:
            list_bullet_style = styles.add_style('CustomListBullet', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            list_bullet_style = styles['CustomListBullet']
        list_bullet_style.base_style = styles['BodyText'] # Basato su BodyText per coerenza
        list_bullet_style.paragraph_format.first_line_indent = Inches(0) # Rimuovi rientro se ListBullet lo gestisce
        # Nota: la formattazione specifica del punto elenco (es. •) è meglio gestirla con add_paragraph(style='List Bullet')

        try:
            list_number_style = styles.add_style('CustomListNumber', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            list_number_style = styles['CustomListNumber']
        list_number_style.base_style = styles['BodyText']
        list_number_style.paragraph_format.first_line_indent = Inches(0)
        # Nota: la numerazione è meglio gestirla con add_paragraph(style='List Number')

        # Stile per le tabelle (se si vuole uno stile personalizzato di base)
        try:
            table_style = styles.add_style('CustomTableGrid', WD_STYLE_TYPE.TABLE)
            # table_style.base_style = styles['TableGrid'] # Non si può basare uno stile tabella su un altro così facilmente
            # Configura bordi, font, allineamento per le celle della tabella qui se necessario
            # Esempio: table_style.font.name = 'Times New Roman'
            # table_style.font.size = Pt(10)
        except ValueError:
            pass # Lo stile tabella esiste già o non si vuole personalizzare oltre 'Table Grid'

    
    def _add_company_header(self, doc, data):
        """Aggiunge l'intestazione della società utilizzando stili definiti."""
        # Nome società - stile principale
        company_name = doc.add_paragraph(data.get('denominazione', 'N/A').upper(), style='CompanyHeader')
        
        # Sede legale - stile secondario
        sede_text = f"Sede in {data.get('sede_legale', 'N/A')}"
        p_sede = doc.add_paragraph(sede_text, style='CompanyHeader')
        p_sede.runs[0].font.size = Pt(10)
        p_sede.runs[0].font.bold = False
        p_sede.paragraph_format.space_before = Pt(0)
        p_sede.paragraph_format.space_after = Pt(0)

        # Capitale sociale
        capitale_deliberato = str(data.get('capitale_deliberato', 'N/A')).strip()
        capitale_versato = str(data.get('capitale_versato', 'N/A')).strip()
        capitale_sottoscritto = str(data.get('capitale_sottoscritto', 'N/A')).strip()
        
        if capitale_versato == capitale_deliberato and capitale_deliberato != 'N/A':
            capitale_text = f"Capitale sociale Euro {capitale_deliberato} i.v."
        elif capitale_deliberato != 'N/A':
            capitale_text = f"Capitale sociale Euro {capitale_deliberato} i.v."
        else:
            capitale_text = "Capitale sociale: N/A"
        
        p_capitale = doc.add_paragraph(capitale_text, style='CompanyHeader')
        p_capitale.runs[0].font.size = Pt(10)
        p_capitale.runs[0].font.bold = False
        p_capitale.paragraph_format.space_before = Pt(0)
        p_capitale.paragraph_format.space_after = Pt(0)

        # Codice fiscale
        cf_piva_text = f"Codice fiscale: {data.get('codice_fiscale', 'N/A')}"
        p_cf = doc.add_paragraph(cf_piva_text, style='CompanyHeader')
        p_cf.runs[0].font.size = Pt(10)
        p_cf.runs[0].font.bold = False
        p_cf.paragraph_format.space_before = Pt(0)
        p_cf.paragraph_format.space_after = Pt(12)
        
        # Riga separatrice
        hr_paragraph = doc.add_paragraph()
        hr_run = hr_paragraph.add_run('_' * 80)
        hr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr_paragraph.paragraph_format.space_before = Pt(6)
        hr_paragraph.paragraph_format.space_after = Pt(18)
    
    def _add_verbale_title(self, doc, data):
        """Aggiunge il titolo del verbale utilizzando stili definiti."""
        doc.add_paragraph("VERBALE DI ASSEMBLEA DEI SOCI", style='VerbaleTitle')
        
        tipo_assemblea = data.get('tipo_assemblea', 'Ordinaria').upper()
        doc.add_paragraph(f"({tipo_assemblea})", style='VerbaleSubtitle')
        
        data_assemblea_str = data.get('data_assemblea', date.today()).strftime('%d/%m/%Y')
        # Applica lo stile 'VerbaleSubtitle' e poi sovrascrivi se necessario per la data
        p_data = doc.add_paragraph(style='VerbaleSubtitle') 
        p_data.text = f"del {data_assemblea_str}"
        # Lo stile VerbaleSubtitle ha già italic = True. Se vuoi bold, aggiungilo.
        p_data.runs[0].font.bold = True # Data in grassetto come richiesto
        p_data.paragraph_format.space_before = Pt(0)
    
    def _add_opening_section(self, doc, data):
        """Aggiunge la sezione di apertura utilizzando lo stile BodyText."""
        # doc.add_paragraph() # Spazio gestito da space_before/after degli stili
        
        data_str = data.get('data_assemblea', date.today()).strftime('%d/%m/%Y')
        ora_str = data.get('ora_inizio', 'HH:MM') # Placeholder se non fornita
        luogo = data.get('luogo_assemblea', data.get('sede_legale', 'N/A'))
        
        text = f"In data {data_str}, alle ore {ora_str}, presso {luogo}, "
        if data.get('audioconferenza', False):
            text += "e con la possibilità di partecipazione mediante mezzi di telecomunicazione che ne assicurino l'identificazione, "
        text += f"si è riunita l'Assemblea {data.get('tipo_assemblea', 'Ordinaria')} dei Soci della Società {data.get('denominazione', 'N/A')}, "
        text += f"{data.get('tipo_convocazione', 'regolarmente convocata')} per discutere e deliberare sul seguente:"
        
        doc.add_paragraph(text, style='BodyText')
    
    def _add_participants_section(self, doc, data):
        """Aggiunge la sezione partecipanti utilizzando stili definiti."""
        # Ordine del giorno
        doc.add_paragraph("ORDINE DEL GIORNO", style='SectionHeader')
        
        punti_odg = data.get('punti_ordine_giorno', [])
        if not punti_odg:
            doc.add_paragraph("Nessun punto all'ordine del giorno specificato.", style='BodyText')
        else:
            for i, punto in enumerate(punti_odg, 1):
                punto_clean = str(punto).strip()
                # Assicurati che la numerazione sia corretta o aggiungila se mancante
                if not re.match(r"^\d+\.\s*", punto_clean): 
                    punto_clean = f"{i}. {punto_clean}"
                p_punto = doc.add_paragraph(style='BodyText')
                p_punto.text = punto_clean
                p_punto.paragraph_format.left_indent = Inches(0.5)
                p_punto.paragraph_format.space_after = Pt(3)
        
        # Presidenza
        presidente = data.get('presidente', 'N/A')
        ruolo_presidente = data.get('ruolo_presidente', 'Amministratore Unico')
        articolo_statuto_presidenza = data.get('articolo_statuto_presidenza', '[...]')
        text_presidenza = f"Assume la presidenza ai sensi dell'art. {articolo_statuto_presidenza} dello statuto sociale il Sig. {presidente} {ruolo_presidente}, il quale dichiara e constata:"
        doc.add_paragraph(text_presidenza, style='BodyText')
        
        # Dichiarazioni del presidente
        dichiarazioni = [
            "- che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. [...] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza;",
            "- che sono presenti/partecipano all'assemblea:"
        ]
        
        for i, dich in enumerate(dichiarazioni, 1):
            p_dich = doc.add_paragraph(style='BodyText')
            p_dich.text = f"{i} {dich}"
            p_dich.paragraph_format.left_indent = Inches(0.25)
            p_dich.paragraph_format.space_after = Pt(3)
        
        # Nomina segretario
        segretario = data.get('segretario', 'N/A')
        if segretario and segretario != 'N/A':
            text_segretario = f"I presenti all'unanimità chiamano a fungere da segretario il signor {segretario}, che accetta l'incarico."
            doc.add_paragraph(text_segretario, style='BodyText')
        else:
            doc.add_paragraph("Il Presidente svolge anche le funzioni di Segretario.", style='BodyText')
    
    def _add_preliminary_statements(self, doc, data):
        """Aggiunge le constatazioni preliminari utilizzando stili definiti."""
        doc.add_paragraph("SOCI PRESENTI E RAPPRESENTATI", style='SectionHeader')
        
        soci_presenti = [s for s in data.get('soci', []) if s.get('nome', '').strip() and s.get('presente', True)]
        if not soci_presenti:
            doc.add_paragraph("Nessun socio risulta presente o rappresentato.", style='BodyText')
        else:
            totale_capitale_rappresentato_perc = 0.0
            # Esempio di calcolo più robusto per il capitale, assumendo che 'quota_euro' sia un numero
            # totale_capitale_rappresentato_euro = sum(float(str(s.get('quota_euro', '0')).replace(',', '.')) for s in soci_presenti)
            # capitale_sociale_totale_euro = float(str(data.get('capitale_deliberato', '1')).replace(',', '.')) # Evita divisione per zero
            # if capitale_sociale_totale_euro > 0:
            #     totale_capitale_rappresentato_perc = (totale_capitale_rappresentato_euro / capitale_sociale_totale_euro) * 100

            for socio in soci_presenti:
                nome = socio.get('nome', 'N/A')
                quota_perc_str = str(socio.get('quota_percentuale', '0')).replace('%', '').replace(',', '.')
                try:
                    quota_perc_val = float(quota_perc_str)
                    totale_capitale_rappresentato_perc += quota_perc_val
                except ValueError:
                    quota_perc_val = 0.0 # o gestisci l'errore
                
                quota_euro = socio.get('quota_euro', 'N/A')
                tipo_partecipazione = socio.get('tipo_partecipazione', 'Diretto')
                delegato = socio.get('delegato', '').strip()
                tipo_soggetto = socio.get('tipo_soggetto', 'Persona Fisica')
                rappresentante_legale = socio.get('rappresentante_legale', '').strip()
                
                desc_socio = f"{nome}"
                if tipo_soggetto == 'Società':
                    desc_socio += " (società)"
                    if rappresentante_legale:
                        desc_socio += f", legalmente rappresentata da {rappresentante_legale}"
                
                if tipo_partecipazione == 'Delegato' and delegato:
                    desc_socio = f"{delegato} (in qualità di delegato di {desc_socio})"
                
                p_socio = doc.add_paragraph(style='BodyText')
                p_socio.text = f"il Sig. {desc_socio} socio recante una quota pari a nominali euro [{quota_euro}] pari al {quota_perc_val:.3f}% del Capitale Sociale"
                p_socio.paragraph_format.left_indent = Inches(0.25)
                p_socio.paragraph_format.space_after = Pt(3)

            p_totale = doc.add_paragraph(style='BodyText')
            p_totale.text = f"che gli interventi sono legittimati alla presente assemblea;"
            p_totale.paragraph_format.space_before = Pt(6)
            p_totale.paragraph_format.space_after = Pt(6)
            
            p_totale2 = doc.add_paragraph(style='BodyText')
            p_totale2.text = f"che tutti gli intervenuti dichiarano di aver ricevuto l'avviso di convocazione nei termini di legge e di essere edotti sugli argomenti posti all'ordine del giorno."
            p_totale2.paragraph_format.space_after = Pt(6)
            # Aggiungere qui la verifica della validità dell'assemblea (maggioranze costitutive)
            # Esempio: if totale_capitale_rappresentato_perc >= data.get('maggioranza_costitutiva_perc', 50.0):
            # doc.add_paragraph("L'Assemblea è validamente costituita per deliberare sugli argomenti posti all'ordine del giorno.", style='BodyText')
            # else: 
            # doc.add_paragraph("ATTENZIONE: L'Assemblea potrebbe non essere validamente costituita.", style='BodyText') # Aggiungere logica per quorum

        # Presidente identifica tutti i presenti
        p_presidente = doc.add_paragraph(style='BodyText')
        p_presidente.text = f"Il Presidente identifica tutti i soggetti collegati e accerta che la loro partecipazione sia conforme alle previsioni statutarie e di legge, quindi dichiara l'assemblea validamente costituita e atta a deliberare sugli argomenti posti all'ordine del giorno."
        p_presidente.paragraph_format.space_before = Pt(6)
        p_presidente.paragraph_format.space_after = Pt(6)
        
        # Altri presenti (Sindaci, Revisore, etc.)
        altri_presenti_list = []
        if data.get('collegio_sindacale_presente', False):
            altri_presenti_list.append("i membri del Collegio Sindacale") # Potrebbe essere più dettagliato con i nomi
        if data.get('revisore_legale_presente', False):
            altri_presenti_list.append("il Revisore Legale dei Conti") # Nome del revisore
        # Aggiungere altri eventuali presenti da 'data'
        
        if altri_presenti_list:
            doc.add_paragraph("ALTRI PRESENTI", style='SectionHeader')
            for altro in altri_presenti_list:
                doc.add_paragraph(f"• {altro}", style='ListBullet')
    
    def _add_bilancio_discussion(self, doc, data):
        """Aggiunge la discussione sull'approvazione del bilancio e destinazione risultato."""
        doc.add_paragraph("DISCUSSIONE E DELIBERAZIONI SUI PUNTI ALL'ORDINE DEL GIORNO", style='SectionHeader')
        doc.add_paragraph("Il Presidente dichiara aperta la discussione sui singoli punti all'ordine del giorno.", style='BodyText')
        
        # Esempio per il primo punto (Bilancio)
        # Questo dovrebbe essere generalizzato o gestito da template specifici se l'OdG è variabile
        doc.add_paragraph("1. Approvazione del bilancio d'esercizio al [Data Chiusura Bilancio]", style='SectionHeader') # Sottotitolo per il punto
        # doc.add_paragraph("PRIMO PUNTO ALL'ORDINE DEL GIORNO", style='SectionHeader') # Alternativa

        data_chiusura_bilancio = data.get('data_chiusura_bilancio', data.get('data_assemblea', date.today())).strftime('%d/%m/%Y')
        
        text_illustrazione = f"Il Presidente illustra ai presenti il progetto di Bilancio d'esercizio chiuso al {data_chiusura_bilancio}, "
        text_illustrazione += "composto da Stato Patrimoniale, Conto Economico e Nota Integrativa. "
        if data.get('relazione_gestione_presente', False):
            text_illustrazione += "Viene altresì illustrata la Relazione sulla Gestione. "
        if data.get('documenti_allegati_bilancio', True): # Assumiamo che siano allegati o disponibili
            text_illustrazione += "Tali documenti, già messi a disposizione dei soci nei termini di legge, vengono allegati al presente verbale sotto la lettera [Lettera Allegato Bilancio]. "
        else:
            text_illustrazione += "Tali documenti sono stati messi a disposizione dei soci nei termini di legge. "
        doc.add_paragraph(text_illustrazione, style='BodyText')

        if data.get('parere_collegio_sindacale_presente', False):
            text_parere_cs = "Viene data lettura della Relazione del Collegio Sindacale, anch'essa allegata al presente verbale sotto la lettera [Lettera Allegato Relazione CS]."
            if data.get('parere_collegio_sindacale_favorevole', True):
                text_parere_cs += " Il Collegio Sindacale ha espresso parere favorevole all'approvazione del bilancio."
            else:
                text_parere_cs += " Il Collegio Sindacale ha espresso osservazioni in merito all'approvazione del bilancio, come da relazione."
            doc.add_paragraph(text_parere_cs, style='BodyText')
        
        if data.get('parere_revisore_presente', False):
            text_parere_rev = "Viene data lettura della Relazione del Revisore Legale dei Conti, allegata sotto la lettera [Lettera Allegato Relazione Revisore]."
            # Aggiungere dettagli sul parere del revisore
            doc.add_paragraph(text_parere_rev, style='BodyText')

        text_discussione = "Dopo ampia ed esauriente discussione, durante la quale vengono forniti tutti i chiarimenti richiesti dai soci, "
        modalita_voto = "per alzata di mano" if data.get('voto_palese', True) else "a scrutinio segreto"
        text_discussione += f"il Presidente pone in votazione la proposta di approvazione del bilancio. Si procede alla votazione {modalita_voto}."
        doc.add_paragraph(text_discussione, style='BodyText')
        
        # Esito votazione bilancio
        esito_votazione_bilancio = data.get('esito_votazione_bilancio', 'approvato all\'unanimità')
        # Qui si potrebbe aggiungere logica per dettagliare voti favorevoli, contrari, astenuti
        text_delibera_bilancio = f"L'Assemblea dei Soci, con voti {esito_votazione_bilancio},"
        doc.add_paragraph(text_delibera_bilancio, style='BodyText')
        doc.add_paragraph("DELIBERA", style='SectionHeader') # Usare uno stile specifico per la parola DELIBERA
        
        doc.add_paragraph(f"1. di approvare il Bilancio d'esercizio della Società chiuso al {data_chiusura_bilancio}, comprensivo di Stato Patrimoniale, Conto Economico e Nota Integrativa, così come presentato dall'organo amministrativo e corredato dalle relazioni dell'organo di controllo e del revisore legale, ove presenti.", style='ListNumber')
        
        # Secondo punto - Destinazione risultato
        doc.add_paragraph("2. Destinazione del risultato d'esercizio", style='SectionHeader')
        
        risultato_esercizio_val = data.get('risultato_esercizio_valore', '0.00')
        risultato_esercizio_tipo = "utile" if float(str(risultato_esercizio_val).replace(',','.')) >= 0 else "perdita"
        
        text_proposta_risultato = f"Il Presidente illustra quindi la proposta dell'organo amministrativo circa la destinazione dell'{risultato_esercizio_tipo} d'esercizio, pari a Euro {risultato_esercizio_val}."
        doc.add_paragraph(text_proposta_risultato, style='BodyText')
        
        proposte_destinazione = data.get('destinazioni_risultato', [])
        if proposte_destinazione:
            doc.add_paragraph("La proposta prevede di destinare tale risultato come segue:", style='BodyText')
            for dest in proposte_destinazione:
                if str(dest).strip():
                    doc.add_paragraph(f"- {str(dest).strip()};", style='ListBullet')
        else:
            doc.add_paragraph("Non vi sono proposte specifiche per la destinazione del risultato.", style='BodyText')

        text_discussione_risultato = f"Dopo breve discussione, il Presidente pone in votazione la proposta di destinazione del risultato d'esercizio. Si procede alla votazione {modalita_voto}."
        doc.add_paragraph(text_discussione_risultato, style='BodyText')
        
        esito_votazione_risultato = data.get('esito_votazione_risultato', 'approvato all\'unanimità')
        text_delibera_risultato = f"L'Assemblea dei Soci, con voti {esito_votazione_risultato},"
        doc.add_paragraph(text_delibera_risultato, style='BodyText')
        doc.add_paragraph("DELIBERA", style='SectionHeader')
        
        if proposte_destinazione:
            for i, dest in enumerate(proposte_destinazione, 1):
                 doc.add_paragraph(f"{i}. di destinare l'{risultato_esercizio_tipo} d'esercizio come segue: {str(dest).strip()};", style='ListNumber')
        else: # Esempio se non ci sono proposte specifiche, magari si riporta a nuovo o si copre la perdita
            if risultato_esercizio_tipo == "utile":
                doc.add_paragraph(f"1. di riportare a nuovo l'{risultato_esercizio_tipo} d'esercizio pari a Euro {risultato_esercizio_val}.", style='ListNumber')
            else:
                doc.add_paragraph(f"1. di coprire la {risultato_esercizio_tipo} d'esercizio pari a Euro {risultato_esercizio_val} mediante [specificare come].", style='ListNumber')
        
        # Aggiungere qui la gestione per ALTRI PUNTI ALL'ORDINE DEL GIORNO
        # Si potrebbe iterare su data.get('altri_punti_odg_discussione_delibere', [])
        # Ogni elemento potrebbe essere un dizionario con 'titolo_punto', 'testo_discussione', 'testo_delibera'
    
    def _add_closing_section(self, doc, data):
        """Aggiunge la sezione di chiusura utilizzando stili definiti."""
        doc.add_paragraph("CHIUSURA DELL'ASSEMBLEA", style='SectionHeader')
        
        doc.add_paragraph("Null'altro essendovi da deliberare e nessuno chiedendo ulteriormente la parola, il Presidente dichiara l'ordine del giorno esaurito.", style='BodyText')
        
        doc.add_paragraph("Il presente verbale, composto da [Numero Pagine Verbale] pagine, previa lettura e conferma, viene approvato all'unanimità (o specificare maggioranza/eventuali contrari/astenuti) e sottoscritto dal Presidente e dal Segretario.", style='BodyText')
        
        ora_fine = data.get('ora_fine', 'HH:MM') # Placeholder se non fornita
        doc.add_paragraph(f"L'Assemblea viene quindi dichiarata chiusa alle ore {ora_fine} del giorno {data.get('data_assemblea', date.today()).strftime('%d/%m/%Y')}.", style='BodyText')

        # Aggiunta di Altre Informazioni o Note al documento
        altre_info_note = data.get('altre_informazioni_note', '')
        if altre_info_note and altre_info_note.strip():
            doc.add_paragraph("ALTRE INFORMAZIONI O NOTE", style='SectionHeader')
            doc.add_paragraph(altre_info_note.strip(), style='BodyText')
        
        # Allegati
        allegati = data.get('allegati', [])
        if allegati:
            doc.add_paragraph("ALLEGATI", style='SectionHeader')
            for i, allegato_desc in enumerate(allegati, 1):
                if str(allegato_desc).strip():
                    # Esempio: "Allegato A) Progetto di Bilancio al 31/12/XXXX"
                    # Si potrebbe avere una struttura più dettagliata per gli allegati in 'data'
                    doc.add_paragraph(f"{str(allegato_desc).strip()};", style='ListBullet')
    
    def _add_signatures(self, doc, data):
        """Aggiunge le firme utilizzando una tabella formattata."""
        # Aggiungi spazio prima delle firme
        doc.add_paragraph().paragraph_format.space_before = Pt(36)
        
        table = doc.add_table(rows=2, cols=2)
        table.autofit = False # Permette di controllare le larghezze delle colonne
        table.columns[0].width = Inches(3.0)
        table.columns[1].width = Inches(3.0)
        table.style = 'TableGrid' # Utilizza lo stile di tabella predefinito o uno personalizzato

        # Rimuovi i bordi della tabella se si preferisce solo testo e linee di firma
        # from docx.oxml.ns import qn
        # for row in table.rows:
        #     for cell in row.cells:
        #         tcPr = cell._tc.get_or_add_tcPr()
        #         tcBorders = tcPr.get_or_add_tcBorders()
        #         for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        #             border_el = tcBorders.find(qn(f'w:{border_name}'))
        #             if border_el is not None:
        #                 tcBorders.remove(border_el)
        #             # Per rimuovere completamente, o impostare 'val' a 'nil' o 'none'
            
        # Riga per i ruoli
        cell_pres_role = table.cell(0, 0)
        cell_pres_role.text = "IL PRESIDENTE"
        cell_pres_role.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_pres_role.paragraphs[0].runs[0].font.bold = True
        cell_pres_role.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell_pres_role.paragraphs[0].runs[0].font.size = Pt(11)
        
        cell_secr_role = table.cell(0, 1)
        cell_secr_role.text = "IL SEGRETARIO"
        cell_secr_role.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_secr_role.paragraphs[0].runs[0].font.bold = True
        cell_secr_role.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell_secr_role.paragraphs[0].runs[0].font.size = Pt(11)

        # Riga per i nomi (o spazio per firma)
        # Aggiungi più spazio verticale prima dei nomi/linee di firma
        # Questo si può fare aggiungendo paragrafi vuoti o modificando lo spazio prima/dopo
        # dei paragrafi nelle celle, o l'altezza della riga.
        
        cell_pres_name = table.cell(1, 0)
        # Aggiungi paragrafi per spaziatura e linea di firma
        cell_pres_name.add_paragraph().paragraph_format.space_before = Pt(24) # Spazio per la firma
        p_pres_name = cell_pres_name.add_paragraph(data.get('presidente', '_________________________'))
        p_pres_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_pres_name.runs[0].font.name = 'Times New Roman'
        p_pres_name.runs[0].font.size = Pt(11)
        
        cell_secr_name = table.cell(1, 1)
        cell_secr_name.add_paragraph().paragraph_format.space_before = Pt(24) # Spazio per la firma
        p_secr_name = cell_secr_name.add_paragraph(data.get('segretario', '_________________________'))
        p_secr_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_secr_name.runs[0].font.name = 'Times New Roman'
        p_secr_name.runs[0].font.size = Pt(11)

        # Allinea l'intera tabella al centro della pagina se necessario
        # Questo è più complesso e richiede l'accesso a OXML o l'uso di sezioni.
        # Per ora, il contenuto delle celle è centrato.


# Registra il template nel factory
DocumentTemplateFactory.register_template("verbale_assemblea_template", VerbaleApprovazioneBilancioTemplate)
