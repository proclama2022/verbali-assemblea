"""
Template per Verbale di Assemblea - Nomina del Collegio Sindacale
Questo template è specifico per verbali di nomina del Collegio Sindacale della società.
"""

import sys
import os

# Aggiungi il path della cartella src (relativo alla root del progetto)
current_dir = os.path.dirname(os.path.abspath(__file__))
src_path = os.path.join(os.path.dirname(current_dir), 'src')
if src_path not in sys.path:
    sys.path.append(src_path)

from document_templates import DocumentTemplate, DocumentTemplateFactory
from common_data_handler import CommonDataHandler
from base_verbale_template import BaseVerbaleTemplate
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date, datetime
import streamlit as st
import pandas as pd
import re

class VerbaleNominaCollegioSindacaleTemplate(BaseVerbaleTemplate):
    """Template per Verbale di Assemblea dei Soci - Nomina del Collegio Sindacale"""
    
    def get_template_name(self) -> str:
        return "Nomina del Collegio Sindacale"
    
    def get_required_fields(self) -> list:
        return [
            "denominazione", "sede_legale", "capitale_sociale", "codice_fiscale",
            "data_assemblea", "ora_assemblea", "presidente", "segretario", 
            "soci", "amministratori", "collegio_sindacale"
        ]
    
    def get_form_fields(self, extracted_data: dict) -> dict:
        """Genera i campi del form per Streamlit."""
        # Chiama il metodo della classe base per ottenere i campi comuni
        form_data = super().get_form_fields(extracted_data)

        st.subheader("🏛️ Configurazioni Nomina Collegio Sindacale")

        # Campi specifici del template
        col1, col2 = st.columns(2)
        with col1:
            form_data["tipo_organo_controllo"] = st.selectbox(
                "Tipo Organo di Controllo",
                ["Collegio Sindacale", "Sindaco Unico"],
                key="tipo_organo_controllo_nomina_cs"
            )
        with col2:
            form_data['durata_incarico'] = st.selectbox(
                "Durata dell'incarico",
                ["Tre esercizi", "Un esercizio", "Altro"],
                key="durata_incarico_cs"
            )
            
        # Dati dei nuovi sindaci
        st.subheader("Nuovi Sindaci da Nominare")
        if form_data.get("tipo_organo_controllo") == "Collegio Sindacale":
            num_sindaci = 3
            st.info("Configurazione per Collegio Sindacale (1 Presidente, 2 Sindaci effettivi)")
        else:
            num_sindaci = 1
            st.info("Configurazione per Sindaco Unico")
            
        nuovi_sindaci = []
        for i in range(num_sindaci):
            ruolo = "Presidente" if i == 0 and num_sindaci > 1 else ("Sindaco Effettivo" if num_sindaci > 1 else "Sindaco Unico")
            st.markdown(f"**{ruolo}**")
            
            col1, col2 = st.columns(2)
            with col1:
                nome = st.text_input(f"Nome e Cognome", key=f"sindaco_nome_{i}_nomina_cs")
            with col2:
                cf = st.text_input(f"Codice Fiscale", key=f"sindaco_cf_{i}_nomina_cs")
            
            if nome:
                nuovi_sindaci.append({
                    "nome": nome,
                    "codice_fiscale": cf,
                    "ruolo": ruolo
                })
        form_data['nuovi_sindaci'] = nuovi_sindaci
        
        # Compenso
        st.subheader("💰 Compenso Organo di Controllo")
        form_data['compenso_previsto'] = st.checkbox("È previsto un compenso?", value=True, key="compenso_previsto_cs")
        if form_data.get('compenso_previsto'):
            form_data['compenso_annuo_sindaci'] = st.text_input("Compenso annuo lordo (€)", "0,00", key="compenso_annuo_cs")

        return form_data
    
    def show_preview(self, form_data: dict):
        """Mostra l'anteprima del documento fuori dal form"""
        # Sezione Anteprima
        st.markdown("---")
        st.subheader("👁️ Anteprima Documento")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col1:
            show_preview = st.checkbox("Mostra anteprima", value=False, key="preview_checkbox")
        
        if show_preview:
            try:
                # Debug dettagliato: mostra i dati estratti
                with st.expander("Debug: Dati estratti dal form"):
                    st.json(form_data)
                    
                    # Debug specifico per i soci
                    if 'soci' in form_data and form_data['soci']:
                        st.subheader("Debug Soci:")
                        for i, socio in enumerate(form_data['soci']):
                            st.write(f"Socio {i+1}:")
                            st.write(f"  - Nome: {socio.get('nome', 'N/A')}")
                            st.write(f"  - Tipo Soggetto: {socio.get('tipo_soggetto', 'N/A')}")
                            st.write(f"  - Tipo Partecipazione: {socio.get('tipo_partecipazione', 'N/A')}")
                            st.write(f"  - Delegato: {socio.get('delegato', 'N/A')}")
                            st.write(f"  - Rappresentante Legale: {socio.get('rappresentante_legale', 'N/A')}")
                            st.write(f"  - Presente: {socio.get('presente', 'N/A')}")
                            st.write("---")
                
                preview_text = self._generate_preview_text(form_data)
                st.text_area(
                    "Contenuto del verbale:",
                    value=preview_text,
                    height=600,
                    key="preview_text_nomina_collegio"
                )
            except Exception as e:
                st.error(f"Errore nell'anteprima: {e}")
    
    def _generate_preview_text(self, data: dict) -> str:
        """Genera il testo di anteprima del verbale"""
        
        # Estrazione dati di base
        denominazione = data.get('denominazione', '[DENOMINAZIONE]')
        sede = data.get('sede_legale', '[SEDE]')
        capitale = data.get('capitale_sociale', '[CAPITALE]')
        cf = data.get('codice_fiscale', '[CODICE FISCALE]')
        
        data_assemblea = data.get('data_assemblea', date.today())
        ora_assemblea = data.get('ora_assemblea', '10:00')
        presidente = data.get('presidente', '[PRESIDENTE]')
        segretario = data.get('segretario', '[SEGRETARIO]')
        
        # Dati collegio sindacale
        collegio_members = data.get('collegio_sindacale', [])
        # Se per compatibilità il campo contiene un boolean invece di una lista, usa il campo "sindaci"
        if not isinstance(collegio_members, list):
            collegio_members = data.get('sindaci', [])
        durata_incarico = data.get('durata_incarico', '[DURATA INCARICO]')
        compenso = data.get('compenso_complessivo', '[COMPENSO]')
        motivo_nomina = data.get('motivo_nomina', '[MOTIVO NOMINA]')
        
        if isinstance(data_assemblea, str):
            data_str = data_assemblea
        else:
            data_str = data_assemblea.strftime('%d/%m/%Y')
        
        # Generazione testo anteprima
        text = f"""{denominazione.upper()}
Sede in {sede}
Capitale sociale Euro {capitale} i.v.
Codice fiscale: {cf}

Verbale di assemblea dei soci
del {data_str}

Oggi {data_str} alle ore {ora_assemblea} presso la sede sociale {sede}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:

Ordine del giorno
{', '.join(data.get('punti_ordine_giorno', ['nomina del Collegio Sindacale della società']))}

Assume la presidenza ai sensi dell'art. [...] dello statuto sociale il Sig. {presidente} {data.get('ruolo_presidente', 'Amministratore Unico')}, il quale dichiara e constata:

1 - che l'assemblea risulta {data.get('tipo_assemblea', 'regolarmente convocata')}
"""

        # Aggiungi partecipazione in audioconferenza se prevista
        if data.get('modalita_partecipazione', False):
            text += "2 - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. [...] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza\n"
            next_num = 3
        else:
            next_num = 2
        
        # Aggiungi sezione partecipanti
        text += f"{next_num} - che sono presenti/partecipano all'assemblea:\n"
        text += f"l'{data.get('ruolo_presidente', 'Amministratore Unico')} nella persona del suddetto Presidente Sig. {presidente}\n"
        
        # Aggiungi soci
        soci_presenti = data.get('soci_presenti', [])
        soci_assenti = data.get('soci_assenti', [])

        # Fallback per mantenere compatibilità se le nuove chiavi non ci sono
        if not soci_presenti and not soci_assenti and 'soci' in data:
            soci_presenti = [s for s in data.get('soci', []) if s.get('presente', True)]
            soci_assenti = [s for s in data.get('soci', []) if not s.get('presente', True)]

        if soci_presenti:
            total_quota_euro = 0.0
            total_quota_percentuale = 0.0
            capitale_sociale_float = 0.0
            try:
                capitale_raw = str(data.get('capitale_sociale', '0')).replace('.', '').replace(',', '.')
                capitale_sociale_float = float(capitale_raw)
            except ValueError:
                pass

            for socio in soci_presenti:
                if isinstance(socio, dict):
                    try:
                        quota_euro_str = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                        total_quota_euro += float(quota_euro_str)
                    except (ValueError, TypeError):
                        pass

                    perc_raw = socio.get('quota_percentuale', '')
                    if perc_raw:
                        try:
                            total_quota_percentuale += float(str(perc_raw).replace('%', '').replace(',', '.'))
                        except (ValueError, TypeError):
                            pass
                    else:
                        try:
                            euro_val = float(str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.'))
                            if capitale_sociale_float > 0:
                                total_quota_percentuale += (euro_val / capitale_sociale_float) * 100
                        except (ValueError, TypeError):
                            pass

            formatted_total_quota_euro = CommonDataHandler.format_currency(total_quota_euro)
            formatted_total_quota_percentuale = CommonDataHandler.format_percentage(total_quota_percentuale)

            text += f"nonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_total_quota_euro} pari al {formatted_total_quota_percentuale} del Capitale Sociale:\n"
            for socio in soci_presenti:
                if isinstance(socio, dict):
                    nome = socio.get('nome', '[Nome Socio]')
                    quota_euro = CommonDataHandler.format_currency(socio.get('quota_euro', '0'))
                    
                    quota_perc = socio.get('quota_percentuale', '')
                    if not quota_perc:
                        try:
                            euro_val = float(str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.'))
                            quota_perc = CommonDataHandler.format_percentage((euro_val / capitale_sociale_float) * 100) if capitale_sociale_float > 0 else '[%]'
                        except (ValueError, TypeError):
                            quota_perc = '[%]'
                    else:
                        quota_perc = CommonDataHandler.clean_percentage(quota_perc)

                    tipo_soggetto = socio.get('tipo_soggetto', 'Persona Fisica')
                    tipo_partecipazione = socio.get('tipo_partecipazione', 'Diretto')
                    delegato = socio.get('delegato', '').strip()
                    rappresentante_legale = socio.get('rappresentante_legale', '').strip()

                    linea_socio = ""
                    if tipo_partecipazione == 'Delegato' and delegato:
                        if tipo_soggetto == 'Società':
                            linea_socio = f"il Sig. {delegato} delegato della società {nome} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                        else:
                            linea_socio = f"il Sig. {delegato} delegato del socio Sig. {nome} recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                    else:
                        if tipo_soggetto == 'Società':
                            if rappresentante_legale:
                                linea_socio = f"la società {nome}, rappresentata dal Sig. {rappresentante_legale}, socia recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                            else:
                                linea_socio = f"la società {nome} socia recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                        else:
                            linea_socio = f"il Sig. {nome} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                    
                    text += f"{linea_socio}\n"

        if soci_assenti:
            text += "\nRisultano invece assenti i seguenti soci:\n"
            for socio in soci_assenti:
                if isinstance(socio, dict) and socio.get('nome'):
                    text += f"- {socio.get('nome')}\n"

        text += f"""
{next_num+1} - che gli intervenuti sono legittimati alla presente assemblea;
{next_num+2} - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno.

I presenti all'unanimità chiamano a fungere da segretario il signor {segretario}, che accetta l'incarico.

Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante.

Il Presidente constata e fa constatare che l'assemblea risulta {data.get('tipo_assemblea', 'regolarmente convocata')} e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.

Si passa quindi allo svolgimento dell'ordine del giorno.

*     *     *

Il Presidente informa l'assemblea che si rende necessaria la nomina del Collegio Sindacale poiché {motivo_nomina}.

Il Presidente ricorda all'assemblea quanto previsto dall'art. 2477 del Codice Civile e dall'atto costitutivo della società.

Prende la parola il socio sig. {data.get('socio_proponente', '[SOCIO PROPONENTE]')} che propone di affidare il controllo legale dei conti ad un collegio sindacale composto dai Sigg. [nomi]. Ai sensi dell'art. 2400, ultimo comma del Codice Civile, prima dell'accettazione dell'incarico, i candidati hanno reso noti all'assemblea gli incarichi di amministrazione e di controllo da essi ricoperti presso altre società, mediante dichiarazioni scritte che resteranno depositate agli atti societari.

Segue breve discussione tra i soci al termine della quale si passa alla votazione con voto palese in forza della quale il Presidente constata che, all'unanimità, l'assemblea

DELIBERA:

di nominare un Collegio Sindacale, composto di tre membri effettivi e due supplenti, che dureranno in carica {durata_incarico}, nelle persone dei signori:
"""

        # Aggiungi membri del collegio
        for i, membro in enumerate(collegio_members):
            nome = membro.get('nome', '[NOME]')
            nato_a = membro.get('nato_a', '[LUOGO NASCITA]')
            nato_il = membro.get('nato_il', '[DATA NASCITA]')
            if isinstance(nato_il, date):
                nato_il_str = nato_il.strftime('%d/%m/%Y')
            else:
                nato_il_str = str(nato_il)
            residente = membro.get('residente', '[LUOGO RESIDENZA]')
            cf_membro = membro.get('codice_fiscale', '[CODICE FISCALE]')
            
            albo_data = membro.get('albo_data_gu', '[DATA GU]')
            if isinstance(albo_data, date):
                albo_data_str = albo_data.strftime('%d/%m/%Y')
            else:
                albo_data_str = str(albo_data)
                
            albo_num = membro.get('albo_num_gu', '[NUM GU]')
            albo_serie = membro.get('albo_serie_gu', '[SERIE GU]')
            ruolo = membro.get('ruolo', '[RUOLO]')
            
            text += f"{nome} nato a {nato_a} il {nato_il_str}, residente in {residente}, codice fiscale {cf_membro}, Revisore Contabile pubblicato sulla G.U. in data {albo_data_str} N. {albo_num}, {albo_serie} serie speciale; {ruolo};\n"

        text += f"""
di corrispondere ai membri effettivi del Collegio Sindacale un compenso annuo complessivo pari a euro {compenso}
"""

        # Aggiunta presenza membri se applicabile
        if data.get('membri_presenti', False):
            qualita = data.get('qualita_presenza', 'invitati')
            text += f"""
I Sigg. [nomi], presenti in assemblea in qualità di {qualita} accettano l'incarico e ringraziano l'assemblea per la fiducia accordata.
"""
        else:
            text += f"""
[L'accettazione della carica da parte dei neo-nominati potrà avvenire successivamente]
"""

        # Nota sul controllo contabile
        if data.get('controllo_contabile_collegio', True):
            text += f"""
[Verificare se l'atto costitutivo non dispone diversamente, il controllo contabile è esercitato dal collegio sindacale].
"""

        text += """
*     *     *

Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.

Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo.

L'assemblea viene sciolta alle ore [...].


Il Presidente                    Il Segretario
_________________            _________________
"""

        return text
    
    def generate_document(self, data: dict) -> Document:
        """Genera il documento Word utilizzando i dati forniti"""
        import os
        
        # Utilizza il template .docx esistente per mantenere la formattazione
        template_path = os.path.join(os.path.dirname(__file__), 'template.docx')
        
        try:
            if os.path.exists(template_path):
                doc = Document(template_path)
                # Rimuovi il contenuto esistente del template mantenendo gli stili
                for paragraph in doc.paragraphs[:]:
                    p = paragraph._element
                    p.getparent().remove(p)
            else:
                doc = Document()
                # Setup stili solo se non usiamo template
                self._setup_document_styles(doc)
        except Exception as e:
            # Fallback a documento vuoto se il template non può essere caricato
            doc = Document()
            self._setup_document_styles(doc)
        
        # Imposta stili del documento
        self._setup_document_styles(doc)
        
        # Aggiungi intestazione azienda
        self._add_company_header(doc, data)
        
        # Aggiungi titolo verbale
        self._add_verbale_title(doc, data)
        
        # Aggiungi sezione di apertura
        self._add_opening_section(doc, data)
        
        # Aggiungi sezione partecipanti
        self._add_participants_section(doc, data)
        
        # Aggiungi dichiarazioni preliminari
        self._add_preliminary_statements(doc, data)
        
        # Aggiungi discussione nomina collegio sindacale
        self._add_nomina_collegio_discussion(doc, data)
        
        # Aggiungi sezione di chiusura
        self._add_closing_section(doc, data)
        
        # Aggiungi firme
        self._add_signatures(doc, data)
        
        return doc
    
    def _setup_document_styles(self, doc):
        """Imposta gli stili del documento"""
        styles = doc.styles
        
        # Stile per l'intestazione
        if 'Intestazione' not in [s.name for s in styles]:
            header_style = styles.add_style('Intestazione', WD_STYLE_TYPE.PARAGRAPH)
            header_style.font.name = 'Times New Roman'
            header_style.font.size = Pt(12)
            header_style.font.bold = True
            header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Stile per il titolo principale
        if 'Titolo Principale' not in [s.name for s in styles]:
            title_style = styles.add_style('Titolo Principale', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Times New Roman'
            title_style.font.size = Pt(14)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Stile per il testo normale
        normal_style = styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.line_spacing = 1.15
    
    def _add_company_header(self, doc, data):
        """Aggiungi intestazione dell'azienda"""
        # Denominazione
        p = doc.add_paragraph()
        p.style = 'Intestazione'
        run = p.add_run(data.get('denominazione', '[DENOMINAZIONE]').upper())
        run.bold = True
        
        # Sede
        p = doc.add_paragraph()
        p.style = 'Intestazione'
        p.add_run(f"Sede in {data.get('sede_legale', '[SEDE]')}")
        
        # Capitale sociale
        p = doc.add_paragraph()
        p.style = 'Intestazione'
        p.add_run(f"Capitale sociale Euro {data.get('capitale_sociale', '[CAPITALE]')} i.v.")
        
        # Codice fiscale
        p = doc.add_paragraph()
        p.style = 'Intestazione'
        p.add_run(f"Codice fiscale: {data.get('codice_fiscale', '[CODICE FISCALE]')}")
        
        # Spazio
        doc.add_paragraph()
    
    def _add_verbale_title(self, doc, data):
        """Aggiungi titolo del verbale"""
        p = doc.add_paragraph()
        p.style = 'Titolo Principale'
        p.add_run("Verbale di assemblea dei soci")
        
        # Data
        data_assemblea = data.get('data_assemblea', date.today())
        if isinstance(data_assemblea, str):
            data_str = data_assemblea
        else:
            data_str = data_assemblea.strftime('%d/%m/%Y')
        
        p = doc.add_paragraph()
        p.style = 'Titolo Principale'
        p.add_run(f"del {data_str}")
        
        # Spazio
        doc.add_paragraph()
    
    def _add_opening_section(self, doc, data):
        """Aggiungi sezione di apertura"""
        data_assemblea = data.get('data_assemblea', date.today())
        if isinstance(data_assemblea, str):
            data_str = data_assemblea
        else:
            data_str = data_assemblea.strftime('%d/%m/%Y')
        
        ora = data.get('ora_assemblea', '10:00')
        sede = data.get('sede_legale', '[SEDE]')
        
        p = doc.add_paragraph()
        p.add_run(f"Oggi {data_str} alle ore {ora} presso la sede sociale {sede}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:")
        
        # Ordine del giorno
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Ordine del giorno")
        run.bold = True
        
        punti_odg = data.get('punti_ordine_giorno', ['nomina del Collegio Sindacale della società'])
        for punto in punti_odg:
            p = doc.add_paragraph()
            p.add_run(punto)
        
        doc.add_paragraph()
    
    def _add_participants_section(self, doc, data):
        """Aggiungi sezione partecipanti"""
        presidente = data.get('presidente', '[PRESIDENTE]')
        ruolo_presidente = data.get('ruolo_presidente', 'Amministratore Unico')
        
        # Presidenza
        p = doc.add_paragraph()
        p.add_run(f"Assume la presidenza ai sensi dell'art. [...] dello statuto sociale il Sig. {presidente} {ruolo_presidente}, il quale dichiara e constata:")
        
        # Dichiarazioni
        counter = 1
        
        # Tipo assemblea
        p = doc.add_paragraph()
        p.add_run(f"{counter} - che l'assemblea risulta {data.get('tipo_assemblea', 'regolarmente convocata')}")
        counter += 1
        
        # Audioconferenza se prevista
        if data.get('modalita_partecipazione', False):
            p = doc.add_paragraph()
            p.add_run(f"{counter} - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. [...] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza")
            counter += 1
        
        # Presenti
        p = doc.add_paragraph()
        p.add_run(f"{counter} - che sono presenti/partecipano all'assemblea:")
        
        p = doc.add_paragraph()
        p.add_run(f"l'{data.get('ruolo_presidente', 'Amministratore Unico')} nella persona del suddetto Presidente Sig. {presidente}")
        
        # Soci
        soci_presenti = data.get('soci_presenti', [])
        soci_assenti = data.get('soci_assenti', [])

        # Fallback per mantenere compatibilità se le nuove chiavi non ci sono
        if not soci_presenti and not soci_assenti and 'soci' in data:
            soci_presenti = [s for s in data.get('soci', []) if s.get('presente', True)]
            soci_assenti = [s for s in data.get('soci', []) if not s.get('presente', True)]

        if soci_presenti:
            total_quota_euro = 0.0
            total_quota_percentuale = 0.0
            capitale_sociale_float = 0.0
            try:
                capitale_raw = str(data.get('capitale_sociale', '0')).replace('.', '').replace(',', '.')
                capitale_sociale_float = float(capitale_raw)
            except ValueError:
                pass

            for socio in soci_presenti:
                if isinstance(socio, dict):
                    try:
                        quota_euro_str = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                        total_quota_euro += float(quota_euro_str)
                    except (ValueError, TypeError):
                        pass

                    perc_raw = socio.get('quota_percentuale', '')
                    if perc_raw:
                        try:
                            total_quota_percentuale += float(str(perc_raw).replace('%', '').replace(',', '.'))
                        except (ValueError, TypeError):
                            pass
                    else:
                        try:
                            euro_val = float(str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.'))
                            if capitale_sociale_float > 0:
                                total_quota_percentuale += (euro_val / capitale_sociale_float) * 100
                        except (ValueError, TypeError):
                            pass

            formatted_total_quota_euro = CommonDataHandler.format_currency(total_quota_euro)
            formatted_total_quota_percentuale = CommonDataHandler.format_percentage(total_quota_percentuale)

            p = doc.add_paragraph()
            p.add_run(f"nonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_total_quota_euro} pari al {formatted_total_quota_percentuale} del Capitale Sociale:")
            
            for socio in soci_presenti:
                nome = socio.get('nome', '[NOME SOCIO]')
                quota_euro = CommonDataHandler.format_currency(socio.get('quota_euro', '0'))
                
                quota_perc = socio.get('quota_percentuale', '')
                if not quota_perc:
                    try:
                        euro_val = float(str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.'))
                        quota_perc = CommonDataHandler.format_percentage((euro_val / capitale_sociale_float) * 100) if capitale_sociale_float > 0 else '[%]'
                    except (ValueError, TypeError):
                        quota_perc = '[%]'
                else:
                    quota_perc = CommonDataHandler.clean_percentage(quota_perc)

                tipo_soggetto = socio.get('tipo_soggetto', 'Persona Fisica')
                tipo_partecipazione = socio.get('tipo_partecipazione', 'Diretto')
                delegato = socio.get('delegato', '').strip()
                rappresentante_legale = socio.get('rappresentante_legale', '').strip()
                
                p = doc.add_paragraph()
                
                linea_socio = ""
                if tipo_partecipazione == 'Delegato' and delegato:
                    if tipo_soggetto == 'Società':
                        linea_socio = f"il Sig. {delegato} delegato della società {nome} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                    else:
                        linea_socio = f"il Sig. {delegato} delegato del socio Sig. {nome} recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                else:
                    if tipo_soggetto == 'Società':
                        if rappresentante_legale:
                            linea_socio = f"la società {nome}, rappresentata dal Sig. {rappresentante_legale}, socia recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                        else:
                            linea_socio = f"la società {nome} socia recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                    else:
                        linea_socio = f"il Sig. {nome} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                p.add_run(linea_socio)

        if soci_assenti:
            p = doc.add_paragraph()
            p.add_run("Risultano invece assenti i seguenti soci:")
            for socio in soci_assenti:
                if isinstance(socio, dict) and socio.get('nome'):
                    p = doc.add_paragraph(f"- {socio.get('nome')}")
        
        counter += 1
        
        # Altre dichiarazioni
        p = doc.add_paragraph()
        p.add_run(f"{counter} - che gli intervenuti sono legittimati alla presente assemblea;")
        counter += 1
        
        p = doc.add_paragraph()
        p.add_run(f"{counter} - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno.")
        
        doc.add_paragraph()
    
    def _add_preliminary_statements(self, doc, data):
        """Aggiungi dichiarazioni preliminari"""
        segretario = data.get('segretario', '[SEGRETARIO]')
        
        p = doc.add_paragraph()
        p.add_run(f"I presenti all'unanimità chiamano a fungere da segretario il signor {segretario}, che accetta l'incarico.")
        
        p = doc.add_paragraph()
        p.add_run("Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante.")
        
        p = doc.add_paragraph()
        tipo_assemblea = data.get('tipo_assemblea', 'regolarmente convocata')
        p.add_run(f"Il Presidente constata e fa constatare che l'assemblea risulta {tipo_assemblea} e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.")
        
        p = doc.add_paragraph()
        p.add_run("Si passa quindi allo svolgimento dell'ordine del giorno.")
        
        # Separatore
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("*     *     *")
        
        doc.add_paragraph()
    
    def _add_nomina_collegio_discussion(self, doc, data):
        """Aggiungi sezione discussione nomina collegio sindacale"""
        motivo_nomina = data.get('motivo_nomina', '[MOTIVO NOMINA]')
        socio_proponente = data.get('socio_proponente', '[SOCIO PROPONENTE]')
        collegio_members = data.get('collegio_sindacale', [])
        # Se per compatibilità il campo contiene un boolean invece di una lista, usa il campo "sindaci"
        if not isinstance(collegio_members, list):
            collegio_members = data.get('sindaci', [])
        durata_incarico = data.get('durata_incarico', '[DURATA INCARICO]')
        compenso = data.get('compenso_complessivo', '[COMPENSO]')
        
        if isinstance(data_assemblea, str):
            data_str = data_assemblea
        else:
            data_str = data_assemblea.strftime('%d/%m/%Y')
        
        # Generazione testo anteprima
        text = f"""{denominazione.upper()}
Sede in {sede}
Capitale sociale Euro {capitale} i.v.
Codice fiscale: {cf}

Verbale di assemblea dei soci
del {data_str}

Oggi {data_str} alle ore {ora_assemblea} presso la sede sociale {sede}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:

Ordine del giorno
{', '.join(data.get('punti_ordine_giorno', ['nomina del Collegio Sindacale della società']))}

Assume la presidenza ai sensi dell'art. [...] dello statuto sociale il Sig. {presidente} {data.get('ruolo_presidente', 'Amministratore Unico')}, il quale dichiara e constata:

1 - che l'assemblea risulta {data.get('tipo_assemblea', 'regolarmente convocata')}
"""

        # Aggiungi partecipazione in audioconferenza se prevista
        if data.get('modalita_partecipazione', False):
            text += "2 - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. [...] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza\n"
            next_num = 3
        else:
            next_num = 2
        
        # Aggiungi sezione partecipanti
        text += f"{next_num} - che sono presenti/partecipano all'assemblea:\n"
        text += f"l'{data.get('ruolo_presidente', 'Amministratore Unico')} nella persona del suddetto Presidente Sig. {presidente}\n"
        
        # Aggiungi soci
        soci_presenti = data.get('soci_presenti', [])
        soci_assenti = data.get('soci_assenti', [])

        # Fallback per mantenere compatibilità se le nuove chiavi non ci sono
        if not soci_presenti and not soci_assenti and 'soci' in data:
            soci_presenti = [s for s in data.get('soci', []) if s.get('presente', True)]
            soci_assenti = [s for s in data.get('soci', []) if not s.get('presente', True)]

        if soci_presenti:
            total_quota_euro = 0.0
            total_quota_percentuale = 0.0
            capitale_sociale_float = 0.0
            try:
                capitale_raw = str(data.get('capitale_sociale', '0')).replace('.', '').replace(',', '.')
                capitale_sociale_float = float(capitale_raw)
            except ValueError:
                pass

            for socio in soci_presenti:
                if isinstance(socio, dict):
                    try:
                        quota_euro_str = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                        total_quota_euro += float(quota_euro_str)
                    except (ValueError, TypeError):
                        pass

                    perc_raw = socio.get('quota_percentuale', '')
                    if perc_raw:
                        try:
                            total_quota_percentuale += float(str(perc_raw).replace('%', '').replace(',', '.'))
                        except (ValueError, TypeError):
                            pass
                    else:
                        try:
                            euro_val = float(str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.'))
                            if capitale_sociale_float > 0:
                                total_quota_percentuale += (euro_val / capitale_sociale_float) * 100
                        except (ValueError, TypeError):
                            pass

            formatted_total_quota_euro = CommonDataHandler.format_currency(total_quota_euro)
            formatted_total_quota_percentuale = CommonDataHandler.format_percentage(total_quota_percentuale)

            text += f"nonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_total_quota_euro} pari al {formatted_total_quota_percentuale} del Capitale Sociale:\n"
            for socio in soci_presenti:
                if isinstance(socio, dict):
                    nome = socio.get('nome', '[NOME SOCIO]')
                    quota_euro = CommonDataHandler.format_currency(socio.get('quota_euro', '0'))
                    
                    quota_perc = socio.get('quota_percentuale', '')
                    if not quota_perc:
                        try:
                            euro_val = float(str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.'))
                            quota_perc = CommonDataHandler.format_percentage((euro_val / capitale_sociale_float) * 100) if capitale_sociale_float > 0 else '[%]'
                        except (ValueError, TypeError):
                            quota_perc = '[%]'
                    else:
                        quota_perc = CommonDataHandler.clean_percentage(quota_perc)

                    tipo_soggetto = socio.get('tipo_soggetto', 'Persona Fisica')
                    tipo_partecipazione = socio.get('tipo_partecipazione', 'Diretto')
                    delegato = socio.get('delegato', '').strip()
                    rappresentante_legale = socio.get('rappresentante_legale', '').strip()

                    linea_socio = ""
                    if tipo_partecipazione == 'Delegato' and delegato:
                        if tipo_soggetto == 'Società':
                            linea_socio = f"il Sig. {delegato} delegato della società {nome} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                        else:
                            linea_socio = f"il Sig. {delegato} delegato del socio Sig. {nome} recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                    else:
                        if tipo_soggetto == 'Società':
                            if rappresentante_legale:
                                linea_socio = f"la società {nome}, rappresentata dal Sig. {rappresentante_legale}, socia recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                            else:
                                linea_socio = f"la società {nome} socia recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                        else:
                            linea_socio = f"il Sig. {nome} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                    
                    text += f"{linea_socio}\n"

        if soci_assenti:
            text += "\nRisultano invece assenti i seguenti soci:\n"
            for socio in soci_assenti:
                if isinstance(socio, dict) and socio.get('nome'):
                    text += f"- {socio.get('nome')}\n"

        text += f"""
{next_num+1} - che gli intervenuti sono legittimati alla presente assemblea;
{next_num+2} - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno.

I presenti all'unanimità chiamano a fungere da segretario il signor {segretario}, che accetta l'incarico.

Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante.

Il Presidente constata e fa constatare che l'assemblea risulta {data.get('tipo_assemblea', 'regolarmente convocata')} e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.

Si passa quindi allo svolgimento dell'ordine del giorno.

*     *     *

Il Presidente informa l'assemblea che si rende necessaria la nomina del Collegio Sindacale poiché {motivo_nomina}.

Il Presidente ricorda all'assemblea quanto previsto dall'art. 2477 del Codice Civile e dall'atto costitutivo della società.

Prende la parola il socio sig. {socio_proponente} che propone di affidare il controllo legale dei conti ad un collegio sindacale composto dai Sigg. [nomi]. Ai sensi dell'art. 2400, ultimo comma del Codice Civile, prima dell'accettazione dell'incarico, i candidati hanno reso noti all'assemblea gli incarichi di amministrazione e di controllo da essi ricoperti presso altre società, mediante dichiarazioni scritte che resteranno depositate agli atti societari.

Segue breve discussione tra i soci al termine della quale si passa alla votazione con voto palese in forza della quale il Presidente constata che, all'unanimità, l'assemblea")
        
        # DELIBERA
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("D E L I B E R A:")
        run.bold = True
        
        doc.add_paragraph()
        
        # Nomina collegio
        durata_incarico = data.get('durata_incarico', '[DURATA INCARICO]')
        
        p = doc.add_paragraph()
        p.add_run(f"di nominare un Collegio Sindacale, composto di tre membri effettivi e due supplenti, che dureranno in carica {durata_incarico}, nelle persone dei signori:")
        
        # Elenco membri
        for membro in collegio_members:
            nome = membro.get('nome', '[NOME]')
            nato_a = membro.get('nato_a', '[LUOGO NASCITA]')
            nato_il = membro.get('nato_il', '[DATA NASCITA]')
            if isinstance(nato_il, date):
                nato_il_str = nato_il.strftime('%d/%m/%Y')
            else:
                nato_il_str = str(nato_il)
            
            residente = membro.get('residente', '[LUOGO RESIDENZA]')
            cf_membro = membro.get('codice_fiscale', '[CODICE FISCALE]')
            
            albo_data = membro.get('albo_data_gu', '[DATA GU]')
            if isinstance(albo_data, date):
                albo_data_str = albo_data.strftime('%d/%m/%Y')
            else:
                albo_data_str = str(albo_data)
                
            albo_num = membro.get('albo_num_gu', '[NUM GU]')
            albo_serie = membro.get('albo_serie_gu', '[SERIE GU]')
            ruolo = membro.get('ruolo', '[RUOLO]')
            
            p = doc.add_paragraph()
            p.add_run(f"{nome} nato a {nato_a} il {nato_il_str}, residente in {residente}, codice fiscale {cf_membro}, Revisore Contabile pubblicato sulla G.U. in data {albo_data_str} N. {albo_num}, {albo_serie} serie speciale; {ruolo};")
        
        # Compenso
        compenso = data.get('compenso_complessivo', '[COMPENSO]')
        p = doc.add_paragraph()
        p.add_run(f"di corrispondere ai membri effettivi del Collegio Sindacale un compenso annuo complessivo pari a euro {compenso}")
        
        doc.add_paragraph()
        
        # Accettazione
        if data.get('membri_presenti', False):
            qualita = data.get('qualita_presenza', 'invitati')
            p = doc.add_paragraph()
            p.add_run(f"I Sigg. [nomi], presenti in assemblea in qualità di {qualita} accettano l'incarico e ringraziano l'assemblea per la fiducia accordata.")
        else:
            p = doc.add_paragraph()
            p.add_run("[L'accettazione della carica da parte dei neo-nominati potrà avvenire successivamente]")
        
        # Nota controllo contabile
        if data.get('controllo_contabile_collegio', True):
            p = doc.add_paragraph()
            p.add_run("[Verificare se l'atto costitutivo non dispone diversamente, il controllo contabile è esercitato dal collegio sindacale].")
        
        # Separatore
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("*     *     *")
        
        doc.add_paragraph()
    
    def _add_closing_section(self, doc, data):
        """Aggiungi sezione di chiusura"""
        p = doc.add_paragraph()
        p.add_run("Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.")
        
        p = doc.add_paragraph()
        p.add_run("Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo.")
        
        p = doc.add_paragraph()
        p.add_run("L'assemblea viene sciolta alle ore [...].")
        
        doc.add_paragraph()
        doc.add_paragraph()
    
    def _add_signatures(self, doc, data):
        """Aggiungi sezione firme"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Crea una tabella per le firme
        table = doc.add_table(rows=2, cols=2)
        table.autofit = False
        
        # Intestazioni
        table.cell(0, 0).text = "Il Presidente"
        table.cell(0, 1).text = "Il Segretario"
        
        # Linee per le firme
        table.cell(1, 0).text = "_________________"
        table.cell(1, 1).text = "_________________"
        
        # Centra la tabella
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Registra il template
DocumentTemplateFactory.register_template("nomina_collegio_sindacale", VerbaleNominaCollegioSindacaleTemplate)