"""
Template per Verbale di Assemblea - Distribuzione Dividendi
Questo template è specifico per verbali di distribuzione dividendi ai soci.
"""

import sys
import os

# Aggiungi il path della cartella src (relativo alla root del progetto)
current_dir = os.path.dirname(os.path.abspath(__file__))
src_path = os.path.join(os.path.dirname(current_dir), 'src')
if src_path not in sys.path:
    sys.path.append(src_path)

from document_templates import DocumentTemplate, DocumentTemplateFactory
from common_data_handler import CommonDataHandler
from base_verbale_template import BaseVerbaleTemplate
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date
import streamlit as st
import pandas as pd
import re

class VerbaleDividendiTemplate(BaseVerbaleTemplate):
    """Template per Verbale di Assemblea dei Soci - Distribuzione Dividendi"""
    
    def get_template_name(self) -> str:
        return "Distribuzione Dividendi"
    
    def get_required_fields(self) -> list:
        return [
            "denominazione", "sede_legale", "capitale_sociale", "codice_fiscale",
            "data_assemblea", "ora_assemblea", "presidente", "segretario", 
            "soci", "amministratori", "importo_dividendi", "socio_proponente"
        ]
    
    def get_form_fields(self, extracted_data: dict) -> dict:
        """Genera i campi del form per Streamlit utilizzando il CommonDataHandler."""
        # Chiama il metodo della classe base per ottenere i campi comuni
        form_data = super().get_form_fields(extracted_data)

        # Campi specifici per distribuzione dividendi
        st.subheader("💰 Configurazioni Specifiche Distribuzione Dividendi")
        
        col1, col2 = st.columns(2)
        with col1:
            form_data["ruolo_presidente"] = st.selectbox("Ruolo del presidente", 
                                                        CommonDataHandler.get_standard_ruoli_presidente(),
                                                        key="ruolo_presidente_dividendi")
            form_data["importo_dividendi"] = st.text_input("Importo totale dividendi (€)", 
                                                          "0,00",
                                                          help="Importo complessivo da distribuire",
                                                          key="importo_dividendi")
        with col2:
            form_data["modalita_assemblea"] = st.selectbox("Modalità assemblea", 
                                                          ["In presenza", "Audioconferenza", "Mista"],
                                                          key="modalita_assemblea_dividendi")
            soci_options = [s.get("nome", "") for s in form_data.get("soci", [])]
            form_data["socio_proponente"] = st.selectbox("Socio proponente", 
                                                         soci_options,
                                                         help="Nome del socio che propone la distribuzione",
                                                         key="socio_proponente_dividendi")

        # Ordine del giorno specifico per dividendi
        st.subheader("📋 Ordine del Giorno")
        default_odg = "Proposta di distribuzione dei dividendi"
        form_data["ordine_giorno"] = st.text_area("Ordine del giorno", 
                                                 default_odg, height=80, key="odg_dividendi")
        
        # Sezione dettagli distribuzione
        st.subheader("💸 Dettagli Distribuzione")
        
        col1, col2 = st.columns(2)
        with col1:
            form_data["tipo_distribuzione"] = st.selectbox("Tipo di distribuzione", 
                                                          ["Utili pregressi", "Riserve disponibili", "Utile dell'esercizio"],
                                                          key="tipo_distribuzione_dividendi")
            form_data["modalita_ripartizione"] = st.selectbox("Modalità ripartizione", 
                                                             ["Proporzionale alla quota", "Secondo atto costitutivo", "Personalizzata"],
                                                             key="modalita_ripartizione_dividendi")
        with col2:
            form_data["prelievo_riserve"] = st.checkbox("Prelievo da fondo di riserva", key="prelievo_riserve_dividendi")
            if form_data["prelievo_riserve"]:
                form_data["nome_riserva"] = st.text_input("Nome del fondo di riserva", "Riserva utili", key="nome_riserva_dividendi")
        
        # Calcolo automatico dividendi per socio
        st.subheader("📊 Calcolo Dividendi per Socio")
        
        if form_data.get("importo_dividendi") and form_data.get("soci"):
            try:
                importo_totale_str = str(form_data["importo_dividendi"]).replace(".", "").replace(",", ".")
                importo_totale = float(importo_totale_str)
                
                # Mostra calcolo per ogni socio
                soci_dividendi = []
                for socio in form_data.get("soci", []):
                    if socio.get("quota_percentuale"):
                        try:
                            perc_str = str(socio["quota_percentuale"]).replace("%", "").replace(",", ".")
                            perc = float(perc_str)
                            dividendo_socio = (importo_totale * perc) / 100
                            soci_dividendi.append({
                                "nome": socio.get("nome", ""),
                                "quota_perc": f"{perc:.2f}%",
                                "dividendo": f"{dividendo_socio:,.2f} €".replace(",", "X").replace(".", ",").replace("X", ".")
                            })
                        except (ValueError, TypeError):
                            pass
                
                if soci_dividendi:
                    st.write("**Calcolo automatico dividendi:**")
                    df_dividendi = pd.DataFrame(soci_dividendi)
                    st.dataframe(df_dividendi, use_container_width=True, hide_index=True)
                    form_data["calcolo_dividendi"] = soci_dividendi
                    
            except (ValueError, TypeError):
                st.warning("⚠️ Inserisci un importo valido per visualizzare il calcolo")
        
        # Modalità di voto
        st.subheader("🗳️ Modalità di Voto")
        
        col1, col2 = st.columns(2)
        with col1:
            form_data["tipo_voto"] = st.selectbox("Esito del voto", 
                                                 ["Unanimità", "Maggioranza con contrari", "Maggioranza con astenuti"],
                                                 key="tipo_voto_dividendi")
        with col2:
            if form_data["tipo_voto"] != "Unanimità":
                form_data["voti_contrari"] = st.text_input("Nomi voti contrari (separati da virgola)", "", key="voti_contrari_dividendi")
                if form_data["tipo_voto"] == "Maggioranza con astenuti":
                    form_data["astensioni"] = st.text_input("Nomi astensioni (separati da virgola)", "", key="astensioni_dividendi")
        
        # Traduzione (se necessaria)
        st.subheader("🌐 Traduzione")
        form_data["richiede_traduzione"] = st.checkbox("È necessaria traduzione per partecipanti stranieri", key="traduzione_dividendi")
        if form_data["richiede_traduzione"]:
            col1, col2 = st.columns(2)
            with col1:
                form_data["partecipante_straniero"] = st.text_input("Nome partecipante straniero", key="part_straniero_dividendi")
                form_data["lingua_straniera"] = st.selectbox("Lingua di traduzione", 
                                                           ["Inglese", "Francese", "Tedesco", "Spagnolo", "Altro"],
                                                           key="lingua_dividendi")
            with col2:
                if form_data["lingua_straniera"] == "Altro":
                    form_data["lingua_altro"] = st.text_input("Specifica lingua", key="lingua_altro_dividendi")
        
        return form_data
    
    def show_preview(self, form_data: dict):
        """Mostra l'anteprima del documento fuori dal form"""
        # Sezione Anteprima
        st.markdown("---")
        st.subheader("👁️ Anteprima Documento")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col1:
            show_preview = st.checkbox("Mostra anteprima", value=False, key="preview_checkbox")
        
        with col2:
            if show_preview:
                st.info("💡 L'anteprima si aggiorna automaticamente con i dati inseriti sopra")
        
        if show_preview:
            try:
                preview_text = self._generate_preview_text(form_data)
                # Anteprima in text_area per permettere copia/modifica come negli altri verbali
                st.text_area("", preview_text, height=600)
            except Exception as e:
                st.error(f"Errore nell'anteprima: {e}")
    
    def _format_date(self, value):
        """Ritorna la data in formato gg/mm/aaaa se value è una data, altrimenti la stringa così com'è."""
        from datetime import date as _date
        if isinstance(value, _date):
            return value.strftime('%d/%m/%Y')
        return str(value)

    def _format_time(self, value):
        """Ritorna l'orario in formato HH:MM se value è un time, altrimenti la stringa così com'è."""
        from datetime import time as _time
        if isinstance(value, _time):
            return value.strftime('%H:%M')
        return str(value)

    def _get_capitale_sociale(self, data):
        """Restituisce la stringa del capitale sociale scegliendo tra i campi disponibili."""
        # Preferenza: deliberato -> sottoscritto -> versato -> capitale_sociale
        for key in ["capitale_deliberato", "capitale_sottoscritto", "capitale_versato", "capitale_sociale"]:
            value = data.get(key)
            if value:
                return str(value)
        return "[CAPITALE]"

    def _generate_preview_text(self, data: dict) -> str:
        """Genera il testo di anteprima del verbale"""
        
        # Determina la frase di convocazione (regolarmente convocata vs totalitaria)
        convocazione_phrase = "totalitaria" if "totalitaria" in str(data.get("tipo_convocazione", "")).lower() else "regolarmente convocata"
        
        # Formattazioni utili
        data_str = self._format_date(data.get('data_assemblea', '[DATA]'))
        ora_str = self._format_time(data.get('ora_inizio', data.get('ora_assemblea', '[ORA]')))

        # Header aziendale
        preview = f"""
**{data.get('denominazione', '[DENOMINAZIONE]')}**

Sede in {data.get('sede_legale', '[SEDE LEGALE]')}
Capitale sociale Euro {self._get_capitale_sociale(data)} i.v.
Codice fiscale: {data.get('codice_fiscale', '[CODICE FISCALE]')}

---

# Verbale di assemblea dei soci
del {data_str}

---

Oggi {data_str} alle ore {ora_str} presso la sede sociale {data.get('sede_legale', '[SEDE]')}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:

## Ordine del giorno
{data.get('ordine_giorno', 'Proposta di distribuzione dei dividendi')}

---

Assume la presidenza ai sensi dell'art. [...] dello statuto sociale il Sig. **{data.get('presidente', '[PRESIDENTE]')}** {data.get('ruolo_presidente', 'Amministratore Unico')}, il quale dichiara e constata:

1. che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. [...] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza

2. che sono presenti/partecipano all'assemblea:
   - l'{data.get('ruolo_presidente', 'Amministratore Unico')} nella persona del suddetto Presidente Sig. **{data.get('presidente', '[PRESIDENTE]')}**
"""
        
        # Collegio sindacale se presente
        if data.get('include_collegio_sindacale', False):
            tipo_organo_controllo = data.get('tipo_organo_controllo', 'Collegio Sindacale')
            
            if tipo_organo_controllo == 'Collegio Sindacale':
                sindaci = data.get('sindaci', [])
                sindaci_presenti = [s for s in sindaci if s.get('presente')]
                if sindaci_presenti:
                    preview += "\n   - per il Collegio Sindacale:\n"
                    for sindaco in sindaci_presenti:
                        carica = sindaco.get('carica', 'Sindaco Effettivo')
                        nome_sindaco = sindaco.get('nome', '')
                        if nome_sindaco:
                            preview += f"     - il Dott. {nome_sindaco} - {carica}\n"
            else: # Sindaco Unico
                sindaci = data.get('sindaci', [])
                if sindaci and sindaci[0].get('nome'):
                    sindaco_unico_nome = sindaci[0].get('nome')
                    preview += f"\n   - il Sindaco Unico nella persona del Sig. {sindaco_unico_nome}"

        # Sezione partecipanti
        soci_presenti = data.get('soci_presenti', [])
        soci_assenti = data.get('soci_assenti', [])
        
        # Fallback per mantenere compatibilità
        if not soci_presenti and not soci_assenti and 'soci' in data:
            soci_presenti = [s for s in data.get('soci', []) if s.get('presente', True)]
            soci_assenti = [s for s in data.get('soci', []) if not s.get('presente', True)]

        if soci_presenti:
            # Calcola totali euro e percentuale solo sui presenti
            totale_euro = 0.0
            totale_perc = 0.0
            capitale_raw = str(data.get('capitale_sociale', '0')).replace('.', '').replace(',', '.')
            try:
                capitale_float = float(capitale_raw)
            except ValueError:
                capitale_float = 0.0

            for socio in soci_presenti:
                if isinstance(socio, dict):
                    euro_raw = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                    try:
                        totale_euro += float(euro_raw or 0)
                    except ValueError:
                        pass
                    perc_raw = socio.get('quota_percentuale', '')
                    if perc_raw:
                        try:
                            totale_perc += float(str(perc_raw).replace('%', '').replace(',', '.'))
                        except ValueError:
                            pass
                    else:
                        try:
                            euro_val = float(euro_raw or 0)
                            if capitale_float:
                                totale_perc += euro_val / capitale_float * 100
                        except ValueError:
                            pass

            formatted_euro = CommonDataHandler.format_currency(totale_euro)
            formatted_perc = CommonDataHandler.format_percentage(totale_perc)

            preview += f"\n   - nonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_euro} pari al {formatted_perc} del Capitale Sociale:\n"
            for socio in soci_presenti:
                if not isinstance(socio, dict) or not socio.get('nome'):
                    continue
                nome = socio.get('nome')
                quota_euro = CommonDataHandler.format_currency(socio.get('quota_euro', '0'))
                quota_perc = socio.get('quota_percentuale', '')
                if not quota_perc:
                    try:
                        euro_val = float(str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.'))
                        quota_perc = CommonDataHandler.format_percentage(euro_val / capitale_float * 100) if capitale_float else '[%]'
                    except ValueError:
                        quota_perc = '[%]'
                else:
                    quota_perc = CommonDataHandler.clean_percentage(quota_perc)

                tipo_sogg = socio.get('tipo_soggetto', 'Persona Fisica')
                tipo_part = socio.get('tipo_partecipazione', 'Diretto')
                delegato = socio.get('delegato', '').strip()
                rappresentante = socio.get('rappresentante_legale', '').strip()

                if tipo_part == 'Delegato' and delegato:
                    if tipo_sogg == 'Società':
                        line = f"il Sig. {delegato} delegato della società {nome} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                    else:
                        line = f"il Sig. {delegato} delegato del socio Sig. {nome} recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                else:
                    if tipo_sogg == 'Società':
                        if rappresentante:
                            line = f"la società {nome}, rappresentata dal Sig {rappresentante}, socia recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                        else:
                            line = f"la società {nome} socia recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                    else:
                        line = f"il Sig. {nome} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"

                preview += f"     - {line}\n"
        
        if soci_assenti:
            preview += "\n   - risultano invece assenti i seguenti soci:\n"
            for socio in soci_assenti:
                if isinstance(socio, dict) and socio.get('nome'):
                    preview += f"     - il Sig. {socio.get('nome')}\n"
        
        preview += f"""
3. che gli intervenuti sono legittimati alla presente assemblea;
4. che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno.

I presenti all'unanimità chiamano a fungere da segretario il signor **{data.get('segretario', '[SEGRETARIO]')}**, che accetta l'incarico.

Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante.
"""
        
        # Sezione traduzione se necessaria
        if data.get('richiede_traduzione'):
            lingua = data.get('lingua_straniera', 'inglese').lower()
            if lingua == 'altro':
                lingua = data.get('lingua_altro', '[LINGUA]').lower()
            
            preview += f"""
In particolare, preso atto che il Sig. **{data.get('partecipante_straniero', '[PARTECIPANTE]')}** non conosce la lingua italiana ma dichiara di conoscere la lingua {lingua}, il Presidente dichiara che provvederà a tradurre dall'italiano al {lingua} (e viceversa) gli interventi dei partecipanti alla discussione nonché a tradurre dall'italiano al {lingua} il verbale che sarà redatto al termine della riunione.
"""
        
        preview += f"""
Il Presidente constata e fa constatare che l'assemblea risulta {convocazione_phrase} e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.

Si passa quindi allo svolgimento dell'ordine del giorno.

---

*     *     *

Il Presidente relaziona l'assemblea dei soci sulla situazione finanziaria della società e illustra le voci che ne costituiscono il patrimonio netto, anche in relazione alla loro possibilità di distribuzione ai soci.

Prende la parola il socio Sig. **{data.get('socio_proponente', '[SOCIO PROPONENTE]')}** che propone di distribuire {data.get('tipo_distribuzione', 'utili pregressi').lower()} ai soci per l'importo complessivo di euro **{data.get('importo_dividendi', '[IMPORTO]')}**.

Quindi si passa alla votazione con voto palese in forza della quale il Presidente constata che"""
        
        # Gestione esito voto
        tipo_voto = data.get('tipo_voto', 'Unanimità')
        if tipo_voto == "Unanimità":
            preview += ", all'unanimità"
        elif tipo_voto == "Maggioranza con contrari":
            contrari = data.get('voti_contrari', '[CONTRARI]')
            preview += f", con il voto contrario dei Sigg. {contrari}"
        elif tipo_voto == "Maggioranza con astenuti":
            contrari = data.get('voti_contrari', '')
            astenuti = data.get('astensioni', '[ASTENUTI]')
            if contrari:
                preview += f", con il voto contrario dei Sigg. {contrari} e l'astensione dei Sigg. {astenuti}"
            else:
                preview += f", con l'astensione dei Sigg. {astenuti}"
        
        preview += ", l'assemblea\n\n**d e l i b e r a:**\n\n"
        
        # Delibera
        modalita_ripartizione = data.get('modalita_ripartizione', 'Proporzionale alla quota')
        if modalita_ripartizione == "Proporzionale alla quota":
            ripartizione_text = "proporzionalmente alla quota di partecipazione"
        elif modalita_ripartizione == "Secondo atto costitutivo":
            ripartizione_text = "nelle proporzioni previste dall'atto costitutivo"
        else:
            ripartizione_text = "nelle modalità specificate"
        
        preview += f"di procedere alla distribuzione di un dividendo complessivo di euro **{data.get('importo_dividendi', '[IMPORTO]')}**, da ripartirsi fra i soci {ripartizione_text}"
        
        # Eventuale prelievo da riserve
        if data.get('prelievo_riserve'):
            nome_riserva = data.get('nome_riserva', 'fondo di riserva')
            preview += f" mediante prelievo del rispettivo importo dal {nome_riserva} in quanto disponibile"
        
        preview += "."
        
        # Calcolo dividendi per socio se disponibile
        if data.get('calcolo_dividendi'):
            preview += "\n\n**Dettaglio dividendi per socio:**\n"
            for calc in data['calcolo_dividendi']:
                preview += f"- {calc['nome']}: {calc['dividendo']} ({calc['quota_perc']})\n"
        
        preview += f"""

---

## Chiusura dell'Assemblea

Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.

Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo.

L'assemblea viene sciolta alle ore **{data.get('ora_chiusura', '[ORA CHIUSURA]')}**.

---

**Il Presidente**                    **Il Segretario**

_________________                    _________________

{data.get('presidente', '[PRESIDENTE]')}                      {data.get('segretario', '[SEGRETARIO]')}
"""
        
        return preview
    
    def generate_document(self, data: dict) -> Document:
        """Genera il documento Word del verbale di distribuzione dividendi"""
        import os
        
        # Utilizza il template .docx esistente per mantenere la formattazione
        template_path = os.path.join(os.path.dirname(__file__), 'template.docx')
        
        try:
            if os.path.exists(template_path):
                doc = Document(template_path)
                # Rimuovi il contenuto esistente del template mantenendo gli stili
                for paragraph in doc.paragraphs[:]:
                    p = paragraph._element
                    p.getparent().remove(p)
            else:
                doc = Document()
                # Setup stili solo se non usiamo template
                self._setup_document_styles(doc)
        except Exception as e:
            # Fallback a documento vuoto se il template non può essere caricato
            doc = Document()
            self._setup_document_styles(doc)
        
        # Imposta gli stili del documento
        self._setup_document_styles(doc)
        
        # Header aziendale
        self._add_company_header(doc, data)
        
        # Titolo del verbale
        self._add_verbale_title(doc, data)
        
        # Sezione di apertura
        self._add_opening_section(doc, data)
        
        # Sezione partecipanti
        self._add_participants_section(doc, data)
        
        # Dichiarazioni preliminari
        self._add_preliminary_statements(doc, data)
        
        # Discussione distribuzione dividendi
        self._add_dividendi_discussion(doc, data)
        
        # Sezione di chiusura
        self._add_closing_section(doc, data)
        
        # Firme
        self._add_signatures(doc, data)
        
        return doc
    
    def _setup_document_styles(self, doc):
        """Imposta gli stili per il documento"""
        # First call parent's method to setup all base styles
        super()._setup_document_styles(doc)
        
        # Stile per il titolo principale
        title_style = doc.styles.add_style('VerbaleTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Times New Roman'
        title_font.size = Pt(16)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
        
        # Stile per i sottotitoli
        subtitle_style = doc.styles.add_style('VerbaleSubtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_font = subtitle_style.font
        subtitle_font.name = 'Times New Roman'
        subtitle_font.size = Pt(12)
        subtitle_font.bold = True
        subtitle_style.paragraph_format.space_before = Pt(12)
        subtitle_style.paragraph_format.space_after = Pt(6)
        
        # Stile per il testo normale
        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = 1.15
    
    def _add_company_header(self, doc, data):
        """Aggiunge l'header con i dati dell'azienda"""
        # Nome azienda
        company_p = doc.add_paragraph(data.get('denominazione', '[DENOMINAZIONE]'))
        company_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in company_p.runs:
            run.font.bold = True
            run.font.size = Pt(14)
        
        # Dati aziendali
        doc.add_paragraph()
        info_p = doc.add_paragraph()
        info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_p.add_run(f"Sede in {data.get('sede_legale', '[SEDE LEGALE]')}\n")
        info_p.add_run(f"Capitale sociale Euro {self._get_capitale_sociale(data)} i.v.\n")
        info_p.add_run(f"Codice fiscale: {data.get('codice_fiscale', '[CODICE FISCALE]')}")
        
        doc.add_paragraph()
    
    def _add_verbale_title(self, doc, data):
        """Aggiunge il titolo del verbale"""
        title_p = doc.add_paragraph("Verbale di assemblea dei soci", style='VerbaleTitle')
        data_p = doc.add_paragraph(f"del {self._format_date(data.get('data_assemblea', '[DATA]'))}", style='VerbaleTitle')
        doc.add_paragraph()
    
    def _add_opening_section(self, doc, data):
        """Aggiunge la sezione di apertura"""
        opening_text = f"""Oggi {self._format_date(data.get('data_assemblea', '[DATA]'))} alle ore {self._format_time(data.get('ora_inizio', data.get('ora_assemblea', '[ORA]')))} presso la sede sociale {data.get('sede_legale', '[SEDE]')}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:"""
        
        doc.add_paragraph(opening_text)
        doc.add_paragraph()
        
        # Ordine del giorno
        odg_title = doc.add_paragraph("Ordine del giorno", style='VerbaleSubtitle')
        odg_p = doc.add_paragraph(data.get('ordine_giorno', 'Proposta di distribuzione dei dividendi'))
        doc.add_paragraph()
    
    def _add_participants_section(self, doc, data):
        """Aggiunge la sezione dei partecipanti"""
        presidente_text = f"""Assume la presidenza ai sensi dell'art. [...] dello statuto sociale il Sig. {data.get('presidente', '[PRESIDENTE]')} {data.get('ruolo_presidente', 'Amministratore Unico')}, il quale dichiara e constata:"""
        
        doc.add_paragraph(presidente_text)
        doc.add_paragraph()
        
        # Dichiarazioni del presidente
        doc.add_paragraph("1 - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. [...] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza")
        doc.add_paragraph()
        
        participants_p = doc.add_paragraph("2 - che sono presenti/partecipano all'assemblea:")
        participants_p.add_run(f"\nl'{data.get('ruolo_presidente', 'Amministratore Unico')} nella persona del suddetto Presidente Sig. {data.get('presidente', '[PRESIDENTE]')}")
        
        # Aggiungi Collegio Sindacale se presente
        if data.get('include_collegio_sindacale', False):
            tipo_organo_controllo = data.get('tipo_organo_controllo', 'Collegio Sindacale')
            
            if tipo_organo_controllo == 'Collegio Sindacale':
                sindaci = data.get('sindaci', [])
                sindaci_presenti = [s for s in sindaci if s.get('presente')]
                if sindaci_presenti:
                    participants_p.add_run("\nper il Collegio Sindacale:")
                    for sindaco in sindaci_presenti:
                        carica = sindaco.get('carica', 'Sindaco Effettivo')
                        nome_sindaco = sindaco.get('nome', '')
                        if nome_sindaco:
                            participants_p.add_run(f"\nil Dott. {nome_sindaco} - {carica}")
            else: # Sindaco Unico
                sindaci = data.get('sindaci', [])
                if sindaci and sindaci[0].get('nome'):
                    sindaco_unico_nome = sindaci[0].get('nome')
                    participants_p.add_run(f"\nil Sindaco Unico nella persona del Sig. {sindaco_unico_nome}")

        # Aggiungi soci
        soci_presenti = data.get('soci_presenti', [])
        soci_assenti = data.get('soci_assenti', [])
        
        # Fallback per mantenere compatibilità
        if not soci_presenti and not soci_assenti and 'soci' in data:
            soci_presenti = [s for s in data.get('soci', []) if s.get('presente', True)]
            soci_assenti = [s for s in data.get('soci', []) if not s.get('presente', True)]

        if soci_presenti:
            # Calcola totali euro e percentuale
            totale_euro = 0.0
            totale_perc = 0.0
            capitale_raw = str(data.get('capitale_sociale', '0')).replace('.', '').replace(',', '.')
            try:
                capitale_float = float(capitale_raw)
            except ValueError:
                capitale_float = 0.0

            for socio in soci_presenti:
                if isinstance(socio, dict):
                    euro_raw = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                    try:
                        totale_euro += float(euro_raw or 0)
                    except ValueError:
                        pass
                    perc_raw = socio.get('quota_percentuale', '')
                    if perc_raw:
                        try:
                            totale_perc += float(str(perc_raw).replace('%', '').replace(',', '.'))
                        except ValueError:
                            pass
                    else:
                        try:
                            euro_val = float(euro_raw or 0)
                            if capitale_float:
                                totale_perc += euro_val / capitale_float * 100
                        except ValueError:
                            pass

            formatted_euro = CommonDataHandler.format_currency(totale_euro)
            formatted_perc = CommonDataHandler.format_percentage(totale_perc)

            participants_p.add_run(f"\nnonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_euro} pari al {formatted_perc} del Capitale Sociale:")

            for socio in soci_presenti:
                if not isinstance(socio, dict) or not socio.get('nome'):
                    continue
                nome = socio.get('nome')
                quota_euro = CommonDataHandler.format_currency(socio.get('quota_euro', '0'))
                quota_perc = socio.get('quota_percentuale', '')
                if not quota_perc:
                    try:
                        euro_val = float(str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.'))
                        quota_perc = CommonDataHandler.format_percentage(euro_val / capitale_float * 100) if capitale_float else '[%]'
                    except ValueError:
                        quota_perc = '[%]'
                else:
                    quota_perc = CommonDataHandler.clean_percentage(quota_perc)

                tipo_sogg = socio.get('tipo_soggetto', 'Persona Fisica')
                tipo_part = socio.get('tipo_partecipazione', 'Diretto')
                delegato = socio.get('delegato', '').strip()
                rappresentante = socio.get('rappresentante_legale', '').strip()

                if tipo_part == 'Delegato' and delegato:
                    if tipo_sogg == 'Società':
                        line = f"il Sig. {delegato} delegato della società {nome} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                    else:
                        line = f"il Sig. {delegato} delegato del socio Sig. {nome} recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                else:
                    if tipo_sogg == 'Società':
                        if rappresentante:
                            line = f"la società {nome}, rappresentata dal Sig {rappresentante}, socia recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                        else:
                            line = f"la società {nome} socia recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                    else:
                        line = f"il Sig. {nome} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"

                participants_p.add_run(f"\n{line}")
        
        if soci_assenti:
            assenti_p = doc.add_paragraph("   - risultano invece assenti i seguenti soci:")
            for socio in soci_assenti:
                if isinstance(socio, dict) and socio.get('nome'):
                    assenti_p.add_run(f"\n     - il Sig. {socio.get('nome')}")

        doc.add_paragraph()
    
    def _add_preliminary_statements(self, doc, data):
        """Aggiunge le dichiarazioni preliminari"""
        doc.add_paragraph("3 - che gli intervenuti sono legittimati alla presente assemblea;")
        doc.add_paragraph("4 - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno.")
        doc.add_paragraph()
        
        segretario_text = f"I presenti all'unanimità chiamano a fungere da segretario il signor {data.get('segretario', '[SEGRETARIO]')}, che accetta l'incarico."
        doc.add_paragraph(segretario_text)
        doc.add_paragraph()
        
        identification_text = "Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante."
        doc.add_paragraph(identification_text)
        
        # Sezione traduzione se necessaria
        if data.get('richiede_traduzione'):
            lingua = data.get('lingua_straniera', 'inglese').lower()
            if lingua == 'altro':
                lingua = data.get('lingua_altro', '[LINGUA]').lower()
            
            translation_text = f"In particolare, preso atto che il Sig. {data.get('partecipante_straniero', '[PARTECIPANTE]')} non conosce la lingua italiana ma dichiara di conoscere la lingua {lingua}, il Presidente dichiara che provvederà a tradurre dall'italiano al {lingua} (e viceversa) gli interventi dei partecipanti alla discussione nonché a tradurre dall'italiano al {lingua} il verbale che sarà redatto al termine della riunione."
            doc.add_paragraph(translation_text)
        
        doc.add_paragraph()
        convocazione_phrase = "totalitaria" if "totalitaria" in str(data.get("tipo_convocazione", "")).lower() else "regolarmente convocata"
        validity_text = f"Il Presidente constata e fa constatare che l'assemblea risulta {convocazione_phrase} e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno."
        doc.add_paragraph(validity_text)
        doc.add_paragraph()
        
        doc.add_paragraph("Si passa quindi allo svolgimento dell'ordine del giorno.")
        doc.add_paragraph("*     *     *")
        doc.add_paragraph()
    
    def _add_dividendi_discussion(self, doc, data):
        """Aggiunge la discussione della distribuzione dividendi"""
        # Relazione del presidente
        president_report = "Il Presidente relaziona l'assemblea dei soci sulla situazione finanziaria della società e illustra le voci che ne costituiscono il patrimonio netto, anche in relazione alla loro possibilità di distribuzione ai soci."
        doc.add_paragraph(president_report)
        doc.add_paragraph()
        
        # Proposta del socio
        proposal_text = f"Prende la parola il socio Sig. {data.get('socio_proponente', '[SOCIO PROPONENTE]')} che propone di distribuire {data.get('tipo_distribuzione', 'utili pregressi').lower()} ai soci per l'importo complessivo di euro {data.get('importo_dividendi', '[IMPORTO]')}."
        doc.add_paragraph(proposal_text)
        doc.add_paragraph()
        
        # Votazione
        vote_intro = "Quindi si passa alla votazione con voto palese in forza della quale il Presidente constata che"
        
        # Gestione esito voto
        tipo_voto = data.get('tipo_voto', 'Unanimità')
        if tipo_voto == "Unanimità":
            vote_intro += ", all'unanimità"
        elif tipo_voto == "Maggioranza con contrari":
            contrari = data.get('voti_contrari', '[CONTRARI]')
            vote_intro += f", con il voto contrario dei Sigg. {contrari}"
        elif tipo_voto == "Maggioranza con astenuti":
            contrari = data.get('voti_contrari', '')
            astenuti = data.get('astensioni', '[ASTENUTI]')
            if contrari:
                vote_intro += f", con il voto contrario dei Sigg. {contrari} e l'astensione dei Sigg. {astenuti}"
            else:
                vote_intro += f", con l'astensione dei Sigg. {astenuti}"
        
        vote_intro += ", l'assemblea"
        doc.add_paragraph(vote_intro)
        doc.add_paragraph()
        
        # Delibera
        delibera_title = doc.add_paragraph("d e l i b e r a:")
        for run in delibera_title.runs:
            run.font.bold = True
        doc.add_paragraph()
        
        # Contenuto delibera
        modalita_ripartizione = data.get('modalita_ripartizione', 'Proporzionale alla quota')
        if modalita_ripartizione == "Proporzionale alla quota":
            ripartizione_text = "proporzionalmente alla quota di partecipazione"
        elif modalita_ripartizione == "Secondo atto costitutivo":
            ripartizione_text = "nelle proporzioni previste dall'atto costitutivo"
        else:
            ripartizione_text = "nelle modalità specificate"
        
        delibera_content = f"di procedere alla distribuzione di un dividendo complessivo di euro {data.get('importo_dividendi', '[IMPORTO]')}, da ripartirsi fra i soci {ripartizione_text}"
        
        # Eventuale prelievo da riserve
        if data.get('prelievo_riserve'):
            nome_riserva = data.get('nome_riserva', 'fondo di riserva')
            delibera_content += f" mediante prelievo del rispettivo importo dal {nome_riserva} in quanto disponibile"
        
        delibera_content += "."
        doc.add_paragraph(delibera_content)
        doc.add_paragraph()
        doc.add_paragraph("*     *     *")
        doc.add_paragraph()
    
    def _add_closing_section(self, doc, data):
        """Aggiunge la sezione di chiusura"""
        closing_text = "Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola."
        doc.add_paragraph(closing_text)
        doc.add_paragraph()
        
        final_approval_text = "Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo."
        doc.add_paragraph(final_approval_text)
        doc.add_paragraph()
        
        end_time_text = f"L'assemblea viene sciolta alle ore {data.get('ora_chiusura', '[ORA CHIUSURA]')}."
        doc.add_paragraph(end_time_text)
        doc.add_paragraph()
    
    def _add_signatures(self, doc, data):
        """Aggiunge le firme"""
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Tabella per le firme usando il metodo sicuro
        signature_table = self._create_table_with_style(doc, rows=3, cols=2)
        
        # Prima riga - titoli
        signature_table.cell(0, 0).text = "Il Presidente"
        signature_table.cell(0, 1).text = "Il Segretario"
        
        # Seconda riga - spazio per firme
        signature_table.cell(1, 0).text = "_________________"
        signature_table.cell(1, 1).text = "_________________"
        
        # Terza riga - nomi
        signature_table.cell(2, 0).text = data.get('presidente', '[PRESIDENTE]')
        signature_table.cell(2, 1).text = data.get('segretario', '[SEGRETARIO]')
        
        # Centra la tabella
        for row in signature_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Registra il template
DocumentTemplateFactory.register_template("dividendi", VerbaleDividendiTemplate)