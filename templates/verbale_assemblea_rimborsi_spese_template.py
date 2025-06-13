"""
Template per Verbale di Assemblea - Riconoscimento di rimborsi spese all'organo amministrativo e autorizzazione all'utilizzo di veicoli personali
"""

import sys
import os

# Aggiungi il path della cartella src
current_dir = os.path.dirname(os.path.abspath(__file__))
src_path = os.path.join(os.path.dirname(current_dir), 'src')
if src_path not in sys.path:
    sys.path.append(src_path)

from document_templates import DocumentTemplate, DocumentTemplateFactory
from common_data_handler import CommonDataHandler
from base_verbale_template import BaseVerbaleTemplate
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date
import streamlit as st

class VerbaleRimborsiSpeseTemplate(BaseVerbaleTemplate):
    """Template per Verbale di Assemblea - Riconoscimento rimborsi spese organo amministrativo"""
    
    def get_template_name(self) -> str:
        return "Riconoscimento rimborsi spese all'organo amministrativo"
    
    def get_required_fields(self) -> list:
        return [
            "denominazione", "sede_legale", "capitale_sociale", "codice_fiscale",
            "data_assemblea", "ora_assemblea", "presidente", "segretario", 
            "soci", "amministratori_beneficiari"
        ]
    
    def get_form_fields(self, extracted_data: dict) -> dict:
        """Genera i campi del form per Streamlit"""
        form_data = {}
        
        # Dati azienda standardizzati
        form_data.update(CommonDataHandler.extract_and_populate_company_data(extracted_data))
        
        # Dati assemblea standardizzati  
        form_data.update(CommonDataHandler.extract_and_populate_assembly_data(extracted_data))
        
        # Configurazioni specifiche
        st.subheader("💰 Configurazioni Specifiche Rimborsi Spese")
        
        col1, col2 = st.columns(2)
        with col1:
            form_data["ruolo_presidente"] = st.selectbox("Ruolo del presidente", 
                                                        ["Amministratore Unico", 
                                                         "Presidente del Consiglio di Amministrazione",
                                                         "Altro (come da statuto)"])
            form_data["tipo_organo"] = st.selectbox("Tipo di organo amministrativo", 
                                                   ["Amministratore Unico", 
                                                    "Consiglio di Amministrazione"])
        with col2:
            form_data["include_collegio_sindacale"] = st.checkbox("Include Collegio Sindacale", value=False)
            form_data["include_revisore"] = st.checkbox("Include revisore contabile", value=False)
        
        # Partecipanti standardizzati
        participants_data = CommonDataHandler.extract_and_populate_participants_data(
            extracted_data, 
            unique_key_suffix="rimborsi_spese"
        )
        form_data.update(participants_data)
        
        # Amministratori beneficiari
        st.subheader("👥 Amministratori Beneficiari dei Rimborsi")
        
        if form_data["tipo_organo"] == "Amministratore Unico":
            num_amministratori = 1
            st.info("Amministratore Unico")
        else:
            num_amministratori = st.number_input("Numero membri CdA", min_value=1, max_value=10, value=3)
        
        amministratori_beneficiari = []
        for i in range(num_amministratori):
            if num_amministratori == 1:
                st.write("**Amministratore Unico**")
            else:
                ruolo = "Presidente" if i == 0 else f"Consigliere {i}"
                st.write(f"**{ruolo}**")
            
            col1, col2 = st.columns(2)
            with col1:
                nome = st.text_input(f"Nome e Cognome", key=f"amm_nome_{i}",
                                   placeholder="es. Sig. Mario Rossi")
            with col2:
                titolo = st.text_input(f"Titolo", key=f"amm_titolo_{i}",
                                     placeholder="es. Sig., Dott., Ing.",
                                     value="Sig.")
            
            if nome:
                amministratori_beneficiari.append({
                    "nome": nome,
                    "titolo": titolo,
                    "ruolo": "Amministratore Unico" if num_amministratori == 1 else ("Presidente" if i == 0 else "Consigliere")
                })
        
        form_data["amministratori_beneficiari"] = amministratori_beneficiari
        
        # Configurazioni rimborsi
        st.subheader("🚗 Configurazioni Rimborsi")
        
        col1, col2 = st.columns(2)
        with col1:
            form_data["include_veicoli_personali"] = st.checkbox("Autorizza uso veicoli personali", value=True)
            form_data["include_tariffe_aci"] = st.checkbox("Usa tariffe ACI", value=True)
        with col2:
            form_data["include_statuto_riferimento"] = st.checkbox("Riferimento articolo statuto", value=True)
            if form_data["include_statuto_riferimento"]:
                form_data["articolo_statuto"] = st.text_input("Articolo statuto", 
                                                             placeholder="es. art. 15")
        
        # Tipi di rimborso
        st.subheader("💵 Tipi di Rimborso")
        
        form_data["tipo_rimborso_trasferte"] = st.selectbox(
            "Tipo rimborso trasferte",
            ["Analitico (a piè di lista)", "Forfettario", "Misto"]
        )
        
        # Importi forfettari (valori standard 2024)
        st.subheader("📊 Importi Forfettari Standard")
        
        col1, col2 = st.columns(2)
        with col1:
            st.write("**Italia**")
            form_data["forfait_completo_italia"] = st.number_input("Forfait completo Italia", value=46.48, step=0.01)
            form_data["forfait_parziale_italia"] = st.number_input("Forfait parziale Italia", value=30.99, step=0.01)
            form_data["forfait_minimo_italia"] = st.number_input("Forfait minimo Italia", value=15.49, step=0.01)
        
        with col2:
            st.write("**Estero**")
            form_data["forfait_completo_estero"] = st.number_input("Forfait completo Estero", value=77.47, step=0.01)
            form_data["forfait_parziale_estero"] = st.number_input("Forfait parziale Estero", value=51.65, step=0.01)
            form_data["forfait_minimo_estero"] = st.number_input("Forfait minimo Estero", value=25.82, step=0.01)
        
        # Esclusioni e limitazioni
        st.subheader("⚠️ Esclusioni e Limitazioni")
        
        form_data["escludi_domicilio_sede"] = st.checkbox("Escludi trasferimenti domicilio-sede", value=True)
        form_data["limite_deducibilita_fiscale"] = st.checkbox("Limite deducibilità fiscale", value=True)
        
        # Documentazione richiesta
        st.subheader("📄 Documentazione Richiesta")
        form_data["richiedi_documentazione"] = st.checkbox("Richiedi documentazione per rimborsi", value=True)
        
        # Votazione
        st.subheader("🗳️ Configurazioni Votazione")
        
        col1, col2 = st.columns(2)
        with col1:
            form_data["votazione_unanime"] = st.checkbox("Votazione unanime", value=True)
            if not form_data["votazione_unanime"]:
                form_data["voti_contrari"] = st.text_input("Voti contrari", 
                                                          placeholder="es. Sig. Rossi, Sig. Verdi")
                form_data["astensioni"] = st.text_input("Astensioni", 
                                                       placeholder="es. Sig. Neri")
        
        with col2:
            form_data["include_traduzione"] = st.checkbox("Include sezione traduzione", value=False)
            form_data["include_presenti_qualita"] = st.checkbox("Include sezione 'presenti in qualità di'", value=False)
        
        # Traduzione se necessaria
        if form_data["include_traduzione"]:
            st.subheader("🌐 Traduzione")
            col1, col2 = st.columns(2)
            with col1:
                form_data["persona_traduzione"] = st.text_input("Persona che necessita traduzione")
            with col2:
                form_data["lingua_traduzione"] = st.selectbox("Lingua", 
                                                             ["Inglese", "Francese", "Tedesco", "Spagnolo"])
        
        # Note aggiuntive
        st.subheader("📝 Note Aggiuntive")
        form_data["note_aggiuntive"] = st.text_area("Note aggiuntive", 
                                                  placeholder="Eventuali note specifiche sui rimborsi...",
                                                  height=80)
        
        return form_data
    
    def show_preview(self, form_data: dict):
        """Mostra l'anteprima del documento"""
        st.markdown("---")
        st.subheader("👁️ Anteprima Documento")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col1:
            show_preview = st.checkbox("Mostra anteprima", value=False, key="preview_checkbox_rimborsi")
        
        with col2:
            if show_preview:
                st.info("💡 L'anteprima si aggiorna automaticamente")
        
        if show_preview:
            try:
                preview_text = self._generate_preview_text(form_data)
                st.text(preview_text)
            except Exception as e:
                st.error(f"Errore nell'anteprima: {e}")
    
    def _generate_preview_text(self, data: dict) -> str:
        """Genera il testo di anteprima"""
        try:
            # Header
            amministratori = data.get('amministratori_beneficiari', [])
            nomi_amm = ', '.join([a.get('nome', '') for a in amministratori if a.get('nome')])
            
            header = f"""{data.get('denominazione', '[Denominazione]')}
Sede in {data.get('sede_legale', '[Sede]')}
Capitale sociale Euro {data.get('capitale_sociale', '[Capitale]')} i.v.
Codice fiscale: {data.get('codice_fiscale', '[CF]')}

Verbale di assemblea dei soci
del {data.get('data_assemblea', '[Data]').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[Data]'}

Ordine del giorno
Riconoscimento di rimborsi spese all'organo amministrativo e autorizzazione all'utilizzo di veicoli personali

Il Presidente propone all'assemblea dei soci di deliberare in merito ai rimborsi spese spettanti ai componenti l'organo amministrativo della società.

d e l i b e r a:

di riconoscere agli Amministratori in carica Sigg. {nomi_amm} il rimborso delle spese incontrate per ragione del proprio ufficio.

Gli Amministratori sono autorizzati all'utilizzo dei propri veicoli ai fini aziendali."""
            
            return header
            
        except Exception as e:
            return f"Errore nella generazione dell'anteprima: {str(e)}"
    
    def generate_document(self, data: dict) -> Document:
        """Genera il documento Word"""
        doc = Document()
        
        # Setup stili
        self._setup_document_styles(doc)
        
        # Header azienda
        self._add_company_header(doc, data)
        
        # Titolo verbale
        self._add_verbale_title(doc, data)
        
        # Sezioni del verbale
        self._add_opening_section(doc, data)
        self._add_participants_section(doc, data)
        self._add_preliminary_statements(doc, data)
        self._add_rimborsi_discussion(doc, data)
        self._add_closing_section(doc, data)
        self._add_signatures(doc, data)
        
        return doc
    
    def _setup_document_styles(self, doc):
        """Imposta gli stili del documento"""
        styles = doc.styles
        
        if 'Titolo Principale' not in [s.name for s in styles]:
            title_style = styles.add_style('Titolo Principale', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Times New Roman'
            title_style.font.size = Pt(14)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        normal_style = styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal_style.paragraph_format.line_spacing = 1.15
    
    def _add_company_header(self, doc, data):
        """Aggiunge header azienda"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data.get('denominazione', '[DENOMINAZIONE]'))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Sede in {data.get('sede_legale', '[SEDE]')}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Capitale sociale Euro {data.get('capitale_sociale', '[CAPITALE]')} i.v.")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Codice fiscale: {data.get('codice_fiscale', '[CF]')}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        doc.add_paragraph()
    
    def _add_verbale_title(self, doc, data):
        """Aggiunge titolo verbale"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Verbale di assemblea dei soci")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        
        data_assemblea = data.get('data_assemblea')
        if hasattr(data_assemblea, 'strftime'):
            data_str = data_assemblea.strftime('%d/%m/%Y')
        else:
            data_str = '[DATA]'
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"del {data_str}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        doc.add_paragraph()
    
    def _add_opening_section(self, doc, data):
        """Aggiunge sezione apertura"""
        data_assemblea = data.get('data_assemblea')
        ora_assemblea = data.get('ora_assemblea')
        
        if hasattr(data_assemblea, 'strftime'):
            data_str = data_assemblea.strftime('%d/%m/%Y')
        else:
            data_str = '[DATA]'
        
        if hasattr(ora_assemblea, 'strftime'):
            ora_str = ora_assemblea.strftime('%H:%M')
        else:
            ora_str = '[ORA]'
        
        text = f"Oggi {data_str} alle ore {ora_str} presso la sede sociale {data.get('sede_legale', '[SEDE]')}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:"
        
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Ordine del giorno")
        run.font.bold = True
        
        p = doc.add_paragraph("Riconoscimento di rimborsi spese all'organo amministrativo e autorizzazione all'utilizzo di veicoli personali")
    
    def _add_participants_section(self, doc, data):
        """Aggiunge sezione partecipanti"""
        doc.add_paragraph()
        
        ruolo_presidente = data.get('ruolo_presidente', 'Amministratore Unico')
        presidente = data.get('presidente', '[PRESIDENTE]')
        
        text = f"Assume la presidenza ai sensi dell'art. […] dello statuto sociale il Sig. {presidente} {ruolo_presidente}, il quale dichiara e constata:"
        
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        p = doc.add_paragraph("1 - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. […] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza")
        
        p = doc.add_paragraph("2 - che sono presenti/partecipano all'assemblea:")
        
        # Organo amministrativo
        if data.get('tipo_organo') == "Amministratore Unico":
            p = doc.add_paragraph(f"l'Amministratore Unico nella persona del suddetto Presidente Sig. {presidente}")
        else:
            p = doc.add_paragraph("[oppure")
            p = doc.add_paragraph("per il Consiglio di Amministrazione:")
            amministratori = data.get('amministratori_beneficiari', [])
            for amm in amministratori:
                nome = amm.get('nome', '[Nome]')
                p = doc.add_paragraph(f"il Sig. {nome}")
            p = doc.add_paragraph("assente giustificato il Sig. […] il quale ha tuttavia rilasciato apposita dichiarazione scritta, conservata agli atti della Società, dalla quale risulta che il medesimo è stato informato su tutti gli argomenti posti all'ordine del giorno e che lo stesso non si oppone alla trattazione degli stessi]")
        
        # Collegio sindacale se presente
        if data.get('include_collegio_sindacale', False):
            p = doc.add_paragraph("[eventualmente")
            p = doc.add_paragraph("per il Collegio Sindacale")
            p = doc.add_paragraph("il Dott. […]")
            p = doc.add_paragraph("il Dott. […]")
            p = doc.add_paragraph("il Dott. […]")
            p = doc.add_paragraph("[oppure")
            p = doc.add_paragraph("il Sindaco Unico nella persona del Sig. […]]]")
        
        # Revisore se presente
        if data.get('include_revisore', False):
            p = doc.add_paragraph("[eventualmente, se invitato")
            p = doc.add_paragraph("il revisore contabile Dott. […] <oppure il dott. […] in rappresentanza della società di revisione incaricata del controllo contabile>")
        
        # Soci
        soci = data.get('soci', [])
        totale_quote_euro = 0.0
        totale_quote_perc = 0.0

        for socio in soci:
            if isinstance(socio, dict):
                try:
                    quota_euro_str = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                    totale_quote_euro += float(quota_euro_str)
                except ValueError:
                    pass  # Ignora valori non numerici
                try:
                    quota_perc_str = str(socio.get('quota_percentuale', '0')).replace('.', '').replace(',', '.')
                    totale_quote_perc += float(quota_perc_str)
                except ValueError:
                    pass  # Ignora valori non numerici

        # Formattazione per l'italiano
        totale_quote_euro_formatted = f"{totale_quote_euro:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        totale_quote_perc_formatted = f"{totale_quote_perc:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

        p = doc.add_paragraph(f"nonché i seguenti soci o loro rappresentanti, [eventualmente così come iscritti a libro soci e] recanti complessivamente una quota pari a nominali euro {totale_quote_euro_formatted} pari al {totale_quote_perc_formatted}% del Capitale Sociale:")
        
        for socio in soci:
            if isinstance(socio, dict):
                nome = socio.get('nome', '[Nome Socio]')
                quota_value = socio.get('quota_euro', '')
                percentuale_value = socio.get('quota_percentuale', '')
                
                # Gestione robusta dei valori nulli o vuoti
                quota = '[Quota]' if quota_value is None or str(quota_value).strip() == '' else str(quota_value).strip()
                percentuale = '[%]' if percentuale_value is None or str(percentuale_value).strip() == '' else str(percentuale_value).strip()
                p = doc.add_paragraph(f"il Sig. {nome} socio [oppure delegato del socio Sig. […]] recante una quota pari a nominali euro {quota} pari al {percentuale}% del Capitale Sociale")
        
        p = doc.add_paragraph("2 - che gli intervenuti sono legittimati alla presente assemblea;")
        p = doc.add_paragraph("3 - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno.")
        
        # Presenti in qualità di
        if data.get('include_presenti_qualita', False):
            p = doc.add_paragraph("[eventualmente")
            p = doc.add_paragraph("4 - che sono altresì presenti, in qualità di […]:")
            p = doc.add_paragraph("il Sig. […]")
            p = doc.add_paragraph("il Sig. […]]")
    
    def _add_preliminary_statements(self, doc, data):
        """Aggiunge dichiarazioni preliminari"""
        doc.add_paragraph()
        
        segretario = data.get('segretario', '[SEGRETARIO]')
        p = doc.add_paragraph(f"I presenti all'unanimità chiamano a fungere da segretario il signor {segretario}, che accetta l'incarico.")
        
        p = doc.add_paragraph("Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante.")
        
        # Traduzione se necessaria
        if data.get('include_traduzione', False):
            persona_traduzione = data.get('persona_traduzione', '[Nome]')
            lingua = data.get('lingua_traduzione', 'inglese').lower()
            p = doc.add_paragraph(f"[eventualmente")
            p = doc.add_paragraph(f"In particolare, preso atto che il Sig. {persona_traduzione} non conosce la lingua italiana ma dichiara di conoscere la lingua {lingua}, il Presidente dichiara che provvederà a tradurre dall'italiano al {lingua} (e viceversa) gli interventi dei partecipanti alla discussione nonché a tradurre dall'italiano al {lingua} il verbale che sarà redatto al termine della riunione.]")
        
        p = doc.add_paragraph("Il Presidente constata e fa constatare che l'assemblea risulta regolarmente convocata [oppure totalitaria] e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.")
        
        p = doc.add_paragraph("Si passa quindi allo svolgimento dell'ordine del giorno.")
        
        # Separatore
        p = doc.add_paragraph("*     *     *")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_rimborsi_discussion(self, doc, data):
        """Aggiunge discussione rimborsi"""
        doc.add_paragraph()
        
        p = doc.add_paragraph("Il Presidente propone all'assemblea dei soci di deliberare in merito ai rimborsi spese spettanti ai componenti l'organo amministrativo della società, con specifico riferimento anche all'eventualità che gli stessi utilizzino, per assolvere al loro mandato, anche veicoli di loro proprietà.")
        
        p = doc.add_paragraph("Segue breve discussione tra i soci al termine della quale si passa alla votazione con voto palese in forza della quale il Presidente constata che, all'unanimità [oppure con il voto contrario dei Sigg. […] e <eventualmente l'astensione dei Sigg. […]>], l'assemblea")
        
        p = doc.add_paragraph("delibera")
        run = p.runs[0]
        run.font.bold = True
        run.font.underline = True
        
        # Deliberazione principale
        amministratori = data.get('amministratori_beneficiari', [])
        nomi_amm = ', '.join([a.get('nome', '') for a in amministratori if a.get('nome')])
        
        statuto_ref = ""
        if data.get('include_statuto_riferimento', False):
            art_statuto = data.get('articolo_statuto', '[…]')
            statuto_ref = f", così come peraltro già espressamente previsto dall'{art_statuto} dello statuto sociale vigente"
        
        p = doc.add_paragraph(f"di riconoscere agli Amministratori in carica Sigg. {nomi_amm} il rimborso delle spese incontrate per ragione del proprio ufficio{statuto_ref}.")
        
        # Autorizzazione veicoli personali
        if data.get('include_veicoli_personali', True):
            p = doc.add_paragraph(f"In particolare gli Amministratori Sigg. {nomi_amm} sono autorizzati all'utilizzo dei propri veicoli ai fini aziendali.")
            
            if data.get('include_tariffe_aci', True):
                p = doc.add_paragraph("Il rimborso delle spese dei viaggi effettuati con il proprio veicolo sarà erogato sulla base del costo di percorrenza chilometrico desumibile dalle tariffe ACI, nei limiti del costo fiscalmente deducibile dal reddito generato dalla società.")
            
            if data.get('escludi_domicilio_sede', True):
                p = doc.add_paragraph("I costi per i trasferimenti dal proprio domicilio alla sede sociale non saranno rimborsati.")
        
        # Tipo di rimborso
        tipo_rimborso = data.get('tipo_rimborso_trasferte', 'analitico')
        p = doc.add_paragraph(f"Il rimborso delle spese per trasferte fuori del Comune sede di lavoro potrà essere {tipo_rimborso.lower()}.")
        
        p = doc.add_paragraph("I medesimi amministratori hanno diritto al rimborso analitico delle spese sostenute durante la trasferta, in ragione del proprio ufficio.")
        
        p = doc.add_paragraph("Inoltre, per le trasferte fuori dal comune in cui ha sede l'attività:")
        
        # Importi forfettari
        forfait_completo_ita = data.get('forfait_completo_italia', 46.48)
        forfait_completo_est = data.get('forfait_completo_estero', 77.47)
        forfait_parziale_ita = data.get('forfait_parziale_italia', 30.99)
        forfait_parziale_est = data.get('forfait_parziale_estero', 51.65)
        forfait_minimo_ita = data.get('forfait_minimo_italia', 15.49)
        forfait_minimo_est = data.get('forfait_minimo_estero', 25.82)
        
        p = doc.add_paragraph(f"in caso di rimborso non documentato delle spese di vitto e di alloggio spetterà all'Amministratore un importo forfettario giornaliero di Euro {forfait_completo_ita:.2f} per trasferte in Italia e di Euro {forfait_completo_est:.2f} per trasferte all'Estero;")
        
        p = doc.add_paragraph(f"in caso di rimborso analitico delle sole spese di vitto o di alloggio spetterà all'Amministratore un importo forfettario giornaliero di Euro {forfait_parziale_ita:.2f} per trasferte in Italia e di Euro {forfait_parziale_est:.2f} per trasferte all'Estero;")
        
        p = doc.add_paragraph(f"in caso di rimborso analitico delle spese di vitto e di alloggio spetterà all'Amministratore un importo forfettario giornaliero di Euro {forfait_minimo_ita:.2f} per trasferte in Italia e di Euro {forfait_minimo_est:.2f} per trasferte all'Estero.")
        
        # Documentazione richiesta
        if data.get('richiedi_documentazione', True):
            p = doc.add_paragraph("Il rimborso delle spese incontrate per ragione del proprio ufficio è subordinato alla presentazione di una nota riepilogativa corredata da idonea documentazione.")
        
        # Separatore
        p = doc.add_paragraph("*     *     *")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_closing_section(self, doc, data):
        """Aggiunge sezione chiusura"""
        doc.add_paragraph()
        
        p = doc.add_paragraph("Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.")
        
        p = doc.add_paragraph("Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo [eventualmente unitamente a quanto allegato].")
        
        ora_chiusura = data.get('ora_chiusura', '[ORA]')
        p = doc.add_paragraph(f"L'assemblea viene sciolta alle ore {ora_chiusura}.")
    
    def _add_signatures(self, doc, data):
        """Aggiunge firme"""
        # Usa la tabella di firme standardizzata
        table = self._add_signature_table(doc, data)
        
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Registra il template
DocumentTemplateFactory.register_template('verbale_assemblea_rimborsi_spese', VerbaleRimborsiSpeseTemplate)