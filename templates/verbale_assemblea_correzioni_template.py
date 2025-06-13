"""
Template per Verbale di Assemblea - Correzioni/Precisazioni
Questo template è specifico per verbali di correzione/precisazione di verbali precedenti.
"""

import sys
import os

# Aggiungi il path della cartella src (relativo alla root del progetto)
current_dir = os.path.dirname(os.path.abspath(__file__))
src_path = os.path.join(os.path.dirname(current_dir), 'src')
if src_path not in sys.path:
    sys.path.append(src_path)

from document_templates import DocumentTemplate, DocumentTemplateFactory
from common_data_handler import CommonDataHandler
from base_verbale_template import BaseVerbaleTemplate
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date
import streamlit as st
import pandas as pd
import re

class VerbaleCorrezioniTemplate(BaseVerbaleTemplate):
    """Template per Verbale di Assemblea dei Soci - Correzioni/Precisazioni"""
    
    def get_template_name(self) -> str:
        return "Correzioni/Precisazioni Verbale"
    
    def get_required_fields(self) -> list:
        return [
            "denominazione", "sede_legale", "capitale_sociale", "codice_fiscale",
            "data_assemblea", "ora_assemblea", "presidente", "segretario", 
            "soci", "amministratori", "data_verbale_precedente", "correzioni"
        ]
    
    def get_form_fields(self, extracted_data: dict) -> dict:
        """Genera i campi del form per Streamlit utilizzando il CommonDataHandler"""
        form_data = {}
        
        # Utilizza il CommonDataHandler per i dati standard
        # Dati azienda standardizzati
        form_data.update(CommonDataHandler.extract_and_populate_company_data(extracted_data))
        
        # Dati assemblea standardizzati
        form_data.update(CommonDataHandler.extract_and_populate_assembly_data(extracted_data))
        
        # Campi specifici per correzioni/precisazioni
        st.subheader("📋 Configurazioni Specifiche Correzioni/Precisazioni")
        
        col1, col2 = st.columns(2)
        with col1:
            form_data["ruolo_presidente"] = st.selectbox("Ruolo del presidente", 
                                                        CommonDataHandler.get_standard_ruoli_presidente())
            form_data["data_verbale_precedente"] = st.date_input("Data del verbale da correggere", 
                                                               value=CommonDataHandler.get_default_date_chiusura())
        with col2:
            form_data["tipo_correzione"] = st.selectbox("Tipo di correzione", 
                                                       ["Errore materiale", "Precisazione", "Integrazione"])
            form_data["modalita_assemblea"] = st.selectbox("Modalità assemblea", 
                                                          ["In presenza", "Audioconferenza", "Mista"])
        
        # Partecipanti standardizzati usando il CommonDataHandler
        participants_data = CommonDataHandler.extract_and_populate_participants_data(
            extracted_data, 
            unique_key_suffix="correzioni"
        )
        form_data.update(participants_data)
        
        # Ordine del giorno specifico per correzioni
        st.subheader("📋 Ordine del Giorno")
        default_odg = f"Precisazioni relative al verbale di assemblea del {form_data['data_verbale_precedente'].strftime('%d/%m/%Y')}"
        form_data["ordine_giorno"] = st.text_area("Ordine del giorno", 
                                                 default_odg, height=80)
        
        # Sezione correzioni
        st.subheader("✏️ Correzioni da Apportare")
        
        # Gestione dinamica delle correzioni
        if 'num_correzioni' not in st.session_state:
            st.session_state.num_correzioni = 1
        
        col1, col2 = st.columns([3, 1])
        with col1:
            st.write("Specifica le correzioni da apportare al verbale precedente:")
        with col2:
            # Use number input instead of buttons to avoid form conflicts
            new_num = st.number_input("Numero correzioni", 
                                    min_value=1, 
                                    max_value=10, 
                                    value=st.session_state.num_correzioni,
                                    key="num_correzioni_input")
            if new_num != st.session_state.num_correzioni:
                st.session_state.num_correzioni = new_num
        
        correzioni = []
        for i in range(st.session_state.num_correzioni):
            st.write(f"**Correzione {i+1}:**")
            col1, col2 = st.columns(2)
            with col1:
                testo_errato = st.text_input(f"Testo errato {i+1}", 
                                           key=f"testo_errato_{i}",
                                           placeholder="Inserisci il testo da correggere...")
            with col2:
                testo_corretto = st.text_input(f"Testo corretto {i+1}", 
                                             key=f"testo_corretto_{i}",
                                             placeholder="Inserisci il testo corretto...")
            
            if testo_errato and testo_corretto:
                correzioni.append({
                    "testo_errato": testo_errato,
                    "testo_corretto": testo_corretto
                })
        
        form_data["correzioni"] = correzioni
        
        # Delibera della precedente assemblea
        st.subheader("📝 Delibera dell'Assemblea Precedente")
        form_data["delibera_precedente"] = st.text_area("Descrivi brevemente cosa fu deliberato nell'assemblea precedente", 
                                                       height=100,
                                                       placeholder="Es: ha deliberato l'approvazione del bilancio...")
        
        # Traduzione (se necessaria)
        st.subheader("🌐 Traduzione")
        form_data["richiede_traduzione"] = st.checkbox("È necessaria traduzione per partecipanti stranieri")
        if form_data["richiede_traduzione"]:
            col1, col2 = st.columns(2)
            with col1:
                form_data["partecipante_straniero"] = st.text_input("Nome partecipante straniero")
                form_data["lingua_straniera"] = st.selectbox("Lingua di traduzione", 
                                                           ["Inglese", "Francese", "Tedesco", "Spagnolo", "Altro"])
            with col2:
                if form_data["lingua_straniera"] == "Altro":
                    form_data["lingua_altro"] = st.text_input("Specifica lingua")
        
        return form_data
    
    def show_preview(self, form_data: dict):
        """Mostra l'anteprima del documento fuori dal form"""
        # Sezione Anteprima
        st.markdown("---")
        st.subheader("👁️ Anteprima Documento")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col1:
            show_preview = st.checkbox("Mostra anteprima", value=False, key="preview_checkbox")
        
        with col2:
            if show_preview:
                st.info("💡 L'anteprima si aggiorna automaticamente con i dati inseriti sopra")
        
        if show_preview:
            try:
                preview_text = self._generate_preview_text(form_data)
                # Anteprima NON modificabile, stessa UX del template completo
                st.text_area("", preview_text, height=600, disabled=True)
            except Exception as e:
                st.error(f"Errore nell'anteprima: {e}")
    
    def _generate_preview_text(self, data: dict) -> str:
        """Genera il testo di anteprima del verbale"""
        
        # Header aziendale
        preview = f"""
**{data.get('denominazione', '[DENOMINAZIONE]')}**

Sede in {data.get('sede_legale', '[SEDE LEGALE]')}
Capitale sociale Euro {data.get('capitale_sociale', '[CAPITALE]')} i.v.
Codice fiscale: {data.get('codice_fiscale', '[CODICE FISCALE]')}

---

# Verbale di assemblea dei soci
del {data.get('data_assemblea', '[DATA]')}

---

Oggi {data.get('data_assemblea', '[DATA]')} alle ore {data.get('ora_assemblea', '[ORA]')} presso la sede sociale {data.get('sede_legale', '[SEDE]')}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:

## Ordine del giorno
{data.get('ordine_giorno', '[ORDINE DEL GIORNO]')}

---

Assume la presidenza ai sensi dell'art. [...] dello statuto sociale il Sig. **{data.get('presidente', '[PRESIDENTE]')}** {data.get('ruolo_presidente', 'Amministratore Unico')}, il quale dichiara e constata:

1. che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. [...] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza

2. che sono presenti/partecipano all'assemblea:
   - l'{data.get('ruolo_presidente', 'Amministratore Unico')} nella persona del suddetto Presidente Sig. **{data.get('presidente', '[PRESIDENTE]')}**
"""
        
        # Sezione partecipanti
        soci = data.get('soci', [])
        if soci:
            # Calcola totali euro e percentuale
            totale_euro = 0.0
            totale_perc = 0.0
            capitale_raw = str(data.get('capitale_sociale', '0')).replace('.', '').replace(',', '.')
            try:
                capitale_float = float(capitale_raw)
            except ValueError:
                capitale_float = 0.0

            for socio in soci:
                if isinstance(socio, dict):
                    # somma euro
                    euro_raw = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                    try:
                        totale_euro += float(euro_raw or 0)
                    except ValueError:
                        pass
                    # somma percentuale se presente, altrimenti calcola
                    perc_raw = socio.get('quota_percentuale', '')
                    if perc_raw:
                        try:
                            totale_perc += float(str(perc_raw).replace('%', '').replace(',', '.'))
                        except ValueError:
                            pass
                    else:
                        try:
                            euro_val = float(euro_raw or 0)
                            if capitale_float:
                                totale_perc += euro_val / capitale_float * 100
                        except ValueError:
                            pass

            formatted_euro = CommonDataHandler.format_currency(totale_euro)
            formatted_perc = CommonDataHandler.format_percentage(totale_perc)

            preview += f"\n   - nonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_euro} pari al {formatted_perc} del Capitale Sociale:\n"
            for socio in soci:
                if not isinstance(socio, dict) or not socio.get('nome'):
                    continue
                nome = socio.get('nome')
                quota_euro = CommonDataHandler.format_currency(socio.get('quota_euro', '0'))
                quota_perc = socio.get('quota_percentuale', '')
                if not quota_perc:
                    # calcola
                    try:
                        euro_val = float(str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.'))
                        quota_perc = CommonDataHandler.format_percentage(euro_val / capitale_float * 100) if capitale_float else '[%]'
                    except ValueError:
                        quota_perc = '[%]'
                else:
                    quota_perc = CommonDataHandler.clean_percentage(quota_perc)

                tipo_sogg = socio.get('tipo_soggetto', 'Persona Fisica')
                tipo_part = socio.get('tipo_partecipazione', 'Diretto')
                delegato = socio.get('delegato', '').strip()
                rappresentante = socio.get('rappresentante_legale', '').strip()

                if tipo_part == 'Delegato' and delegato:
                    if tipo_sogg == 'Società':
                        line = f"il Sig. {delegato} delegato della società {nome} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                    else:
                        line = f"il Sig. {delegato} delegato del socio Sig. {nome} recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                else:
                    if tipo_sogg == 'Società':
                        if rappresentante:
                            line = f"la società {nome}, rappresentata dal Sig {rappresentante}, socia recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                        else:
                            line = f"la società {nome} socia recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                    else:
                        line = f"il Sig. {nome} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"

                preview += f"     - {line}\n"
        
        preview += f"""
3. che gli intervenuti sono legittimati alla presente assemblea;
4. che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno.

I presenti all'unanimità chiamano a fungere da segretario il signor **{data.get('segretario', '[SEGRETARIO]')}**, che accetta l'incarico.

Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante.
"""
        
        # Sezione traduzione se necessaria
        if data.get('richiede_traduzione'):
            lingua = data.get('lingua_straniera', 'inglese').lower()
            if lingua == 'altro':
                lingua = data.get('lingua_altro', '[LINGUA]').lower()
            
            preview += f"""
In particolare, preso atto che il Sig. **{data.get('partecipante_straniero', '[PARTECIPANTE]')}** non conosce la lingua italiana ma dichiara di conoscere la lingua {lingua}, il Presidente dichiara che provvederà a tradurre dall'italiano al {lingua} (e viceversa) gli interventi dei partecipanti alla discussione nonché a tradurre dall'italiano al {lingua} il verbale che sarà redatto al termine della riunione.
"""
        
        preview += f"""
Il Presidente constata e fa constatare che l'assemblea risulta regolarmente convocata e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.

Si passa quindi allo svolgimento dell'ordine del giorno.

---

*     *     *

Il Presidente ricorda agli intervenuti che l'assemblea dei soci riunitasi lo scorso **{data.get('data_verbale_precedente', '[DATA PRECEDENTE]')}** ha deliberato **{data.get('delibera_precedente', '[DELIBERA]')}**; purtroppo, causa un errore materiale, il verbale della suddetta assemblea riporta i termini errati:
"""
        
        # Sezione correzioni
        correzioni = data.get('correzioni', [])
        if correzioni:
            for i, correzione in enumerate(correzioni, 1):
                preview += f"\n**Correzione {i}:**\n"
                preview += f"- **Testo errato:** {correzione['testo_errato']}\n"
                preview += f"- **Testo corretto:** {correzione['testo_corretto']}\n"
        else:
            preview += "\n- **Testo errato:** [TESTO DA CORREGGERE]\n"
            preview += "- **Testo corretto:** [TESTO CORRETTO]\n"
        
        preview += f"""

Il Presidente invita pertanto a correggere il verbale stesso.

L'Assemblea prende atto delle dichiarazioni del Presidente.

Viene quindi corretto il suddetto verbale dell'assemblea dei soci del **{data.get('data_verbale_precedente', '[DATA PRECEDENTE]')}** e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo che viene allegato al presente verbale per la sua trascrizione sul libro sociale.

---

## Chiusura dell'Assemblea

Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.

Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo.

L'assemblea viene sciolta alle ore **{data.get('ora_chiusura', '[ORA CHIUSURA]')}**.

---

**Il Presidente**                    **Il Segretario**

_________________                    _________________

{data.get('presidente', '[PRESIDENTE]')}                      {data.get('segretario', '[SEGRETARIO]')}
"""
        
        return preview
    
    def generate_document(self, data: dict) -> Document:
        """Genera il documento Word del verbale di correzioni. Se l'utente ha modificato l'anteprima viene usato il testo personalizzato."""

        # Genera il documento con la struttura automatica in stile "completo"
        doc = Document()

        # Imposta gli stili del documento
        self._setup_document_styles(doc)

        # Header aziendale
        self._add_company_header(doc, data)

        # Titolo del verbale
        self._add_verbale_title(doc, data)

        # Sezione di apertura
        self._add_opening_section(doc, data)

        # Sezione partecipanti
        self._add_participants_section(doc, data)

        # Dichiarazioni preliminari
        self._add_preliminary_statements(doc, data)

        # Discussione delle correzioni
        self._add_correzioni_discussion(doc, data)

        # Sezione di chiusura
        self._add_closing_section(doc, data)

        # Firme
        self._add_signatures(doc, data)

        return doc
    
    def _setup_document_styles(self, doc):
        """Imposta gli stili per il documento"""
        # Stile per il titolo principale
        title_style = doc.styles.add_style('VerbaleTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Times New Roman'
        title_font.size = Pt(16)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
        
        # Stile per i sottotitoli
        subtitle_style = doc.styles.add_style('VerbaleSubtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_font = subtitle_style.font
        subtitle_font.name = 'Times New Roman'
        subtitle_font.size = Pt(12)
        subtitle_font.bold = True
        subtitle_style.paragraph_format.space_before = Pt(12)
        subtitle_style.paragraph_format.space_after = Pt(6)
        
        # Stile per il testo normale
        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = 1.15
    
    def _add_company_header(self, doc, data):
        """Aggiunge l'header con i dati dell'azienda"""
        # Nome azienda
        company_p = doc.add_paragraph(data.get('denominazione', '[DENOMINAZIONE]'))
        company_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in company_p.runs:
            run.font.bold = True
            run.font.size = Pt(14)
        
        # Dati aziendali
        doc.add_paragraph()
        info_p = doc.add_paragraph()
        info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_p.add_run(f"Sede in {data.get('sede_legale', '[SEDE LEGALE]')}\n")
        info_p.add_run(f"Capitale sociale Euro {data.get('capitale_sociale', '[CAPITALE]')} i.v.\n")
        info_p.add_run(f"Codice fiscale: {data.get('codice_fiscale', '[CODICE FISCALE]')}")
        
        doc.add_paragraph()
    
    def _add_verbale_title(self, doc, data):
        """Aggiunge il titolo del verbale"""
        title_p = doc.add_paragraph("Verbale di assemblea dei soci", style='VerbaleTitle')
        data_p = doc.add_paragraph(f"del {data.get('data_assemblea', '[DATA]')}", style='VerbaleTitle')
        doc.add_paragraph()
    
    def _add_opening_section(self, doc, data):
        """Aggiunge la sezione di apertura"""
        opening_text = f"""Oggi {data.get('data_assemblea', '[DATA]')} alle ore {data.get('ora_assemblea', '[ORA]')} presso la sede sociale {data.get('sede_legale', '[SEDE]')}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:"""
        
        doc.add_paragraph(opening_text)
        doc.add_paragraph()
        
        # Ordine del giorno
        odg_title = doc.add_paragraph("Ordine del giorno", style='VerbaleSubtitle')
        odg_p = doc.add_paragraph(data.get('ordine_giorno', '[ORDINE DEL GIORNO]'))
        doc.add_paragraph()
    
    def _add_participants_section(self, doc, data):
        """Aggiunge la sezione dei partecipanti"""
        presidente_text = f"""Assume la presidenza ai sensi dell'art. [...] dello statuto sociale il Sig. {data.get('presidente', '[PRESIDENTE]')} {data.get('ruolo_presidente', 'Amministratore Unico')}, il quale dichiara e constata:"""
        
        doc.add_paragraph(presidente_text)
        doc.add_paragraph()
        
        # Dichiarazioni del presidente
        doc.add_paragraph("1 - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. [...] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza")
        doc.add_paragraph()
        
        participants_p = doc.add_paragraph("2 - che sono presenti/partecipano all'assemblea:")
        participants_p.add_run(f"\nl'{data.get('ruolo_presidente', 'Amministratore Unico')} nella persona del suddetto Presidente Sig. {data.get('presidente', '[PRESIDENTE]')}")
        
        # Aggiungi soci
        soci = data.get('soci', [])
        if soci:
            # Calcola totali euro e percentuale
            totale_euro = 0.0
            totale_perc = 0.0
            capitale_raw = str(data.get('capitale_sociale', '0')).replace('.', '').replace(',', '.')
            try:
                capitale_float = float(capitale_raw)
            except ValueError:
                capitale_float = 0.0

            for socio in soci:
                if isinstance(socio, dict):
                    # somma euro
                    euro_raw = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                    try:
                        totale_euro += float(euro_raw or 0)
                    except ValueError:
                        pass
                    # somma percentuale se presente, altrimenti calcola
                    perc_raw = socio.get('quota_percentuale', '')
                    if perc_raw:
                        try:
                            totale_perc += float(str(perc_raw).replace('%', '').replace(',', '.'))
                        except ValueError:
                            pass
                    else:
                        try:
                            euro_val = float(euro_raw or 0)
                            if capitale_float:
                                totale_perc += euro_val / capitale_float * 100
                        except ValueError:
                            pass

            formatted_euro = CommonDataHandler.format_currency(totale_euro)
            formatted_perc = CommonDataHandler.format_percentage(totale_perc)

            participants_p.add_run(f"\nnonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_euro} pari al {formatted_perc} del Capitale Sociale:")

            # Elenco dettagliato
            for socio in soci:
                if not isinstance(socio, dict) or not socio.get('nome'):
                    continue
                nome = socio.get('nome')
                quota_euro = CommonDataHandler.format_currency(socio.get('quota_euro', '0'))
                quota_perc = socio.get('quota_percentuale', '')
                if not quota_perc:
                    # calcola
                    try:
                        euro_val = float(str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.'))
                        quota_perc = CommonDataHandler.format_percentage(euro_val / capitale_float * 100) if capitale_float else '[%]'
                    except ValueError:
                        quota_perc = '[%]'
                else:
                    quota_perc = CommonDataHandler.clean_percentage(quota_perc)

                tipo_sogg = socio.get('tipo_soggetto', 'Persona Fisica')
                tipo_part = socio.get('tipo_partecipazione', 'Diretto')
                delegato = socio.get('delegato', '').strip()
                rappresentante = socio.get('rappresentante_legale', '').strip()

                if tipo_part == 'Delegato' and delegato:
                    if tipo_sogg == 'Società':
                        line = f"il Sig. {delegato} delegato della società {nome} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                    else:
                        line = f"il Sig. {delegato} delegato del socio Sig. {nome} recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                else:
                    if tipo_sogg == 'Società':
                        if rappresentante:
                            line = f"la società {nome}, rappresentata dal Sig {rappresentante}, socia recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                        else:
                            line = f"la società {nome} socia recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"
                    else:
                        line = f"il Sig. {nome} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale"

                participants_p.add_run(f"\n{line}")
        
        doc.add_paragraph()
    
    def _add_preliminary_statements(self, doc, data):
        """Aggiunge le dichiarazioni preliminari"""
        doc.add_paragraph("3 - che gli intervenuti sono legittimati alla presente assemblea;")
        doc.add_paragraph("4 - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno.")
        doc.add_paragraph()
        
        segretario_text = f"I presenti all'unanimità chiamano a fungere da segretario il signor {data.get('segretario', '[SEGRETARIO]')}, che accetta l'incarico."
        doc.add_paragraph(segretario_text)
        doc.add_paragraph()
        
        identification_text = "Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante."
        doc.add_paragraph(identification_text)
        
        # Sezione traduzione se necessaria
        if data.get('richiede_traduzione'):
            lingua = data.get('lingua_straniera', 'inglese').lower()
            if lingua == 'altro':
                lingua = data.get('lingua_altro', '[LINGUA]').lower()
            
            translation_text = f"In particolare, preso atto che il Sig. {data.get('partecipante_straniero', '[PARTECIPANTE]')} non conosce la lingua italiana ma dichiara di conoscere la lingua {lingua}, il Presidente dichiara che provvederà a tradurre dall'italiano al {lingua} (e viceversa) gli interventi dei partecipanti alla discussione nonché a tradurre dall'italiano al {lingua} il verbale che sarà redatto al termine della riunione."
            doc.add_paragraph(translation_text)
        
        doc.add_paragraph()
        validity_text = "Il Presidente constata e fa constatare che l'assemblea risulta regolarmente convocata e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno."
        doc.add_paragraph(validity_text)
        doc.add_paragraph()
        
        doc.add_paragraph("Si passa quindi allo svolgimento dell'ordine del giorno.")
        doc.add_paragraph("*     *     *")
        doc.add_paragraph()
    
    def _add_correzioni_discussion(self, doc, data):
        """Aggiunge la discussione delle correzioni"""
        intro_text = f"Il Presidente ricorda agli intervenuti che l'assemblea dei soci riunitasi lo scorso {data.get('data_verbale_precedente', '[DATA PRECEDENTE]')} ha deliberato {data.get('delibera_precedente', '[DELIBERA]')}; purtroppo, causa un errore materiale, il verbale della suddetta assemblea riporta i termini errati"
        
        # Gestione correzioni
        correzioni = data.get('correzioni', [])
        if correzioni:
            for correzione in correzioni:
                intro_text += f" [{correzione['testo_errato']}] invece dei corretti [{correzione['testo_corretto']}]"
        else:
            intro_text += " [TESTO ERRATO] invece dei corretti [TESTO CORRETTO]"
        
        intro_text += "."
        doc.add_paragraph(intro_text)
        doc.add_paragraph()
        
        doc.add_paragraph("Il Presidente invita pertanto a correggere il verbale stesso.")
        doc.add_paragraph()
        doc.add_paragraph("L'Assemblea prende atto delle dichiarazioni del Presidente.")
        doc.add_paragraph()
        
        approval_text = f"Viene quindi corretto il suddetto verbale dell'assemblea dei soci del {data.get('data_verbale_precedente', '[DATA PRECEDENTE]')} e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo che viene allegato al presente verbale per la sua trascrizione sul libro sociale."
        doc.add_paragraph(approval_text)
        doc.add_paragraph()
        doc.add_paragraph("*     *     *")
        doc.add_paragraph()
    
    def _add_closing_section(self, doc, data):
        """Aggiunge la sezione di chiusura"""
        closing_text = "Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola."
        doc.add_paragraph(closing_text)
        doc.add_paragraph()
        
        final_approval_text = "Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo."
        doc.add_paragraph(final_approval_text)
        doc.add_paragraph()
        
        end_time_text = f"L'assemblea viene sciolta alle ore {data.get('ora_chiusura', '[ORA CHIUSURA]')}."
        doc.add_paragraph(end_time_text)
        doc.add_paragraph()
    
    def _add_signatures(self, doc, data):
        """Aggiunge le firme"""
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Tabella per le firme usando il metodo sicuro
        signature_table = self._create_table_with_style(doc, rows=3, cols=2)
        
        # Prima riga - titoli
        signature_table.cell(0, 0).text = "Il Presidente"
        signature_table.cell(0, 1).text = "Il Segretario"
        
        # Seconda riga - spazio per firme
        signature_table.cell(1, 0).text = "_________________"
        signature_table.cell(1, 1).text = "_________________"
        
        # Terza riga - nomi
        signature_table.cell(2, 0).text = data.get('presidente', '[PRESIDENTE]')
        signature_table.cell(2, 1).text = data.get('segretario', '[SEGRETARIO]')
        
        # Centra la tabella
        for row in signature_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Registra il template
DocumentTemplateFactory.register_template("correzioni", VerbaleCorrezioniTemplate)