"""
Template per Verbale di Assemblea Irregolare dei Soci
Questo template è specifico per verbali di assemblee irregolari per mancanza del numero legale.
"""

import sys
import os

# Aggiungi il path della cartella src (relativo alla root del progetto)
current_dir = os.path.dirname(os.path.abspath(__file__))
src_path = os.path.join(os.path.dirname(current_dir), 'src')
if src_path not in sys.path:
    sys.path.append(src_path)

from document_templates import DocumentTemplate, DocumentTemplateFactory
from common_data_handler import CommonDataHandler
from base_verbale_template import BaseVerbaleTemplate
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date
import streamlit as st
import pandas as pd
import re

class VerbaleAssembleaIrregolareTemplate(BaseVerbaleTemplate):
    """Template per Verbale di Assemblea Irregolare dei Soci"""
    
    def get_template_name(self) -> str:
        return "Verbale Assemblea Irregolare"
    
    def get_required_fields(self) -> list:
        return [
            "denominazione", "sede_legale", "capitale_sociale", "codice_fiscale",
            "data_assemblea", "ora_assemblea", "presidente", "segretario", 
            "soci", "amministratori", "ordine_giorno"
        ]
    
    def get_form_fields(self, extracted_data: dict) -> dict:
        """Genera i campi del form per Streamlit."""
        # Chiama il metodo della classe base per ottenere i campi comuni
        form_data = super().get_form_fields(extracted_data)

        st.subheader("🚨 Configurazioni Assemblea Irregolare / Totalitaria")

        # Campi specifici del template
        form_data['dettagli_convocazione'] = st.text_area(
            "Dettagli sulla mancata convocazione o irregolarità",
            "L'assemblea si è riunita in forma totalitaria...",
            key="dettagli_convocazione_irregolare",
            height=100
        )
        
        # Sezione per la rinuncia ai termini di convocazione
        st.subheader("✍️ Rinuncia ai Termini")
        form_data['rinuncia_termini'] = st.checkbox(
            "I presenti dichiarano di rinunciare ai termini di convocazione",
            value=True,
            key="rinuncia_termini_irregolare"
        )
        
        if form_data.get('rinuncia_termini'):
            form_data['motivo_rinuncia'] = st.text_input(
                "Motivo della rinuncia (opzionale)",
                "per urgenti motivi gestionali",
                key="motivo_rinuncia_irregolare"
            )

        # Ordine del giorno
        st.subheader("📋 Ordine del Giorno")
        form_data['ordine_del_giorno'] = st.text_area(
            "Ordine del giorno trattato",
            "1. ...\n2. ...",
            key="odg_irregolare",
            height=150
        )
        
        return form_data
    
    def show_preview(self, form_data: dict):
        """Mostra un'anteprima del documento che verrà generato"""
        st.subheader("📄 Anteprima Documento")
        
        preview_text = self._generate_preview_text(form_data)
        st.text_area("", value=preview_text, height=600, disabled=True)
    
    def _generate_preview_text(self, data: dict) -> str:
        """Genera il testo di anteprima del documento"""
        lines = []
        
        # Header
        denominazione = data.get('denominazione', '[DENOMINAZIONE]')
        sede = data.get('sede_legale', '[SEDE]')
        
        # Gestione capitale sociale con tre campi
        capitale_deliberato = str(data.get('capitale_deliberato', '10.000,00')).strip()
        capitale_versato = str(data.get('capitale_versato', '10.000,00')).strip()
        capitale_sottoscritto = str(data.get('capitale_sottoscritto', '10.000,00')).strip()
        
        # Gestione "i.v." (interamente versato) - solo se versato = deliberato
        if capitale_versato == capitale_deliberato:
            capitale_text = f"Capitale sociale Euro {capitale_deliberato} i.v."
        else:
            capitale_text = f"Capitale sociale Euro deliberato: {capitale_deliberato}, sottoscritto: {capitale_sottoscritto}, versato: {capitale_versato}"
        
        cf = data.get('codice_fiscale', '[CF]')
        
        lines.extend([
            f"{denominazione}",
            f"Sede in {sede}",
            capitale_text,
            f"Codice fiscale: {cf}",
            "",
            "Verbale di assemblea dei soci",
            f"del {data.get('data_assemblea', date.today()).strftime('%d/%m/%Y')}",
            "",
        ])
        
        # Apertura assemblea
        ora_inizio = data.get('ora_inizio', '09:00')
        luogo = data.get('luogo_assemblea', sede)
        lines.extend([
            f"Oggi {data.get('data_assemblea', date.today()).strftime('%d/%m/%Y')} alle ore {ora_inizio} presso {luogo}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:",
            "",
            "Ordine del giorno"
        ])
        
        # Ordine del giorno
        for punto in data.get('ordine_giorno', []):
            lines.append(f"[{punto}]")
        
        lines.append("")
        
        # Presidente
        presidente = data.get('presidente', '[PRESIDENTE]')
        ruolo = data.get('ruolo_presidente', 'Amministratore Unico')
        
        lines.extend([
            f"Assume la presidenza ai sensi dell'art. [...] dello statuto sociale il Sig. {presidente} {ruolo}", 
            # Aggiungi il testo condizionale per AU
            f"{'[oppure Presidente del Consiglio di Amministrazione o altro (come da statuto)]' if ruolo == 'Amministratore Unico' else ''}, il quale dichiara e constata:",
            "",
            "1 - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. [...] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza" if data.get('audioconferenza') else "1 - che l'assemblea è regolarmente convocata",
            "",
            "2 - che sono presenti/partecipano all'assemblea:",
            f"     l'{ruolo} nella persona del suddetto Presidente Sig. {presidente}",
            ""
        ])
        
        # Amministratori
        amministratori_presenti = [a for a in data.get('amministratori', []) if a.get('presente', True) and a.get('nome', '').strip()]
        if len(amministratori_presenti) > 1:
            lines.append("per il Consiglio di Amministrazione:")
            for amm in amministratori_presenti:
                if amm.get('nome') != presidente:  # Non ripetere il presidente
                    lines.append(f"il Sig {amm.get('nome', '')}")
        
        # Collegio sindacale
        if data.get('collegio_sindacale'):
            lines.extend([
                "per il Collegio Sindacale",
                "il Dott. [NOME_SINDACO_1]",
                "il Dott. [NOME_SINDACO_2]",
                "il Dott. [NOME_SINDACO_3]"
            ])
        
        # Revisore
        if data.get('revisore'):
            nome_revisore = data.get('nome_revisore', '[NOME_REVISORE]')
            lines.append(f"il revisore contabile Dott. {nome_revisore}")
        
        lines.append("")
        
        # Calcola soci presenti e totali
        # Fallback per la retrocompatibilità
        if 'soci_presenti' not in data:
            data['soci_presenti'] = [s for s in data.get('soci', []) if s.get('presente', True)]
            data['soci_assenti'] = [s for s in data.get('soci', []) if not s.get('presente', True)]

        soci_presenti = data.get('soci_presenti', [])
        soci_assenti = data.get('soci_assenti', [])
        
        totale_quote_euro = 0
        totale_quote_perc = 0
        
        for socio in soci_presenti:
            if socio.get('nome', '').strip():
                quota_perc_str = str(socio.get('quota_percentuale', '0')).replace('%', '').replace(',', '.').strip()
                quota_euro_str = str(socio.get('quota_euro', '0')).replace('€', '', 1).replace('Euro', '', 1).replace('.', '').replace(',', '.').strip()
                
                # Calcola i totali
                try:
                    if quota_euro_str:
                        totale_quote_euro += float(quota_euro_str)
                except (ValueError, TypeError):
                    pass
                
                try:
                    if quota_perc_str:
                        totale_quote_perc += float(quota_perc_str)
                except (ValueError, TypeError):
                    pass
        
        # Gestione soci presenti
        if soci_presenti:
            # Formatta i totali per la visualizzazione
            formatted_total_quota_euro = f"{totale_quote_euro:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            formatted_total_quota_percentuale = f"{totale_quote_perc:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

            lines.append(f"nonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_total_quota_euro} pari al {formatted_total_quota_percentuale}% del Capitale Sociale:")
            
            for socio_info in soci_presenti:
                nome = socio_info.get('nome', '[Nome]')
                quota_euro = socio_info.get('quota_euro', '0')
                quota_perc = socio_info.get('quota_percentuale', '0%')
                tipo = socio_info.get('tipo_partecipazione', 'Diretto')
                delegato = socio_info.get('delegato', '')
                tipo_soggetto = socio_info.get('tipo_soggetto', 'Persona Fisica')
                rappresentante_legale = socio_info.get('rappresentante_legale', '')
                
                socio_line = self._format_socio_line(nome, quota_euro, quota_perc, tipo, delegato, tipo_soggetto, rappresentante_legale)
                lines.append(socio_line)
        
        # Aggiunge soci assenti se presenti
        if soci_assenti:
            lines.append("\nRisultano invece assenti i seguenti soci:")
            for socio in soci_assenti:
                lines.append(f"- Sig. {socio.get('nome', '[Nome]')}")
        
        lines.append("")
        
        # Conclusione e irregolarità
        percentuale_presente = data.get('percentuale_presente', '40')
        if soci_presenti:
            percentuale_effettiva = totale_quote_perc if totale_quote_perc > 0 else float(percentuale_presente)
            lines.extend([
                f"Il Presidente constata e fa constatare che l'assemblea risulta regolarmente convocata ma che sono presenti soci rappresentanti soltanto il {percentuale_effettiva:.1f}% del capitale sociale; dichiara pertanto che l'Assemblea deve considerarsi irregolarmente costituita per mancanza del numero legale.",
                "",
                "Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo.",
            ])
        else:
            lines.extend([
                "Il Presidente constata e fa constatare che l'assemblea risulta regolarmente convocata ma che non sono presenti soci; dichiara pertanto che l'Assemblea deve considerarsi deserta per mancanza totale di partecipazione.",
                "",
                "Viene quindi redatto il presente verbale di assemblea deserta.",
            ])
        
        lines.append("")
        lines.append(f"L'assemblea viene sciolta alle ore {data.get('ora_fine', '09:30')}.")
        
        return '\n'.join(lines)
    
    def generate_document(self, data: dict) -> Document:
        """Genera il documento Word"""
        doc = Document()
        self._setup_document_styles(doc)
        self._add_company_header(doc, data)
        self._add_verbale_title(doc, data)
        self._add_opening_section(doc, data)
        self._add_participants_section(doc, data)
        self._add_preliminary_statements(doc, data)
        self._add_irregularity_section(doc, data)
        self._add_closing_section(doc, data)
        self._add_signatures(doc, data)
        
        return doc
    
    def _setup_document_styles(self, doc):
        """Configura gli stili del documento"""
        # First call parent's method to setup all base styles
        super()._setup_document_styles(doc)
        
        # Imposta font di default
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
    
    def _add_company_header(self, doc, data):
        """Aggiunge l'intestazione della società"""
        denominazione = data.get('denominazione', '[DENOMINAZIONE]')
        sede = data.get('sede_legale', '[SEDE]')
        
        # Gestione capitale sociale con tre campi
        capitale_deliberato = str(data.get('capitale_deliberato', '10.000,00')).strip()
        capitale_versato = str(data.get('capitale_versato', '10.000,00')).strip()
        capitale_sottoscritto = str(data.get('capitale_sottoscritto', '10.000,00')).strip()
        
        # Gestione "i.v." (interamente versato) - solo se versato = deliberato
        if capitale_versato == capitale_deliberato:
            capitale_text = f"Capitale sociale Euro {capitale_deliberato} i.v."
        else:
            capitale_text = f"Capitale sociale Euro deliberato: {capitale_deliberato}, sottoscritto: {capitale_sottoscritto}, versato: {capitale_versato}"
        
        cf = data.get('codice_fiscale', '[CF]')
        
        # Denominazione centrata e in grassetto
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(denominazione)
        run.font.bold = True
        run.font.size = Pt(14)
        
        # Sede
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Sede in {sede}")
        run.font.size = Pt(12)
        
        # Capitale sociale
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(capitale_text)
        run.font.size = Pt(12)
        
        # Codice fiscale
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Codice fiscale: {cf}")
        run.font.size = Pt(12)
        
        doc.add_paragraph()
    
    def _add_verbale_title(self, doc, data):
        """Aggiunge il titolo del verbale"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Verbale di assemblea dei soci")
        run.font.bold = True
        run.font.size = Pt(14)
        
        data_assemblea = data.get('data_assemblea', date.today())
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"del {data_assemblea.strftime('%d/%m/%Y')}")
        run.font.bold = True
        run.font.size = Pt(14)
        
        doc.add_paragraph()
    
    def _add_opening_section(self, doc, data):
        """Aggiunge la sezione di apertura"""
        data_assemblea = data.get('data_assemblea', date.today())
        ora_inizio = data.get('ora_inizio', '09:00')
        luogo = data.get('luogo_assemblea', data.get('sede_legale', '[SEDE]'))
        
        text = f"Oggi {data_assemblea.strftime('%d/%m/%Y')} alle ore {ora_inizio} presso {luogo}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:"
        doc.add_paragraph(text)
        
        # Ordine del giorno
        p = doc.add_paragraph("Ordine del giorno")
        run = p.runs[0]
        run.font.bold = True
        
        for punto in data.get('ordine_giorno', []):
            doc.add_paragraph(f"[{punto}]")
        
        doc.add_paragraph()
    
    def _add_participants_section(self, doc, data):
        """Aggiunge la sezione dei partecipanti"""
        presidente = data.get('presidente', '[PRESIDENTE]')
        ruolo = data.get('ruolo_presidente', 'Amministratore Unico')
        
        text = f"Assume la presidenza ai sensi dell'art. [...] dello statuto sociale il Sig. {presidente} {ruolo}, il quale dichiara e constata:"
        if ruolo == "Amministratore Unico":
            text += " [oppure Presidente del Consiglio di Amministrazione o altro (come da statuto)]"
        text += ", che l'assemblea è regolarmente convocata secondo le modalità statutarie"
        
        doc.add_paragraph(text)
        doc.add_paragraph()
        
        # Dichiarazioni
        dichiarazioni = []
        if data.get('audioconferenza'):
            dichiarazioni.append("che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. [...] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza")
        else:
            dichiarazioni.append("che l'assemblea è regolarmente convocata")
        
        dichiarazioni.append(f"che sono presenti/partecipano all'assemblea:\nl' {ruolo} nella persona del suddetto Presidente Sig. {presidente}")
        
        for i, dich in enumerate(dichiarazioni, 1):
            doc.add_paragraph(f"{i} - {dich}")
        
        # Amministratori
        amministratori_presenti = [a for a in data.get('amministratori', []) if a.get('presente', True) and a.get('nome', '').strip()]
        if len(amministratori_presenti) > 1:
            doc.add_paragraph()
            doc.add_paragraph("per il Consiglio di Amministrazione:")
            for amm in amministratori_presenti:
                if amm.get('nome') != presidente:  # Non ripetere il presidente
                    doc.add_paragraph(f"il Sig {amm.get('nome', '')}")
        
        # Collegio sindacale
        if data.get('collegio_sindacale'):
            doc.add_paragraph()
            doc.add_paragraph("per il Collegio Sindacale")
            doc.add_paragraph("il Dott. [NOME_SINDACO_1]")
            doc.add_paragraph("il Dott. [NOME_SINDACO_2]")
            doc.add_paragraph("il Dott. [NOME_SINDACO_3]")
        
        # Revisore
        if data.get('revisore'):
            doc.add_paragraph()
            nome_revisore = data.get('nome_revisore', '[NOME_REVISORE]')
            doc.add_paragraph(f"il revisore contabile Dott. {nome_revisore}")
        
        doc.add_paragraph()
    
    def _add_preliminary_statements(self, doc, data):
        """Aggiunge le constatazioni preliminari"""
        # Fallback per la retrocompatibilità
        if 'soci_presenti' not in data:
            data['soci_presenti'] = [s for s in data.get('soci', []) if s.get('presente', True)]
            data['soci_assenti'] = [s for s in data.get('soci', []) if not s.get('presente', True)]

        soci_presenti = data.get('soci_presenti', [])
        soci_assenti = data.get('soci_assenti', [])
        
        totale_quote_euro = 0
        totale_quote_perc = 0
        
        for socio in soci_presenti:
            if socio.get('nome', '').strip():
                quota_perc_str = str(socio.get('quota_percentuale', '0')).replace('%', '').replace(',', '.').strip()
                quota_euro_str = str(socio.get('quota_euro', '0')).replace('€', '', 1).replace('Euro', '', 1).replace('.', '').replace(',', '.').strip()
                
                # Calcola i totali
                try:
                    if quota_euro_str:
                        totale_quote_euro += float(quota_euro_str)
                except (ValueError, TypeError):
                    pass
                
                try:
                    if quota_perc_str:
                        totale_quote_perc += float(quota_perc_str)
                except (ValueError, TypeError):
                    pass
        
        # Gestione soci presenti
        if soci_presenti:
            formatted_total_quota_euro = f"{totale_quote_euro:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            formatted_total_quota_percentuale = f"{totale_quote_perc:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            
            self.add_paragraph_with_font(doc, f"nonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_total_quota_euro} pari al {formatted_total_quota_percentuale}% del Capitale Sociale:", size=Pt(12), font_name='Times New Roman', space_before=Pt(12))
            
            for socio_info in soci_presenti:
                nome = socio_info.get('nome', '[Nome]')
                quota_euro = socio_info.get('quota_euro', '0')
                quota_perc = socio_info.get('quota_percentuale', '0%')
                tipo = socio_info.get('tipo_partecipazione', 'Diretto')
                delegato = socio_info.get('delegato', '')
                tipo_soggetto = socio_info.get('tipo_soggetto', 'Persona Fisica')
                rappresentante_legale = socio_info.get('rappresentante_legale', '')
                
                socio_line = self._format_socio_line(nome, quota_euro, quota_perc, tipo, delegato, tipo_soggetto, rappresentante_legale)
                self.add_paragraph_with_font(doc, socio_line, size=Pt(12), font_name='Times New Roman', left_indent=Inches(0.5))
        
        # Aggiunge soci assenti se presenti
        if soci_assenti:
            self.add_paragraph_with_font(doc, "Risultano invece assenti i seguenti soci:", size=Pt(12), font_name='Times New Roman', space_before=Pt(12))
            for socio in soci_assenti:
                self.add_paragraph_with_font(doc, f"- Sig. {socio.get('nome', '[Nome]')}", size=Pt(12), font_name='Times New Roman', left_indent=Inches(0.5))
        
        doc.add_paragraph()
    
    def _add_irregularity_section(self, doc, data):
        """Aggiunge la sezione di constatazione dell'irregolarità"""
        # Fallback per la retrocompatibilità
        if 'soci_presenti' not in data:
            data['soci_presenti'] = [s for s in data.get('soci', []) if s.get('presente', True)]
            data['soci_assenti'] = [s for s in data.get('soci', []) if not s.get('presente', True)]

        soci_presenti = data.get('soci_presenti', [])
        soci_assenti = data.get('soci_assenti', [])
        
        totale_quote_perc = 0
        for socio in soci_presenti:
            try:
                quota_perc_str = str(socio.get('quota_percentuale', '0')).replace('%', '').replace(',', '.').strip()
                if quota_perc_str:
                    totale_quote_perc += float(quota_perc_str)
            except (ValueError, TypeError):
                pass
        
        percentuale_effettiva = totale_quote_perc if totale_quote_perc > 0 else float(data.get('percentuale_presente', '0'))
        convocazione = "regolarmente convocata" if data.get('tipo_convocazione') == "regolarmente convocata" else "convocata"
        
        text = f"Il Presidente constata e fa constatare che l'assemblea risulta {convocazione} ma che sono presenti soci rappresentanti soltanto il {percentuale_effettiva:.1f}% del capitale sociale; dichiara pertanto che l'Assemblea deve considerarsi irregolarmente costituita per mancanza del numero legale."
        doc.add_paragraph(text)
        doc.add_paragraph()
    
    def _add_closing_section(self, doc, data):
        """Aggiunge la sezione di chiusura"""
        # Fallback per la retrocompatibilità
        if 'soci_presenti' not in data:
            data['soci_presenti'] = [s for s in data.get('soci', []) if s.get('presente', True)]
            data['soci_assenti'] = [s for s in data.get('soci', []) if not s.get('presente', True)]

        soci_presenti = data.get('soci_presenti', [])
        soci_assenti = data.get('soci_assenti', [])
        
        if soci_presenti:
            # Assemblea irregolare (con soci presenti ma sotto il numero legale)
            doc.add_paragraph("Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo.")
        else:
            # Assemblea deserta (nessun socio presente)
            doc.add_paragraph("Viene quindi redatto il presente verbale di assemblea deserta.")
        
        ora_fine = data.get('ora_fine', '09:30')
        doc.add_paragraph(f"L'assemblea viene sciolta alle ore {ora_fine}.")
        doc.add_paragraph()
    
    def _add_signatures(self, doc, data):
        """Aggiunge le firme"""
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Tabella per le firme usando il metodo sicuro
        table = self._create_table_with_style(doc, rows=3, cols=2)
        
        # Intestazioni
        table.cell(0, 0).text = "IL PRESIDENTE"
        table.cell(0, 1).text = "IL SEGRETARIO"
        
        # Nomi
        presidente = data.get('presidente', '[PRESIDENTE]')
        segretario = data.get('segretario', '[SEGRETARIO]')
        table.cell(1, 0).text = presidente
        table.cell(1, 1).text = segretario
        
        # Linee per le firme
        table.cell(2, 0).text = "_" * 30
        table.cell(2, 1).text = "_" * 30
        
        # Centra la tabella
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


# Registra il template nel factory
DocumentTemplateFactory.register_template("verbale_assemblea_irregolare", VerbaleAssembleaIrregolareTemplate)