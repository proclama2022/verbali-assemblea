"""
Template per Verbale di Assemblea - Nomina Amministratore Unico
Questo template è specifico per verbali di nomina dell'amministratore unico della società.
"""

import sys
import os

# Aggiungi il path della cartella src (relativo alla root del progetto)
current_dir = os.path.dirname(os.path.abspath(__file__))
src_path = os.path.join(os.path.dirname(current_dir), 'src')
if src_path not in sys.path:
    sys.path.append(src_path)

from document_templates import DocumentTemplate, DocumentTemplateFactory
from common_data_handler import CommonDataHandler
from base_verbale_template import BaseVerbaleTemplate
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date
import streamlit as st
import pandas as pd
import re

class VerbaleAmministratoreUnicoTemplate(BaseVerbaleTemplate):
    """Template per Verbale di Assemblea dei Soci - Nomina Amministratore Unico"""
    
    def get_template_name(self) -> str:
        return "Nomina Amministratore Unico"
    
    def get_required_fields(self) -> list:
        return [
            "denominazione", "sede_legale", "capitale_sociale", "codice_fiscale",
            "data_assemblea", "ora_assemblea", "presidente", "segretario", 
            "soci", "amministratore_unico"
        ]
    
    def get_form_fields(self, extracted_data: dict) -> dict:
        """Genera i campi del form per Streamlit."""
        # Chiama il metodo della classe base per ottenere i campi comuni
        form_data = super().get_form_fields(extracted_data)

        st.subheader("👤 Configurazioni Nomina Amministratore Unico")

        # Campi specifici del template
        col1, col2 = st.columns(2)
        with col1:
            form_data['motivo_nomina'] = st.selectbox(
                "Motivo della nomina",
                ["Dimissioni organo precedente", "Scadenza mandato", "Prima nomina", "Altro"],
                key="motivo_nomina_au"
            )
        with col2:
            form_data['durata_incarico'] = st.selectbox(
                "Durata dell'incarico",
                ["A tempo indeterminato", "Tre esercizi", "Un esercizio", "Altro"],
                key="durata_incarico_au"
            )

        st.subheader("Dati del nuovo Amministratore Unico")
        # Dati anagrafici
        col1, col2 = st.columns(2)
        with col1:
            nuovo_amministratore_nome = st.text_input("Nome e Cognome", key="nuovo_admin_nome_au")
            nuovo_amministratore_cf = st.text_input("Codice Fiscale", key="nuovo_admin_cf_au")
        with col2:
            nuovo_amministratore_luogo_nascita = st.text_input("Luogo di nascita", key="nuovo_admin_luogo_nascita_au")
            nuovo_amministratore_data_nascita = st.date_input("Data di nascita", value=None, key="nuovo_admin_data_nascita_au")
        nuovo_amministratore_domicilio = st.text_input("Domicilio", key="nuovo_admin_domicilio_au")

        # Raggruppa i dati dell'amministratore in un dizionario
        form_data['amministratore_unico'] = {
            "nome": nuovo_amministratore_nome,
            "codice_fiscale": nuovo_amministratore_cf,
            "luogo_nascita": nuovo_amministratore_luogo_nascita,
            "data_nascita": nuovo_amministratore_data_nascita,
            "domicilio": nuovo_amministratore_domicilio,
        }
        
        # Compenso
        st.subheader("💰 Compenso Amministratore")
        form_data['include_compensi'] = st.checkbox("È previsto un compenso?", value=True, key="include_compensi_au")
        if form_data.get('include_compensi'):
            col1, col2 = st.columns(2)
            with col1:
                form_data['compenso_annuo'] = st.text_input(
                    "Importo compenso annuo lordo (€)",
                    "0,00",
                    key="compenso_annuo_au"
                )
            with col2:
                form_data['rimborso_spese'] = st.checkbox(
                    "Includi rimborso spese documentate",
                    value=True,
                    key="rimborso_spese_au"
                )

            form_data['modalita_liquidazione'] = st.selectbox(
                "Modalità liquidazione",
                ["Periodicamente", "Annuale", "A fine mandato"],
                key="modalita_liquidazione_au"
            )
        
        # Opzioni avanzate che erano presenti in versioni precedenti e sono utili alla generazione
        st.subheader("⚙️ Opzioni Avanzate del Verbale")
        col1, col2 = st.columns(2)
        with col1:
            form_data["articolo_statuto_presidenza"] = st.text_input("Articolo statuto per presidenza", "15", key="art_presidenza_au")
            form_data["articolo_statuto_compensi"] = st.text_input("Articolo statuto per compensi", "20", key="art_compensi_au")
        with col2:
             form_data["socio_proponente"] = st.text_input("Socio proponente la nomina", "tutti i soci", key="socio_proponente_au")

        return form_data
    
    def show_preview(self, form_data: dict):
        """Mostra l'anteprima del documento fuori dal form"""
        # Sezione Anteprima
        st.markdown("---")
        st.subheader("👁️ Anteprima Documento")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col1:
            show_preview = st.checkbox("Mostra anteprima", value=False, key="preview_checkbox_admin_unico")
        
        with col2:
            if show_preview:
                st.info("💡 L'anteprima si aggiorna automaticamente con i dati inseriti sopra")
        
        if show_preview:
            try:
                # Debug dettagliato: mostra i dati estratti
                with st.expander("Debug: Dati estratti dal form"):
                    st.json(form_data)
                    
                    # Debug specifico per i soci
                    if 'soci' in form_data and form_data['soci']:
                        st.subheader("Debug Soci:")
                        for i, socio in enumerate(form_data['soci']):
                            st.write(f"Socio {i+1}:")
                            st.write(f"  - Nome: {socio.get('nome', 'N/A')}")
                            st.write(f"  - Tipo Soggetto: {socio.get('tipo_soggetto', 'N/A')}")
                            st.write(f"  - Tipo Partecipazione: {socio.get('tipo_partecipazione', 'N/A')}")
                            st.write(f"  - Delegato: {socio.get('delegato', 'N/A')}")
                            st.write(f"  - Rappresentante Legale: {socio.get('rappresentante_legale', 'N/A')}")
                            st.write(f"  - Presente: {socio.get('presente', 'N/A')}")
                            st.write("---")
                
                preview_text = self._generate_preview_text(form_data)
                st.text_area(
                    "Contenuto del verbale:",
                    value=preview_text,
                    height=600,
                    key="preview_text_amministratore_unico"
                )
            except Exception as e:
                st.error(f"Errore nell'anteprima: {e}")
                st.exception(e)
    
    def _generate_preview_text(self, data: dict) -> str:
        """Genera il testo di anteprima del verbale"""
        try:
            # Header dell'azienda con gestione dati mancanti
            denominazione = data.get('denominazione', '[Denominazione]')
            sede_legale = data.get('sede_legale', '[Sede]')
            
            # Gestione capitale sociale - usa il capitale versato se disponibile, altrimenti quello deliberato
            capitale_sociale_raw = data.get('capitale_versato') or data.get('capitale_deliberato') or data.get('capitale_sociale', '[Capitale]')
            capitale_sociale = CommonDataHandler.format_currency(capitale_sociale_raw)
            
            codice_fiscale = data.get('codice_fiscale', '[CF]')
            data_assemblea = data.get('data_assemblea', '[Data]')
            ora_assemblea = data.get('ora_assemblea', '[Ora]')
            
            # Formatta le date se presenti
            data_assemblea_str = data_assemblea.strftime('%d/%m/%Y') if hasattr(data_assemblea, 'strftime') else data_assemblea
            ora_assemblea_str = ora_assemblea.strftime('%H:%M') if hasattr(ora_assemblea, 'strftime') else ora_assemblea
            
            header = f"""{denominazione}
Sede in {sede_legale}
Capitale sociale Euro {capitale_sociale} i.v.
Codice fiscale: {codice_fiscale}

Verbale di assemblea dei soci
del {data_assemblea_str}

Oggi {data_assemblea_str} alle ore {ora_assemblea_str} presso la sede sociale {sede_legale}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:

Ordine del giorno
1. nomina dell'amministratore della società"""
            
            if data.get('include_compensi', True):
                header += "\n2. attribuzione di compensi all'amministratore della società"
            
            # Presidenza
            presidente = data.get('presidente', '[Presidente]')
            ruolo_presidente = data.get('ruolo_presidente', 'Amministratore Unico')
            articolo_statuto_presidenza = data.get('articolo_statuto_presidenza', '15')
            
            presidente_section = f"""
Assume la presidenza ai sensi dell'art. {articolo_statuto_presidenza} dello statuto sociale il Sig. {presidente} {ruolo_presidente}, il quale dichiara e constata:"""

            # Audioconferenza se inclusa
            if data.get('include_audioconferenza', True):
                articolo_statuto_audioconferenza = data.get('articolo_statuto_audioconferenza', '16')
                presidente_section += f"""

1 - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. {articolo_statuto_audioconferenza} dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza"""

            presidente_line = None
            for amm in data.get('amministratori', []):
                nome_amm = amm.get('nome', '').strip().lower()
                if nome_amm == str(presidente).strip().lower():
                    if amm.get('presente', True):
                        presidente_line = (
                            f"l'Amministratore Unico nella persona del suddetto Presidente Sig. {presidente}"
                        )
                    elif amm.get('assente_giustificato', False):
                        presidente_line = (
                            f"assente giustificato il Sig. {presidente} Amministratore Unico il quale ha tuttavia rilasciato apposita dichiarazione scritta"
                        )
                    else:
                        presidente_line = f"assente il Sig. {presidente} Amministratore Unico"
                    break

            # Se non è stato trovato alcun amministratore coincidente con il presidente,
            # controlliamo l'eventuale dizionario `amministratore_unico` presente nei dati
            if not presidente_line and 'amministratore_unico' in data:
                admin_unico = data.get('amministratore_unico', {})
                nome_admin = str(admin_unico.get('nome', '')).strip().lower()
                if nome_admin == str(presidente).strip().lower():
                    if admin_unico.get('presente', True):
                        presidente_line = (
                            f"l'Amministratore Unico nella persona del suddetto Presidente Sig. {presidente}"
                        )
                    elif admin_unico.get('assente_giustificato', False):
                        presidente_line = (
                            f"assente giustificato il Sig. {presidente} Amministratore Unico il quale ha tuttavia rilasciato apposita dichiarazione scritta"
                        )
                    else:
                        presidente_line = f"assente il Sig. {presidente} Amministratore Unico"

            if not presidente_line:
                presidente_line = (
                    f"l'Amministratore Unico nella persona del suddetto Presidente Sig. {presidente}"
                )

            presidente_section += f"""

2 - che sono presenti/partecipano all'assemblea:
{presidente_line}"""
            
            # Collegio sindacale se presente
            if data.get('include_collegio_sindacale', False):
                tipo_organo_controllo = data.get('tipo_organo_controllo', 'Collegio Sindacale')
                
                if tipo_organo_controllo == 'Collegio Sindacale':
                    sindaci = data.get('sindaci', [])
                    sindaci_presenti = [s for s in sindaci if s.get('presente')]
                    if sindaci_presenti:
                        presidente_section += "\n\nper il Collegio Sindacale:"
                        for sindaco in sindaci_presenti:
                            carica = sindaco.get('carica', 'Sindaco Effettivo')
                            nome_sindaco = sindaco.get('nome', '')
                            if nome_sindaco:
                                presidente_section += f"\nil Dott. {nome_sindaco} - {carica}"
                else: # Sindaco Unico
                    sindaci = data.get('sindaci', [])
                    if sindaci and sindaci[0].get('nome'):
                        sindaco_unico_nome = sindaci[0].get('nome')
                        presidente_section += f"\n\nil Sindaco Unico nella persona del Sig. {sindaco_unico_nome}"
            
            # Revisore se presente
            if data.get('include_revisore', False):
                tipo_revisore = data.get('tipo_revisore', 'Revisore contabile')
                nome_revisore = data.get('nome_revisore', '')
                
                if nome_revisore:
                    if tipo_revisore == 'Revisore contabile':
                        presidente_section += f"\n\nil revisore contabile Dott. {nome_revisore}"
                    else:
                        presidente_section += f"\n\nil dott. {nome_revisore} in rappresentanza della società di revisione incaricata del controllo contabile"
            
            # Soci presenti
            soci_presenti = data.get('soci_presenti', [])
            soci_assenti = data.get('soci_assenti', [])

            # Fallback per mantenere compatibilità se le nuove chiavi non ci sono
            if not soci_presenti and not soci_assenti and 'soci' in data:
                soci_presenti = [s for s in data.get('soci', []) if s.get('presente', True)]
                soci_assenti = [s for s in data.get('soci', []) if not s.get('presente', True)]
            
            total_quota_euro = 0.0
            total_quota_percentuale = 0.0
            
            # Calcola totali
            capitale_sociale_float = 0.0
            try:
                capitale_per_calcoli = data.get('capitale_versato') or data.get('capitale_deliberato') or data.get('capitale_sociale', '0')
                capitale_sociale_str = str(capitale_per_calcoli).replace('.', '').replace(',', '.')
                capitale_sociale_float = float(capitale_sociale_str)
            except ValueError:
                pass
                
            for socio in soci_presenti:
                if isinstance(socio, dict):
                    try:
                        quota_euro_str = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                        if quota_euro_str and quota_euro_str != '0':
                            total_quota_euro += float(quota_euro_str)
                    except ValueError:
                        pass
                    
                    quota_percentuale_raw = socio.get('quota_percentuale', '')
                    if quota_percentuale_raw and str(quota_percentuale_raw).strip() != '':
                        try:
                            perc_clean = str(quota_percentuale_raw).replace('%', '').replace(',', '.').strip()
                            total_quota_percentuale += float(perc_clean)
                        except ValueError:
                            pass
                    else:
                        try:
                            quota_euro_val = float(str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.'))
                            if quota_euro_val > 0 and capitale_sociale_float > 0:
                                perc_individuale = (quota_euro_val / capitale_sociale_float * 100)
                                total_quota_percentuale += perc_individuale
                        except ValueError:
                            pass

            # Formatta i totali per la visualizzazione
            formatted_total_quota_euro = CommonDataHandler.format_currency(total_quota_euro)
            formatted_total_quota_percentuale = CommonDataHandler.format_percentage(total_quota_percentuale)

            # Sezione soci - inserita dopo il presidente e prima della sezione partecipanti
            soci_section = ""
            if soci_presenti:
                soci_section = f"""
nonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_total_quota_euro} pari al {formatted_total_quota_percentuale} del Capitale Sociale:"""
                
                for socio in soci_presenti:
                    if isinstance(socio, dict):
                        nome = socio.get('nome', '[Nome Socio]')
                        
                        quota_euro_raw = socio.get('quota_euro', '')
                        quota_percentuale_raw = socio.get('quota_percentuale', '')
                        
                        if not quota_euro_raw or str(quota_euro_raw).strip() == '':
                            quota = '[Quota]'
                        else:
                            quota = CommonDataHandler.format_currency(quota_euro_raw)
                        
                        if quota_percentuale_raw and str(quota_percentuale_raw).strip() != '':
                            try:
                                perc_clean = str(quota_percentuale_raw).replace('%', '').replace(',', '.').strip()
                                perc_val = float(perc_clean)
                                percentuale = CommonDataHandler.format_percentage(perc_val)
                            except ValueError:
                                percentuale = str(quota_percentuale_raw)
                                if not percentuale.endswith('%'):
                                    percentuale += '%'
                        else:
                            quota_euro_val = 0.0
                            try:
                                quota_euro_val = float(str(quota_euro_raw).replace('.', '').replace(',', '.'))
                            except ValueError:
                                pass
                            
                            if quota_euro_val > 0 and capitale_sociale_float > 0:
                                quota_percentuale_individuale = (quota_euro_val / capitale_sociale_float * 100)
                                percentuale = CommonDataHandler.format_percentage(quota_percentuale_individuale)
                            else:
                                percentuale = '[Percentuale]'
                        
                        tipo_soggetto = socio.get('tipo_soggetto', 'Persona Fisica')
                        tipo_partecipazione = socio.get('tipo_partecipazione', 'Diretta')
                        
                        if tipo_partecipazione == 'Delegato':
                            delegato = socio.get('delegato', '')
                            if delegato:
                                soci_section += f"\nil Sig {nome} delegato del socio Sig {delegato} recante una quota pari a nominali euro {quota} pari al {percentuale} del Capitale Sociale"
                            else:
                                soci_section += f"\nil Sig {nome} socio recante una quota pari a nominali euro {quota} pari al {percentuale} del Capitale Sociale"
                        else:
                            soci_section += f"\nil Sig {nome} socio recante una quota pari a nominali euro {quota} pari al {percentuale} del Capitale Sociale"

            if soci_assenti:
                soci_section += "\n\nRisultano invece assenti i seguenti soci:"
                for socio in soci_assenti:
                    if isinstance(socio, dict) and socio.get('nome'):
                        soci_section += f"\n- il Sig. {socio.get('nome')}"

            # Concludi la sezione dei partecipanti
            partecipanti_section = """

3 - che gli intervenuti sono legittimati alla presente assemblea;
4 - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno."""
            
            # Segretario
            segretario = data.get('segretario', '[SEGRETARIO]')
            segretario_section = f"""
I presenti all'unanimità chiamano a fungere da segretario il signor {segretario}, che accetta l'incarico.

Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante."""
            
            # Tipo assemblea
            tipo_assemblea = data.get('tipo_assemblea', 'regolarmente convocata')
            
            segretario_section += f"""

Il Presidente constata e fa constatare che l'assemblea risulta {tipo_assemblea} e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.

Si passa quindi allo svolgimento dell'ordine del giorno.

*     *     *"""
            
            # Discussione nomina amministratore unico
            motivo_nomina = data.get('motivo_nomina', 'Dimissioni dell\'organo in carica')
            admin_unico = data.get('amministratore_unico', {})
            nome_admin = admin_unico.get('nome', '')
            socio_proponente = data.get('socio_proponente', '')
            
            nomina_section = f"""
Il Presidente informa l'assemblea che si rende necessaria la nomina di un nuovo organo amministrativo per {motivo_nomina.lower()}. 

Il Presidente ricorda all'assemblea quanto previsto dall'art. 2475 del Codice Civile e dall'atto costitutivo della società."""

            if socio_proponente:
                nomina_section += f"""

Prende la parola il socio sig. {socio_proponente} che propone di nominare Amministratore Unico della società il sig. {nome_admin}, dando evidenza della comunicazione scritta con cui il candidato, prima di accettare l'eventuale nomina, ha dichiarato:"""
            else:
                nomina_section += f"""

Il Presidente propone di nominare Amministratore Unico della società il sig. {nome_admin}, dando evidenza della comunicazione scritta con cui il candidato, prima di accettare l'eventuale nomina, ha dichiarato:"""
                
            nomina_section += """

l'insussistenza a suo carico di cause di ineleggibilità alla carica di amministratore di società ed in particolare di non essere stato dichiarato interdetto, inabilitato o fallito e di non essere stato condannato ad una pena che importa l'interdizione, anche temporanea, dai pubblici uffici o l'incapacità ad esercitare uffici direttivi. 

l'insussistenza a suo carico di interdizioni dal ruolo di amministratore adottate da una Stato membro dell'Unione Europea."""
            
            # Compensi se inclusi
            compensi_section = ""
            if data.get('include_compensi', True):
                articolo_statuto_compensi = data.get('articolo_statuto_compensi', '20')
                compensi_section = f"""

Il Presidente invita anche l'assemblea a deliberare il compenso da attribuire all'organo amministrativo che verrà nominato, ai sensi dell'art. {articolo_statuto_compensi} dello statuto sociale."""
            
            # Deliberazione
            durata_incarico = data.get('durata_incarico', 'A tempo indeterminato fino a revoca o dimissioni')
            compenso_annuo = data.get('compenso_annuo', '0,00')
            tipo_votazione = data.get('tipo_votazione', 'Unanimità')
            
            deliberazione_section = """
Segue breve discussione tra i soci al termine della quale si passa alla votazione con voto palese in forza della quale il Presidente constata che,"""
            
            if tipo_votazione == 'Unanimità':
                deliberazione_section += " all'unanimità,"
            else:
                contrari = data.get('contrari', '')
                astenuti = data.get('astenuti', '')
                
                if contrari:
                    deliberazione_section += f" con il voto contrario dei Sigg. {contrari}"
                    if astenuti:
                        deliberazione_section += f" e l'astensione dei Sigg. {astenuti},"
                    else:
                        deliberazione_section += ","
                elif astenuti:
                    deliberazione_section += f" con l'astensione dei Sigg. {astenuti},"
            
            # Recupera tutti i dati dell'amministratore
            amministratore_unico = data.get('amministratore_unico', {})
            nome_admin = amministratore_unico.get('nome', '[Nome Amministratore]')
            cf_admin = amministratore_unico.get('codice_fiscale', '[CF]')
            nato_a = amministratore_unico.get('luogo_nascita', '[Luogo di Nascita]')
            data_nascita_raw = amministratore_unico.get('data_nascita')
            nato_il = data_nascita_raw.strftime('%d/%m/%Y') if hasattr(data_nascita_raw, 'strftime') else '[Data di Nascita]'
            domicilio = amministratore_unico.get('domicilio', '[Domicilio]')

            deliberazione_section += f"""

l'assemblea

d e l i b e r a:

di nominare quale Amministratore Unico della società il Sig. {nome_admin}, nato a {nato_a} il {nato_il}, codice fiscale {cf_admin} e residente in {domicilio}, il quale, presente all'assemblea, dichiara di accettare la carica e di non trovarsi in alcuna delle cause di ineleggibilità o di incompatibilità previste dalla legge e dallo statuto sociale.

che l'amministratore resti in carica {durata_incarico.lower()}"""

            if data.get('include_compensi', True):
                rimborso_text = " oltre al rimborso delle spese sostenute in ragione del suo ufficio" if data.get('rimborso_spese', True) else ""
                modalita_liquidazione = data.get('modalita_liquidazione', 'periodicamente')
                
                deliberazione_section += f"""

di attribuire all'amministratore unico testè nominato il compenso annuo ed omnicomprensivo pari a nominali euro {compenso_annuo} al lordo di ritenute fiscali e previdenziali{rimborso_text}. Il compenso verrà liquidato {modalita_liquidazione.lower()}, in ragione della permanenza in carica."""
            
            # Accettazione
            qualifica = admin_unico.get('qualifica', 'socio')
            accettazione_section = f"""

Il sig. {nome_admin}, presente in assemblea in qualità di {qualifica.lower()}, accetta l'incarico e ringrazia l'assemblea per la fiducia accordata."""
            
            # Chiusura
            ora_chiusura = data.get('ora_chiusura', data.get('ora_assemblea', '[Ora]'))
            ora_chiusura_str = ora_chiusura.strftime('%H:%M') if hasattr(ora_chiusura, 'strftime') else ora_chiusura
            
            chiusura_section = f"""

*     *     *

Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola. 

Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo. 

L'assemblea viene sciolta alle ore {ora_chiusura_str}.


Il Presidente                    Il Segretario
{presidente}            {segretario}"""
            
            # Combina tutte le sezioni
            full_text = (header + presidente_section + soci_section + partecipanti_section + 
                        segretario_section + nomina_section + compensi_section + 
                        deliberazione_section + accettazione_section + chiusura_section)
            
            return full_text
            
        except Exception as e:
            return f"Errore nella generazione dell'anteprima: {str(e)}"
    
    def generate_document(self, data: dict) -> Document:
        """Genera il documento Word del verbale"""
        import os
        
        # Utilizza il template .docx esistente per mantenere la formattazione
        template_path = os.path.join(os.path.dirname(__file__), 'template.docx')
        
        try:
            if os.path.exists(template_path):
                doc = Document(template_path)
                # Rimuovi il contenuto esistente del template mantenendo gli stili
                for paragraph in doc.paragraphs[:]:
                    p = paragraph._element
                    p.getparent().remove(p)
                # Setup stili del documento anche quando usiamo template
                self._setup_document_styles(doc)
            else:
                doc = Document()
                # Setup stili del documento
                self._setup_document_styles(doc)
        except Exception as e:
            # Fallback a documento vuoto se il template non può essere caricato
            doc = Document()
            self._setup_document_styles(doc)
        
        # Assicurati che la variabile soci sia definita per evitare errori
        if 'soci' not in data:
            data['soci'] = []
        
        # Aggiungi header azienda
        self._add_company_header(doc, data)
        
        # Aggiungi titolo verbale
        self._add_verbale_title(doc, data)
        
        # Aggiungi sezione di apertura
        self._add_opening_section(doc, data)
        
        # Aggiungi partecipanti
        self._add_participants_section(doc, data)
        
        # Aggiungi discussione nomina amministratore unico
        self._add_nomination_discussion(doc, data)
        
        # Aggiungi sezione di chiusura
        self._add_closing_section(doc, data)
        
        # Aggiungi firme
        self._add_signatures(doc, data)
        
        return doc

    def _setup_document_styles(self, doc):
        """Imposta gli stili del documento"""
        styles = doc.styles
        
        # Stile per il titolo della società
        if 'TitoloSocieta' not in styles:
            title_style = styles.add_style('TitoloSocieta', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Times New Roman'
            title_style.font.size = Pt(14)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(6)
        
        # Stile per il corpo del testo
        if 'BodyText' not in styles:
            body_style = styles.add_style('BodyText', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = 'Times New Roman'
            body_style.font.size = Pt(12)
            body_style.paragraph_format.space_after = Pt(6)
            body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _add_company_header(self, doc, data):
        """Aggiunge l'header con i dati della società"""
        try:
            p = doc.add_paragraph(style='TitoloSocieta')
        except KeyError:
            p = doc.add_paragraph()
        
        run = p.add_run(data.get('denominazione', '[DENOMINAZIONE]'))
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        try:
            p = doc.add_paragraph(style='TitoloSocieta')
        except KeyError:
            p = doc.add_paragraph()
        
        run = p.add_run(f"Sede in {data.get('sede_legale', '[SEDE]')}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        try:
            p = doc.add_paragraph(style='TitoloSocieta')
        except KeyError:
            p = doc.add_paragraph()
        
        capitale_sociale = CommonDataHandler.format_currency(data.get('capitale_sociale', '[CAPITALE]'))
        run = p.add_run(f"Capitale sociale Euro {capitale_sociale} i.v.")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        try:
            p = doc.add_paragraph(style='TitoloSocieta')
        except KeyError:
            p = doc.add_paragraph()
        
        run = p.add_run(f"Codice fiscale: {data.get('codice_fiscale', '[CF]')}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_verbale_title(self, doc, data):
        """Aggiunge il titolo del verbale"""
        doc.add_paragraph()
        
        try:
            p = doc.add_paragraph(style='TitoloSocieta')
        except KeyError:
            p = doc.add_paragraph()
        
        run = p.add_run("Verbale di assemblea dei soci")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        data_str = data.get('data_assemblea').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[DATA]'
        
        try:
            p = doc.add_paragraph(style='TitoloSocieta')
        except KeyError:
            p = doc.add_paragraph()
        
        run = p.add_run(f"del {data_str}")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_opening_section(self, doc, data):
        """Aggiunge la sezione di apertura del verbale"""
        doc.add_paragraph()
        
        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()
        
        # Format data assemblea
        data_val = data.get('data_assemblea', '[DATA]')
        if hasattr(data_val, 'strftime'):
            data_str = data_val.strftime('%d/%m/%Y')
        else:
            data_str = str(data_val)
        
        # Format ora assemblea - usa ora_chiusura se non c'è ora_assemblea
        ora_val = data.get('ora_assemblea') or data.get('ora_chiusura', '[ORA]')
        if hasattr(ora_val, 'strftime'):
            ora_str = ora_val.strftime('%H:%M')
        else:
            ora_str = str(ora_val) or '[ORA]'
        
        sede = data.get('sede_legale', '[SEDE]')
        
        opening_text = f"Oggi {data_str} alle ore {ora_str} presso la sede sociale {sede}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:"
        
        run = p.add_run(opening_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        # Ordine del giorno
        doc.add_paragraph()
        
        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()
        
        run = p.add_run("Ordine del giorno")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()
        
        run = p.add_run("1. nomina dell'amministratore della società")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        if data.get('include_compensi', True):
            try:
                p = doc.add_paragraph(style='BodyText')
            except KeyError:
                p = doc.add_paragraph()
            
            run = p.add_run("2. attribuzione di compensi all'amministratore della società")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    def _add_participants_section(self, doc, data):
        """Aggiunge la sezione dei partecipanti"""
        doc.add_paragraph()
        
        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()

        presidente = data.get('presidente', '[PRESIDENTE]')
        ruolo_presidente = data.get('ruolo_presidente', 'Amministratore Unico')
        art_pres = data.get('articolo_statuto_presidenza', '15')

        intro_text = (
            f"Assume la presidenza ai sensi dell'art. {art_pres} dello statuto sociale "
            f"il Sig. {presidente} {ruolo_presidente}, il quale dichiara e constata:"
        )
        run = p.add_run(intro_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        bullet_num = 1
        if data.get('include_audioconferenza', True):
            art_audio = data.get('articolo_statuto_audioconferenza', '16')
            try:
                p = doc.add_paragraph(style='BodyText')
            except KeyError:
                p = doc.add_paragraph()
            run = p.add_run(
                f"{bullet_num} - che (come indicato anche nell'avviso di convocazione ed in conformità "
                f"alle previsioni dell'art. {art_audio} dello statuto sociale) l'intervento all'assemblea "
                "può avvenire anche in audioconferenza"
            )
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            bullet_num += 1

        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()
        run = p.add_run(f"{bullet_num} - che sono presenti/partecipano all'assemblea:")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        presidente_line = None
        for amm in data.get('amministratori', []):
            nome_amm = amm.get('nome', '').strip().lower()
            if nome_amm == str(presidente).strip().lower():
                if amm.get('presente', True):
                    presidente_line = f"l'{ruolo_presidente} nella persona del suddetto Presidente Sig. {presidente}"
                elif amm.get('assente_giustificato', False):
                    presidente_line = (
                        f"assente giustificato il Sig. {presidente} {ruolo_presidente} "
                        "il quale ha tuttavia rilasciato apposita dichiarazione scritta"
                    )
                else:
                    presidente_line = f"assente il Sig. {presidente} {ruolo_presidente}"
                break

        if not presidente_line and 'amministratore_unico' in data:
            admin_unico = data.get('amministratore_unico', {})
            nome_admin = str(admin_unico.get('nome', '')).strip().lower()
            if nome_admin == str(presidente).strip().lower():
                if admin_unico.get('presente', True):
                    presidente_line = f"l'{ruolo_presidente} nella persona del suddetto Presidente Sig. {presidente}"
                elif admin_unico.get('assente_giustificato', False):
                    presidente_line = (
                        f"assente giustificato il Sig. {presidente} {ruolo_presidente} "
                        "il quale ha tuttavia rilasciato apposita dichiarazione scritta"
                    )
                else:
                    presidente_line = f"assente il Sig. {presidente} {ruolo_presidente}"

        if not presidente_line:
            presidente_line = f"l'{ruolo_presidente} nella persona del suddetto Presidente Sig. {presidente}"

        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()
        run = p.add_run(presidente_line)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        if data.get('include_collegio_sindacale', False):
            tipo_organo_controllo = data.get('tipo_organo_controllo', 'Collegio Sindacale')
            if tipo_organo_controllo == 'Collegio Sindacale':
                sindaci = [s for s in data.get('sindaci', []) if s.get('presente')]
                if sindaci:
                    p = doc.add_paragraph("per il Collegio Sindacale:")
                    for s in sindaci:
                        doc.add_paragraph(f"- {s.get('nome', '[NOME]')} {s.get('carica', '')}")
            else:
                sindaci = data.get('sindaci', [])
                if sindaci and sindaci[0].get('nome'):
                    doc.add_paragraph(f"il Sindaco Unico {sindaci[0].get('nome')}")

        if data.get('include_revisore', False):
            nome_rev = data.get('nome_revisore', '[NOME REVISORE]')
            doc.add_paragraph(f"il revisore contabile Dott. {nome_rev}")

        soci = data.get('soci', [])
        if soci:
            self._add_soci_section(doc, data)

    def _add_soci_section(self, doc, data):
        """Aggiunge la sezione dettagliata dei soci"""
        soci_presenti = data.get('soci_presenti', [])
        soci_assenti = data.get('soci_assenti', [])

        # Fallback per mantenere compatibilità se le nuove chiavi non ci sono
        if not soci_presenti and not soci_assenti and 'soci' in data:
            soci_presenti = [s for s in data.get('soci', []) if s.get('presente', True)]
            soci_assenti = [s for s in data.get('soci', []) if not s.get('presente', True)]
        
        if not soci_presenti:
            return
        
        # Calcola totali
        total_quota_euro = 0.0
        total_quota_percentuale = 0.0
        
        # Parse capitale_sociale as float
        capitale_sociale_float = 0.0
        try:
            capitale_per_calcoli = data.get('capitale_versato') or data.get('capitale_deliberato') or data.get('capitale_sociale', '0')
            capitale_sociale_str = str(capitale_per_calcoli).replace('.', '').replace(',', '.')
            capitale_sociale_float = float(capitale_sociale_str)
        except ValueError:
            pass
        
        for socio in soci_presenti:
            if isinstance(socio, dict):
                # Calcola totale quote in euro
                try:
                    quota_euro_str = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                    if quota_euro_str and quota_euro_str != '0':
                        total_quota_euro += float(quota_euro_str)
                except ValueError:
                    pass
                
                # Calcola totale percentuali
                quota_percentuale_raw = socio.get('quota_percentuale', '')
                if quota_percentuale_raw and str(quota_percentuale_raw).strip() != '':
                    try:
                        perc_clean = str(quota_percentuale_raw).replace('%', '').replace(',', '.').strip()
                        total_quota_percentuale += float(perc_clean)
                    except ValueError:
                        pass
                else:
                    # Calcola basandosi sulla quota euro
                    try:
                        quota_euro_val = float(str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.'))
                        if quota_euro_val > 0 and capitale_sociale_float > 0:
                            perc_individuale = (quota_euro_val / capitale_sociale_float * 100)
                            total_quota_percentuale += perc_individuale
                    except ValueError:
                        pass
        
        # Formatta i totali
        formatted_total_quota_euro = CommonDataHandler.format_currency(total_quota_euro)
        formatted_total_quota_percentuale = CommonDataHandler.format_percentage(total_quota_percentuale)
        
        # Intestazione soci
        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()
        
        run = p.add_run(f"nonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_total_quota_euro} pari al {formatted_total_quota_percentuale} del Capitale Sociale:")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        # Lista dei soci
        for socio in soci_presenti:
            if isinstance(socio, dict):
                nome = socio.get('nome', '[Nome Socio]')
                
                # Gestione quote
                quota_euro_raw = socio.get('quota_euro', '')
                quota_percentuale_raw = socio.get('quota_percentuale', '')
                
                # Formatta quota in euro
                if not quota_euro_raw or str(quota_euro_raw).strip() == '':
                    quota = '[Quota]'
                else:
                    quota = CommonDataHandler.format_currency(quota_euro_raw)
                
                # Gestione percentuale
                if quota_percentuale_raw and str(quota_percentuale_raw).strip() != '':
                    try:
                        perc_clean = str(quota_percentuale_raw).replace('%', '').replace(',', '.').strip()
                        perc_val = float(perc_clean)
                        percentuale = CommonDataHandler.format_percentage(perc_val)
                    except ValueError:
                        percentuale = str(quota_percentuale_raw)
                        if not percentuale.endswith('%'):
                            percentuale += '%'
                else:
                    # Calcola la percentuale
                    quota_euro_val = 0.0
                    try:
                        quota_euro_val = float(str(quota_euro_raw).replace('.', '').replace(',', '.'))
                    except ValueError:
                        pass
                    
                    if quota_euro_val > 0 and capitale_sociale_float > 0:
                        quota_percentuale_individuale = (quota_euro_val / capitale_sociale_float * 100)
                        percentuale = CommonDataHandler.format_percentage(quota_percentuale_individuale)
                    else:
                        percentuale = '[Percentuale]'
                
                tipo_soggetto = socio.get('tipo_soggetto', 'Persona Fisica')
                
                # Crea il paragrafo per il socio
                try:
                    p = doc.add_paragraph(style='BodyText')
                except KeyError:
                    p = doc.add_paragraph()
                
                if tipo_soggetto == 'Società':
                    socio_text = f"la società {nome} socio recante una quota pari a nominali euro {quota} pari al {percentuale} del Capitale Sociale"
                else:
                    socio_text = f"il Sig {nome} socio recante una quota pari a nominali euro {quota} pari al {percentuale} del Capitale Sociale"
                
                run = p.add_run(socio_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        
        if soci_assenti:
            p = doc.add_paragraph()
            p.add_run("Risultano invece assenti i seguenti soci:")
            for socio in soci_assenti:
                if isinstance(socio, dict) and socio.get('nome'):
                    p = doc.add_paragraph(f"- il Sig. {socio.get('nome')}", style='BodyText')

    def _add_nomination_discussion(self, doc, data):
        """Aggiunge la discussione sulla nomina e la delibera."""
        
        # --- Introduzione alla discussione ---
        p = doc.add_paragraph(style='BodyText')
        p.add_run("*     *     *").alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        motivo_nomina = data.get('motivo_nomina', 'dimissioni dell\'organo in carica')
        
        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"Il Presidente informa l'assemblea che si rende necessaria la nomina di un nuovo organo amministrativo per {motivo_nomina.lower()}.")
        
        p = doc.add_paragraph(style='BodyText')
        p.add_run("Il Presidente ricorda all'assemblea quanto previsto dall'art. 2475 del Codice Civile e dall'atto costitutivo della società.")
        
        # --- Proposta di nomina ---
        socio_proponente = data.get('socio_proponente', 'Il Presidente')
        amministratore_unico = data.get('amministratore_unico', {})
        nome_admin = amministratore_unico.get('nome', '[Nome Amministratore]')
        
        p = doc.add_paragraph(style='BodyText')
        if socio_proponente.lower() != 'il presidente':
            run = p.add_run(f"Prende la parola il socio sig. {socio_proponente} che propone di nominare Amministratore Unico della società il sig. ")
        else:
            run = p.add_run("Il Presidente propone di nominare Amministratore Unico della società il sig. ")

        run = p.add_run(nome_admin)
        run.bold = True
        p.add_run(", dando evidenza della comunicazione scritta con cui il candidato, prima di accettare l'eventuale nomina, ha dichiarato:")
        
        # Dichiarazioni di insussistenza
        p = doc.add_paragraph("l'insussistenza a suo carico di cause di ineleggibilità...", style='BodyText')
        p.paragraph_format.left_indent = Inches(0.5)
        
        p = doc.add_paragraph("l'insussistenza a suo carico di interdizioni dal ruolo...", style='BodyText')
        p.paragraph_format.left_indent = Inches(0.5)

        # --- Discussione compenso (se presente) ---
        if data.get('include_compensi', True):
            articolo_statuto_compensi = data.get('articolo_statuto_compensi', '20')
            p = doc.add_paragraph(style='BodyText')
            p.add_run(f"Il Presidente invita anche l'assemblea a deliberare il compenso da attribuire all'organo amministrativo che verrà nominato, ai sensi dell'art. {articolo_statuto_compensi} dello statuto sociale.")

        # --- Votazione e delibera ---
        p = doc.add_paragraph(style='BodyText')
        p.add_run("Segue breve discussione tra i soci al termine della quale si passa alla votazione con voto palese in forza della quale il Presidente constata che, all'unanimità,")
        
        p = doc.add_paragraph("l'assemblea", style='BodyText')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = doc.add_paragraph(style='BodyText')
        run = p.add_run("d e l i b e r a:")
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # --- Punto 1: Nomina Amministratore con dati completi ---
        p = doc.add_paragraph(style='BodyText')
        p.add_run("di nominare quale Amministratore Unico della società il Sig. ")
        
        cf_admin = amministratore_unico.get('codice_fiscale', '[CF]')
        nato_a = amministratore_unico.get('luogo_nascita', '[Luogo di Nascita]')
        data_nascita_raw = amministratore_unico.get('data_nascita')
        nato_il = data_nascita_raw.strftime('%d/%m/%Y') if hasattr(data_nascita_raw, 'strftime') else '[Data di Nascita]'
        domicilio = amministratore_unico.get('domicilio', '[Domicilio]')
        
        run = p.add_run(nome_admin)
        run.bold = True
        p.add_run(f", nato a {nato_a} il {nato_il}, codice fiscale {cf_admin} e residente in {domicilio}, il quale, presente all'assemblea, dichiara di accettare la carica e di non trovarsi in alcuna delle cause di ineleggibilità o di incompatibilità previste dalla legge e dallo statuto sociale.")

        # --- Punto 2: Durata incarico ---
        durata_incarico = data.get('durata_incarico', 'A tempo indeterminato')
        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"che l'amministratore resti in carica {durata_incarico.lower()}")

        # --- Punto 3: Compenso (se presente) ---
        if data.get('include_compensi', True):
            compenso_annuo = CommonDataHandler.format_currency(data.get('compenso_annuo', '0,00'))
            rimborso_text = " oltre al rimborso delle spese sostenute in ragione del suo ufficio" if data.get('rimborso_spese', True) else ""
            
            modalita = data.get('modalita_liquidazione', 'Periodicamente').lower()

            p = doc.add_paragraph(style='BodyText')
            p.add_run(
                f"di attribuire all'amministratore unico testè nominato il compenso annuo ed omnicomprensivo pari a nominali euro {compenso_annuo} al lordo di ritenute fiscali e previdenziali{rimborso_text}. "
                f"Il compenso verrà liquidato {modalita}, in ragione della permanenza in carica."
            )

        # --- Accettazione finale ---
        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"Il sig. {nome_admin}, presente in assemblea, accetta l'incarico e ringrazia l'assemblea per la fiducia accordata.")

    def _add_closing_section(self, doc, data):
        """Aggiunge la sezione di chiusura"""
        p = doc.add_paragraph()
        
        # Ora di chiusura (usa ora_assemblea se non specificata)
        ora_val = data.get('ora_chiusura', data.get('ora_assemblea', '[ORA]'))
        if hasattr(ora_val, 'strftime'):
            ora_str = ora_val.strftime('%H:%M')
        else:
            ora_str = str(ora_val) or '[ORA]'
        
        run = p.add_run(f"L'assemblea viene sciolta alle ore {ora_str}.")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    def _add_signatures(self, doc, data):
        """Aggiunge le firme"""
        doc.add_paragraph()
        doc.add_paragraph()
        
        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()
        
        run = p.add_run("Il Presidente                    Il Segretario")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()
        
        presidente = data.get('presidente', '[PRESIDENTE]')
        segretario = data.get('segretario', '[SEGRETARIO]')
        run = p.add_run(f"{presidente}            {segretario}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# Registrazione del template nel factory
DocumentTemplateFactory.register_template('verbale_assemblea_amministratore_unico_template', VerbaleAmministratoreUnicoTemplate)
