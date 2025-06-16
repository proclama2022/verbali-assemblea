"""
Template per verbale di assemblea - Nomina Consiglio di Amministrazione
"""

import sys
import os
import pandas as pd

# Aggiungi il path della cartella src (relativo alla root del progetto)
current_dir = os.path.dirname(os.path.abspath(__file__))
src_path = os.path.join(os.path.dirname(current_dir), 'src')
if src_path not in sys.path:
    sys.path.append(src_path)

from document_templates import DocumentTemplate
from document_templates import DocumentTemplateFactory
from base_verbale_template import BaseVerbaleTemplate
from common_data_handler import CommonDataHandler
import streamlit as st
from datetime import datetime, date
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

class VerbaleConsiglioAmministrazioneTemplate(BaseVerbaleTemplate):
    """Template per verbale di nomina Consiglio di Amministrazione"""
    
    def get_template_name(self) -> str:
        return "Nomina Consiglio di Amministrazione"
    
    def get_required_fields(self) -> list:
        return ['denominazione', 'sede_legale', 'capitale_sociale', 'codice_fiscale', 
                'data_assemblea', 'ora_inizio', 'presidente', 'segretario', 'consiglieri', 'soci']
    
    def get_form_fields(self, extracted_data: dict) -> dict:
        """Genera i campi del form per Streamlit."""
        # Chiama il metodo della classe base per ottenere i campi comuni
        form_data = super().get_form_fields(extracted_data)

        st.subheader("🏛️ Configurazioni Consiglio di Amministrazione")

        # Sezioni opzionali del verbale
        col1, col2, col3 = st.columns(3)
        with col1:
            form_data['nomina_presidente'] = st.checkbox("Nomina Presidente", value=True, key="nomina_presidente_cda")
        with col2:
            form_data['attribuzione_poteri'] = st.checkbox("Attribuzione Poteri", value=True, key="attribuzione_poteri_cda")
        with col3:
            form_data['determinazione_compensi'] = st.checkbox("Determinazione Compensi", value=True, key="determinazione_compensi_cda")

        # Dettagli condizionali
        if form_data.get('nomina_presidente'):
            st.subheader("Nomina del Presidente del CdA")
            # Assumendo che gli amministratori siano già in form_data dalla classe base
            admin_names = [a.get('nome', '') for a in form_data.get('amministratori', []) if a.get('nome')]
            form_data['presidente_eletto'] = st.selectbox(
                "Seleziona il Presidente da eleggere",
                admin_names,
                key="presidente_eletto_cda"
            )

        if form_data.get('attribuzione_poteri'):
            st.subheader("Attribuzione dei Poteri")
            form_data['poteri_presidente'] = st.text_area(
                "Poteri attribuiti al Presidente",
                "Al Presidente del Consiglio di Amministrazione sono attribuiti i poteri di...",
                height=150,
                key="poteri_presidente_cda"
            )

        if form_data.get('determinazione_compensi'):
            st.subheader("Determinazione dei Compensi")
            form_data['compenso_cda'] = st.text_input(
                "Compenso annuo lordo per il CdA (€)",
                "0,00",
                key="compenso_cda"
            )
        
        return form_data
    
    def show_preview(self, form_data: dict):
        """Mostra l'anteprima del verbale"""
        st.subheader("📄 Anteprima Verbale")
        
        preview_text = self._generate_preview_text(form_data)
        
        # Mostra anteprima in un container scrollabile
        st.text_area("Anteprima del documento:", 
                    value=preview_text, 
                    height=400, 
                    disabled=True)
    
    def _generate_preview_text(self, data: dict) -> str:
        """Genera il testo di anteprima del verbale"""
        try:
            # Header azienda
            header = f"""{data.get('denominazione', '[DENOMINAZIONE SOCIETÀ]')}
Sede in {data.get('sede_legale', '[SEDE]')}
Capitale sociale Euro {data.get('capitale_sociale', '[CAPITALE]')} i.v.
Codice fiscale: {data.get('codice_fiscale', '[CODICE FISCALE]')}

Verbale di assemblea dei soci
del {data.get('data_assemblea', '[DATA]').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[DATA]'}

Oggi {data.get('data_assemblea', '[DATA]').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[DATA]'} alle ore {data.get('ora_inizio', '[ORA]')} presso la sede sociale {data.get('sede_legale', '[SEDE]')}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:

Ordine del giorno
• nomina del Consiglio di Amministrazione della società"""
            
            if data.get('include_compensi', True):
                header += "\n• attribuzione di compensi al Consiglio di Amministrazione della società"
            
            # --- Costruzione elenco soci prima della sezione presidente ---
            soci_presenti = data.get('soci_presenti', [])
            soci_assenti = data.get('soci_assenti', [])

            # Fallback
            if not soci_presenti and not soci_assenti and 'soci' in data:
                soci_presenti = [s for s in data.get('soci', []) if s.get('presente', True)]
                soci_assenti = [s for s in data.get('soci', []) if not s.get('presente', True)]

            soci_listing_lines = []
            if soci_presenti:
                for socio in soci_presenti:
                    if isinstance(socio, dict) and socio.get('nome'):
                        quota_euro = socio.get('quota_euro', '')
                        quota_perc = socio.get('quota_percentuale', '')
                        quota_text = ""
                        if quota_euro:
                            quota_text += f" euro {quota_euro}"
                        if quota_perc:
                            quota_text += f" pari al {quota_perc}%"
                        soci_listing_lines.append(f"- {socio.get('nome')} ({socio.get('tipo_soggetto', 'PF')}){quota_text}")

            soci_listing = "\n".join(soci_listing_lines) if soci_listing_lines else "- Nessun socio presente"
            
            if soci_assenti:
                soci_listing += "\n\nSoci assenti:"
                for socio in soci_assenti:
                    if isinstance(socio, dict) and socio.get('nome'):
                        soci_listing += f"\n- {socio.get('nome')}"
            
            # ---- Calcolo totali quota euro e percentuale ----
            total_quota_euro_val = 0.0
            total_quota_perc_val = 0.0

            # Capitale sociale float (serve per calcolare percentuali se non presenti)
            capitale_raw = data.get('capitale_sociale', '').replace('.', '').replace(',', '.')
            try:
                capitale_sociale_float = float(capitale_raw) if capitale_raw else 0.0
            except ValueError:
                capitale_sociale_float = 0.0

            for socio in soci_presenti:
                if isinstance(socio, dict):
                    # totale euro
                    quota_euro_raw = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                    try:
                        total_quota_euro_val += float(quota_euro_raw) if quota_euro_raw else 0.0
                    except ValueError:
                        pass

                    # totale percentuale
                    quota_perc_raw = socio.get('quota_percentuale', '')
                    if quota_perc_raw:
                        try:
                            perc_clean = str(quota_perc_raw).replace('%', '').replace(',', '.').strip()
                            total_quota_perc_val += float(perc_clean)
                        except ValueError:
                            pass
                    else:
                        # se manca percentuale calcola da euro
                        try:
                            quota_euro_val = float(quota_euro_raw) if quota_euro_raw else 0.0
                            if capitale_sociale_float > 0 and quota_euro_val > 0:
                                total_quota_perc_val += (quota_euro_val / capitale_sociale_float * 100)
                        except ValueError:
                            pass

            # Formatta totali
            total_quota_euro = CommonDataHandler.format_currency(total_quota_euro_val)
            total_quota_perc = CommonDataHandler.format_percentage(total_quota_perc_val)
            
            # Sezione presidente con soci già inseriti
            presidente_section = f"""

Assume la presidenza ai sensi dell'art. […] dello statuto sociale il Sig. {data.get('presidente', '[PRESIDENTE]')} Amministratore Unico [oppure Presidente del Consiglio di Amministrazione o altro (come da statuto)], il quale dichiara e constata:

1 - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. […] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza

2 - che sono presenti/partecipano all'assemblea:
l'Amministratore Unico nella persona del suddetto Presidente Sig. {data.get('presidente', '[PRESIDENTE]')}"""
            
            # Collegio sindacale se presente
            if data.get('include_collegio_sindacale', False):
                tipo_organo_controllo = data.get('tipo_organo_controllo', 'Collegio Sindacale')
                
                if tipo_organo_controllo == 'Collegio Sindacale':
                    sindaci_list = data.get('sindaci', [])
                    sindaci_presenti = [s for s in sindaci_list if s.get('presente')]
                    if sindaci_presenti:
                        presidente_section += "\n\nper il Collegio Sindacale:"
                        for sindaco in sindaci_presenti:
                            carica = sindaco.get('carica', 'Sindaco Effettivo')
                            nome_sindaco = sindaco.get('nome', '')
                            if nome_sindaco:
                                presidente_section += f"\nil Dott. {nome_sindaco} - {carica}"
                else: # Sindaco Unico
                    sindaci_list = data.get('sindaci', [])
                    if sindaci_list and sindaci_list[0].get('nome'):
                        sindaco_unico_nome = sindaci_list[0].get('nome')
                        presidente_section += f"\n\nil Sindaco Unico nella persona del Sig. {sindaco_unico_nome}"

            presidente_section += f"""

nonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {total_quota_euro} pari al {total_quota_perc} del Capitale Sociale:
{soci_listing}

2 - che gli intervenuti sono legittimati alla presente assemblea;
3 - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno.

I presenti all'unanimità chiamano a fungere da segretario il signor {data.get('segretario', '[SEGRETARIO]')}, che accetta l'incarico.

Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante.

Il Presidente constata e fa constatare che l'assemblea risulta regolarmente convocata [oppure totalitaria] e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.

Si passa quindi allo svolgimento dell'ordine del giorno.

*     *     *"""
            
            # Discussione nomina
            motivo_nomina = data.get('motivo_nomina', 'Dimissioni dell\'organo in carica')
            
            nomina_section = f"""

Il Presidente informa l'assemblea che si rende necessaria la nomina di un nuovo organo amministrativo [{motivo_nomina.lower()}].

Il Presidente ricorda all'assemblea quanto previsto dall'art. 2475 del Codice Civile e dall'atto costitutivo della società [verificare quanto previsto dall'atto costitutivo in tema di amministrazione].

Prende la parola il socio sig. […] che propone di affidare l'amministrazione della società ad un Consiglio di Amministrazione di {len(data.get('consiglieri', []))} membri e composto dai Sigg.:"""
            
            # Lista consiglieri
            consiglieri = data.get('consiglieri', [])
            for cons in consiglieri:
                if cons.get('nome'):
                    data_nascita_str = cons.get('data_nascita').strftime('%d/%m/%Y') if hasattr(cons.get('data_nascita'), 'strftime') else '[Data nascita]'
                    nomina_section += f"\n- {cons.get('nome', '[Nome]')}, nato a {cons.get('luogo_nascita', '[Luogo nascita]')} il {data_nascita_str}, C.F. {cons.get('codice_fiscale', '[CF]')}, residente in {cons.get('residenza', '[Residenza]')}"
            
            nomina_section += """\n\ndando evidenza della comunicazione scritta con cui i candidati, prima di accettare l'eventuale nomina, hanno dichiarato:
• l'insussistenza a loro carico di cause di ineleggibilità alla carica di amministratore di società ed in particolare di non essere stati dichiarati interdetti, inabilitati o falliti e di non essere stati condannati ad una pena che importa l'interdizione, anche temporanea, dai pubblici uffici o l'incapacità ad esercitare uffici direttivi.
• l'insussistenza a loro carico di interdizioni dal ruolo di amministratore adottate da una Stato membro dell'Unione Europea.

[verificare che l'atto costitutivo non preveda ulteriori requisiti per l'assunzione della carica e quanto previsto da leggi speciali in relazione all'esercizio di particolari attività] [se esiste il collegio sindacale o il revisore, verificare eventuali incompatibilità con i neo amministratori]."""
            
            # Compensi
            compensi_section = ""
            if data.get('include_compensi', True):
                compensi_section = """

Il Presidente invita anche l'assemblea a deliberare il compenso da attribuire all'organo amministrativo che verrà nominato, ai sensi dell'art. […] dello statuto sociale."""
            
            # Deliberazione
            deliberazione_section = f"""

Segue breve discussione tra i soci al termine della quale si passa alla votazione con voto palese in forza della quale il Presidente constata che, all'unanimità [oppure con il voto contrario dei Sigg. […] e [eventualmente l'astensione dei Sigg. […]]], l'assemblea

d e l i b e r a:

di affidare l'amministrazione della società ad un Consiglio di Amministrazione composto da {len(consiglieri)} membri che resterà in carica {data.get('durata_incarico', 'a tempo indeterminato fino a revoca o dimissioni').lower()} [verificare che l'atto costitutivo non preveda una durata massima per l'incarico]

di nominare Consiglieri di Amministrazione della società i Sigg.:"""
            
            for cons in consiglieri:
                if cons.get('nome'):
                    data_nascita_str = cons.get('data_nascita').strftime('%d/%m/%Y') if hasattr(cons.get('data_nascita'), 'strftime') else '[Data nascita]'
                    deliberazione_section += f"\n- {cons.get('nome', '[Nome]')}, nato a {cons.get('luogo_nascita', '[Luogo nascita]')} il {data_nascita_str}, C.F. {cons.get('codice_fiscale', '[CF]')}, residente in {cons.get('residenza', '[Residenza]')}"
            
            # Presidente CdA
            if data.get('presidente_cda_option') == "Nomina diretta in assemblea" and data.get('presidente_cda'):
                deliberazione_section += f"\n\ndi nominare Presidente del Consiglio di Amministrazione della società il Sig. {data.get('presidente_cda')} ai sensi dell'art. […] dello statuto sociale"
            else:
                deliberazione_section += "\n\ndi rimandare al Consiglio di Amministrazione testé deliberato la nomina del Presidente ai sensi dell'art. […] dello statuto sociale"
            
            # Compensi nella deliberazione
            if data.get('include_compensi', True):
                compenso = data.get('compenso_annuo', '0,00')
                rimborso_text = " oltre al rimborso delle spese sostenute dai consiglieri in ragione del loro ufficio" if data.get('rimborso_spese', True) else ""
                deliberazione_section += f"\n\ndi attribuire all'organo amministrativo testè nominato il compenso annuo ed omnicomprensivo pari a nominali euro {compenso} al lordo di ritenute fiscali e previdenziali{rimborso_text}. L'organo amministrativo delibererà in merito alla ripartizione del compenso tra i suoi membri, anche in considerazione dei compiti e delle deleghe che verranno attribuite a ciascun consigliere."
            
            # Accettazione
            nomi_cons = [c.get('nome', '[Nome]') for c in consiglieri if c.get('nome')]
            if len(nomi_cons) > 1:
                nomi_str = " e ".join(nomi_cons)
                accettazione_section = f"\n\nI sigg. {nomi_str}, presenti in assemblea in qualità di [indicare (socio, amministratore uscente, invitato o altro)] accettano l'incarico e ringraziano l'assemblea per la fiducia accordata."
            else:
                accettazione_section = f"\n\nIl sig. {nomi_cons[0] if nomi_cons else '[Nome]'}, presente in assemblea accetta l'incarico e ringrazia l'assemblea per la fiducia accordata."
            
            # Chiusura
            chiusura_section = f"""

*     *     *

Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.

Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo [eventualmente unitamente a quanto allegato].

L'assemblea viene sciolta alle ore {data.get('ora_chiusura', '[ORA]')}.


Il Presidente                    Il Segretario
{data.get('presidente', '[PRESIDENTE]')}            {data.get('segretario', '[SEGRETARIO]')}"""
            
            # Combina tutte le sezioni
            full_text = (header + presidente_section + nomina_section + compensi_section + 
                        deliberazione_section + accettazione_section + chiusura_section)
            
            return full_text
            
        except Exception as e:
            return f"Errore nella generazione dell'anteprima: {str(e)}"
    
    def generate_document(self, data: dict) -> Document:
        """Genera il documento Word del verbale con formattazione simile agli altri template."""
        import os

        # Se esiste un template .docx di base, prova a caricarlo per mantenere stili
        template_path = os.path.join(os.path.dirname(__file__), 'template.docx')

        try:
            if os.path.exists(template_path):
                doc = Document(template_path)
                # Svuota il contenuto preservando gli stili
                for paragraph in doc.paragraphs[:]:
                    paragraph._element.getparent().remove(paragraph._element)
                # Configura stili base definiti nella superclass
                self._setup_document_styles(doc)
            else:
                doc = Document()
                self._setup_document_styles(doc)
        except Exception:
            doc = Document()
            self._setup_document_styles(doc)

        # Garantisce che existano campi lista soci
        if 'soci' not in data:
            data['soci'] = []

        # Header, titolo, apertura, partecipanti, discussione, chiusura, firme
        self._add_company_header(doc, data)
        self._add_verbale_title(doc, data)
        self._add_opening_section(doc, data)
        self._add_participants_section(doc, data)
        self._add_nomination_discussion(doc, data)
        self._add_closing_section(doc, data)
        self._add_signatures(doc, data)

        return doc

    # --------------------- Sezioni helper per il documento ---------------------

    def _add_opening_section(self, doc, data):
        """Sezione iniziale con data, ora e sede e ordine del giorno"""
        doc.add_paragraph()

        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()

        # Data e ora formattate
        data_val = data.get('data_assemblea', '[DATA]')
        data_str = data_val.strftime('%d/%m/%Y') if hasattr(data_val, 'strftime') else str(data_val)

        ora_val = data.get('ora_inizio', '[ORA]')
        ora_str = ora_val.strftime('%H:%M') if hasattr(ora_val, 'strftime') else str(ora_val)

        sede = data.get('sede_legale', '[SEDE]')

        opening_text = (f"Oggi {data_str} alle ore {ora_str} presso la sede sociale {sede}, "
                        "si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:")

        run = p.add_run(opening_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # Ordine del giorno
        doc.add_paragraph()
        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()
        run = p.add_run("Ordine del giorno")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # punti
        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()
        run = p.add_run("1. nomina del Consiglio di Amministrazione della società")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        if data.get('include_compensi', True):
            try:
                p = doc.add_paragraph(style='BodyText')
            except KeyError:
                p = doc.add_paragraph()
            run = p.add_run("2. attribuzione di compensi al Consiglio di Amministrazione della società")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    def _add_participants_section(self, doc, data):
        """Sezione presidente e partecipanti, inclusi soci"""
        doc.add_paragraph()
        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()

        presidente = data.get('presidente', '[PRESIDENTE]')
        run = p.add_run(
            f"Assume la presidenza ai sensi dell'art. […] dello statuto sociale il Sig. {presidente} Amministratore Unico [oppure Presidente del Consiglio di Amministrazione o altro], il quale dichiara e constata:")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # Sezione soci dettagliata
        soci = data.get('soci', [])
        if soci:
            self._add_soci_section(doc, data)

    def _add_soci_section(self, doc, data):
        """Lista soci con quote e totali"""
        soci_presenti = data.get('soci_presenti', [])
        soci_assenti = data.get('soci_assenti', [])

        # Fallback
        if not soci_presenti and not soci_assenti and 'soci' in data:
            soci_presenti = [s for s in data.get('soci', []) if s.get('presente', True)]
            soci_assenti = [s for s in data.get('soci', []) if not s.get('presente', True)]

        if not soci_presenti:
            return

        # Calcola totali
        total_euro = 0.0
        total_perc = 0.0
        capitale_raw = str(data.get('capitale_sociale', '0')).replace('.', '').replace(',', '.')
        try:
            capitale_float = float(capitale_raw)
        except ValueError:
            capitale_float = 0.0

        for socio in soci_presenti:
            if isinstance(socio, dict):
                # euro
                euro_raw = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                try:
                    total_euro += float(euro_raw or 0)
                except ValueError:
                    pass

                # percentuale
                perc_raw = socio.get('quota_percentuale', '')
                if perc_raw:
                    try:
                        total_perc += float(str(perc_raw).replace('%', '').replace(',', '.'))
                    except ValueError:
                        pass
                else:
                    try:
                        euro_val = float(euro_raw or 0)
                        if capitale_float > 0 and euro_val > 0:
                            total_perc += euro_val / capitale_float * 100
                    except ValueError:
                        pass

        formatted_euro = CommonDataHandler.format_currency(total_euro)
        formatted_perc = CommonDataHandler.format_percentage(total_perc)

        # Paragrafo totali
        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()
        run = p.add_run(
            f"nonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_euro} pari al {formatted_perc} del Capitale Sociale:")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # Lista soci
        for socio in soci_presenti:
            if isinstance(socio, dict) and socio.get('nome'):
                quota_euro = CommonDataHandler.format_currency(socio.get('quota_euro', '0'))
                quota_perc = socio.get('quota_percentuale', '')
                if not quota_perc:
                    # calcola
                    try:
                        euro_val = float(str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.'))
                        quota_perc = CommonDataHandler.format_percentage(euro_val / capitale_float * 100) if capitale_float else '[%]'
                    except ValueError:
                        quota_perc = '[%]'
                else:
                    quota_perc = CommonDataHandler.clean_percentage(quota_perc)

                tipo_sogg = socio.get('tipo_soggetto', 'Persona Fisica')
                tipo_part = socio.get('tipo_partecipazione', 'Diretto')
                delegato = socio.get('delegato', '').strip()
                rappresentante = socio.get('rappresentante_legale', '').strip()

                # Costruisci descrizione socio con deleghe/rappr. legale
                if tipo_part == 'Delegato' and delegato:
                    # Caso delegato
                    if tipo_sogg == 'Società':
                        descr_line = (f"il Sig {delegato} delegato della società {socio.get('nome')} socio "
                                      f"recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale")
                    else:
                        descr_line = (f"il Sig {delegato} delegato del socio Sig {socio.get('nome')} "
                                      f"recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale")
                else:
                    # Partecipazione diretta
                    if tipo_sogg == 'Società':
                        if rappresentante:
                            descr_line = (f"la società {socio.get('nome')}, rappresentata dal Sig {rappresentante}, socia "
                                          f"recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale")
                        else:
                            descr_line = (f"la società {socio.get('nome')} socia recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale")
                    else:
                        descr_line = (f"il Sig {socio.get('nome')} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_perc} del Capitale Sociale")

                try:
                    p = doc.add_paragraph(style='BodyText')
                except KeyError:
                    p = doc.add_paragraph()
                run = p.add_run(descr_line)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        
        if soci_assenti:
            doc.add_paragraph()
            p = doc.add_paragraph("Risultano invece assenti i seguenti soci:", style='BodyText')
            for socio in soci_assenti:
                 if isinstance(socio, dict) and socio.get('nome'):
                    p = doc.add_paragraph(f"- {socio.get('nome')}", style='BodyText')

    def _add_nomination_discussion(self, doc, data):
        """Discussione sulla nomina del CdA e deliberazione"""
        doc.add_paragraph()

        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()

        motivo = data.get('motivo_nomina', 'dimissioni dell\'organo in carica').lower()
        membri = len(data.get('consiglieri', []))
        run = p.add_run(
            f"Il Presidente informa l'assemblea che si rende necessaria la nomina di un nuovo organo amministrativo [{motivo}]. "
            f"Propone di affidare l'amministrazione della società ad un Consiglio di Amministrazione di {membri} membri.")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # Elenco consiglieri
        consiglieri = data.get('consiglieri', [])
        for cons in consiglieri:
            if cons.get('nome'):
                try:
                    p = doc.add_paragraph(style='BodyText')
                except KeyError:
                    p = doc.add_paragraph()
                data_nascita = cons.get('data_nascita')
                data_str = data_nascita.strftime('%d/%m/%Y') if hasattr(data_nascita, 'strftime') else '[Data nascita]'
                run = p.add_run(f"- {cons.get('nome')} nato a {cons.get('luogo_nascita', '')} il {data_str}, C.F. {cons.get('codice_fiscale', '')}, residente in {cons.get('residenza', '')}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

        # Compensi
        if data.get('include_compensi', True):
            try:
                p = doc.add_paragraph(style='BodyText')
            except KeyError:
                p = doc.add_paragraph()
            run = p.add_run("L'assemblea delibera inoltre di attribuire all'organo amministrativo il compenso annuo previsto.")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    def _add_closing_section(self, doc, data):
        """Sezione di chiusura con ora scioglimento"""
        doc.add_paragraph()
        try:
            p = doc.add_paragraph(style='BodyText')
        except KeyError:
            p = doc.add_paragraph()

        ora_val = data.get('ora_chiusura', '[ORA]')
        ora_str = ora_val.strftime('%H:%M') if hasattr(ora_val, 'strftime') else str(ora_val)

        run = p.add_run(f"Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola. L'assemblea viene sciolta alle ore {ora_str}.")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# Registrazione del template nel factory per renderlo selezionabile nell'app
DocumentTemplateFactory.register_template('verbale_assemblea_consiglio_amministrazione', VerbaleConsiglioAmministrazioneTemplate)
