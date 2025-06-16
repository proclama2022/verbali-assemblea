"""
Template Completo per Verbale di Assemblea dei Soci
Questo template include tutte le opzioni possibili per verbali di assemblea.
"""

import sys
import os

# Aggiungi il path della cartella src (relativo alla root del progetto)
current_dir = os.path.dirname(os.path.abspath(__file__))
src_path = os.path.join(os.path.dirname(current_dir), 'src')
if src_path not in sys.path:
    sys.path.append(src_path)

from document_templates import DocumentTemplate, DocumentTemplateFactory
from common_data_handler import CommonDataHandler
from base_verbale_template import BaseVerbaleTemplate
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date
import streamlit as st
import pandas as pd
import re

class VerbaleAssembleaCompletoTemplate(BaseVerbaleTemplate):
    """Template Completo per Verbale di Assemblea dei Soci"""
    
    def get_template_name(self) -> str:
        return "Verbale Assemblea Approvazione Bilancio"
    
    def get_required_fields(self) -> list:
        return [
            "denominazione", "sede_legale", "capitale_sociale", "codice_fiscale",
            "data_assemblea", "ora_assemblea", "presidente", "segretario", 
            "soci", "amministratori"
        ]
    
    def get_form_fields(self, extracted_data: dict) -> dict:
        """Genera i campi del form per Streamlit."""
        # Chiama il metodo della classe base per ottenere i campi comuni
        form_data = super().get_form_fields(extracted_data)

        st.subheader("📑 Configurazioni Specifiche del Verbale Completo")

        # Sezioni opzionali
        col1, col2, col3 = st.columns(3)
        with col1:
            form_data['include_nomina_cda'] = st.checkbox("Include Nomina CdA", key="include_nomina_cda_completo")
            form_data['include_nomina_revisore'] = st.checkbox("Include Nomina Revisore", key="include_nomina_revisore_completo")
        with col2:
            form_data['include_approvazione_bilancio'] = st.checkbox("Include Approvazione Bilancio", key="include_bilancio_completo")
            form_data['include_distribuzione_utili'] = st.checkbox("Include Distribuzione Utili", key="include_utili_completo")
        with col3:
            form_data['include_ratifica_operato'] = st.checkbox("Include Ratifica Operato", key="include_ratifica_completo")
            form_data['altre_delibere'] = st.checkbox("Include Altre Delibere", key="include_altre_delibere_completo")

        # Campi condizionali basati sulle selezioni
        if form_data.get('include_approvazione_bilancio'):
            st.subheader("📊 Dati Bilancio")
            # Aggiungere qui i campi specifici per l'approvazione del bilancio
            form_data['esercizio_bilancio'] = st.text_input("Esercizio di riferimento del bilancio", key="esercizio_bilancio_completo")

        if form_data.get('include_distribuzione_utili'):
            st.subheader("💰 Dati Distribuzione Utili")
            # Aggiungere qui i campi specifici per la distribuzione degli utili
            form_data['importo_utili'] = st.text_input("Importo utili da distribuire (€)", "0,00", key="importo_utili_completo")

        if form_data.get('altre_delibere'):
            st.subheader("✍️ Altre Delibere")
            form_data['testo_altre_delibere'] = st.text_area("Testo per altre delibere", height=150, key="testo_altre_delibere_completo")

        return form_data

    def show_preview(self, form_data: dict):
        """Mostra l'anteprima del documento"""
        st.subheader("📄 Anteprima Verbale")
        preview_text = self._generate_preview_text(form_data)
        st.text_area("", preview_text, height=600, disabled=True)

    def _generate_preview_text(self, data: dict) -> str:
        """Genera il testo di anteprima del verbale"""
        
        # Header
        preview = f"""
{data.get('denominazione', '[DENOMINAZIONE]')}
Sede in {data.get('sede_legale', '[SEDE]')}
Capitale sociale Euro {data.get('capitale_sociale', '[CAPITALE]')} i.v.
Codice fiscale: {data.get('codice_fiscale', '[CF]')}

Verbale di assemblea dei soci
del {data.get('data_assemblea', '[DATA]').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[DATA]'}

Oggi {data.get('data_assemblea', '[DATA]').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[DATA]'} alle ore {data.get('ora_inizio', '[ORA]')} presso {data.get('luogo_assemblea', 'la sede sociale')}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:

ORDINE DEL GIORNO
"""
        
        # Ordine del giorno
        for i, punto in enumerate(data.get('punti_ordine_giorno', []), 1):
            preview += f"{i}. {punto}\n"
        
        # Presidente
        presidente = data.get('presidente', '[PRESIDENTE]')
        if data.get('tipo_amministrazione') == "Amministratore Unico":
            ruolo_presidente = "Amministratore Unico"
        else:
            # Trova il ruolo del presidente negli amministratori
            ruolo_presidente = "Presidente del Consiglio di Amministrazione"
            for amm in data.get('amministratori', []):
                if amm.get('nome') == presidente and 'Presidente' in amm.get('carica', ''):
                    ruolo_presidente = amm.get('carica')
                    break
        
        preview += f"""
Assume la presidenza ai sensi dell'art. […] dello statuto sociale il Sig. {presidente} {ruolo_presidente}, il quale dichiara e constata:

1 - che"""
        
        if data.get('audioconferenza'):
            preview += " (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. […] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza"
        
        preview += f"""
2 - che sono presenti/partecipano all'assemblea:

{ruolo_presidente} nella persona del suddetto Presidente Sig. {presidente}
"""
        
        # Amministratori
        if data.get('tipo_amministrazione') == "Consiglio di Amministrazione":
            preview += "\nper il Consiglio di Amministrazione:\n"
            for amm in data.get('amministratori', []):
                if amm.get('presente'):
                    nome_amm = amm.get('nome', '')
                    if nome_amm != presidente:
                        preview += f"il Sig. {nome_amm} - {amm.get('carica', '')}\n"
                elif amm.get('assente_giustificato'):
                    nome_amm = amm.get('nome', '')
                    if nome_amm != presidente:
                        preview += f"assente giustificato il Sig. {nome_amm} il quale ha tuttavia rilasciato apposita dichiarazione scritta\n"
        
        # Collegio Sindacale
        if data.get('has_collegio_sindacale'):
            if data.get('tipo_collegio') == "Sindaco Unico":
                preview += f"\nil Sindaco Unico nella persona del Sig. {data.get('sindaco_unico', '')}\n"
            else:
                preview += "\nper il Collegio Sindacale:\n"
                for sindaco in data.get('sindaci', []):
                    if sindaco.get('presente'):
                        preview += f"il Dott. {sindaco.get('nome', '')} - {sindaco.get('carica', '')}\n"
        
        # Revisore
        if data.get('has_revisore'):
            preview += f"\nil revisore contabile Dott. {data.get('nome_revisore', '')}\n"
        
        # Altri partecipanti
        if data.get('has_altri_partecipanti'):
            for altro in data.get('altri_partecipanti', []):
                if altro.get('nome'):
                    preview += f"\nil Sig. {altro.get('nome', '')} in qualità di {altro.get('qualita', '')}\n"
        
        # Soci
        soci_presenti = data.get('soci_presenti', [])
        soci_assenti = data.get('soci_assenti', [])

        # Fallback per mantenere compatibilità se le nuove chiavi non ci sono
        if not soci_presenti and not soci_assenti and 'soci' in data:
            soci_presenti = [s for s in data.get('soci', []) if s.get('presente', True)]
            soci_assenti = [s for s in data.get('soci', []) if not s.get('presente', True)]

        total_quota_euro = 0.0
        total_quota_percentuale = 0.0

        for socio in soci_presenti:
            if isinstance(socio, dict):
                try:
                    # Rimuovi punti e sostituisci virgola con punto per la conversione a float
                    quota_euro_str = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                    total_quota_euro += float(quota_euro_str)
                except ValueError:
                    pass # Ignora valori non numerici
                try:
                    quota_percentuale_str = str(socio.get('quota_percentuale', '0')).replace('.', '').replace(',', '.')
                    total_quota_percentuale += float(quota_percentuale_str)
                except ValueError:
                    pass # Ignora valori non numerici
        
        # Formatta i totali per la visualizzazione
        formatted_total_quota_euro = f"{total_quota_euro:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') # Formato italiano
        formatted_total_quota_percentuale = f"{total_quota_percentuale:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') # Formato italiano

        if soci_presenti:
            preview += f"\nnonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_total_quota_euro} pari al {formatted_total_quota_percentuale}% del Capitale Sociale:\n"
            
            for socio in soci_presenti:
                nome = socio.get('nome', '')
                quota_euro = socio.get('quota_euro', '')
                quota_percentuale = socio.get('quota_percentuale', '')
                tipo_partecipazione = socio.get('tipo_partecipazione', 'Diretto')
                delegato = socio.get('delegato', '').strip()
                tipo_soggetto = socio.get('tipo_soggetto', 'Persona Fisica')
                rappresentante_legale = socio.get('rappresentante_legale', '').strip()
                
                # Gestione completa: delegati e rappresentanti legali
                if tipo_partecipazione == 'Delegato' and delegato:
                    if tipo_soggetto == 'Società':
                        preview += f"il Sig. {delegato} delegato della società {nome}"
                        if rappresentante_legale:
                            preview += f" (nella persona del legale rappresentante {rappresentante_legale})"
                        preview += f" recante una quota pari a nominali euro {quota_euro} pari al {quota_percentuale}% del Capitale Sociale\n"
                    else:
                        preview += f"il Sig. {delegato} delegato del socio {nome} recante una quota pari a nominali euro {quota_euro} pari al {quota_percentuale}% del Capitale Sociale\n"
                else:
                    if tipo_soggetto == 'Società':
                        if rappresentante_legale:
                            preview += f"la società {nome} nella persona del legale rappresentante {rappresentante_legale} recante una quota pari a nominali euro {quota_euro} pari al {quota_percentuale}% del Capitale Sociale\n"
                        else:
                            preview += f"la società {nome} recante una quota pari a nominali euro {quota_euro} pari al {quota_percentuale}% del Capitale Sociale\n"
                    else:
                        preview += f"il Sig. {nome} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_percentuale}% del Capitale Sociale\n"
        
        if soci_assenti:
            preview += "\nRisultano invece assenti i seguenti soci:\n"
            for socio in soci_assenti:
                if isinstance(socio, dict) and socio.get('nome'):
                    preview += f"- {socio.get('nome')}\n"

        preview += """
2 - che gli intervenuti sono legittimati alla presente assemblea;
3 - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno.

I presenti all'unanimità chiamano a fungere da segretario il signor {}, che accetta l'incarico.

Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante.
""".format(data.get('segretario', '[SEGRETARIO]'))
        
        # Lingua straniera
        if data.get('lingua_straniera'):
            preview += f"""
In particolare, preso atto che il Sig. {data.get('nome_straniero', '')} non conosce la lingua italiana ma dichiara di conoscere la lingua {data.get('lingua_conosciuta', 'inglese')}, il Presidente dichiara che provvederà a tradurre dall'italiano all'{data.get('lingua_conosciuta', 'inglese')} (e viceversa) gli interventi dei partecipanti alla discussione nonché a tradurre dall'italiano all'{data.get('lingua_conosciuta', 'inglese')} il verbale che sarà redatto al termine della riunione.
"""
        
        preview += f"""
Il Presidente constata e fa constatare che l'assemblea risulta {data.get('tipo_convocazione', 'regolarmente convocata')} e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.

Si passa quindi allo svolgimento dell'ordine del giorno.

*     *     *

In relazione al primo punto il presidente legge il bilancio al {data.get('data_chiusura', '[DATA]').strftime('%d/%m/%Y') if hasattr(data.get('data_chiusura'), 'strftime') else '[DATA]'} composto da stato patrimoniale, conto economico e nota integrativa (allegati di seguito al presente verbale).
"""
        
        # Relazione Collegio Sindacale
        if data.get('has_collegio_sindacale'):
            if data.get('tipo_collegio') == "Sindaco Unico":
                preview += f"Prende la parola il Sindaco Unico che legge la relazione al Bilancio chiuso al {data.get('data_chiusura', '[DATA]').strftime('%d/%m/%Y') if hasattr(data.get('data_chiusura'), 'strftime') else '[DATA]'} (allegata di seguito al presente verbale).\n"
            else:
                preview += f"Prende la parola il Presidente del Collegio Sindacale che legge la relazione del collegio sindacale al Bilancio chiuso al {data.get('data_chiusura', '[DATA]').strftime('%d/%m/%Y') if hasattr(data.get('data_chiusura'), 'strftime') else '[DATA]'} (allegata di seguito al presente verbale).\n"
        
        # Relazione Revisore
        if data.get('has_revisore'):
            preview += f"Prende infine la parola il Dott. {data.get('nome_revisore', '')} che legge la relazione del revisore contabile (allegata di seguito al presente verbale).\n"
        
        preview += f"""
Segue breve discussione tra i soci al termine della quale si passa alla votazione con voto {'palese' if data.get('voto_palese') else 'segreto'} in forza della quale il Presidente constata che, {data.get('esito_votazione', 'all\'unanimità')} l'assemblea

DELIBERA

l'approvazione del bilancio di esercizio chiuso al {data.get('data_chiusura', '[DATA]').strftime('%d/%m/%Y') if hasattr(data.get('data_chiusura'), 'strftime') else '[DATA]'} e dei relativi documenti che lo compongono.

*     *     *
"""
        
        # Ripianamento perdite
        if data.get('has_ripianamento'):
            preview += f"""
In relazione al secondo punto posto all'ordine del giorno, il Presidente"""
            
            if data.get('has_collegio_sindacale'):
                preview += ", sentito il parere favorevole del Collegio Sindacale,"
            
            preview += f""" propone all'assemblea di ripianare la perdita mediante la rinuncia al rimborso di una corrispondente quota del credito vantato dai soci nei confronti della società a titolo di {data.get('tipo_credito', 'finanziamento infruttifero soci')}."

Segue breve discussione tra i soci al termine della quale si passa alla votazione con voto {'palese' if data.get('voto_palese') else 'segreto'} in forza della quale il Presidente constata che, {data.get('esito_votazione', 'all\'unanimità')} l'assemblea

DELIBERA

di approvare la proposta del Presidente e di ripianare, con effetto dalla data odierna, la perdita di euro {data.get('importo_perdita', '')} mediante la rinuncia al rimborso di una corrispondente quota del credito vantato dai soci nei confronti della società a titolo di {data.get('tipo_credito', 'finanziamento infruttifero soci')}.

Le rinunce al rimborso avverranno in proporzione alla quota di {data.get('tipo_credito', 'finanziamento infruttifero soci')} attualmente in essere e quindi:
"""
            
            for rinuncia in data.get('rinunce_soci', []):
                if rinuncia.get('socio'):
                    preview += f"socio {rinuncia.get('socio', '')} per euro {rinuncia.get('importo_rinuncia', '')} pari al {rinuncia.get('percentuale', '')}% della perdita da ripianare;\n"
            
            preview += f"""
Con effetto dalla data odierna, il residuo credito verso i soci a titolo di {data.get('tipo_credito', 'finanziamento infruttifero soci')} continuerà ad essere regolato dalle condizioni previgenti, ma solo per il residuo importo (al netto dell'avvenuta rinuncia) qui riepilogato:
"""
            
            for rinuncia in data.get('rinunce_soci', []):
                if rinuncia.get('socio'):
                    preview += f"socio {rinuncia.get('socio', '')} per euro {rinuncia.get('residuo', '')};\n"
            
            preview += """
Ciascun socio rilascia seduta stante al Presidente una dichiarazione sostitutiva di atto notorio che attesta il valore fiscale del credito a cui ha testé rinunciato.
Le dichiarazioni rilasciate dai soci verranno conservate agli atti della società.

*     *     *
"""
        
        preview += f"""
Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.

Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea {data.get('esito_votazione', 'all\'unanimità')}, con voto {'palese' if data.get('voto_palese') else 'segreto'}, ne approva il testo{'unitamente a quanto allegato' if data.get('documenti_allegati') else ''}.

L'assemblea viene sciolta alle ore {data.get('ora_fine', '[ORA]')}.


Il Presidente                           Il Segretario

_____________________                  _____________________
{data.get('presidente', '[PRESIDENTE]')}                         {data.get('segretario', '[SEGRETARIO]')}
"""
        
        return preview

    def generate_document(self, data: dict) -> Document:
        """Genera il documento Word del verbale"""
        import os
        
        # Utilizza il template .docx esistente per mantenere la formattazione
        template_path = os.path.join(os.path.dirname(__file__), 'template.docx')
        
        try:
            if os.path.exists(template_path):
                doc = Document(template_path)
                # Rimuovi il contenuto esistente del template mantenendo gli stili
                for paragraph in doc.paragraphs[:]:
                    p = paragraph._element
                    p.getparent().remove(p)
                # Imposta gli stili anche quando usiamo un template
                self._setup_document_styles(doc)
            else:
                doc = Document()
                self._setup_document_styles(doc)
        except Exception as e:
            # Fallback a documento vuoto se il template non può essere caricato
            doc = Document()
            self._setup_document_styles(doc)
        
        # Header società
        self._add_company_header(doc, data)
        
        # Titolo verbale
        self._add_verbale_title(doc, data)
        
        # Sezione di apertura
        self._add_opening_section(doc, data)
        
        # Partecipanti
        self._add_participants_section(doc, data)
        
        # Dichiarazioni preliminari
        self._add_preliminary_statements(doc, data)
        
        # Discussione punti OdG
        self._add_discussion_section(doc, data)
        
        # Ripianamento perdite se presente
        if data.get('has_ripianamento'):
            self._add_ripianamento_section(doc, data)
        
        # Chiusura
        self._add_closing_section(doc, data)
        
        # Firme
        self._add_signatures(doc, data)
        
        return doc

    def _setup_document_styles(self, doc):
        """Imposta gli stili del documento in modo completo"""
        styles = doc.styles
        
        # Stile per il titolo della società
        if 'TitoloSocieta' not in styles:
            title_style = styles.add_style('TitoloSocieta', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Times New Roman'
            title_style.font.size = Pt(14)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(6)
        
        # Stile per il titolo del verbale
        if 'TitoloVerbale' not in styles:
            verbale_title_style = styles.add_style('TitoloVerbale', WD_STYLE_TYPE.PARAGRAPH)
            verbale_title_style.font.name = 'Times New Roman'
            verbale_title_style.font.size = Pt(16)
            verbale_title_style.font.bold = True
            verbale_title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            verbale_title_style.paragraph_format.space_before = Pt(18)
            verbale_title_style.paragraph_format.space_after = Pt(18)
        
        # Stile per gli heading 1
        if 'Heading1' not in styles:
            heading1_style = styles.add_style('Heading1', WD_STYLE_TYPE.PARAGRAPH)
            heading1_style.font.name = 'Times New Roman'
            heading1_style.font.size = Pt(14)
            heading1_style.font.bold = True
            heading1_style.paragraph_format.space_before = Pt(12)
            heading1_style.paragraph_format.space_after = Pt(6)
            heading1_style.paragraph_format.keep_with_next = True
        
        # Stile per gli heading 2
        if 'Heading2' not in styles:
            heading2_style = styles.add_style('Heading2', WD_STYLE_TYPE.PARAGRAPH)
            heading2_style.font.name = 'Times New Roman'
            heading2_style.font.size = Pt(13)
            heading2_style.font.bold = True
            heading2_style.paragraph_format.space_before = Pt(6)
            heading2_style.paragraph_format.space_after = Pt(6)
            heading2_style.paragraph_format.keep_with_next = True
        
        # Stile per il corpo del testo
        if 'BodyText' not in styles:
            body_style = styles.add_style('BodyText', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = 'Times New Roman'
            body_style.font.size = Pt(12)
            body_style.paragraph_format.space_after = Pt(6)
            body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Stile per le firme
        if 'Firma' not in styles:
            firma_style = styles.add_style('Firma', WD_STYLE_TYPE.PARAGRAPH)
            firma_style.font.name = 'Times New Roman'
            firma_style.font.size = Pt(12)
            firma_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            firma_style.paragraph_format.space_before = Pt(24)

    def _add_company_header(self, doc, data):
        """Aggiunge l'header con i dati della società"""
        p = doc.add_paragraph(style='TitoloSocieta')
        p.add_run(data.get('denominazione', '[DENOMINAZIONE]')).bold = True
        
        p = doc.add_paragraph(style='TitoloSocieta')
        p.add_run(f"Sede in {data.get('sede_legale', '[SEDE]')}")
        
        p = doc.add_paragraph(style='TitoloSocieta')
        p.add_run(f"Capitale sociale Euro {data.get('capitale_sociale', '[CAPITALE]')} i.v.")
        
        p = doc.add_paragraph(style='TitoloSocieta')
        p.add_run(f"Codice fiscale: {data.get('codice_fiscale', '[CF]')}")

    def _add_verbale_title(self, doc, data):
        """Aggiunge il titolo del verbale"""
        doc.add_paragraph()
        p = doc.add_paragraph(style='TitoloVerbale')
        p.add_run("Verbale di assemblea dei soci").bold = True
        
        data_str = data.get('data_assemblea').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[DATA]'
        p = doc.add_paragraph(style='TitoloVerbale')
        p.add_run(f"del {data_str}").bold = True

    def _add_opening_section(self, doc, data):
        """Aggiunge la sezione di apertura"""
        doc.add_paragraph()
        
        data_str = data.get('data_assemblea').strftime('%d/%m/%Y') if hasattr(data.get('data_assemblea'), 'strftime') else '[DATA]'
        luogo = data.get('luogo_assemblea', 'la sede sociale')
        ora = data.get('ora_inizio', '[ORA]')
        
        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"Oggi {data_str} alle ore {ora} presso {luogo}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:")
        
        doc.add_paragraph()
        p = doc.add_paragraph(style='Heading1')
        p.add_run("ORDINE DEL GIORNO")
        
        for i, punto in enumerate(data.get('punti_ordine_giorno', []), 1):
            p = doc.add_paragraph(style='BodyText')
            p.add_run(f"{i}. {punto}")

    def _add_participants_section(self, doc, data):
        """Aggiunge la sezione dei partecipanti"""
        doc.add_paragraph()
        
        presidente = data.get('presidente', '[PRESIDENTE]')
        if data.get('tipo_amministrazione') == "Amministratore Unico":
            ruolo_presidente = "Amministratore Unico"
        else:
            ruolo_presidente = "Presidente del Consiglio di Amministrazione"
            for amm in data.get('amministratori', []):
                if amm.get('nome') == presidente and 'Presidente' in amm.get('carica', ''):
                    ruolo_presidente = amm.get('carica')
                    break
        
        p = doc.add_paragraph()
        p.add_run(f"Assume la presidenza ai sensi dell'art. […] dello statuto sociale il Sig. {presidente} {ruolo_presidente}, il quale dichiara e constata:")

    def _add_preliminary_statements(self, doc, data):
        """Aggiunge le dichiarazioni preliminari"""
        # Punto 1 - Audioconferenza
        p = doc.add_paragraph()
        text = "1 - che"
        if data.get('audioconferenza'):
            text += " (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. […] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza"
        p.add_run(text)
        
        # Punto 2 - Partecipanti
        p = doc.add_paragraph()
        p.add_run("2 - che sono presenti/partecipano all'assemblea:")
        
        # Presidente
        presidente = data.get('presidente', '[PRESIDENTE]')
        ruolo_presidente = "Amministratore Unico" if data.get('tipo_amministrazione') == "Amministratore Unico" else "Presidente del Consiglio di Amministrazione"
        for amm in data.get('amministratori', []):
            if amm.get('nome') == presidente and 'Presidente' in amm.get('carica', ''):
                ruolo_presidente = amm.get('carica')
                break
                
        p = doc.add_paragraph()
        p.add_run(f"{ruolo_presidente} nella persona del suddetto Presidente Sig. {presidente}")
        
        # Amministratori
        if data.get('tipo_amministrazione') == "Consiglio di Amministrazione":
            p = doc.add_paragraph()
            p.add_run("per il Consiglio di Amministrazione:")
            
            for amm in data.get('amministratori', []):
                nome_amm = amm.get('nome', '')
                if nome_amm != presidente and amm.get('presente', False):
                    p = doc.add_paragraph()
                    p.add_run(f"il Sig. {nome_amm} - {amm.get('carica', '')}")
                elif nome_amm != presidente and amm.get('assente_giustificato', False):
                    p = doc.add_paragraph()
                    p.add_run(f"assente giustificato il Sig. {nome_amm} il quale ha tuttavia rilasciato apposita dichiarazione scritta")
        
        # Collegio Sindacale
        if data.get('has_collegio_sindacale'):
            if data.get('tipo_collegio') == "Sindaco Unico":
                p = doc.add_paragraph()
                p.add_run(f"il Sindaco Unico nella persona del Sig. {data.get('sindaco_unico', '')}")
            else:
                p = doc.add_paragraph()
                p.add_run("per il Collegio Sindacale:")
                
                for sindaco in data.get('sindaci', []):
                    if sindaco.get('presente', False):
                        p = doc.add_paragraph()
                        p.add_run(f"il Dott. {sindaco.get('nome', '')} - {sindaco.get('carica', '')}")
        
        # Altri partecipanti
        if data.get('has_altri_partecipanti'):
            for altro in data.get('altri_partecipanti', []):
                if altro.get('nome'):
                    p = doc.add_paragraph()
                    p.add_run(f"il Sig. {altro.get('nome', '')} in qualità di {altro.get('qualita', '')}")
        
        # Calcola il totale delle quote dei soci presenti
        soci_presenti = data.get('soci_presenti', [])
        soci_assenti = data.get('soci_assenti', [])

        # Fallback per mantenere compatibilità se le nuove chiavi non ci sono
        if not soci_presenti and not soci_assenti and 'soci' in data:
            soci_presenti = [s for s in data.get('soci', []) if s.get('presente', True)]
            soci_assenti = [s for s in data.get('soci', []) if not s.get('presente', True)]

        total_quota_euro = 0.0
        total_quota_percentuale = 0.0

        for socio in soci_presenti:
            if isinstance(socio, dict):
                try:
                    # Rimuovi punti e sostituisci virgola con punto per la conversione a float
                    quota_euro_str = str(socio.get('quota_euro', '0')).replace('.', '').replace(',', '.')
                    total_quota_euro += float(quota_euro_str)
                except ValueError:
                    pass # Ignora valori non numerici
                try:
                    quota_percentuale_str = str(socio.get('quota_percentuale', '0')).replace('.', '').replace(',', '.')
                    total_quota_percentuale += float(quota_percentuale_str)
                except ValueError:
                    pass # Ignora valori non numerici
        
        # Formatta i totali per la visualizzazione
        formatted_total_quota_euro = f"{total_quota_euro:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') # Formato italiano
        formatted_total_quota_percentuale = f"{total_quota_percentuale:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') # Formato italiano
        
        # Soci
        if soci_presenti:
            p = doc.add_paragraph()
            p.add_run(f"nonché i seguenti soci o loro rappresentanti, recanti complessivamente una quota pari a nominali euro {formatted_total_quota_euro} pari al {formatted_total_quota_percentuale}% del Capitale Sociale:")
            
            for socio in soci_presenti:
                nome = socio.get('nome', '')
                quota_euro = socio.get('quota_euro', '')
                quota_percentuale = socio.get('quota_percentuale', '')
                tipo_partecipazione = socio.get('tipo_partecipazione', 'Diretto')
                delegato = socio.get('delegato', '').strip()
                tipo_soggetto = socio.get('tipo_soggetto', 'Persona Fisica')
                rappresentante_legale = socio.get('rappresentante_legale', '').strip()
                
                p = doc.add_paragraph()
                
                # Gestione completa: delegati e rappresentanti legali
                if tipo_partecipazione == 'Delegato' and delegato:
                    if tipo_soggetto == 'Società':
                        text = f"il Sig. {delegato} delegato della società {nome}"
                        if rappresentante_legale:
                            text += f" (nella persona del legale rappresentante {rappresentante_legale})"
                        text += f" recante una quota pari a nominali euro {quota_euro} pari al {quota_percentuale}% del Capitale Sociale"
                    else:
                        text = f"il Sig. {delegato} delegato del socio {nome} recante una quota pari a nominali euro {quota_euro} pari al {quota_percentuale}% del Capitale Sociale"
                else:
                    if tipo_soggetto == 'Società':
                        if rappresentante_legale:
                            text = f"la società {nome} nella persona del legale rappresentante {rappresentante_legale} recante una quota pari a nominali euro {quota_euro} pari al {quota_percentuale}% del Capitale Sociale"
                        else:
                            text = f"la società {nome} recante una quota pari a nominali euro {quota_euro} pari al {quota_percentuale}% del Capitale Sociale"
                    else:
                        text = f"il Sig. {nome} socio recante una quota pari a nominali euro {quota_euro} pari al {quota_percentuale}% del Capitale Sociale"
                
                p.add_run(text)

        if soci_assenti:
            p = doc.add_paragraph()
            p.add_run("Risultano invece assenti i seguenti soci:")
            for socio in soci_assenti:
                if isinstance(socio, dict) and socio.get('nome'):
                    p = doc.add_paragraph()
                    p.add_run(f"- {socio.get('nome')}")
        
        # Dichiarazioni finali
        p = doc.add_paragraph()
        p.add_run("3 - che gli intervenuti sono legittimati alla presente assemblea;")
        
        p = doc.add_paragraph()
        p.add_run("4 - che tutti gli intervenuti si dichiarano edotti sugli argomenti posti all'ordine del giorno.")
        
        # Segretario
        p = doc.add_paragraph()
        segretario = data.get('segretario', '')
        if segretario:
            p.add_run(f"I presenti all'unanimità chiamano a fungere da segretario il signor {segretario}, che accetta l'incarico.")
        else:
            p.add_run("I presenti all'unanimità chiamano a fungere da segretario un membro dell'assemblea, che accetta l'incarico.")
        
        # Identificazione partecipanti
        p = doc.add_paragraph()
        p.add_run("Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante.")
        
        # Lingua straniera
        if data.get('lingua_straniera'):
            p = doc.add_paragraph()
            p.add_run(f"In particolare, preso atto che il Sig. {data.get('nome_straniero', '')} non conosce la lingua italiana ma dichiara di conoscere la lingua {data.get('lingua_conosciuta', 'inglese')}, il Presidente dichiara che provvederà a tradurre dall'italiano all'{data.get('lingua_conosciuta', 'inglese')} (e viceversa) gli interventi dei partecipanti alla discussione nonché a tradurre dall'italiano all'{data.get('lingua_conosciuta', 'inglese')} il verbale che sarà redatto al termine della riunione.")
        
        # Validità dell'assemblea
        p = doc.add_paragraph()
        p.add_run(f"Il Presidente constata e fa constatare che l'assemblea risulta {data.get('tipo_convocazione', 'regolarmente convocata')} e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.")
        
        p = doc.add_paragraph()
        p.add_run("Si passa quindi allo svolgimento dell'ordine del giorno.")

    def _add_discussion_section(self, doc, data):
        """Aggiunge la discussione dei punti all'OdG"""
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("*     *     *").bold = True
        
        data_chiusura = data.get('data_chiusura').strftime('%d/%m/%Y') if hasattr(data.get('data_chiusura'), 'strftime') else '[DATA]'
        
        p = doc.add_paragraph()
        p.add_run(f"In relazione al primo punto il presidente legge il bilancio al {data_chiusura} composto da stato patrimoniale, conto economico e nota integrativa (allegati di seguito al presente verbale).")

    def _add_ripianamento_section(self, doc, data):
        """Aggiunge la sezione del ripianamento perdite"""
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("*     *     *").bold = True
        
        p = doc.add_paragraph()
        text = "In relazione al secondo punto posto all'ordine del giorno, il Presidente"
        if data.get('has_collegio_sindacale'):
            text += ", sentito il parere favorevole del Collegio Sindacale,"
        text += f" propone all'assemblea di ripianare la perdita mediante la rinuncia al rimborso di una corrispondente quota del credito vantato dai soci nei confronti della società a titolo di {data.get('tipo_credito', 'finanziamento infruttifero soci')}."
        p.add_run(text)
        
        p = doc.add_paragraph()
        text = f"Segue breve discussione tra i soci al termine della quale si passa alla votazione con voto {'palese' if data.get('voto_palese') else 'segreto'} in forza della quale il Presidente constata che, {data.get('esito_votazione', 'all\'unanimità')} l'assemblea"
        p.add_run(text)
        
        p = doc.add_paragraph()
        p.add_run("DELIBERA").bold = True
        
        p = doc.add_paragraph()
        text = f"di approvare la proposta del Presidente e di ripianare, con effetto dalla data odierna, la perdita di euro {data.get('importo_perdita', '')} mediante la rinuncia al rimborso di una corrispondente quota del credito vantato dai soci nei confronti della società a titolo di {data.get('tipo_credito', 'finanziamento infruttifero soci')}."
        p.add_run(text)
        
        p = doc.add_paragraph()
        text = f"Le rinunce al rimborso avverranno in proporzione alla quota di {data.get('tipo_credito', 'finanziamento infruttifero soci')} attualmente in essere e quindi:"
        p.add_run(text)
        
        for rinuncia in data.get('rinunce_soci', []):
            if rinuncia.get('socio'):
                p = doc.add_paragraph()
                text = f"socio {rinuncia.get('socio', '')} per euro {rinuncia.get('importo_rinuncia', '')} pari al {rinuncia.get('percentuale', '')}% della perdita da ripianare;"
                p.add_run(text)
        
        p = doc.add_paragraph()
        text = f"Con effetto dalla data odierna, il residuo credito verso i soci a titolo di {data.get('tipo_credito', 'finanziamento infruttifero soci')} continuerà ad essere regolato dalle condizioni previgenti, ma solo per il residuo importo (al netto dell'avvenuta rinuncia) qui riepilogato:"
        p.add_run(text)
        
        for rinuncia in data.get('rinunce_soci', []):
            if rinuncia.get('socio'):
                p = doc.add_paragraph()
                text = f"socio {rinuncia.get('socio', '')} per euro {rinuncia.get('residuo', '')};"
                p.add_run(text)
        
        p = doc.add_paragraph()
        text = "Ciascun socio rilascia seduta stante al Presidente una dichiarazione sostitutiva di atto notorio che attesta il valore fiscale del credito a cui ha testé rinunciato."
        p.add_run(text)
        
        p = doc.add_paragraph()
        text = "Le dichiarazioni rilasciate dai soci verranno conservate agli atti della società."
        p.add_run(text)
        
        p = doc.add_paragraph()
        p.add_run("*     *     *").bold = True

    def _add_closing_section(self, doc, data):
        """Aggiunge la sezione di chiusura"""
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.")
        
        esito = data.get('esito_votazione', 'all\'unanimità')
        voto = 'palese' if data.get('voto_palese') else 'segreto'
        allegati = 'unitamente a quanto allegato' if data.get('documenti_allegati') else ''
        
        p = doc.add_paragraph()
        p.add_run(f"Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea {esito}, con voto {voto}, ne approva il testo {allegati}.")
        
        ora_fine = data.get('ora_fine', '[ORA]')
        p = doc.add_paragraph()
        p.add_run(f"L'assemblea viene sciolta alle ore {ora_fine}.")

    def _add_signatures(self, doc, data):
        """Aggiunge le firme con formattazione professionale"""
        # Aggiunge un'interruzione di pagina prima delle firme
        doc.add_page_break()
        
        # Aggiunge le firme con stile dedicato
        p = doc.add_paragraph(style='Firma')
        p.add_run("_____________________")
        
        p = doc.add_paragraph(style='Firma')
        p.add_run(data.get('presidente', '[PRESIDENTE]'))
        
        p = doc.add_paragraph(style='Firma')
        p.add_run("Il Presidente")
        
        doc.add_paragraph()
        
        p = doc.add_paragraph(style='Firma')
        p.add_run("_____________________")
        
        # Usa il segretario specificato o un placeholder se non disponibile
        segretario = data.get('segretario', '[SEGRETARIO]')
        
        p = doc.add_paragraph(style='Firma')
        p.add_run(segretario)
        
        p = doc.add_paragraph(style='Firma')
        p.add_run("Il Segretario")
        
        # Aggiunge la data di generazione
        doc.add_paragraph()
        p = doc.add_paragraph(style='BodyText')
        p.add_run(f"Generato il {date.today().strftime('%d/%m/%Y')}")

# Registra il template
DocumentTemplateFactory.register_template('verbale_assemblea_approvazione_bilancio', VerbaleAssembleaCompletoTemplate)
