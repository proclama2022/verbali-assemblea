import streamlit as st
import os
from dotenv import load_dotenv
from mistralai import Mistral
from docx import Document
from datetime import datetime, date
import PyPDF2
import json
import io
import pandas as pd
from packaging import version
import re

# Load environment variables
load_dotenv()
api_key = os.environ.get("MISTRAL_API_KEY")

# Inizializza il client Mistral
client = Mistral(api_key=api_key)

# Controllo versione Streamlit per data_editor avanzato
if version.parse(st.__version__) < version.parse("1.25.0"):
    st.error("La funzione 'st.data_editor' con 'num_rows' richiede Streamlit >= 1.25.0. Aggiorna Streamlit con 'pip install --upgrade streamlit'.")

def extract_visura_info(text):
    """Estrae le informazioni rilevanti dalla visura camerale."""
    info = {
        "denominazione": "",
        "sede_legale": "",
        "pec": "",
        "codice_fiscale": "",
        "forma_giuridica": "",
        "rappresentante": "",
        "capitale_sociale": "",
        "soci": [],
        "amministratori": [],
        "sindaci": [],
    }
    prompt = f"""Estrai le seguenti informazioni dalla visura camerale, rispondendo SOLO con un dizionario JSON:
    - denominazione (nome completo dell'azienda)
    - sede_legale (indirizzo completo)
    - pec (indirizzo PEC)
    - codice_fiscale
    - forma_giuridica
    - rappresentante (nome del rappresentante legale)
    - capitale_sociale (se presente)
    - soci: lista di oggetti con chiavi 'nome', 'quota_percentuale', 'quota_euro' (es: [{{"nome": "Mario Rossi", "quota_percentuale": "50%", "quota_euro": "5000"}}, ...])
    - amministratori: lista di oggetti con chiavi 'nome', 'carica' (es: [{{"nome": "Mario Rossi", "carica": "Amministratore Unico"}}, ...])
    - sindaci: lista di oggetti con chiavi 'nome', 'carica' (es: [{{"nome": "Luca Bianchi", "carica": "Presidente Collegio Sindacale"}}, ...])

    Testo della visura:
    {text}
    
    Rispondi SOLO con il dizionario JSON, senza altro testo."""
    messages = [
        {"role": "user", "content": prompt}
    ]
    chat_response = client.chat.complete(
        model="mistral-small-latest",
        messages=messages,
        temperature=0
    )
    try:
        response_text = chat_response.choices[0].message.content
        json_start = response_text.find('{')
        json_end = response_text.rfind('}')
        if json_start != -1 and json_end != -1:
            json_string = response_text[json_start : json_end + 1]
            extracted_info = json.loads(json_string)
            info.update(extracted_info)
        else:
            st.error("Impossibile trovare un blocco JSON valido nella risposta dell'API.")
            st.text(response_text)
    except json.JSONDecodeError as e:
        st.error(f"Errore nel decodificare la risposta JSON estratta: {e}")
        st.text(json_string)
    except Exception as e:
        st.error(f"Errore nell'estrazione delle informazioni (generico): {e}")
        st.text(response_text)
    return info

def compile_verbale(info):
    """Compila il verbale con le informazioni estratte."""
    doc = Document("templates/template.docx")
    
    # Sostituisci i campi nel template
    for paragraph in doc.paragraphs:
        text = paragraph.text
        text = text.replace("[DENOMINAZIONE]", info["denominazione"])
        text = text.replace("[SEDE]", info["sede_legale"])
        text = text.replace("[CAPITALE_SOCIALE]", info["capitale_sociale"])
        text = text.replace("[CODICE_FISCALE]", info["codice_fiscale"])
        text = text.replace("[DATA]", datetime.now().strftime("%d/%m/%Y"))
        paragraph.text = text
    
    # Salva il documento compilato
    output_path = "verbale_compilato.docx"
    doc.save(output_path)
    return output_path

# Funzione per estrarre testo con PyPDF2
def estrai_testo_pypdf2(pdf_bytes):
    testo = ""
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                testo += page_text + "\n"
    except Exception as e:
        testo = ""
    return testo

def genera_testo_verbale(
    denominazione,
    sede,
    capitale_sociale,
    codice_fiscale,
    data,
    ora,
    luogo,
    presidente,
    segretario,
    ordine_del_giorno,
    altri_presenti,
    tipo_convocazione,
    testo_libero=None
):
    if testo_libero:
        return testo_libero
    return f"""
{denominazione}
Sede in {sede}
Capitale sociale Euro {capitale_sociale} i.v.
Codice fiscale: {codice_fiscale}

Verbale di assemblea dei soci
del {data}

Oggi {data} alle ore {ora} presso la sede sociale {luogo}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:
Ordine del giorno
{ordine_del_giorno}
Assume la presidenza ai sensi dell'art. […] dello statuto sociale il Sig. {presidente}, il quale dichiara e constata:
1 - che l'intervento all'assemblea può avvenire anche in audioconferenza
2 - che sono presenti/partecipano all'assemblea: {altri_presenti}
I presenti all'unanimità chiamano a fungere da segretario il signor {segretario}, che accetta l'incarico.
Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante.
Il Presidente constata e fa constatare che l'assemblea risulta {tipo_convocazione} e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.\nSi passa quindi allo svolgimento dell'ordine del giorno.\n*     *     *\n"
In relazione al primo punto il presidente legge il bilancio al {data} composto da stato patrimoniale, conto economico e nota integrativa (allegati di seguito al presente verbale).
Segue breve discussione tra i soci al termine della quale si passa alla votazione con voto palese in forza della quale il Presidente constata che, all'unanimità, l'assemblea delibera l'approvazione del bilancio di esercizio chiuso al {data} e dei relativi documenti che lo compongono.
*     *     *
In relazione al secondo punto posto all'ordine del giorno, il Presidente propone all'assemblea di così destinare il risultato d'esercizio:
[……………………………………]
[……………………………………]
[……………………………………]
*     *     *
Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.
Viene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo.
L'assemblea viene sciolta alle ore {ora}.
"""

# Streamlit UI
st.title("Generatore Verbali da Visura Camerale")

# Spiegazione dell'app
st.write("""
Questa applicazione ti permette di:
1. Caricare una visura camerale
2. Estrarre automaticamente le informazioni rilevanti
3. Generare un verbale di assemblea pre-compilato
""")

# File uploader per la visura
uploaded_file = st.file_uploader("Carica la visura camerale (testo o PDF)", type=["txt", "pdf"])

if uploaded_file is not None:
    # Determina il tipo di file e leggi il contenuto
    file_type = uploaded_file.type
    visura_text = ""
    testo_pypdf2 = ""
    testo_ocr = ""
    scelta = None

    if file_type == "text/plain":
        visura_text = uploaded_file.getvalue().decode("utf-8")
    elif file_type == "application/pdf":
        pdf_bytes = uploaded_file.getvalue()
        # 1. Estrazione PyPDF2
        testo_pypdf2 = estrai_testo_pypdf2(pdf_bytes)
        # 2. Estrazione OCR Mistral
        st.info("Elaborazione del file PDF tramite OCR di Mistral...")
        try:
            uploaded_pdf = client.files.upload(
                file={
                    "file_name": uploaded_file.name,
                    "content": pdf_bytes
                },
                purpose="ocr"
            )
            ocr_response = client.ocr.process(
                model="mistral-ocr-latest",
                document={
                    "type": "document_url",
                    "document_url": client.files.get_signed_url(file_id=uploaded_pdf.id).url
                }
            )
            testo_ocr = "".join(page.markdown for page in ocr_response.pages)
            st.success("OCR completato.")
        except Exception as e:
            st.error(f"Errore nell'elaborazione del file PDF con Mistral OCR: {e}")
            testo_ocr = None
        # Mostra entrambi i testi
        if testo_pypdf2.strip():
            st.subheader("Testo estratto con PyPDF2 (PDF digitale)")
            st.text_area("Testo PyPDF2", testo_pypdf2, height=200)
        if testo_ocr and testo_ocr.strip():
            st.subheader("Testo estratto con Mistral OCR")
            st.text_area("Testo OCR Mistral", testo_ocr, height=200)
        # Scelta
        opzioni = []
        if testo_pypdf2.strip():
            opzioni.append("PyPDF2")
        if testo_ocr and testo_ocr.strip():
            opzioni.append("Mistral OCR")
        if opzioni:
            scelta = st.radio("Quale testo vuoi usare per l'estrazione AI?", opzioni)
            if scelta == "PyPDF2":
                visura_text = testo_pypdf2
            elif scelta == "Mistral OCR":
                visura_text = testo_ocr
        else:
            st.error("Nessun testo estratto dal PDF.")
            visura_text = None

    # Mostra il testo scelto (se presente)
    if visura_text is not None and visura_text.strip() != "":
        with st.expander("Visualizza testo della visura estratto (usato per AI)"):
            st.text(visura_text)
        
        # Initialize session state for info extraction
        if 'info_extracted' not in st.session_state:
            st.session_state['info_extracted'] = False
        if 'info' not in st.session_state:
            st.session_state['info'] = {}

        # Bottone per estrarre le informazioni
        if st.button("Estrai informazioni e genera verbale"):
            with st.spinner("Estrazione informazioni in corso..."):
                st.session_state.info = extract_visura_info(visura_text)
            st.session_state.info_extracted = True

        # Display fields after extraction
        if st.session_state.info_extracted:
            info = st.session_state.info
            st.subheader("Modifica e completa i dati del verbale (avanzato)")
            denominazione = st.text_input("Denominazione", info.get("denominazione", ""))
            sede = st.text_input("Sede legale", info.get("sede_legale", ""))
            capitale_sociale = st.text_input("Capitale sociale", info.get("capitale_sociale", ""))
            codice_fiscale = st.text_input("Codice fiscale", info.get("codice_fiscale", ""))
            data = st.date_input("Data assemblea")
            # Data di chiusura del bilancio, default 31/12 anno precedente
            default_chiusura = date(data.year - 1, 12, 31)
            data_chiusura = st.date_input("Data chiusura bilancio", default_chiusura)
            ora = st.text_input("Ora assemblea", "09:00")
            luogo = st.text_input("Luogo assemblea", sede)
            tipo_assemblea = st.selectbox("Tipo di assemblea", ["Ordinaria", "Straordinaria"]) 
            ruolo_presidente = st.selectbox("Ruolo del presidente", ["Amministratore Unico", "Presidente CdA", "Altro"]) 
            esito_votazione = st.selectbox("Esito votazione", ["Approvato all'unanimità", "Approvato a maggioranza", "Respinto", "Altro"])
            collegio_sindacale = st.checkbox("Collegio sindacale presente")
            revisore = st.checkbox("Revisore presente")
            if revisore:
                revisore_nome = st.text_input("Nome del revisore contabile", "")
            else:
                revisore_nome = ""
            # Tabelle dinamiche per soci, amministratori, sindaci
            df_soci = pd.DataFrame(info.get("soci", []))
            st.markdown("**Soci e quote**")
            df_soci = st.data_editor(df_soci, num_rows="dynamic")
            lista_soci = df_soci.to_dict("records")
            st.markdown("**Amministratori**")
            df_amm = pd.DataFrame(info.get("amministratori", []))
            df_amm = st.data_editor(df_amm, num_rows="dynamic")
            lista_amm = df_amm.to_dict("records")
            st.markdown("**Sindaci**")
            df_sindaci = pd.DataFrame(info.get("sindaci", []))
            df_sindaci = st.data_editor(df_sindaci, num_rows="dynamic")
            lista_sindaci = df_sindaci.to_dict("records")
            # Lista dinamica ordine del giorno
            st.markdown("**Punti all'ordine del giorno** (aggiungi uno per riga)")
            # Ordine del giorno con data bilancio preimpostata
            default_odg = f"Approvazione del Bilancio al {data_chiusura.strftime('%d/%m/%Y')} e dei documenti correlati;\nDelibere consequenziali."
            punti_odg = st.text_area("Ordine del giorno", default_odg)
            lista_odg = [p.strip() for p in punti_odg.split("\n") if p.strip()]
            altri_presenti = st.text_area("Altri presenti", "")
            # Selezione dinamica Presidente e Segretario
            if lista_amm:
                options_pres = [a.get('nome', '') for a in lista_amm]
                presidente = st.selectbox("Seleziona Presidente", options_pres)
            else:
                presidente = st.text_input("Nome presidente", info.get("rappresentante", ""))
            # Seleziona Segretario tra soci e amministratori
            options_sec = [item.get('nome', '') for item in lista_soci + lista_amm]
            if options_sec:
                segretario = st.selectbox("Seleziona Segretario (soci o amministratori)", options_sec)
            else:
                segretario = st.text_input("Segretario", "")
            # Opzioni per il secondo punto (parere dei Sindaci e destinazione del risultato d'esercizio)
            sentito_parere_sindaci = st.checkbox("Sentito il parere favorevole del Collegio Sindacale per il secondo punto", value=True)
            destinazione_risultato = st.text_area(
                "Destinazione del risultato d'esercizio / ripiano perdite (uno per riga)",
                "[Inserisci destinazione del risultato d'esercizio o ripiano perdite]\n[Inserisci eventuale seconda destinazione]\n[Inserisci eventuale terza destinazione]"
            )
            lista_destinazioni = [p.strip() for p in destinazione_risultato.split("\n") if p.strip()]
            # Selezione del tipo di convocazione (regolarmente convocata o totalitaria)
            tipo_convocazione = st.selectbox(
                "Tipo di convocazione",
                ["regolarmente convocata", "totalitaria"]
            )
            # Input per gli allegati
            allegati = st.text_area(
                "Elenco allegati (uno per riga)",
                ""
            )
            lista_allegati = [a.strip() for a in allegati.split("\n") if a.strip()]

            # Genera testo verbale dinamico secondo schema completo (senza duplicare l'ordine del giorno)
            testo_verbale = f"""
{denominazione}
Sede in {sede}
Capitale sociale Euro {capitale_sociale} i.v.
Codice fiscale: {codice_fiscale}

Verbale di assemblea dei soci
del {data.strftime('%d/%m/%Y')}

Oggi {data.strftime('%d/%m/%Y')} alle ore {ora} presso la sede sociale {luogo}, si è tenuta l'assemblea generale dei soci, per discutere e deliberare sul seguente:
Ordine del giorno:
"""
            # Aggiungo il corpo dinamico del verbale
            for idx, punto in enumerate(lista_odg, 1):
                testo_verbale += f"\n{idx}. {punto}"
            # Inserisco prima le dichiarazioni del presidente e poi elenco i partecipanti
            # Dichiarazioni del presidente
            testo_verbale += f"\n\nAssume la presidenza ai sensi dell'art. […] dello statuto sociale il Sig. {presidente} ({ruolo_presidente}), il quale dichiara e constata:\n"
            testo_verbale += "1 - che (come indicato anche nell'avviso di convocazione ed in conformità alle previsioni dell'art. […] dello statuto sociale) l'intervento all'assemblea può avvenire anche in audioconferenza\n"
            testo_verbale += "2 - che sono presenti/partecipano all'assemblea i soci e gli amministratori come di seguito indicati.\n"
            testo_verbale += f"I presenti all'unanimità chiamano a fungere da segretario il signor {segretario}, che accetta l'incarico.\n"
            # Elenco dei partecipanti (soci e amministratori)
            testo_verbale += "\n\nSoci presenti e quote:\n"
            for socio in lista_soci:
                testo_verbale += f"- {socio.get('nome','')} (Quota: {socio.get('quota_percentuale','')} - Euro {socio.get('quota_euro','')})\n"
            testo_verbale += "\nAmministratori:\n"
            for amm in lista_amm:
                testo_verbale += f"- {amm.get('nome','')} ({amm.get('carica','')})\n"
            testo_verbale += "Il Presidente identifica tutti i partecipanti e si accerta che ai soggetti collegati mediante mezzi di telecomunicazione sia consentito seguire la discussione, trasmettere e ricevere documenti, intervenire in tempo reale, con conferma da parte di ciascun partecipante.\n"
            testo_verbale += f"Il Presidente constata e fa constatare che l'assemblea risulta {tipo_convocazione} e deve ritenersi valida ed atta a deliberare sul citato ordine del giorno.\nSi passa quindi allo svolgimento dell'ordine del giorno.\n*     *     *\n"
            testo_verbale += f"In relazione al primo punto il presidente legge il bilancio al {data_chiusura.strftime('%d/%m/%Y')} composto da stato patrimoniale, conto economico e nota integrativa (allegati di seguito al presente verbale).\n"
            testo_verbale += f"Segue breve discussione tra i soci al termine della quale si passa alla votazione con voto palese in forza della quale il Presidente constata che, {esito_votazione.lower()}, l'assemblea delibera l'approvazione del bilancio di esercizio chiuso al {data_chiusura.strftime('%d/%m/%Y')} e dei relativi documenti che lo compongono.\n*     *     *\n"
            testo_verbale += f"In relazione al secondo punto posto all'ordine del giorno, il Presidente"
            if sentito_parere_sindaci:
                testo_verbale += ", <sentito il parere favorevole del Collegio Sindacale>"
            testo_verbale += ", propone all'assemblea di così destinare il risultato d'esercizio:\n"
            for dest in lista_destinazioni:
                testo_verbale += f"{dest}\n"
            testo_verbale += "*     *     *\n"
            testo_verbale += "Il Presidente constata che l'ordine del giorno è esaurito e che nessuno chiede la parola.\nViene quindi redatto il presente verbale e dopo averne data lettura, il Presidente constata che l'assemblea all'unanimità, con voto palese, ne approva il testo [eventualmente unitamente a quanto allegato].\n"
            testo_verbale += f"L'assemblea viene sciolta alle ore {ora}."
            # Aggiungo elenco allegati se presenti
            if lista_allegati:
                testo_verbale += "\n\nAllegati:\n"
                for allegato in lista_allegati:
                    testo_verbale += f"- {allegato}\n"
            # Inserisco le firme alla fine del verbale
            testo_verbale += "\nIl Presidente:\n__________________________\n\nIl Segretario:\n__________________________\n"
            testo_verbale = st.text_area("Testo verbale finale (modificabile)", testo_verbale, height=500)
            if st.button("Genera e scarica verbale Word"):
                doc = Document()
                # Generazione con formattazione: heading, liste numerate e bullet
                for par in testo_verbale.split("\n"):
                    stripped = par.strip()
                    # Numerazione punti ordine del giorno
                    if re.match(r"^\d+\.\s", stripped):
                        doc.add_paragraph(stripped, style="List Number")
                    # Bullet list (soci, amministratori, allegati)
                    elif stripped.startswith("- "):
                        doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
                    # Sezione con due punti
                    elif stripped.endswith(":"):
                        doc.add_paragraph(stripped, style="Heading 2")
                    else:
                        doc.add_paragraph(stripped)
                output_path = "verbale_compilato.docx"
                doc.save(output_path)
                with open(output_path, "rb") as file:
                    st.download_button(
                        label="Scarica verbale compilato",
                        data=file,
                        file_name="verbale_compilato.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

st.markdown("---")
st.markdown("Nota: questo è un prototipo. L'estrazione strutturata dei dati dalla visura verrà implementata nel prossimo passo.") 